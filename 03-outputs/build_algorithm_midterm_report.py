from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\moniC\project\learn")
TEMPLATE_PATH = Path(r"D:\不知道是啥\算法课期中报告模板.docx")
RESULT_PATH = Path(r"C:\Users\rainhotle\.codex\attachments\a504d31e-179d-4ef5-9fe5-c0f83a7f6c30\pasted-text.txt")
CODE_PATH = ROOT / "03-outputs" / "block-sequence-grand-master_solver_targeted_fix.cpp"
OUTPUT_DOCX = ROOT / "03-outputs" / "algorithm-midterm-report-block-sequence.docx"
OUTPUT_PNG = ROOT / "03-outputs" / "algorithm-midterm-report-score.png"


def set_run_font(run, name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def style_paragraph(paragraph, *, name: str = "宋体", size: float = 10.5,
                    bold: bool = False, after: float = 3.0,
                    first_line_chars: float | None = None) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(0)
    fmt.line_spacing = 1.15
    if first_line_chars is not None:
        fmt.first_line_indent = Pt(first_line_chars * size)
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=bold)


def clear_cell(cell) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")


def add_cell_paragraph(cell, text: str, *, name: str = "宋体",
                       size: float = 10.5, bold: bool = False,
                       after: float = 3.0) -> None:
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        paragraph = cell.paragraphs[0]
    else:
        paragraph = cell.add_paragraph()
    paragraph.text = text
    style_paragraph(paragraph, name=name, size=size, bold=bold, after=after)


def parse_result_summary() -> dict[str, object]:
    outer = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
    text = outer["result"]
    level_scores = [int(m.group(1)) for m in re.finditer(r"LEVEL_END\s+(\d+)", text)]
    final_scores = [int(m.group(1)) for m in re.finditer(r"FINAL_SCORE\s+(\d+)", text)]
    final_score = final_scores[-1] if final_scores else sum(level_scores)
    return {
        "level_scores": level_scores,
        "final_score": final_score,
        "match_id": outer.get("id", ""),
        "game_title": outer.get("game", {}).get("title", ""),
    }


def make_score_image(summary: dict[str, object]) -> None:
    lines = [
        "OJ Highest Verified Result",
        f"Match ID: {summary['match_id']}",
        f"Game: {summary['game_title']}",
        "",
        f"FINAL_SCORE = {summary['final_score']}",
        "",
    ]
    for idx, value in enumerate(summary["level_scores"], start=1):
        lines.append(f"LEVEL_{idx:02d} = {value}")
    lines.extend([
        "",
        "Code appendix in report: targeted_fix version",
        "Verified score source: latest exported OJ result",
    ])

    width, height = 1180, 720
    image = Image.new("RGB", (width, height), "#10151c")
    draw = ImageDraw.Draw(image)

    title_font = ImageFont.truetype(r"C:\Windows\Fonts\msyh.ttc", 34)
    body_font = ImageFont.truetype(r"C:\Windows\Fonts\consola.ttf", 26)
    small_font = ImageFont.truetype(r"C:\Windows\Fonts\consola.ttf", 22)

    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=24,
                           fill="#151d27", outline="#2f435a", width=3)
    y = 70
    for i, line in enumerate(lines):
        if i == 0:
            draw.text((70, y), line, fill="#e8f1ff", font=title_font)
            y += 64
            continue
        font = body_font if "FINAL_SCORE" in line or "LEVEL_" in line or "Match ID" in line else small_font
        color = "#9fd0ff" if "FINAL_SCORE" in line else "#d6deeb"
        if "LEVEL_" in line:
            color = "#cfe8b4"
        draw.text((70, y), line, fill=color, font=font)
        y += 38 if line else 22

    image.save(OUTPUT_PNG)


def fill_cover(doc: Document) -> None:
    lines = {
        1: "课程名称    算法分析与设计                    教师  待补充",
        2: "专业        待补充                          班级  待补充",
        3: "学生姓名    李林浩",
        4: "学生学号    202483290054",
    }
    for idx, text in lines.items():
        if idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        para.text = text
        style_paragraph(para, name="宋体", size=12, after=4)


def build_report_content(doc: Document, summary: dict[str, object]) -> None:
    table = doc.tables[0]

    clear_cell(table.cell(0, 1))
    data_lines = [
        "本策略同时使用了题目给出的静态棋盘信息和可预测的未来掉落信息。",
        "1. 当前棋盘 grid：读取每个格子的颜色、是否为炸弹、是否为 wildcard，用于构建连通块与可行路径。",
        "2. 关卡信息 level、棋盘边长 N、步数 step：决定不同关卡和不同阶段采用不同的搜索参数。",
        "3. 随机种子 seed 生成的 drop_queue 与 queue_ptr：在本地沙盘中精确预览执行某条路径后的新棋盘。",
        "4. 连通块统计信息：包括组件大小、节点度数、边缘深度、剩余出口数和可扩展分支数。",
        "5. 路径评价信息：当前路径得分、未来搜索得分、纯色局面结构分、wildcard 桥点消耗等。",
    ]
    for line in data_lines:
        add_cell_paragraph(table.cell(0, 1), line)

    clear_cell(table.cell(1, 1))
    strategy_lines = [
        "最终代码采用“分关卡 + 分阶段 + 多步预览”的策略框架。",
        "1. 共性框架：先按颜色与 wildcard 构造连通块，再在每个连通块内做 DFS 路径搜索，保留高质量候选，最后进行 3 层左右的未来评估。",
        "2. 第 1/2 关（纯色关）：目标不是只吃眼前高分，而是优先制造更大的后续连通块。针对较大的组件，起点更偏向中心和高连接度节点，避免过早修边角；评价函数同时考虑当前分数、链长偏好和局部未来潜力。",
        "3. 第 3 关（wildcard 关）：在保持一定贪心性的同时，增加 wildcard 桥点保护，尽量避免过早消耗连接度高的 wildcard，从而维持中后期大链能力。",
        "4. 第 4/5 关（炸弹关）：保留更稳健的未来权重，优先使用沙盘预览判断炸弹爆炸收益与后续死局风险，避免因局部贪心破坏整体布局。",
        "5. 针对性修复：本次最终代码把“结构潜力”前移到候选池筛选阶段，使真正有后续价值的铺垫招法能进入最终比较，而不是在早期就被短链高分候选挤掉。",
    ]
    for line in strategy_lines:
        add_cell_paragraph(table.cell(1, 1), line)

    clear_cell(table.cell(2, 1))
    iteration_lines = [
        "本项目不是一次成型，而是围绕 OJ 实测结果持续迭代。",
        "1. high_ceiling：总分 23952。特点是后两关较稳，但前两关偏保守。",
        "2. aggressive：总分 23632。前两关明显提升到 3194/3194，但第 4/5 关被过度激进的候选排序拖累。",
        "3. greedy_hybrid：总分 22864。第 3 关 wildcard 关提升到 4778，但炸弹关损失较大。",
        "4. specialized：总分 24312。开始按关卡拆分策略，后两关恢复稳定，第 1/2 关也超过早期版本。",
        "5. level_phased：总分 24676（已完整验证的最新成绩）。进一步把第 1/2/3 关按前中后期分段处理，第 1/2 关分别达到 3298 和 3222。",
        "6. targeted_fix（本报告附录代码）：基于手工复盘继续修复三个缺陷：纯色关中盘爱吃小块、结构分介入过晚、wildcard 桥点消耗过早。该版本作为下一轮冲分代码保存。",
    ]
    for line in iteration_lines:
        add_cell_paragraph(table.cell(2, 1), line)

    clear_cell(table.cell(3, 1))
    ai_lines = [
        "本项目使用了 AI 编程助手 Codex 进行分析与辅助实现。",
        "1. 读取并比较多次 OJ 原始导出结果，自动提取分关卡成绩和路径长度特征。",
        "2. 对不同版本策略进行代码对比，定位候选筛选、路径扩展顺序、关卡参数设置等差异。",
        "3. 根据手工复盘结论生成针对性修复版代码，并协助整理为实验报告初稿。",
        "4. 人工负责确定优化方向、判断哪些修改真正符合题目特点，并对最终文档内容进行确认。",
    ]
    for line in ai_lines:
        add_cell_paragraph(table.cell(3, 1), line)

    clear_cell(table.cell(4, 1))
    source_lines = [
        "完整源码见附录 A。",
        "附录源码文件名：block-sequence-grand-master_solver_targeted_fix.cpp",
        "说明：报告中的策略描述与附录代码一致；已完整验证的最高分来自同一策略框架下的上一轮稳定版本。"
    ]
    for line in source_lines:
        add_cell_paragraph(table.cell(4, 1), line)

    clear_cell(table.cell(5, 1))
    score_lines = [
        f"当前已完整验证的最高总分：{summary['final_score']}",
        "分关卡成绩：" + "，".join(
            f"第{i}关 {value}" for i, value in enumerate(summary["level_scores"], start=1)
        ),
        "下图根据 OJ 原始导出结果整理，用于报告归档。",
    ]
    for line in score_lines:
        add_cell_paragraph(table.cell(5, 1), line)
    table.cell(5, 1).add_paragraph().add_run().add_picture(str(OUTPUT_PNG), width=Inches(5.7))


def add_appendix(doc: Document) -> None:
    doc.add_page_break()

    heading = doc.add_paragraph()
    heading.add_run("附录 A  最终策略源码（targeted_fix 版本）")
    style_paragraph(heading, name="黑体", size=14, bold=True, after=6)

    note = doc.add_paragraph(
        "说明：以下代码为本次手工复盘后形成的针对性修复版。为便于查阅，附录保留完整源码，并在每行前增加行号。"
    )
    style_paragraph(note, name="宋体", size=10.5, after=8)

    code_text = CODE_PATH.read_text(encoding="utf-8").splitlines()
    for idx, line in enumerate(code_text, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(f"{idx:04d} | {line}")
        set_run_font(run, name="Consolas", size=8.2)


def main() -> None:
    summary = parse_result_summary()
    make_score_image(summary)

    doc = Document(TEMPLATE_PATH)
    fill_cover(doc)
    build_report_content(doc, summary)
    add_appendix(doc)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
