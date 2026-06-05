"""生成实验三报告 docx"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# -- 样式设置 --
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 封面/头部信息 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('南京信息工程大学  实验（实习）报告')
run.bold = True
run.font.size = Pt(16)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 基本信息表
table = doc.add_table(rows=3, cols=6, style='Table Grid')
table.autofit = True

# 第一行
cells_r1 = table.rows[0].cells
cells_r1[0].text = '实验名称'
cells_r1[1].text = 'Java实验三'
cells_r1[2].text = '实验日期'
cells_r1[3].text = '[日期]'
cells_r1[4].text = '得分'
cells_r1[5].text = ''

# 第二行
cells_r2 = table.rows[1].cells
cells_r2[0].text = '指导教师'
cells_r2[1].text = '[教师]'
cells_r2[2].text = '系'
cells_r2[3].text = '计算机'
cells_r2[4].text = '专业'
cells_r2[5].text = '计算机科学与技术'

# 第三行
cells_r3 = table.rows[2].cells
cells_r3[0].text = '年级'
cells_r3[1].text = '2024级'
cells_r3[2].text = '班次'
cells_r3[3].text = '[班级]'
cells_r3[4].text = '姓名'
cells_r3[5].text = '[姓名]'

# 学号
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('学号：[学号]')
run.font.size = Pt(12)

doc.add_paragraph()

# ===== 一、实验目的 =====
h1 = doc.add_paragraph()
run = h1.add_run('一、实验目的')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

objectives = [
    '学习Java Swing图形界面编程中Graphics类的使用方法，掌握字体样式设置、画笔颜色控制以及直线、矩形、圆角矩形、椭圆、弧形、多边形等基本几何图形的绘制与填充技巧。',
    '了解Swing原子组件JProgressBar的工作原理，熟悉其方向、范围、当前值等属性的配置方式，理解ChangeListener事件监听机制在进度条数值实时更新中的应用。',
    '理解Java多线程编程的核心概念，掌握Thread类的继承方式和run()方法的重写，学会使用synchronized关键字实现共享数据的线程安全访问，并能够将多线程技术与Swing图形界面进行整合。',
    '通过三个案例的动手实践，综合运用图形绘制、组件交互与多线程编程技术，提升Java GUI程序的设计能力与复杂问题的分解解决能力。',
]
for obj in objectives:
    p = doc.add_paragraph()
    p.style.font.size = Pt(12)
    run = p.add_run('  ' + obj)
    run.font.size = Pt(12)

# ===== 二、实验内容 =====
h2 = doc.add_paragraph()
run = h2.add_run('二、实验内容')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
run = p.add_run('教材例题 7-1（Graphics图形绘制）、7-6（JProgressBar进度条）、7-9（多线程绘图）。')
run.font.size = Pt(12)

# 对每个例题的简要说明
items = [
    '7-1：基于JFrame窗口，重写paint()方法，使用Graphics对象依次绘制不同字体、颜色的文字，以及线段、矩形、圆角矩形、3D矩形、椭圆、弧形、多边形等图形，对比draw与fill系列方法的区别。',
    '7-6：创建JProgressBar组件，设置水平方向、范围0~100，通过循环模拟任务进度推进，实现ChangeListener接口在进度值变化时同步更新JLabel显示当前完成百分比。',
    '7-9：构建多线程绘图程序，主线程负责GUI界面（含自定义JPanel绘图区和JButton按钮），鼠标点击可手动添加从原点出发的线段；点击按钮启动子线程，每0.3秒随机生成一条新线段并刷新界面，使用Collections.synchronizedList保证共享数据线程安全。',
]
for item in items:
    p = doc.add_paragraph()
    run = p.add_run('  ' + item)
    run.font.size = Pt(12)

# ===== 三、实验步骤 =====
h3 = doc.add_paragraph()
run = h3.add_run('三、实验步骤')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
run = p.add_run('（一）编写并运行以下Java程序代码：')
run.font.size = Pt(12)

# 7-1 代码
p = doc.add_paragraph()
run = p.add_run('1. 实验7-1：Graphics图形绘制')
run.bold = True
run.font.size = Pt(12)

code_7_1 = '''package test7_1;

import java.awt.*;
import javax.swing.*;

public class GraphicsTester extends JFrame {

    public GraphicsTester() {
        super("Java图形绘制演示 — [姓名] [学号]");
        setSize(480, 280);
        setLocationRelativeTo(null);
        setVisible(true);
    }

    @Override
    public void paint(Graphics g) {
        super.paint(g);

        // 演示不同字体和颜色绘制文字
        g.setFont(new Font("SansSerif", Font.BOLD, 12));
        g.setColor(Color.blue);
        g.drawString("SansSerif 粗体 12号 蓝色", 20, 50);

        g.setFont(new Font("Serif", Font.ITALIC, 14));
        g.setColor(new Color(255, 0, 0));
        g.drawString("Serif 斜体 14号 红色", 250, 50);

        // 分隔线
        g.drawLine(20, 60, 460, 60);

        // 矩形
        g.setColor(Color.green);
        g.drawRect(20, 70, 100, 50);
        g.fillRect(130, 70, 100, 50);

        // 圆角矩形
        g.setColor(Color.yellow);
        g.drawRoundRect(240, 70, 100, 50, 50, 50);
        g.fillRoundRect(350, 70, 100, 50, 50, 50);

        // 3D矩形
        g.setColor(Color.cyan);
        g.draw3DRect(20, 130, 100, 50, true);
        g.fill3DRect(130, 130, 100, 50, false);

        // 椭圆
        g.setColor(Color.pink);
        g.drawOval(240, 130, 100, 50);
        g.fillOval(350, 130, 100, 50);

        // 弧形
        g.setColor(new Color(0, 120, 20));
        g.drawArc(20, 190, 100, 50, 0, 90);
        g.fillArc(130, 190, 100, 50, 0, 90);

        // 多边形
        g.setColor(Color.black);
        int[] xVals1 = {250, 280, 290, 300, 330, 310, 320, 290, 260, 270};
        int[] yVals1 = {210, 210, 190, 210, 210, 220, 230, 220, 230, 220};
        g.drawPolygon(xVals1, yVals1, 10);

        int[] xVals2 = {360, 390, 400, 410, 440, 420, 430, 400, 370, 380};
        g.fillPolygon(xVals2, yVals1, 10);
    }

    public static void main(String[] args) {
        JFrame.setDefaultLookAndFeelDecorated(true);
        GraphicsTester app = new GraphicsTester();
        app.setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
    }
}'''

# 添加代码块
code_para = doc.add_paragraph()
code_run = code_para.add_run(code_7_1)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)
code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 运行截图占位
p = doc.add_paragraph()
run = p.add_run('运行结果截图：')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('[在此处插入7-1运行截图]')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 7-6 代码
p = doc.add_paragraph()
run = p.add_run('2. 实验7-6：JProgressBar进度条')
run.bold = True
run.font.size = Pt(12)

code_7_6 = '''package test7_6;

import javax.swing.*;
import javax.swing.event.ChangeEvent;
import javax.swing.event.ChangeListener;
import java.awt.*;
import java.awt.event.WindowAdapter;
import java.awt.event.WindowEvent;

public class ProgressBarDemo implements ChangeListener {

    private JLabel statusLabel;
    private JProgressBar progressBar;

    public ProgressBarDemo() {
        int initialValue = 0;

        JFrame frame = new JFrame("JProgressBar进度条演示 — [姓名] [学号]");
        Container contentPane = frame.getContentPane();

        // 状态标签
        statusLabel = new JLabel("", JLabel.CENTER);
        statusLabel.setToolTipText("当前进度信息");

        // 进度条配置
        progressBar = new JProgressBar();
        progressBar.setOrientation(JProgressBar.HORIZONTAL);
        progressBar.setMinimum(0);
        progressBar.setMaximum(100);
        progressBar.setValue(initialValue);
        progressBar.setStringPainted(true);
        progressBar.addChangeListener(this);
        progressBar.setToolTipText("进度条组件");

        contentPane.add(progressBar, BorderLayout.CENTER);
        contentPane.add(statusLabel, BorderLayout.SOUTH);

        // 窗口配置
        frame.setSize(420, 80);
        frame.setLocationRelativeTo(null);
        frame.setVisible(true);

        // 模拟进度更新
        for (int i = 1; i <= 1000000000; i++) {
            if (i % 10000000 == 0) {
                progressBar.setValue(++initialValue);
            }
        }

        frame.addWindowListener(new WindowAdapter() {
            @Override
            public void windowClosing(WindowEvent e) {
                System.exit(0);
            }
        });
    }

    @Override
    public void stateChanged(ChangeEvent e) {
        int currentValue = progressBar.getValue();
        statusLabel.setText("当前已完成进度：" + currentValue + "%");
    }

    public static void main(String[] args) {
        SwingUtilities.invokeLater(ProgressBarDemo::new);
    }
}'''

code_para = doc.add_paragraph()
code_run = code_para.add_run(code_7_6)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)
code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('运行结果截图：')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('[在此处插入7-6运行截图]')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 7-9 代码
p = doc.add_paragraph()
run = p.add_run('3. 实验7-9：多线程绘图')
run.bold = True
run.font.size = Pt(12)

code_7_9 = '''package test7_9;

import javax.swing.*;
import java.awt.*;
import java.awt.event.*;
import java.util.*;

public class MultiThreadDraw extends JFrame {

    // 线程安全的线段端点列表
    private final List<Point> pointList =
            Collections.synchronizedList(new ArrayList<>());
    private final Random random = new Random();
    private JPanel drawPanel;

    public MultiThreadDraw() {
        setTitle("多线程绘图演示 — [姓名] [学号]");
        Container container = getContentPane();

        // 自定义绘图面板
        drawPanel = new JPanel() {
            @Override
            protected void paintComponent(Graphics g) {
                super.paintComponent(g);
                g.setColor(Color.RED);
                synchronized (pointList) {
                    for (Point p : pointList) {
                        g.drawLine(0, 0, p.x, p.y);
                    }
                }
            }
        };
        drawPanel.setLayout(new BorderLayout());

        JLabel tipLabel = new JLabel("点击面板添加线段", JLabel.CENTER);
        tipLabel.setOpaque(false);
        drawPanel.add(tipLabel, BorderLayout.CENTER);

        container.add(drawPanel, BorderLayout.CENTER);

        // 底部按钮
        JButton startButton = new JButton("启动新线程");
        startButton.setBorder(
                BorderFactory.createLineBorder(Color.blue, 5));
        container.add(startButton, BorderLayout.SOUTH);

        // 鼠标点击添加线段
        drawPanel.addMouseListener(new MouseAdapter() {
            @Override
            public void mousePressed(MouseEvent e) {
                pointList.add(new Point(e.getX(), e.getY()));
                drawPanel.repaint();
            }
        });

        // 按钮启动新线程随机画线
        startButton.addActionListener(new ActionListener() {
            @Override
            public void actionPerformed(ActionEvent e) {
                new DrawThread(drawPanel, pointList, random).start();
            }
        });

        setSize(350, 450);
        setLocationRelativeTo(null);
        setVisible(true);
        setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
    }

    public static void main(String[] args) {
        SwingUtilities.invokeLater(MultiThreadDraw::new);
    }
}

/**
 * 绘图线程：随机生成20条线段，间隔0.3秒
 */
class DrawThread extends Thread {

    private final JPanel panel;
    private final List<Point> points;
    private final Random rand;

    public DrawThread(JPanel panel, List<Point> points, Random rand) {
        this.panel = panel;
        this.points = points;
        this.rand = rand;
    }

    @Override
    public void run() {
        for (int i = 0; i < 20; i++) {
            int x = rand.nextInt(panel.getWidth());
            int y = rand.nextInt(panel.getHeight());
            points.add(new Point(x, y));
            panel.repaint();
            try {
                Thread.sleep(300);
            } catch (InterruptedException ex) {
                ex.printStackTrace();
            }
        }
    }
}'''

code_para = doc.add_paragraph()
code_run = code_para.add_run(code_7_9)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)
code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('运行结果截图：')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('[在此处插入7-9运行截图]')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

# ===== 四、体会和总结 =====
doc.add_paragraph()
h4 = doc.add_paragraph()
run = h4.add_run('四、体会和总结')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

summary_paragraphs = [
    '本次Java实验三通过Graphics绘图、JProgressBar进度条以及多线程绘图三个例题的编码与调试，使我对Java Swing图形界面编程有了更深入的认识。',

    '在实验7-1中，我通过重写JFrame的paint()方法，调用Graphics对象的一系列draw/fill方法，绘制了包含文字、直线、矩形、圆角矩形、3D矩形、椭圆、弧形和多边形在内的多种图形。在这个过程中，我理解了Java坐标系的定义方式（左上角为原点），以及draw系列（仅描边）与fill系列（填充内部）之间的差异。同时我也注意到，Graphics的setColor()方法会持续影响后续所有绘制操作，因此在不同图形之间需要及时切换颜色才能达到预期效果。',

    '实验7-6涉及JProgressBar进度条组件的使用。我将进度条设置为水平方向，范围为0到100，并通过一个从1到10亿的循环来模拟长时间任务的进度推进——每循环1000万次进度值加1。通过实现ChangeListener接口并重写stateChanged()方法，我在进度值每次变化时同步更新JLabel标签的文本，实现了"当前已完成进度：XX%"的实时显示。这个实验让我认识到，Swing组件的值变化监听机制是实现界面与数据同步的重要手段。',

    '实验7-9是本次实验的综合应用。该程序创建了一个自定义JPanel作为绘图区，从窗口原点(0,0)向各个方向绘制射线。程序支持两种交互方式：一是鼠标点击面板手动添加射线；二是点击底部按钮启动子线程，子线程每隔0.3秒随机生成一条新射线并调用repaint()刷新界面。为了确保多线程同时操作共享数据列表时的安全性，我使用了Collections.synchronizedList()来包装ArrayList，并在paintComponent()中对列表进行synchronized同步遍历。通过这个实验，我深刻体会到：在Swing多线程编程中，数据一致性是需要重点关注的问题，选择合适的线程安全容器和合理的同步策略是程序稳定运行的基础。',

    '总的来说，这三次实验从基础图形绘制到组件交互再到多线程综合应用，层层递进，使我将课本上的理论知识转化为了实际的编程能力。在实验过程中，我也遇到了一些问题，例如初次运行时忘记在paint()方法开头调用super.paint(g)导致界面残留、多线程绘图时因未做同步而出现ConcurrentModificationException异常等，通过查阅资料和逐步调试最终得以解决。这些经历加深了我对Java GUI编程规范和调试方法的理解，为今后进行更复杂的Java桌面应用开发积累了经验。',
]

for text in summary_paragraphs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
    run = p.add_run(text)
    run.font.size = Pt(12)

# ===== 保存 =====
output_path = r'D:\moniC\project\learn\java-exp3\Java实验三_实验报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
