"""
Extract paragraphs and tables from a docx file, output structured markdown.

Usage:
  python extract_docx.py <docx_path> [output_md_path]

If output_md_path is omitted, writes to <docx_name>_content.md in the same directory.
"""
import sys
import os
from docx import Document


def extract_docx(docx_path: str) -> str:
    doc = Document(docx_path)

    lines = []
    # Title from first meaningful paragraph
    if doc.paragraphs:
        title = doc.paragraphs[0].text.strip()
        lines.append(f"# {title}")
        lines.append("")

    lines.append("## 头部信息")
    lines.append("- 实验名称：（待填写）")
    lines.append("- 学院/专业/年级/班次/姓名/学号：（待填写）")
    lines.append("")

    current_section = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        style = p.style.name if p.style else ""

        # Detect sections
        if "实验目的" in text:
            current_section = "purpose"
            lines.append("## 一、实验目的")
            continue
        elif "实验内容" in text or "实验步骤" in text:
            current_section = "steps"
            lines.append("## 二、实验内容与步骤")
            lines.append("")
            continue
        elif "实验心得" in text or "实验总结" in text:
            current_section = "summary"
            lines.append("## 三、实验心得")
            lines.append("（待填写）")
            lines.append("")
            continue

        # Skip header lines (university name, experiment name, student info table)
        if style == "Normal" and any(kw in text for kw in ["大学", "实验名称", "学院", "指导教师"]):
            continue

        # Detect numbered questions
        if current_section == "steps":
            # Match patterns like （1）, (1), 1、, 1.
            import re
            qmatch = re.match(r'[（(]\s*(\d+)\s*[）)]', text)
            if qmatch:
                qnum = qmatch.group(1)
                lines.append(f"### 题目 {qnum}：{text}")
                lines.append("")
                continue

            # Also match things like "1、"
            qmatch2 = re.match(r'^(\d+)[、.]', text)
            if qmatch2 and int(qmatch2.group(1)) <= 20:
                lines.append(f"### 题目 {qmatch2.group(1)}：{text}")
                lines.append("")
                continue

        # Purpose items
        if current_section == "purpose":
            # Match patterns like 1、, 2、
            if text and text[0].isdigit():
                lines.append(f"{text}")
                continue

    # Tables
    if doc.tables:
        lines.append("---")
        lines.append("")
        lines.append("## 附录：表格信息")
        lines.append("")
        for i, table in enumerate(doc.tables):
            lines.append(f"### 表格 {i + 1} ({len(table.rows)}行 × {len(table.columns)}列)")
            lines.append("")
            for r, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if r == 0:
                    lines.append("|" + "|".join(["---"] * len(cells)) + "|")
            lines.append("")

    return "\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_path> [output_md_path]")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"Error: file not found: {docx_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        base = os.path.splitext(os.path.basename(docx_path))[0]
        output_path = os.path.join(os.path.dirname(docx_path) or ".", f"{base}_content.md")

    md_content = extract_docx(docx_path)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"Extracted to: {output_path}")
    print(f"  Characters: {len(md_content)}")


if __name__ == "__main__":
    main()
