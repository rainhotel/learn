"""
Generate a filled experiment report docx by inserting SQL and results.

Usage:
  python generate_docx.py <template_docx> <queries_json> [output_path]

queries_json format:
{
  "questions": [
    {
      "num": 1,
      "keyword": "unique text in the question paragraph",
      "has_placeholders": true,
      "sql": "SELECT ...",
      "result": "formatted result text"
    },
    {
      "num": 2,
      "keyword": "unique text in the question paragraph",
      "sql": "SELECT ...",
      "result": "formatted result text"
    },
    {
      "num": 6,
      "keyword": "unique text in the question paragraph",
      "methods": [
        {"label": "方法一：连接查询", "sql": "...", "result": "..."},
        {"label": "方法二：EXISTS嵌套查询", "sql": "...", "result": "..."}
      ]
    }
  ],
  "summary": "实验心得内容..."
}
"""
import sys
import os
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def make_paragraph(text, font_name="宋体", font_size=10.5):
    """Create a w:p XML element with text."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(ref_elem, paragraphs):
    """Insert multiple paragraphs after ref_elem. Returns last inserted."""
    prev = ref_elem
    for text in paragraphs:
        p = make_paragraph(text)
        prev.addnext(p)
        prev = p
    return prev


def find_paragraph(doc, keyword):
    """Find the first paragraph containing keyword."""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i
    return -1


def build_lines(question):
    """Build the list of paragraph texts to insert for a question."""
    lines = []

    # For Q1 with placeholder paragraphs, don't add extra labels
    if question.get("has_placeholders"):
        # Only insert SQL and result content, no labels
        return question.get("sql", "").split("\n"), question.get("result", "").split("\n")

    lines.append("涉及的各基本表原始数据（截图）：[此处插入截图]")
    lines.append("")

    if "methods" in question:
        for i, method in enumerate(question["methods"]):
            if i > 0:
                lines.append("")
            lines.append(f"【{method['label']}】")
            lines.append("SQL查询语句：")
            lines.extend(method["sql"].split("\n"))
            lines.append("")
            lines.append("查询结果（截图）：[此处插入截图]")
            lines.extend(method["result"].split("\n"))
    else:
        lines.append("SQL查询语句：")
        lines.extend(question["sql"].split("\n"))
        lines.append("")
        lines.append("查询结果（截图）：[此处插入截图]")
        lines.extend(question["result"].split("\n"))

    return lines


def generate(template_path, queries_data, output_path):
    doc = Document(template_path)

    questions = queries_data["questions"]

    # Process Q1 first (has special placeholders)
    q1 = next((q for q in questions if q["num"] == 1), None)
    if q1 and q1.get("has_placeholders"):
        # Insert result text after result screenshot placeholder first (to avoid index shift)
        result_idx = find_paragraph(doc, "执行SQL语句后查询结果")
        if result_idx >= 0:
            rlines = q1.get("result", "").split("\n")
            insert_after(doc.paragraphs[result_idx]._p, rlines)

        # Then insert SQL after SQL placeholder
        sql_idx = find_paragraph(doc, "SQL查询语句（文本）")
        if sql_idx >= 0:
            slines = q1.get("sql", "").split("\n")
            insert_after(doc.paragraphs[sql_idx]._p, slines)

    # Process remaining questions from bottom to top (preserve indices)
    other_qs = sorted(
        [q for q in questions if q["num"] != 1],
        key=lambda q: q["num"],
        reverse=True,
    )
    for q in other_qs:
        kw = q.get("keyword", "")
        if not kw:
            # Try to find by question number pattern
            kw = f"（{q['num']}）"

        q_idx = find_paragraph(doc, kw)
        if q_idx < 0:
            print(f"Warning: Q{q['num']} not found with keyword '{kw}'")
            continue

        lines = build_lines(q)
        insert_after(doc.paragraphs[q_idx]._p, lines)
        print(f"Q{q['num']}: inserted {len(lines)} lines")

    # Add summary (实验心得)
    summary_idx = find_paragraph(doc, "实验心得")
    if summary_idx < 0:
        summary_idx = find_paragraph(doc, "三、实验心得")

    if summary_idx >= 0 and "summary" in queries_data:
        summary_text = queries_data["summary"]
        insert_after(doc.paragraphs[summary_idx]._p, summary_text.split("\n"))
        print("Summary: inserted after 实验心得")

    doc.save(output_path)
    print(f"Saved to: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <template_docx> <queries_json> [output_path]")
        sys.exit(1)

    template_path = sys.argv[1]
    queries_path = sys.argv[2]

    if not os.path.exists(template_path):
        print(f"Error: template not found: {template_path}")
        sys.exit(1)
    if not os.path.exists(queries_path):
        print(f"Error: queries JSON not found: {queries_path}")
        sys.exit(1)

    with open(queries_path, "r", encoding="utf-8") as f:
        queries_data = json.load(f)

    if len(sys.argv) >= 4:
        output_path = sys.argv[3]
    else:
        base = os.path.splitext(os.path.basename(template_path))[0]
        output_path = os.path.join(
            os.path.dirname(template_path) or ".", f"{base}_已完成.docx"
        )

    generate(template_path, queries_data, output_path)


if __name__ == "__main__":
    main()
