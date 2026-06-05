# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import copy

SRC = r'D:\moniC\202483290054李林浩-实验2.docx'
DST = r'D:\moniC\202483290054李林浩-实验2.docx'

doc = Document(SRC)

# Identify paragraphs that contain images (drawings)
image_paras = set()
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawings:
            image_paras.add(i)
            break

print(f'Image paragraphs: {sorted(image_paras)}')
print(f'Total paragraphs before edit: {len(doc.paragraphs)}')

def set_para_text(para, text):
    """Replace all text in a paragraph with new text."""
    if not para.runs:
        return
    for run in para.runs:
        run.text = ''
    para.runs[0].text = text

def insert_paragraph_after(paragraph, text):
    """Insert a new paragraph after the given paragraph."""
    new_p = docx.oxml.OxmlElement("w:p")
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        new_pPr = copy.deepcopy(pPr)
        new_p.append(new_pPr)
    r = docx.oxml.OxmlElement("w:r")
    rPr = docx.oxml.OxmlElement("w:rPr")
    r.append(rPr)
    t = docx.oxml.OxmlElement("w:t")
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph.part)

# ============================================================
# STEP 1: Modify text of existing paragraphs (in-place, original indices)
# ============================================================

modifications = {
    10: '1、在GoodsOrder数据库中，利用SQL语句实现下面各类单表查询操作。本实验涵盖基础查询、条件筛选、模糊匹配、聚合统计、分组过滤、排序输出、集合运算等多种查询类型：',
    12: '（1）查询所有商品的详细信息。本查询使用SELECT * 返回GoodsInfo表中所有列、所有行的数据，是最基础的全表查询操作，用于查看商品表的完整内容。',
    18: '（2）查询客户订单中的所有付款方式，并且付款方式不重复显示。本查询使用DISTINCT关键字消除结果集中的重复行，从OrderList表中提取所有不重复的付款方式，便于了解平台支持的支付渠道种类。',
    24: '（3）查询各种球类商品的编号、名称和价格等信息。本查询使用WHERE子句对商品类别进行条件过滤，从GoodsInfo表中筛选出商品类别为体育用品的所有记录，展示球类商品的基本属性和定价。',
    30: '（4）查询食品类的商品编号、商品名称、库存量和到期天数（即距离保质期截止还有多少天）。本查询综合运用列筛选、WHERE条件过滤与日期函数DATEDIFF，通过计算当前日期与保质期之间的天数差，实时反映食品库存的临期情况，对食品类商品的库存管理具有实际参考价值。',
    49: '（5）查询所有女客户以及本月过生日的男客户的详细信息。本查询使用OR组合多个筛选条件：性别为女的客户无条件返回，性别为男的客户则需进一步通过MONTH函数判断其出生月份是否与当前月份相同，实现了基于性别和生日的复合条件过滤。',
    66: '（6）查询上海、江苏南京、河南郑州三地的所有客户信息，并按客户所在省市降序排列，同省市的按编号升序排列。本查询使用IN谓词实现确定集合的多值匹配，配合ORDER BY子句实现双重排序规则（主排序键为所在省市降序，次排序键为客户编号升序），展示了多条件排序的典型应用场景。',
    82: '（7）现针对备注不为空的VIP客户推出促销活动：金牌客户享受8折优惠，银牌客户享受9折优惠，新客户享受9.5折优惠，其他情况无折扣。请查询这些客户的编号、姓名和对应的折扣率。本查询使用CASE WHEN表达式实现多分支条件判断，将备注字段中的客户等级映射为不同的折扣率，同时通过WHERE条件限定VIP标识为真且备注非空，是CASE表达式在业务计算中的典型应用。',
    102: '（8）查询姓张且不为单名的所有客户的编号、姓名、性别和出生日期，并按年纪从小到大排序（即出生日期降序排列）。本查询使用LIKE运算符进行模糊匹配（张%表示以张开头），结合LEN函数排除单名客户（姓名长度≥3，含姓氏），最终按出生日期降序排列，使年龄最小的客户排在前面。',
    120: '（9）统计体育用品类商品的最低单价和最高单价。本查询使用聚合函数MIN和MAX对指定类别商品的单价进行极值统计，快速获取体育用品类商品的价格区间，为定价分析和市场定位提供数据支撑。',
    133: '（10）统计各类商品的最低库存数，统计结果所在列的标题显示为最低库存量。本查询使用GROUP BY子句按商品类别进行分组，结合MIN聚合函数计算每组的最低库存量，并通过AS关键字为结果列指定别名。查询排除了食品类别，聚焦于非食品类商品的库存底数分析。',
    148: '（11）查询非食品类各商品的订购总数量在前3名的商品编号和对应的订购总数量。本查询综合运用了JOIN联表、WHERE过滤、GROUP BY分组、SUM聚合、ORDER BY排序和TOP子句限制返回行数：先通过JOIN关联OrderList与GoodsInfo获取商品的类别信息，排除食品类后按商品编号分组求和，最后用TOP 3取订购总量最高的前三名商品。',
    167: '（12）查询在2020年2月下旬（即2月15日及之后）销售的各商品编号和对应的客户数，并按客户数降序排列、商品编号升序排列。本查询使用YEAR、MONTH、DAY三个日期提取函数配合逻辑表达式构造指定范围的谓词条件，精确锁定2020年2月15日至月底的销售记录，再按商品编号分组统计去重客户数。',
    183: '（13）查询至少购买了两种商品的客户编号。本查询在GROUP BY分组的基础上使用HAVING子句对分组后的结果进行条件过滤（COUNT(商品编号) >= 2）。HAVING与WHERE的区别在于：WHERE在分组前过滤行，HAVING在分组后过滤组，本查询正是HAVING子句的典型应用场景。',
    193: '（14）在送货上门的订单中，按付款方式统计订单总数量大于等于3的付款方式及对应的订单总数，并按订单总数升序排序。本查询将WHERE行级过滤（送货方式=送货上门）、GROUP BY分组统计、HAVING组级过滤（COUNT >= 3）和ORDER BY排序组合在一起，完整展示了SQL逻辑子句的执行顺序。',
    207: '（15）利用集合查询找出没有订购商品的客户编号。本查询使用EXCEPT集合运算符，返回存在于CustomerInfo客户表中但不存在于OrderList订单表中的客户编号，即从未下过订单的客户。EXCEPT运算符自动去重，返回第一个结果集中独有的行，是实现差集查询的标准SQL语法。',
    217: '通过本次数据库单表查询实验，我系统性地掌握了SELECT语句的完整语法体系和各类查询技术。在基础查询方面，熟练运用了SELECT * 全表查询、DISTINCT去重查询以及WHERE条件筛选。在高级查询方面，掌握了LIKE模糊匹配、IN集合谓词、BETWEEN范围谓词以及复合逻辑表达式（AND/OR）的灵活组合。在聚合分析方面，能够运用COUNT、SUM、MIN、MAX、AVG等聚合函数配合GROUP BY和HAVING子句进行数据分组统计与过滤，深刻理解了WHERE与HAVING的执行时机差异。此外，还学习了日期函数（GETDATE、DATEDIFF、MONTH、YEAR等）的实际运用、CASE WHEN条件表达式的业务逻辑映射、ORDER BY多字段排序、TOP子句限制返回行数以及EXCEPT集合差集运算等高级技巧。在实操过程中，我加深了对SQL语句逻辑执行顺序（FROM→WHERE→GROUP BY→HAVING→SELECT→ORDER BY→TOP）的理解，认识到规范书写SQL、精确设定查询条件以及合理使用分组过滤对于获取准确查询结果的重要性。本次实验为后续多表连接查询、子查询以及复杂数据库操作打下了坚实的基础。',
}

for idx, text in modifications.items():
    set_para_text(doc.paragraphs[idx], text)
    print(f'Modified paragraph [{idx}]')

# ============================================================
# STEP 2: Insert supplementary notes in REVERSE order
# (largest index first, so earlier insertions don't shift later ones)
# ============================================================

# These are ORIGINAL paragraph indices (before any insertions)
# Sorted largest-first for safe insertion
supplementary_notes = [
    (215, '【补充说明】EXCEPT、INTERSECT和UNION是SQL中的三种集合运算符，分别对应差集、交集和并集运算。使用时需注意两个查询的列数和数据类型必须匹配。'),
    (206, '【补充说明】SQL逻辑执行顺序与书写顺序不同：先FROM确定数据来源，再WHERE过滤行，然后GROUP BY分组，接着HAVING过滤组，之后才是SELECT选取列和ORDER BY排序。'),
    (192, '【补充说明】HAVING子句中的条件只能使用SELECT列表中出现的列或聚合函数，不能引用未分组的原始列，这是初学者常见的错误点。'),
    (182, '【补充说明】在进行日期范围查询时，需要注意时间部分的边界问题。如果日期列包含时间信息（如datetime类型），WHERE条件需同时考虑时分秒，否则可能遗漏或多余包含边界日期的记录。'),
    (166, '【补充说明】TOP子句是SQL Server特有的语法（MySQL使用LIMIT，Oracle使用ROWNUM），配合ORDER BY使用时，确保返回排序后的前N条记录。'),
    (147, '【补充说明】使用AS为聚合结果列指定别名是一个良好的编程习惯，它使查询结果的可读性大大提高，尤其在多表连接和嵌套查询中尤为重要。'),
    (132, '【补充说明】MIN和MAX函数不仅适用于数值类型，也可用于日期、字符串等可比较的数据类型：日期取最早/最晚，字符串取字母序最小/最大。'),
    (119, '【补充说明】LEN函数返回字符串的字符数（不包括尾部空格），而DATALENGTH返回字节数。对于中文Unicode字符，一个汉字占2个字节，但LEN计为1个字符。'),
    (101, '【补充说明】CASE WHEN表达式是SQL中的条件判断结构，类似于编程语言中的if-else语句，可用于SELECT列表、WHERE条件、ORDER BY等多个子句中。'),
    (81, '【补充说明】IN谓词后面括号内的值列表既可以是显式列举的常量值，也可以是一个子查询的结果集，后者在实际开发中更为灵活。'),
    (65, '【补充说明】MONTH(CURRENT_DATE)中的CURRENT_DATE也可以替换为GETDATE()函数，两者在只取日期部分时效果类似，但GETDATE()同时返回时间部分。'),
    (48, '【补充说明】DATEDIFF函数的第一个参数datepart决定了差值的计算单位，常用值包括day（天）、month（月）、year（年）、hour（小时）等。'),
    (29, '【补充说明】WHERE子句中字符串值N体育用品的前缀N表示Unicode字符串，这是SQL Server中对NVARCHAR/NCHAR类型列进行比较的标准写法。'),
    (23, '【补充说明】DISTINCT作用于所有选择的列，如果SELECT后面有多个列，则组合值完全相同的行才会被去重。'),
    (17, '【补充说明】SELECT * 虽然简单便捷，但在实际开发中建议明确列出所需字段，避免不必要的数据传输和性能开销。'),
]

# Sort by index descending so insertions don't affect each other
supplementary_notes.sort(key=lambda x: x[0], reverse=True)

for img_idx, note_text in supplementary_notes:
    if img_idx in image_paras:
        para = doc.paragraphs[img_idx]  # Safe because we go in reverse
        insert_paragraph_after(para, note_text)
        print(f'Added note after paragraph [{img_idx}]')
    else:
        print(f'WARNING: Para [{img_idx}] is NOT an image, skipping')

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f'\nDocument saved: {DST}')
print(f'Total paragraphs after edit: {len(doc.paragraphs)}')
print('Done!')
