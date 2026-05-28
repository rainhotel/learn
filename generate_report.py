import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---- Page setup: A4 ----
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ---- Helper functions ----
def set_font(run, name_cn='宋体', name_en='Times New Roman', size=Pt(11), bold=False):
    """Set font with both Chinese and English font names."""
    run.font.size = size
    run.bold = bold
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)

def add_paragraph(text, style_name='Normal', font_cn='宋体', font_en='Times New Roman',
                  size=Pt(11), bold=False, alignment=None, spacing_after=Pt(6),
                  first_line_indent=None):
    """Add a paragraph with proper Chinese font handling."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = spacing_after
    pf.space_before = Pt(0)
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    return p

def add_title(text, font_cn='黑体', font_en='Times New Roman', size=Pt(16), bold=True):
    """Add a centered title."""
    return add_paragraph(text, font_cn=font_cn, font_en=font_en, size=size,
                         bold=bold, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         spacing_after=Pt(12))

def add_heading_text(text, font_cn='黑体', size=Pt(11)):
    """Add a bold section heading."""
    return add_paragraph(text, font_cn=font_cn, size=size, bold=True,
                         spacing_after=Pt(4))

def add_body(text, indent=True):
    """Add body text."""
    return add_paragraph(text, size=Pt(10.5),
                         first_line_indent=Cm(0.7) if indent else None,
                         spacing_after=Pt(3))

def add_image(image_path, width=Inches(5.5)):
    """Add an image, centered."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        return p
    else:
        add_body(f'[图片未找到: {image_path}]')
        return None

# ============================================================
# COVER / HEADER
# ============================================================
add_title('南京信息工程大学  实验（实习）报告', size=Pt(18))

# Info line 1
add_paragraph(
    '实验（实习）名称：  频域图像增强              实验（实习）日期：               '
    '得分：        指导教师：',
    size=Pt(10.5), spacing_after=Pt(2))

# Info line 2
add_paragraph(
    '学院：  计算机学院          专业：               年级：           班级：         '
    ' 姓名：             学号：',
    size=Pt(10.5), spacing_after=Pt(12))

# ============================================================
# 一、实验目的
# ============================================================
add_heading_text('一、实验目的：')
add_body('1. 掌握离散傅里叶变换的基本性质；')
add_body('2. 掌握二维图像傅里叶变换的方法和应用；')
add_body('3. 理解二维图像频谱；')
add_body('4. 掌握常用频域图像增强方法。')

# ============================================================
# 二、实验内容
# ============================================================
add_heading_text('二、实验内容：')

# --- 任务1 ---
add_heading_text('任务1：图像生成与FFT分析', font_cn='楷体')
add_body('（1）编程生成图像 f1(x,y)，大小为128×128，服从均匀分布（均值=0，方差=255）。'
         '对 f1 进行快速傅里叶变换（FFT），在同一窗口同时显示原图和 FFT 幅度谱。')
add_body('（2）将 f1(x,y) 顺时针旋转45度得到 f2(x,y)，显示 f2 的 FFT 幅度谱，'
         '并与 f1 的幅度谱进行比较。')

# Add key numeric results
p = doc.add_paragraph()
run = p.add_run('生成结果验证：f1 实际均值 ≈ -0.0220，实际方差 ≈ 254.94。'
                '旋转后频谱同步旋转45°，验证了傅里叶变换的旋转不变性。')
set_font(run, size=Pt(9), bold=False)
p.paragraph_format.first_line_indent = Cm(0.7)

add_image('D:/moniC/project/learn/lab5_output/task1_fft_rotation.png',
          width=Inches(5.2))

# --- 任务2 ---
add_heading_text('任务2：频域低通滤波', font_cn='楷体')
add_body('对图像 Fig0333(a)(test_pattern_blurring_orig).tif 在频域实现低通滤波，'
         '分别使用 Butterworth 低通滤波器（n=2）和高斯低通滤波器，'
         '观察不同截止频率 D₀ = 5, 15, 30, 80, 230 下的滤波效果。')

add_body('Butterworth低通滤波：H(u,v) = 1 / [1 + (D(u,v)/D₀)^(2n)]', indent=False)
add_body('高斯低通滤波：H(u,v) = exp(-D²(u,v) / (2D₀²))', indent=False)

add_body('结果分析：D₀ 越小，图像越模糊（高频分量被大量滤除）；D₀=230 时基本保留所有频率分量，'
         '图像接近原图。Butterworth 滤波器存在振铃现象，高斯滤波器过渡更平滑。')

add_image('D:/moniC/project/learn/lab5_output/task2_lowpass.png',
          width=Inches(5.2))

# --- 任务3 ---
add_heading_text('任务3：高斯噪声去除', font_cn='楷体')
add_body('对测试图像添加不同方差的高斯噪声（σ² = 0.001, 0.005, 0.01），观察噪声图像的频谱，'
         '根据频谱分析选择合适的低通滤波器类型和截止频率去除噪声。')

add_body('分析：高斯噪声在频谱中表现为全频带的白噪声，无明显频率集中。'
         '选择 Butterworth 低通滤波器进行去噪，噪声方差越大，截止频率 D₀ 应越小（滤波更强）。')

add_image('D:/moniC/project/learn/lab5_output/task3_gaussian_noise.png',
          width=Inches(5.2))

# --- 任务4 ---
add_heading_text('任务4：频域高通滤波', font_cn='楷体')
add_body('对测试图像在频域实现高通滤波，分别使用 Butterworth 高通滤波器和'
         '高斯高通滤波器，截止频率 D₀ = 15，对比显示结果。')

add_body('Butterworth高通滤波：H(u,v) = 1 / [1 + (D₀/D(u,v))^(2n)]', indent=False)
add_body('高斯高通滤波：H(u,v) = 1 - exp(-D²(u,v) / (2D₀²))', indent=False)

add_body('结果分析：高通滤波保留边缘和细节信息，但丢失了低频背景。'
         'Butterworth HPF 边缘提取更清晰，高斯 HPF 过渡更平滑。')

add_image('D:/moniC/project/learn/lab5_output/task4_highpass.png',
          width=Inches(5.2))

# --- 任务5 ---
add_heading_text('任务5：高频增强与高频强调滤波', font_cn='楷体')
add_body('在频域实现高频增强和高频强调滤波，截止频率 D₀ = 15。'
         '高频强调滤波公式为 H_hfe(u,v) = a + b × H_hp(u,v)，'
         '其中 a 控制低频分量保留程度，b 控制高频分量增强程度。')

add_body('分别测试四组参数：(a,b) = (0.5, 1.0), (0.5, 2.0), (1.0, 1.5), (0.3, 3.0)。', indent=False)

add_body('结果分析：a 越大，背景保留越完整；b 越大，边缘增强越明显。'
         '当 b > 1 时高频被强调，图像边缘更加突出。')

add_image('D:/moniC/project/learn/lab5_output/task5_hfe.png',
          width=Inches(5.2))

# --- 任务6 ---
add_heading_text('任务6：周期性噪声去除', font_cn='楷体')
add_body('对测试图像叠加正弦周期性噪声（水平和垂直方向），观察其频谱图。'
         '周期性噪声在频谱中表现为成对出现的亮点（冲激信号），'
         '通过设计 Butterworth 陷波滤波器，在噪声频率位置设置阻带，'
         '精确去除噪声分量而不影响其他频率成分。')

add_body('陷波滤波器：H_notch = ∏ H_k(u,v)，其中 H_k 是在噪声频率点 (u_k, v_k) 处'
         '的 Butterworth 带阻滤波器。', indent=False)

add_image('D:/moniC/project/learn/lab5_output/task6_periodic_noise.png',
          width=Inches(5.2))

# ============================================================
# 三、实验要求
# ============================================================
add_heading_text('三、实验要求：')
add_body('1. 使用 Matlab 语言进行编程，实现上述所有功能，程序通过调试并正确运行。')
add_body('2. 实验报告总结实验收获与体会。')
add_body('3. 实验报告控制在 A4 纸 6 页之内，不贴大段代码。')
add_body('4. 成果提交方式：撰写实验报告并打包程序和相关结果。')

# ============================================================
# 四、实验总结与体会
# ============================================================
add_heading_text('四、实验总结与体会：')

conclusions = [
    '通过本次频域图像增强实验，我对离散傅里叶变换（DFT）的基本性质有了更深入的理解。'
    '实验验证了傅里叶变换的旋转性质——空域中图像旋转45°后，其频谱也同步旋转45°，'
    '这说明空域和频域之间存在紧密的几何对应关系。',

    '在低通滤波实验中，我观察到截止频率 D₀ 对滤波效果的决定性影响：D₀ 越小，'
    '高频分量被滤除得越多，图像变得越模糊。Butterworth 滤波器在阶数较高时会产生'
    '"振铃"效应，而高斯滤波器在空域和频域都具有平滑的过渡特性，不会产生振铃现象。'
    '这让我理解了不同滤波器设计之间的权衡——锐截止 vs. 平滑过渡。',

    '在处理高斯噪声时，通过在频域观察噪声图像的频谱可以发现，高斯噪声表现为全频带'
    '的随机干扰。低通滤波虽然能有效去除高频噪声，但不可避免地会损失图像的边缘和细节'
    '信息，这是频域滤波的固有局限性。',

    '高通滤波实验使我认识到，频域滤波可以灵活地选择保留或抑制特定频率分量。'
    '高频强调滤波通过参数 a 和 b 的调节，可以在保留背景信息的同时增强边缘，'
    '这种方法比单纯的高通滤波更加实用，在实际图像增强中有广泛应用。',

    '在处理周期性噪声时，频谱分析的优势尤为明显。正弦噪声在频谱中表现为对称的亮点，'
    '可以通过陷波滤波器进行精确的"点对点"去除，几乎不影响其他频率分量。'
    '这体现了频域处理的独特优势——可以将复杂的空域卷积问题转化为频域的简单乘积运算。',

    '总体而言，本次实验让我掌握了频域图像增强的基本方法，包括低通滤波、高通滤波、'
    '高频强调滤波和陷波滤波等，加深了对傅里叶变换在图像处理中应用的认识。'
    '从频谱的视角理解图像，为后续学习更复杂的图像处理算法打下了良好基础。',
]

for c in conclusions:
    add_body(c)

# ============================================================
# Save
# ============================================================
output_path = 'D:/moniC/project/learn/lab5_实验报告_频域图像增强.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print('Done!')
