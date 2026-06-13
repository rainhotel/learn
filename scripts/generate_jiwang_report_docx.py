from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


WORKSPACE = Path(r"D:\moniC\project\learn")
SOURCE_MD = WORKSPACE / "03-outputs" / "jiwang-keshe-report-draft-v2.md"
OUTPUT_DOCX = WORKSPACE / "03-outputs" / "202483290198_周可名_计网实践.docx"

IMAGE_MAP = {
    "94b6b6ce25c83a9b5da97361680cac02.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\94b6b6ce25c83a9b5da97361680cac02.png"
    ),
    "58f2a2ffb552dab2c200fcb739c1a0d1.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\58f2a2ffb552dab2c200fcb739c1a0d1.png"
    ),
    "c738da036ffbdd69c90eb4ffe318c4e5.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\c738da036ffbdd69c90eb4ffe318c4e5.png"
    ),
    "fa0a3199d75bbffac2180a0beab20e43.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\fa0a3199d75bbffac2180a0beab20e43.png"
    ),
    "0e2baa9be30b7736ae1892a7dd34b838.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\0e2baa9be30b7736ae1892a7dd34b838.png"
    ),
    "2b30e8f0565944630c4c472e316c8913.png": Path(
        r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\2b30e8f0565944630c4c472e316c8913.png"
    ),
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_page_margins(section) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def apply_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph_text(doc: Document, text: str, style: str = "Body") -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if style == "Body":
        apply_font(run, "宋体", 10.5, False)
    elif style == "Abstract":
        apply_font(run, "仿宋", 9, False)
    elif style == "Keyword":
        apply_font(run, "宋体", 9, False)
    elif style == "Code":
        apply_font(run, "Consolas", 9, False)
    else:
        apply_font(run, "宋体", 10.5, False)


def add_heading_text(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_font(run, "黑体", 14, True)
    elif level == 2:
        apply_font(run, "黑体", 12, True)
    else:
        apply_font(run, "黑体", 10.5, True)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="Code")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line.rstrip("\n"))
        apply_font(run, "Consolas", 9, False)


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
    doc.add_paragraph("")


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[截图缺失：{caption}]")
        apply_font(r, "宋体", 10.5, False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    cr = cp.add_run(caption)
    apply_font(cr, "宋体", 10.5, False)


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Body", "Abstract", "Keyword", "Code"]:
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("《计算机网络实践》课程报告")
    apply_font(run, "黑体", 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    run = p.add_run("基于 Packet Tracer 的校园网络设计与基于 Socket 的局域网聊天软件设计与实现")
    apply_font(run, "黑体", 16, True)

    info_lines = [
        "姓名：周可名",
        "学号：202483290198",
        "专业：________________",
        "班级：________________",
        "日期：二〇二六年六月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(line)
        apply_font(run, "宋体", 14, False)

    doc.add_page_break()


def build_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Packet Tracer 的校园网络设计与基于 Socket 的局域网聊天软件设计与实现")
    apply_font(run, "黑体", 16, True)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("周可名")
    apply_font(run, "宋体", 10.5, False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("计算机与软件学院，专业，年级班")
    apply_font(run, "宋体", 9, False)

    doc.add_paragraph("")


def render_markdown(doc: Document, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    image_pending: tuple[Path, str] | None = None
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith(">"):
            stripped = stripped.lstrip(">").strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("## 题目"):
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            # skip title block in md; we build custom one
            while i < len(lines) and not lines[i].startswith("## 摘要"):
                i += 1
            continue

        if stripped.startswith("## 摘要"):
            add_heading_text(doc, "摘要", 1)
            i += 1
            continue

        if stripped.startswith("**关键词：**"):
            add_paragraph_text(doc, stripped.replace("**关键词：**", "关键词："), "Keyword")
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading_text(doc, strip_md(stripped[3:]), 1)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading_text(doc, strip_md(stripped[4:]), 2)
            i += 1
            continue

        if stripped.startswith("#### "):
            add_heading_text(doc, strip_md(stripped[5:]), 3)
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("**[此处插入图"):
            caption = strip_md(stripped).replace("[此处插入", "").replace("]", "").strip()
            image_pending = (Path(), caption)
            i += 1
            continue

        if stripped.startswith("对应文件："):
            filename = stripped.split("：", 1)[1].strip().strip("`")
            if filename in IMAGE_MAP:
                caption = image_pending[1] if image_pending else ""
                add_image(doc, IMAGE_MAP[filename], caption if caption else filename)
            image_pending = None
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="Body")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(strip_md(stripped))
            apply_font(run, "宋体", 10.5, False)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="Body")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run("• " + strip_md(stripped[2:]))
            apply_font(run, "宋体", 10.5, False)
            i += 1
            continue

        add_paragraph_text(doc, strip_md(stripped), "Abstract" if "本课程设计围绕" in stripped else "Body")
        i += 1


def main() -> None:
    doc = Document()
    ensure_styles(doc)
    for section in doc.sections:
        set_page_margins(section)

    build_cover(doc)
    build_title_block(doc)
    render_markdown(doc, SOURCE_MD)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
