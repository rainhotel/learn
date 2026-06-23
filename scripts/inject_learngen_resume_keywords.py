from pathlib import Path
import os
import re
import shutil
import sys
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import build_learngen_business_plan_revised as base_plan  # noqa: E402


OUT = ROOT / "03-outputs" / "learngen-business-plan-resume-enhanced.docx"


MEMBERS = [
    {
        "name": "李林浩",
        "student_id": "202483290054",
        "major": "计算机科学与技术",
        "phone": "13357836327",
        "email": "lilinhao5292006@163.com",
        "role": "项目负责人 / AI Agent 与后端架构负责人",
        "cover_role": "项目负责人、AI 技术路线、后端架构与商业方案统筹",
        "division": "负责项目总体统筹、商业模式完善、AI Agent 技术路线、RAG/知识图谱方案、后端架构和核心产品闭环设计。",
        "result": "将企业级 Agent 平台开发、Spring Boot/Redis/MySQL/Vue、LangChain/LangGraph/RAG、提示词工程和 AI 辅助开发经验转化为项目核心技术壁垒。",
        "profile": "校园软件技术团队负责人，具备企业级 AI Agent 平台开发实习经历，参与 Agent Memory、插件运行时、编程代理集成、飞书会话接入、企业身份与账号绑定、多实例通信可靠性等工程模块；掌握 Spring Boot、Redis、MySQL、Vue、LangChain、LangGraph、RAG、Agent Runtime、上下文工程、提示词工程和 AI 辅助开发；曾负责微信公众号内容管理系统、校园餐饮评价小程序等项目从 0 到 1 落地，兼具技术研发、产品交付和团队协作能力。",
    },
    {
        "name": "周可名",
        "student_id": "202483290198",
        "major": "计算机科学与技术",
        "phone": "__________",
        "email": "",
        "role": "运维与平台可靠性负责人",
        "cover_role": "云原生运维、CI/CD、部署安全与平台可靠性保障",
        "division": "负责自动化部署、Docker 容器化、CI/CD 流水线、数据库版本管理、云服务资源配置、安全组与对象存储权限管理。",
        "result": "将 Jenkins、Docker/Docker Compose、Flyway、Caddy、GitLab、云服务资源管理和 AI 辅助编程能力注入平台工程化落地方案。",
        "profile": "英语六级 701/710，具备扎实计算机基础与较强英文资料阅读能力；熟悉 Docker、Docker Compose、Jenkins、GitLab、Flyway、Caddy、K3s/K8s 等运维与云原生工具；参与校园美食推荐系统和微信公众号管理系统运维，能够搭建自动化发版流程、管理数据库版本、配置网关与 SSL 证书、维护云服务器安全组和对象存储权限，适合承担平台上线、持续交付和稳定性保障工作。",
    },
    {
        "name": "余越",
        "student_id": "202483290188",
        "major": "计算机科学与技术",
        "phone": "15248138746",
        "email": "491278763@qq.com",
        "role": "AI 应用与全栈产品负责人",
        "cover_role": "多智能体应用、全栈开发、产品功能闭环与交互体验",
        "division": "负责 AI 应用功能设计、多智能体内容生成流程、SSE 流式交互、前后端联调、权限体系和用户安全模块设计。",
        "result": "将 Spring AI Alibaba、StateGraph 多智能体编排、React18/Ant Design Pro、Spring Boot、MySQL/Redis、RBAC 权限控制和蓝桥杯省一能力注入产品研发。",
        "profile": "蓝桥杯江苏省 C++ 组省一，获得校级奖学金、三好学生，英语四六级通过；熟悉 Java、Spring Boot、MyBatis/MyBatis-Plus、MySQL、Redis、Vue3、React、TypeScript、RESTful API、RBAC 与 Token 认证机制；参与基于 Spring AI Alibaba 和 StateGraph 的多智能体自媒体文章创作工具，负责选题、标题、大纲、正文、配图等 Agent 编排、SSE 流式输出、React18 + Ant Design Pro 后台和多策略配图系统；参与校园老吃家用户权限与密码管理模块重构，具备 AI 应用工程化、全栈开发和产品体验优化能力。",
    },
]


def get_or_add(parent, tag):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    return node


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False):
    run.font.name = west
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = get_or_add(rpr, "w:rFonts")
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)
    rfonts.set(qn("w:eastAsia"), east)


def rewrite_paragraph(paragraph, text, *, bold=False, east="宋体", size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.text = ""
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, east=east, size=size, bold=bold)


def format_cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, east="宋体", size=10.5, bold=bold)


def resize_body_rows(table, body_count):
    while len(table.rows) - 1 > body_count:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) - 1 < body_count:
        table.add_row()


def fill_table(table, headers, rows, aligns):
    resize_body_rows(table, len(rows))
    for i, header in enumerate(headers):
        format_cell(table.rows[0].cells[i], header, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            format_cell(table.rows[row_idx].cells[col_idx], value, align=aligns[col_idx])


def replace_paragraphs(doc):
    replacements = {
        "联系人：__________    电话：__________    电子邮箱：__________": "联系人：李林浩    电话：13357836327    电子邮箱：lilinhao5292006@163.com",
        "智学派教育科技有限公司（拟）是一家专注于 AI+教育领域的科技创业公司，由高校学生团队发起。公司拟设立于武汉光谷或高校创业孵化器，利用当地高校资源、软件人才、云计算生态和创业政策，降低初期研发和运营成本。": "智学派教育科技有限公司（拟）是一家由南京信息工程大学计算机科学与技术专业学生团队发起的 AI+教育科技创业项目。团队核心成员具备企业级 AI Agent 平台研发、多智能体应用开发、校园软件项目交付、云原生运维、CI/CD 自动化部署和高校场景产品理解能力，能够把“大学生真实学习痛点”与“可落地的软件工程能力”结合起来，推动项目从课程作业方案走向可试点 MVP。",
        "团队成员来自计算机、软件工程、教育技术、管理和财务相关方向，具有贴近大学生用户、理解课程学习痛点、快速迭代产品原型的优势。学生团队虽然资源有限，但更容易进入真实课程群和学习场景，也更能理解同龄用户在考试压力、学习节奏和付费意愿上的特点。": "团队目前由三名计算机科学与技术专业核心成员构成，形成“AI Agent 平台研发 + 多智能体应用开发 + 云原生运维交付”的互补组合。李林浩负责项目统筹、AI Agent/RAG 技术路线、后端架构与产品闭环，具备企业级 Agent 平台实习和多个线上项目交付经验；余越负责 Spring AI 多智能体应用、SSE 流式交互、React 后台和权限安全模块，具备蓝桥杯省一和全栈 AI 应用落地经验；周可名负责容器化部署、CI/CD、云服务资源、安全配置和数据库版本管理，具备运维流水线、Docker Compose、Flyway、Caddy、Jenkins 等工程化能力。三名成员都熟悉 AI 辅助开发，能够在较短周期内完成原型开发、部署上线、反馈迭代和稳定性保障。",
        "团队目前由两名计算机科学与技术专业核心成员构成，形成“AI 平台研发 + 云原生运维交付”的互补组合。李林浩负责项目统筹、AI Agent/RAG 技术路线、后端架构与产品闭环，具备企业级 Agent 平台实习和多个线上项目交付经验；周可名负责容器化部署、CI/CD、云服务资源、安全配置和数据库版本管理，具备运维流水线、Docker Compose、Flyway、Caddy、Jenkins 等工程化能力。两名成员都熟悉 AI 辅助开发，能够在较短周期内完成原型开发、部署上线、反馈迭代和稳定性保障。": "团队目前由三名计算机科学与技术专业核心成员构成，形成“AI Agent 平台研发 + 多智能体应用开发 + 云原生运维交付”的互补组合。李林浩负责项目统筹、AI Agent/RAG 技术路线、后端架构与产品闭环，具备企业级 Agent 平台实习和多个线上项目交付经验；余越负责 Spring AI 多智能体应用、SSE 流式交互、React 后台和权限安全模块，具备蓝桥杯省一和全栈 AI 应用落地经验；周可名负责容器化部署、CI/CD、云服务资源、安全配置和数据库版本管理，具备运维流水线、Docker Compose、Flyway、Caddy、Jenkins 等工程化能力。三名成员都熟悉 AI 辅助开发，能够在较短周期内完成原型开发、部署上线、反馈迭代和稳定性保障。",
        "表 8 初期组织结构与岗位职责": "表 8 核心成员能力画像与岗位匹配",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            if text.startswith("表 8"):
                rewrite_paragraph(paragraph, replacements[text], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                rewrite_paragraph(paragraph, replacements[text])


def update_tables(doc):
    cover_rows = [[m["major"], m["student_id"], m["name"], m["phone"]] for m in MEMBERS]
    fill_table(
        doc.tables[0],
        ["专业班级", "学号", "姓名", "联系电话"],
        cover_rows,
        [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )

    team_rows = [[m["name"], m["role"], m["profile"]] for m in MEMBERS]
    fill_table(
        doc.tables[7],
        ["成员", "团队角色", "能力关键词与项目贡献"],
        team_rows,
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    division_rows = [
        [m["name"], f'{m["major"]} {m["student_id"]}', m["phone"], m["division"], m["result"]]
        for m in MEMBERS
    ]
    fill_table(
        doc.tables[-1],
        ["成员", "专业与学号", "联系电话", "主要分工", "成果说明"],
        division_rows,
        [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )


def force_black_and_white(doc):
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            style.font.color.rgb = RGBColor(0, 0, 0)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                for shd in list(tc_pr.findall(qn("w:shd"))):
                    tc_pr.remove(shd)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def scrub_docx_ooxml(path):
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                xml = data.decode("utf-8", errors="ignore")
                xml = re.sub(r'(<w:color\b[^>]*?w:val=")([^"]+)("[^>]*/>)', r'\g<1>000000\g<3>', xml)
                xml = re.sub(r'<w:shd\b[^>]*/>', '', xml)
                data = xml.encode("utf-8")
            zout.writestr(info, data)
    shutil.move(str(tmp), str(path))


def build_enhanced():
    base_path = base_plan.build()
    doc = Document(base_path)
    replace_paragraphs(doc)
    update_tables(doc)
    force_black_and_white(doc)
    doc.save(OUT)
    scrub_docx_ooxml(OUT)
    return OUT


if __name__ == "__main__":
    print(build_enhanced())
