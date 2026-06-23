from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "03-outputs" / "learngen-business-plan-revised.docx"


def cm_to_dxa(value: float) -> int:
    return int(value / 2.54 * 1440)


def get_or_add(parent, tag):
    found = parent.find(qn(tag))
    if found is None:
        found = OxmlElement(tag)
        parent.append(found)
    return found


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=False):
    run.font.name = west
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = get_or_add(rPr, "w:rFonts")
    rFonts.set(qn("w:ascii"), west)
    rFonts.set(qn("w:hAnsi"), west)
    rFonts.set(qn("w:eastAsia"), east)


def set_style_font(style, east="宋体", west="Times New Roman", size=10.5, bold=False):
    style.font.name = west
    style.font.size = Pt(size)
    style.font.bold = bold
    rPr = style._element.get_or_add_rPr()
    rFonts = get_or_add(rPr, "w:rFonts")
    rFonts.set(qn("w:ascii"), west)
    rFonts.set(qn("w:hAnsi"), west)
    rFonts.set(qn("w:eastAsia"), east)


def format_paragraph(paragraph, *, align=None, first_line=False, before=0, after=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    else:
        fmt.first_line_indent = None
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, *, size=10.5, east="宋体", west="Times New Roman", bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, before=0, after=0):
    p = doc.add_paragraph()
    format_paragraph(p, align=align, first_line=first_line, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, east=east, west=west, size=size, bold=bold)
    return p


def add_center(doc, text, *, size=10.5, east="宋体", bold=False, before=0, after=0):
    return add_text(
        doc,
        text,
        size=size,
        east=east,
        bold=bold,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        before=before,
        after=after,
    )


def add_chapter(doc, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=0, after=8)
    run = p.add_run(title)
    set_run_font(run, east="黑体", west="Times New Roman", size=14, bold=True)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=6, after=0)
    run = p.add_run(title)
    set_run_font(run, east="黑体", west="Times New Roman", size=10.5, bold=True)
    return p


def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=4, after=0)
    run = p.add_run(title)
    set_run_font(run, east="黑体", west="Times New Roman", size=10.5, bold=True)
    return p


def add_caption(doc, text):
    return add_text(
        doc,
        text,
        size=9,
        east="宋体",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        before=3,
        after=3,
    )


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = get_or_add(tcPr, "w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = get_or_add(tcMar, f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = get_or_add(tcPr, "w:tcW")
    tcW.set(qn("w:w"), str(cm_to_dxa(width_cm)))
    tcW.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = get_or_add(tblPr, "w:tblW")
    tblW.set(qn("w:w"), str(sum(cm_to_dxa(w) for w in widths_cm)))
    tblW.set(qn("w:type"), "dxa")
    layout = get_or_add(tblPr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(cm_to_dxa(width)))
        grid.append(col)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tcPr, "w:shd")
    shd.set(qn("w:fill"), fill)


def format_cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    format_paragraph(p, align=align, first_line=False, before=0, after=0, line=1.25)
    run = p.add_run(text)
    set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=bold)


def add_table(doc, caption, headers, rows, widths_cm, alignments=None):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    if alignments is None:
        alignments = [WD_ALIGN_PARAGRAPH.CENTER] * len(headers)
    for i, header in enumerate(headers):
        set_cell_width(table.rows[0].cells[i], widths_cm[i])
        format_cell(table.rows[0].cells[i], header, bold=True, fill="D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths_cm[i])
            format_cell(cells[i], value, align=alignments[i])
    add_text(doc, "", first_line=False, after=2)
    return table


def add_footer(section):
    footer = section.footer.paragraphs[0]
    format_paragraph(footer, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=0, after=0)
    run = footer.add_run("第 ")
    set_run_font(run, east="宋体", size=9)

    run = footer.add_run()
    set_run_font(run, east="宋体", size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = footer.add_run()
    set_run_font(run, east="宋体", size=9)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    run = footer.add_run()
    set_run_font(run, east="宋体", size=9)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run = footer.add_run(" 页")
    set_run_font(run, east="宋体", size=9)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    set_style_font(normal, east="宋体", west="Times New Roman", size=10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, east, size, centered in [
        ("Heading 1", "黑体", 14, True),
        ("Heading 2", "黑体", 10.5, False),
        ("Heading 3", "黑体", 10.5, False),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, east=east, west="Times New Roman", size=size, bold=True)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        if centered:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_footer(section)


def add_cover(doc):
    for _ in range(2):
        add_text(doc, "", first_line=False)
    add_center(doc, "《创新创业基础》课程作业", size=16, east="黑体", bold=True, after=12)
    add_center(doc, "创 业 计 划 书", size=22, east="黑体", bold=True, after=22)
    add_center(doc, "项目名称：智学派（LearnGen）", size=14, east="黑体", bold=True, after=6)
    add_center(doc, "AI 驱动的个性化学习辅助平台", size=12, east="宋体", bold=False, after=24)

    rows = [["", "", "", ""] for _ in range(5)]
    add_table(
        doc,
        "表 1 团队成员信息（提交前填写完整）",
        ["专业班级", "学号", "姓名", "联系电话"],
        rows,
        [4.5, 4.0, 3.5, 5.0],
        [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_text(doc, "", first_line=False, after=8)
    add_center(doc, "团队名称：智学派项目团队", size=10.5, east="宋体", after=3)
    add_center(doc, "编制时间：2026 年 6 月", size=10.5, east="宋体", after=3)
    add_center(doc, "联系人：__________    电话：__________    电子邮箱：__________", size=10.5, east="宋体")
    doc.add_page_break()


def add_integrity_page(doc):
    add_chapter(doc, "原创性与保密说明")
    add_section(doc, "一、原创性说明")
    add_text(doc, "本计划书由智学派项目团队围绕“AI 驱动的个性化学习辅助平台”创业项目独立完成。报告中的行业数据、政策表述和竞品信息来自公开资料整理，团队在使用时已结合项目定位重新分析、归纳和表达。提交前，团队成员应再次核验数据来源、成员信息与分工，确保符合课程关于学术规范和原创性的要求。")
    add_section(doc, "二、保密说明")
    add_text(doc, "本商业计划书仅用于《创新创业基础》课程作业、项目展示和创业训练用途。未经团队许可，不得将其中涉及的产品方案、技术路线、商业模式、财务测算和团队信息用于与本项目无关的商业传播。")
    doc.add_page_break()


def add_toc(doc):
    add_center(doc, "目        录", size=16, east="黑体", bold=True, after=12)
    toc_items = [
        ("第一章  项目简介", 14, 0),
        ("第二章  产品说明", 14, 0),
        ("第三章  运营模式", 14, 0),
        ("第四章  市场分析", 14, 0),
        ("第五章  营销与销售", 14, 0),
        ("第六章  公司概述与创业团队", 14, 0),
        ("第七章  财务分析", 14, 0),
        ("第八章  风险控制与实施保障", 14, 0),
        ("第九章  附录", 14, 0),
        ("团队成员分工表", 12, 0.74),
    ]
    for text, size, indent in toc_items:
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=0, after=4)
        p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(text)
        set_run_font(run, east="宋体", size=size, bold=False)
    doc.add_page_break()


def add_project_intro(doc):
    add_chapter(doc, "第一章  项目简介")
    add_section(doc, "一、项目概况")
    add_text(doc, "智学派（LearnGen）是一个面向大学生的 AI 个性化学习辅助平台，核心目标是让每一位学生都能拥有一位“懂课程、懂考试、懂自己学习节奏”的 AI 学习导师。平台以大语言模型、检索增强生成（RAG）、课程知识图谱和学习诊断算法为技术基础，为大学生提供课程答疑、学习路径规划、薄弱知识点诊断、考前冲刺训练和学习数据看板等服务。")
    add_text(doc, "项目初期聚焦高等数学、线性代数、大学物理、程序设计、数据结构等高校公共基础课和理工科高频课程。这些课程覆盖面广、学习难度高、考试压力集中，是大学生学习痛点最明显、也是 AI 个性化辅导最容易产生效果的场景。")
    add_section(doc, "二、项目背景")
    add_text(doc, "高等教育规模持续扩大，但学生在课堂之外获得个性化指导的机会仍然不足。传统教师答疑受时间和班级规模限制，一对一辅导费用较高，普通录播课和慕课又缺乏针对个人薄弱环节的反馈。与此同时，大学生已经普遍接触 ChatGPT、Kimi、豆包、通义千问等 AI 工具，但通用 AI 在具体课程课件、教师讲义和考试范围上的准确性有限，容易出现泛泛而谈或生成错误答案的问题。")
    add_text(doc, "智学派的机会在于把通用 AI 能力和高校课程场景深度结合：一方面利用大模型的自然语言理解、推理和解释能力，降低学生提问门槛；另一方面通过课程知识库、知识图谱和学习行为数据约束回答范围，让 AI 辅导从“能聊天”走向“能学会”。")
    add_section(doc, "三、项目定位")
    add_text(doc, "智学派不是单纯的搜题工具，也不是传统网课平台，而是“AI 学习操作系统”。它既能回答学生当前的问题，也能告诉学生为什么不会、应该先补什么、下一步怎么学。产品定位为大学生课内学习与考试备考的智能助教，兼顾 C 端个人订阅和 B 端高校教学辅助场景。")
    add_section(doc, "四、核心价值")
    add_text(doc, "对学生而言，智学派提供低成本、随时可用、可持续反馈的学习支持，帮助学生降低挂科风险、提升学习效率和备考信心。对高校而言，平台可形成班级和课程层面的学习数据反馈，帮助教师发现共性薄弱点，辅助教学改进。对投资者而言，项目切入 AI+教育和高校数字化交叉赛道，具备较清晰的用户增长路径、商业化方式和技术壁垒。")


def add_product(doc):
    doc.add_page_break()
    add_chapter(doc, "第二章  产品说明")
    add_section(doc, "一、产品形态")
    add_text(doc, "智学派采用 App、小程序和 Web 端协同的产品形态。学生端强调随时提问、随手拍题、学习路径查看和考前冲刺；教师端和高校管理端强调课程知识库管理、班级学习数据查看、共性问题统计和教学反馈。项目 MVP 阶段优先开发学生端核心功能，随后再逐步开放教师端管理能力。")
    add_section(doc, "二、核心功能")
    add_table(
        doc,
        "表 2 核心功能与实现方式",
        ["功能模块", "用户价值", "技术实现"],
        [
            ["AI 智能答疑", "支持文字、图片和语音提问，给出分步讲解、知识点定位和同类题建议。", "大模型推理、OCR、语音识别、RAG 检索和答案置信度标注。"],
            ["个性化学习路径", "根据课程目标和当前水平生成学习计划，展示先修关系、重点难度和建议学习时长。", "课程知识图谱、学习目标解析、路径规划算法和学习进度追踪。"],
            ["薄弱知识点诊断", "通过练习记录和答疑行为识别薄弱点，生成热力图和补强任务。", "贝叶斯知识追踪模型、错因标签体系和知识点掌握度评估。"],
            ["课件与笔记解析", "上传课件、教材片段和笔记后，自动生成摘要、考点清单和问答索引。", "文档解析、向量化存储、知识片段召回和内容安全过滤。"],
            ["学习数据看板", "展示学习时长、正确率、知识点覆盖率、连续打卡和阶段报告。", "行为埋点、数据分析、可视化报表和匿名同伴对比。"],
        ],
        [3.2, 6.5, 7.3],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "三、技术特点")
    add_subsection(doc, "（一）大模型与 RAG 混合架构")
    add_text(doc, "平台底层接入主流大模型能力，同时建设自有课程知识库。学生提问后，系统先在课程课件、教材摘要、题库解析和知识图谱中检索相关片段，再将检索结果与问题一并交给大模型生成答案。该方案能减少通用大模型在专业课程内容上的幻觉，提升回答与课堂内容的匹配度。")
    add_subsection(doc, "（二）课程知识图谱")
    add_text(doc, "课程知识图谱记录知识点之间的前置依赖、难度层级、常见错误和典型题型。例如学习“特征值与特征向量”前，系统会识别学生是否掌握矩阵运算、行列式、线性相关性等前置知识。图谱不仅服务答疑，也服务学习路径规划和薄弱点诊断。")
    add_subsection(doc, "（三）学习诊断算法")
    add_text(doc, "平台根据学生的答题结果、追问次数、停留时长和复习行为，估计每个知识点的掌握概率。系统会把“不会做题”拆解为概念不清、公式记忆不足、步骤错误、计算粗心等不同错因，并给出针对性的补救建议。")
    add_section(doc, "四、研发方法")
    add_text(doc, "项目研发采用“课程样本小闭环、功能迭代快验证”的方法。第一阶段选取高等数学和线性代数两门课程建设 MVP，完成知识点拆分、课件入库、基础答疑和学习路径生成。第二阶段通过 2 至 3 个班级试点收集真实提问和学习反馈，优化知识库质量和提示词模板。第三阶段扩展到更多公共基础课，并建立教师协作和内容审核机制。")
    add_section(doc, "五、产品路线图")
    add_table(
        doc,
        "表 3 产品路线图",
        ["阶段", "时间", "重点任务", "验收指标"],
        [
            ["MVP 验证期", "0 至 3 个月", "完成高数、线代课程知识库；上线文字问答、学习路径和基础诊断。", "试点用户 1000 人，周留存率超过 25%。"],
            ["校园试点期", "4 至 9 个月", "接入拍照提问、课件解析、教师端数据看板；进入 5 至 10 所高校试点。", "注册用户 5 万人，Pro 付费用户 3000 人以上。"],
            ["规模增长期", "10 至 24 个月", "扩展 30 门课程，建设校园大使体系和 B 端合作流程。", "覆盖 300 所高校，B 端合作 50 所。"],
            ["业务拓展期", "25 至 36 个月", "拓展考研、考公、四六级和职业证书备考场景。", "累计注册用户 300 万人，实现年度盈亏平衡。"],
        ],
        [3.0, 3.0, 7.0, 4.0],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )


def add_operation(doc):
    doc.add_page_break()
    add_chapter(doc, "第三章  运营模式")
    add_section(doc, "一、商业运作逻辑")
    add_text(doc, "智学派采用“C 端 Freemium + B 端高校合作”的双引擎模式。C 端通过免费基础功能快速获取学生用户，再以高频答疑、深度诊断、考前冲刺和课件解析等功能促进付费转化。B 端则面向高校、学院和课程团队提供教学辅助系统、课程知识库建设和学习数据分析服务。")
    add_text(doc, "这种模式能降低单一收入来源带来的风险。学生端有利于形成规模化数据和产品口碑，高校端有利于获得稳定合同收入和场景背书。两者之间形成正循环：高校试点带来种子用户，学生使用数据反过来证明教学辅助价值。")
    add_section(doc, "二、进入市场方式")
    add_text(doc, "项目冷启动从团队所在高校及周边高校切入，优先寻找高等数学、线性代数、程序设计等公共课教师合作，以“AI 助力课程学习试点”的方式进入班级。试点期为学生提供免费账号和课程专属知识库，教师获得班级学习薄弱点报告。通过一个班级、一个课程群、一个学院逐步扩散，降低传统广告获客成本。")
    add_section(doc, "三、盈利模式与定价策略")
    add_table(
        doc,
        "表 4 收入来源与定价设计",
        ["收入来源", "目标客户", "定价方式", "说明"],
        [
            ["C 端 Pro 订阅", "大学生个人用户", "29.9 元／月，199 元／年", "无限答疑、深度诊断、课件解析、考前计划。"],
            ["考前冲刺包", "期末备考用户", "19.9 元／次起", "高频考点、模拟题、错题复盘和冲刺计划。"],
            ["高校合作版", "高校、学院、课程团队", "院系试点 10 至 20 万元／年；校级合作 50 至 100 万元／年", "全校账号、教师端面板、课程知识库和 API 对接。"],
            ["课程内容共建", "教师团队、教研机构", "项目制收费或收益分成", "共建课程知识图谱、题库解析和教学数据报告。"],
        ],
        [3.2, 3.2, 4.4, 6.2],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "四、运营流程")
    add_text(doc, "运营流程分为课程建库、用户导入、学习服务、数据反馈和商业转化五个环节。课程建库阶段由团队对教材、课件和历年题型进行知识点拆分；用户导入阶段通过教师推荐、课程群邀请码和校园大使完成注册；学习服务阶段围绕答疑、路径、诊断和冲刺形成高频使用；数据反馈阶段向教师输出班级薄弱点和教学建议；商业转化阶段则通过 Pro 试用、考前冲刺包和高校续约实现收入。")
    add_section(doc, "五、关键运营指标")
    add_table(
        doc,
        "表 5 关键运营指标",
        ["指标", "目标值", "管理意义"],
        [
            ["次日留存率", "35% 以上", "判断新用户是否在首次使用中感受到价值。"],
            ["周留存率", "25% 以上", "判断产品是否进入真实学习习惯。"],
            ["免费转付费率", "8% 至 15%", "衡量 Pro 功能和考前场景的商业价值。"],
            ["获客成本", "8 至 12 元／人", "控制校园推广和内容投放效率。"],
            ["LTV／CAC", "大于 5", "判断用户生命周期价值是否覆盖获客成本。"],
            ["B 端续约率", "70% 以上", "衡量高校合作服务的稳定性。"],
        ],
        [4.0, 4.0, 9.0],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )


def add_market(doc):
    doc.add_page_break()
    add_chapter(doc, "第四章  市场分析")
    add_section(doc, "一、市场容量与趋势")
    add_text(doc, "我国高等教育在校生规模超过 4000 万人，大学生群体对数字化学习工具接受度高，且具有明确的考试通过、绩点提升、考研备考和技能学习需求。在线教育市场经过多年发展后，单纯录播课和题库产品的增长放缓，而 AI 个性化学习、智能助教和教学数据分析正在成为新的增长方向。")
    add_text(doc, "从政策环境看，教育数字化、人工智能应用和高校教学改革持续受到鼓励。《中国教育现代化 2035》《新一代人工智能发展规划》以及教育强国相关政策均强调用信息技术提升教育质量和促进教育公平。智学派切入的大学课程 AI 学习场景，既符合学生端需求，也符合高校提升教学质量和学习支持能力的方向。")
    add_table(
        doc,
        "表 6 目标市场规模测算",
        ["维度", "保守测算", "说明"],
        [
            ["中国高等教育在校生", "4000 万人以上", "本科、高职高专和研究生构成潜在用户基础。"],
            ["初期核心用户", "约 1500 万人", "优先覆盖理工科公共课学习压力较高的学生。"],
            ["三年注册目标", "300 万人", "对应核心市场约 20% 渗透率。"],
            ["三年 Pro 用户目标", "20 至 30 万人", "按 8% 至 10% 左右付费转化率测算。"],
            ["B 端合作目标", "100 所高校或院系", "以院系试点向校级合作逐步升级。"],
        ],
        [4.8, 4.2, 8.0],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "二、目标客户分析")
    add_text(doc, "第一类客户是在校大学生，尤其是公共基础课压力较大、期末考试需求明确、愿意尝试 AI 工具的学生。他们价格敏感，但只要产品能直接提升学习效率和考试通过率，就愿意为低价订阅或考前冲刺包付费。")
    add_text(doc, "第二类客户是高校教师和课程团队。他们关注学生是否真正理解课程内容，是否能及时发现共性问题，以及是否能减少重复答疑压力。智学派提供的班级薄弱点报告和课程问答数据，可作为教学改进和课程建设的参考。")
    add_text(doc, "第三类客户是高校管理部门和学院。他们关注学业预警、挂科率、教学质量评价和数字化建设成果。B 端合作版可以作为学习支持中心、智慧教学平台或课程改革试点工具的一部分。")
    add_section(doc, "三、竞争对手分析")
    add_table(
        doc,
        "表 7 竞品对比分析",
        ["竞品类型", "代表产品", "优势", "不足", "智学派应对策略"],
        [
            ["通用 AI 工具", "ChatGPT、Kimi、豆包等", "模型能力强，用户认知度高。", "不懂具体课程和教师课件，学习路径与诊断能力弱。", "用课程知识库和知识图谱做垂直化差异。"],
            ["传统在线教育平台", "中国大学 MOOC、智慧树、超星学习通", "课程资源和高校渠道较强。", "互动性弱，学生主动使用意愿不足。", "以 AI 互动和个性化反馈提升学习体验。"],
            ["搜题与题库工具", "作业帮、学小易、考途", "题库丰富，使用习惯成熟。", "偏向给答案，难以证明学生真正学会。", "强调步骤讲解、错因分析和知识点补救。"],
            ["AI 教育公司", "有道 AI、学而思九章等", "教育品牌和内容积累较深。", "多聚焦 K12，高校课程覆盖不足。", "聚焦大学课程，形成错位竞争和先发优势。"],
        ],
        [3.0, 3.2, 3.6, 3.6, 3.6],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "四、竞争策略")
    add_text(doc, "智学派的竞争策略是聚焦、垂直和深度绑定。聚焦是指不一开始做全学段教育，而是先做好大学公共基础课和理工科核心课。垂直是指围绕课程知识图谱、课件检索和考试目标做专门优化，而不是简单套用通用聊天模型。深度绑定是指通过教师试点、课程群运营和高校合作建立场景入口，让产品出现在学生真实学习链路中。")
    add_section(doc, "五、市场进入与增长节奏")
    add_text(doc, "第一年为种子期，目标是覆盖 50 所高校，重点验证产品市场契合度。第二年为增长期，借助校园大使、内容营销和院系合作扩展至 300 所高校。第三年为扩张期，在稳定大学课程场景后，进入考研、考公、四六级和职业证书备考市场，形成第二增长曲线。")


def add_sales(doc):
    doc.add_page_break()
    add_chapter(doc, "第五章  营销与销售")
    add_section(doc, "一、获客策略")
    add_text(doc, "智学派的获客重点不放在大规模硬广，而放在校园场景中的低成本裂变。项目通过课程群邀请码、教师推荐、学习社群分享和校园大使拉新，触达有即时学习需求的学生。对于期末考试、补考、考研基础课等高压场景，产品用“考前两周冲刺计划”“AI 讲懂一类题”等内容吸引用户首次体验。")
    add_section(doc, "二、转化策略")
    add_text(doc, "免费版保留每日基础答疑和有限学习路径，让用户可以低门槛体验核心价值。Pro 版则提供无限答疑、深度诊断报告、课件解析和考前冲刺计划。转化节点主要设置在考试前、连续使用达到一定次数、完成知识点诊断后，以及用户上传课件并需要深度解析时。")
    add_section(doc, "三、留存策略")
    add_text(doc, "留存依赖产品是否持续帮助学生解决学习问题。平台会通过每日学习打卡、薄弱知识点提醒、阶段性复习计划、学习周报和班级匿名对比等方式，让学生形成持续使用习惯。对于长期未使用用户，系统只在考试节点和薄弱点复习节点进行低频提醒，避免过度打扰。")
    add_section(doc, "四、品牌定位")
    add_text(doc, "智学派的品牌口号为“你的 AI 学长，24 小时在线”。品牌调性强调年轻、可靠、陪伴式和懂学生，不制造教育焦虑，而是把复杂课程拆成可理解、可跟进、可复习的学习任务。")
    add_section(doc, "五、销售路径")
    add_text(doc, "B 端销售采取“免费试点、数据证明、院系付费、校级扩展”的路径。试点阶段用一个课程或一个学院证明学习效果，例如答疑响应速度、学生使用率、薄弱点报告质量和教师重复答疑减少情况。形成案例后，再向学院和学校信息化部门推进年度合作。")


def add_company_team(doc):
    doc.add_page_break()
    add_chapter(doc, "第六章  公司概述与创业团队")
    add_section(doc, "一、公司概述")
    add_text(doc, "智学派教育科技有限公司（拟）是一家专注于 AI+教育领域的科技创业公司，由高校学生团队发起。公司拟设立于武汉光谷或高校创业孵化器，利用当地高校资源、软件人才、云计算生态和创业政策，降低初期研发和运营成本。")
    add_section(doc, "二、愿景与使命")
    add_text(doc, "公司愿景是成为大学生首选的 AI 学习平台，让每个学生都拥有自己的 AI 导师。公司使命是利用大模型和教育数据技术降低优质学习支持的门槛，让因材施教从理念走向可规模化落地的产品。")
    add_section(doc, "三、组织结构")
    add_table(
        doc,
        "表 8 初期组织结构与岗位职责",
        ["岗位", "主要职责", "能力要求"],
        [
            ["项目负责人／CEO", "负责战略规划、融资沟通、外部合作、课程作业统筹和路演表达。", "具备组织协调、商业分析和表达能力。"],
            ["技术负责人／CTO", "负责大模型接入、RAG 架构、知识图谱、后端服务和数据安全。", "熟悉软件开发、AI 应用和系统架构。"],
            ["产品负责人／CPO", "负责用户需求、产品原型、交互流程、MVP 迭代和用户反馈分析。", "具备产品设计、用户研究和文档能力。"],
            ["运营负责人／COO", "负责校园推广、社群运营、校园大使、B 端试点和客户服务。", "具备活动组织、渠道沟通和数据运营能力。"],
            ["财务与合规负责人", "负责财务测算、融资预算、知识产权、数据合规和风险控制。", "具备财务分析、规范意识和资料整理能力。"],
        ],
        [3.2, 8.2, 5.6],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "四、团队优势")
    add_text(doc, "团队成员来自计算机、软件工程、教育技术、管理和财务相关方向，具有贴近大学生用户、理解课程学习痛点、快速迭代产品原型的优势。学生团队虽然资源有限，但更容易进入真实课程群和学习场景，也更能理解同龄用户在考试压力、学习节奏和付费意愿上的特点。")
    add_section(doc, "五、外部协作")
    add_text(doc, "项目计划与高校教师、学习社团、创业导师、云服务厂商和法律财务顾问建立合作。高校教师负责课程内容把关，学习社团和校园大使负责用户触达，云服务厂商提供算力与存储支持，法律财务顾问协助知识产权、公司注册、合同和合规建设。")
    add_section(doc, "六、知识产权策略")
    add_text(doc, "公司拟围绕“智学派”“LearnGen”品牌申请商标保护；围绕 App、Web 平台和教师管理后台申请软件著作权；围绕学习路径推荐、课程知识图谱构建和 RAG 课程问答系统申请发明专利或实用新型；课程知识库、提示词模板和学习行为分析模型作为商业秘密管理。")


def add_finance(doc):
    doc.add_page_break()
    add_chapter(doc, "第七章  财务分析")
    add_section(doc, "一、资本需求")
    add_text(doc, "项目启动期计划融资 300 万元人民币，用于完成 MVP 产品研发、课程知识库建设、校园试点推广和基础运营。融资形式可包括团队自筹、大学生创业基金、创业比赛奖金、天使投资和政府创新创业扶持资金。")
    add_table(
        doc,
        "表 9 启动资金使用计划",
        ["用途", "金额（万元）", "占比", "说明"],
        [
            ["产品研发", "150", "50%", "前后端开发、AI 引擎、知识图谱、数据安全和测试。"],
            ["课程内容建设", "45", "15%", "首批课程知识点拆分、题库解析、教师顾问和内容审核。"],
            ["市场推广", "60", "20%", "校园大使、内容营销、试点活动和用户激励。"],
            ["云服务与运维", "30", "10%", "模型 API、GPU 实例、向量数据库、存储和带宽。"],
            ["法务财务与办公", "15", "5%", "公司注册、知识产权、合同、办公和基础行政。"],
        ],
        [4.0, 3.2, 2.4, 7.4],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_section(doc, "二、三年收入预测")
    add_table(
        doc,
        "表 10 三年收入预测",
        ["项目", "第 1 年", "第 2 年", "第 3 年"],
        [
            ["注册用户（万人）", "5", "50", "300"],
            ["Pro 订阅用户（万人）", "0.4", "4", "25"],
            ["C 端订阅收入（万元）", "80", "800", "5000"],
            ["B 端合作院校或院系（个）", "8", "50", "100"],
            ["B 端合作收入（万元）", "240", "2500", "7000"],
            ["其他增值收入（万元）", "0", "150", "800"],
            ["总收入（万元）", "320", "3450", "12800"],
        ],
        [5.0, 4.0, 4.0, 4.0],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_section(doc, "三、三年成本与盈利预测")
    add_table(
        doc,
        "表 11 三年成本与盈利预测",
        ["项目", "第 1 年", "第 2 年", "第 3 年"],
        [
            ["人力成本（万元）", "120", "650", "2400"],
            ["模型与云服务成本（万元）", "60", "520", "2600"],
            ["课程内容与数据建设（万元）", "50", "360", "1200"],
            ["市场推广费用（万元）", "100", "1000", "3600"],
            ["客服、办公及管理费用（万元）", "70", "520", "2100"],
            ["总成本（万元）", "400", "3050", "11900"],
            ["税前利润（万元）", "-80", "400", "900"],
        ],
        [5.0, 4.0, 4.0, 4.0],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_section(doc, "四、关键财务假设")
    add_text(doc, "财务测算基于以下假设：C 端 Pro 用户年均收入约 199 元；B 端合作从院系试点逐步升级到校级合作，平均合同金额随服务深度提升；模型推理成本随用户规模增长而上升，但通过缓存、模型路由和开源模型部署逐步优化；市场费用在第二、三年快速增加，以支持校园规模化扩张。")
    add_section(doc, "五、融资与退出机制")
    add_text(doc, "项目完成种子期验证后，可在第二年寻求 800 至 1500 万元 A 轮融资，用于课程扩张、B 端销售和算法团队建设。退出方式包括被在线教育企业、互联网平台或高校智慧教育服务商并购；若用户规模、收入和盈利能力持续增长，也可在中长期探索独立上市。")


def add_risk(doc):
    doc.add_page_break()
    add_chapter(doc, "第八章  风险控制与实施保障")
    add_section(doc, "一、技术风险")
    add_text(doc, "主要风险是大模型在专业课程场景中生成不准确答案。应对措施包括：使用课程知识库做检索增强；对答案标注来源和置信度；对高风险内容设置人工审核；建立用户反馈纠错机制；在考试题、证明题和实验报告等场景中强调学习辅导而非替代完成。")
    add_section(doc, "二、市场风险")
    add_text(doc, "大学生价格敏感，免费工具较多，付费转化存在不确定性。项目通过 Freemium 模式降低体验门槛，通过考前冲刺和深度诊断强化付费价值，并用 B 端高校合作平衡 C 端收入波动。")
    add_section(doc, "三、竞争风险")
    add_text(doc, "大厂和教育公司可能进入 AI 学习助手赛道。智学派的应对方式是尽早聚焦大学课程细分场景，建设课程级知识图谱和高校试点案例，用快速迭代和贴近学生的产品体验形成差异化。")
    add_section(doc, "四、合规风险")
    add_text(doc, "平台需要遵守《个人信息保护法》《网络安全法》和生成式人工智能相关管理要求。项目将坚持最小必要原则收集数据，敏感数据脱敏存储，用户可删除个人学习记录，教师端和学校端只展示必要的统计结果。")
    add_section(doc, "五、实施保障")
    add_text(doc, "项目实施采用阶段性里程碑管理。每个阶段设置明确验收指标，包括课程知识库完成度、答疑准确率、用户留存率、付费转化率和试点合作数量。团队每两周进行一次产品复盘，每月进行一次财务和运营复盘，确保项目计划能够根据真实反馈持续调整。")


def add_appendix(doc):
    doc.add_page_break()
    add_chapter(doc, "第九章  附录")
    add_section(doc, "一、首批课程知识图谱建设计划")
    add_text(doc, "数学类：高等数学（上、下）、线性代数、概率论与数理统计、离散数学。")
    add_text(doc, "计算机类：程序设计基础、数据结构、操作系统、计算机网络、数据库系统、计算机组成原理、算法设计与分析。")
    add_text(doc, "物理与电子类：大学物理、电路分析、模拟电子技术、数字电子技术、信号与系统。")
    add_text(doc, "经管与外语类：微观经济学、宏观经济学、管理学原理、会计学原理、大学英语四级与六级备考。")
    add_section(doc, "二、MVP 功能清单")
    add_table(
        doc,
        "表 12 MVP 功能清单",
        ["编号", "功能", "优先级", "预估工期"],
        [
            ["F01", "AI 文字问答", "P0", "4 周"],
            ["F02", "用户注册与登录", "P0", "1 周"],
            ["F03", "高数与线代学习路径生成", "P0", "3 周"],
            ["F04", "基础知识图谱入库", "P0", "3 周"],
            ["F05", "RAG 课件检索增强", "P1", "3 周"],
            ["F06", "知识点诊断模型", "P1", "4 周"],
            ["F07", "学习数据看板", "P1", "2 周"],
            ["F08", "拍照提问与 OCR", "P1", "3 周"],
            ["F09", "教师端班级报告", "P2", "4 周"],
        ],
        [2.4, 7.0, 3.0, 4.6],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_section(doc, "三、参考政策与规范")
    add_text(doc, "《中华人民共和国网络安全法》")
    add_text(doc, "《中华人民共和国个人信息保护法》")
    add_text(doc, "《生成式人工智能服务管理暂行办法》")
    add_text(doc, "《中国教育现代化 2035》")
    add_text(doc, "《新一代人工智能发展规划》")
    add_text(doc, "《教育强国建设规划纲要（2024－2035 年）》")


def add_final_work_division(doc):
    doc.add_page_break()
    add_chapter(doc, "团队成员分工表")
    add_text(doc, "本页用于满足课程要求中“最后一页标明每位成员分工”的提交规范。请在提交前将姓名、专业、学号和联系方式补充完整，并确保与封面团队成员信息一致。")
    add_table(
        doc,
        "表 13 团队成员分工",
        ["成员", "专业与学号", "联系电话", "主要分工", "成果说明"],
        [
            ["成员 1", "", "", "项目负责人：整体方案、商业模式、摘要撰写、进度统筹和路演表达。", "完成项目简介、运营模式和终稿统筹。"],
            ["成员 2", "", "", "技术负责人：AI 架构、RAG 流程、知识图谱、MVP 功能设计。", "完成产品说明、技术特点和研发路线。"],
            ["成员 3", "", "", "市场负责人：市场调研、用户画像、竞品分析和营销策略。", "完成市场分析、竞争分析和获客方案。"],
            ["成员 4", "", "", "财务负责人：资金需求、收入预测、成本测算和风险控制。", "完成财务分析、关键假设和风险章节。"],
            ["成员 5", "", "", "文档负责人：资料整理、排版规范、附录和课程提交检查。", "完成格式检查、参考资料和分工整理。"],
        ],
        [2.4, 3.2, 2.8, 5.2, 3.4],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    props = doc.core_properties
    props.title = "智学派 LearnGen 创业计划书"
    props.subject = "创新创业基础课程作业"
    props.keywords = "智学派; LearnGen; AI教育; 创业计划书"

    add_cover(doc)
    add_integrity_page(doc)
    add_toc(doc)
    add_project_intro(doc)
    add_product(doc)
    add_operation(doc)
    add_market(doc)
    add_sales(doc)
    add_company_team(doc)
    add_finance(doc)
    add_risk(doc)
    add_appendix(doc)
    add_final_work_division(doc)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
