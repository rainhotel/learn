from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor


BLACK = RGBColor(0, 0, 0)


def set_run_black(run):
    run.font.color.rgb = BLACK


def set_style_black(style):
    if style.type not in {WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER}:
        return
    style.font.color.rgb = BLACK


def blacken_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))

    for style in doc.styles:
        set_style_black(style)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_black(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_black(run)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                set_run_black(run)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                set_run_black(run)

    doc.save(str(output_path))


if __name__ == "__main__":
    input_path = Path(r"D:\moniC\project\learn\03-outputs\实验5_数据库系统_报告填写稿.docx")
    output_path = Path(r"D:\moniC\project\learn\03-outputs\实验5_数据库系统_报告填写稿_全黑.docx")
    blacken_docx(input_path, output_path)
    print(output_path)
