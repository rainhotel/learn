"""
Generate the properly formatted business plan DOCX following the template formatting rules:
- A4 paper, margins: top/bottom 2.5cm, left/right 2cm
- Line spacing 1.25x
- Body text first-line indent 2 characters (≈ 0.74cm for 小四/12pt)
- Heading hierarchy: 一/二/三 → (一)/(二)/(三) → 1/2/3 → 1.1/1.2
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_line_spacing(paragraph, spacing=1.25):
    """Set line spacing to a multiple (e.g., 1.25x)"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing_elem = pPr.find(qn('w:spacing'))
    if spacing_elem is None:
        spacing_elem = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing_elem)
    spacing_elem.set(qn('w:line'), str(int(spacing * 240)))
    spacing_elem.set(qn('w:lineRule'), 'auto')


def add_heading_custom(doc, text, level=1, font_name='黑体', font_size=None, bold=True):
    """Add a heading with custom formatting matching Chinese document standards."""
    p = doc.add_paragraph()

    # Set font sizes based on heading level
    if font_size is None:
        if level == 0:  # Title
            font_size = Pt(22)
        elif level == 1:  # 一、二、三 chapter titles
            font_size = Pt(16)
        elif level == 2:  # (一)(二)(三) section titles
            font_size = Pt(14)
        elif level == 3:  # 1. 2. 3. subsection
            font_size = Pt(12)
        else:
            font_size = Pt(12)

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold

    # Heading alignment: title centered, others left
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    set_line_spacing(p, 1.25)

    # Add some spacing before/after
    pPr = p._element.get_or_add_pPr()
    spacing_elem = pPr.find(qn('w:spacing'))
    if spacing_elem is None:
        spacing_elem = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing_elem)

    if level == 0:
        spacing_elem.set(qn('w:before'), '480')
        spacing_elem.set(qn('w:after'), '360')
    elif level == 1:
        spacing_elem.set(qn('w:before'), '360')
        spacing_elem.set(qn('w:after'), '200')
    elif level == 2:
        spacing_elem.set(qn('w:before'), '240')
        spacing_elem.set(qn('w:after'), '120')
    else:
        spacing_elem.set(qn('w:before'), '120')
        spacing_elem.set(qn('w:after'), '60')

    return p


def add_body_text(doc, text, indent=True, font_name='宋体', font_size=Pt(12)):
    """Add body text with first-line indent and proper formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

    set_line_spacing(p, 1.25)

    # First-line indent: 2 characters ≈ 24pt for 12pt font, or about 0.84cm
    if indent:
        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(ind)
        # 2 characters at 12pt = 24pt = ~0.85cm
        ind.set(qn('w:firstLine'), '480')  # ~24pt in twips
        ind.set(qn('w:firstLineChars'), '200')  # 2 characters

    return p


def add_bullet_text(doc, text, level=0):
    """Add bullet-point style text."""
    p = doc.add_paragraph()
    prefix = "● " if level == 0 else "  ○ "
    run = p.add_run(prefix + text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)

    set_line_spacing(p, 1.25)

    # Indent bullets
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr.append(ind)
    indent_val = 720 + level * 360
    ind.set(qn('w:left'), str(indent_val))
    ind.set(qn('w:hanging'), '360')

    return p


def add_empty_line(doc):
    """Add an empty line."""
    p = doc.add_paragraph()
    run = p.add_run('')
    run.font.size = Pt(8)
    set_line_spacing(p, 1.0)
    return p


def build_table(doc, headers, rows, col_widths=None):
    """Build a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10)
        run.bold = True
        set_line_spacing(p, 1.25)
        # Gray background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            # Center-align short content, left-align long
            if len(str(text)) < 20:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            set_line_spacing(p, 1.25)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generate_docx():
    doc = Document()

    # ==================== PAGE SETUP ====================
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)  # A4
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ==================== STYLES ====================
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pPr = style._element.get_or_add_pPr()
    sp = parse_xml(f'<w:spacing {nsdecls("w")} w:line="300" w:lineRule="auto"/>')
    pPr.append(sp)

    # ==================== COVER PAGE ====================
    # Add multiple blank lines for centering effect
    for _ in range(6):
        add_empty_line(doc)

    # Company/Project name
    add_heading_custom(doc, '智学派（LearnGen）', level=0, font_size=Pt(26))

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI 驱动的个性化学习辅助平台')
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(89, 89, 89)
    set_line_spacing(p, 1.25)

    add_empty_line(doc)
    add_empty_line(doc)

    add_heading_custom(doc, '商 业 计 划 书', level=0, font_size=Pt(28))

    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)
    add_empty_line(doc)

    # Cover info
    cover_items = [
        '编制时间：2026 年 6 月',
        '联 系 人：___________',
        '电　　话：___________',
        '电子邮箱：___________',
    ]
    for item in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        set_line_spacing(p, 1.5)

    # Page break
    doc.add_page_break()

    # ==================== CONFIDENTIALITY NOTICE ====================
    add_heading_custom(doc, '保密须知', level=1)
    add_body_text(doc, '本商业计划书属于商业机密，所有权属于智学派项目团队。所涉及的内容和资料仅限于已签署投资意向的投资者使用。收到本计划书后，收件人应即刻确认，并遵守以下规定：')
    add_body_text(doc, '（1）若收件人不希望参与本计划书所述项目，请按上述地址尽快将本计划书完整退回；')
    add_body_text(doc, '（2）在没有取得智学派项目团队书面同意前，收件人不得将本计划书全部或部分地予以复制、传递给他人、影印、泄露或散布给他人；')
    add_body_text(doc, '（3）应以对待贵公司机密资料一样的态度对待本计划书所提供的所有商业秘密资料。')
    add_body_text(doc, '本商业计划书不可用作销售报价使用，也不可用作购买时的报价使用。')

    add_empty_line(doc)
    add_body_text(doc, '商业计划书编号：2026-001', indent=False)
    add_body_text(doc, '签　　字：___________', indent=False)
    add_body_text(doc, '公　　司：___________', indent=False)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading_custom(doc, '目　录', level=1)
    add_empty_line(doc)

    toc_entries = [
        '第一章　摘　要',
        '第二章　公司概述',
        '　　一、宗旨与使命',
        '　　二、公司简介',
        '　　三、公司战略',
        '　　四、核心技术',
        '　　五、价值主张',
        '　　六、公司管理',
        '　　七、组织与协作关系',
        '　　八、知识产权策略',
        '　　九、场地与设施',
        '　　十、风险分析',
        '第三章　市场分析',
        '　　一、市场介绍',
        '　　二、目标市场',
        '　　三、顾客购买准则',
        '　　四、销售策略',
        '　　五、市场渗透与增长',
        '第四章　竞争分析',
        '第五章　产品与服务',
        '　　一、核心产品模块',
        '　　二、延展产品方向',
        '　　三、产品路线图',
        '第六章　营销与销售',
        '　　一、获客策略',
        '　　二、用户留存与转化',
        '　　三、品牌定位',
        '第七章　财务计划',
        '　　一、初期融资需求',
        '　　二、收入预测',
        '　　三、成本与利润预测',
        '　　四、关键财务假设',
        '　　五、退出机制',
        '第八章　附　录',
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if entry.startswith('第'):
            run.font.size = Pt(12)
            run.bold = True
        else:
            run.font.size = Pt(11)
        set_line_spacing(p, 1.5)

    doc.add_page_break()

    # ==================== CHAPTER 1: EXECUTIVE SUMMARY ====================
    add_heading_custom(doc, '第一章　摘　要', level=1)
    add_empty_line(doc)

    add_body_text(doc, '智学派（LearnGen）是一个面向大学生的 AI 个性化学习助手，核心功能包括 AI 驱动的学习路径生成、基于课程知识库的智能答疑以及知识点薄弱环节精准诊断。产品以大语言模型（LLM）结合检索增强生成（RAG）技术为基础，自建高校主流课程知识图谱，为学生提供比通用 AI 工具更准确、更懂课程的专属学习辅导。')
    add_body_text(doc, '中国高等教育在校生规模已超 4000 万，在线教育与教育科技市场规模突破 5000 亿元人民币，AI+教育赛道年复合增长率超过 30%。与此同时，传统教育模式下优质师资稀缺、一对一家教费用高昂、慕课/录播课完课率不足 5%，学生高效个性化学习的需求远未被满足。智学派以"Freemium + 高校合作"双引擎模式切入市场：C 端 Pro 版定价 29.9 元/月，B 端高校合作版定价 50-100 万元/校/年。')
    add_body_text(doc, '项目启动资金需求为 300 万元人民币，主要用于产品研发（60%）、市场推广（25%）和运营储备（15%）。预计三年内覆盖 1000 所高校，注册用户突破 300 万，第三年实现盈亏平衡。本轮融资为 Pre-A 轮，拟出让 10%-15% 股权。')

    doc.add_page_break()

    # ==================== CHAPTER 2: COMPANY OVERVIEW ====================
    add_heading_custom(doc, '第二章　公司概述', level=1)
    add_empty_line(doc)

    # 2.1
    add_heading_custom(doc, '一、宗旨与使命', level=2)
    add_body_text(doc, '愿景：成为 Z 世代首选 AI 学习平台，让每个学生都拥有自己的 AI 导师。')
    add_body_text(doc, '使命：利用大模型技术打破教育资源不均的壁垒，为每一位大学生提供普惠且高质量的个性化学习体验，用技术的力量让"因材施教"从理念走向现实。')

    # 2.2
    add_heading_custom(doc, '二、公司简介', level=2)
    add_body_text(doc, '智学派教育科技有限公司（LearnGen EduTech Co., Ltd.）是一家专注 AI+教育领域的科技创业公司，由华中科技大学计算机学院学生团队发起创立。公司致力于将前沿大语言模型技术与高校教学场景深度融合，打造懂课程、懂学生、懂考试的 AI 学习助手。')
    add_body_text(doc, '公司总部拟设在武汉光谷，充分利用光谷科创大走廊的政策扶持、高校人才资源和相对较低的运营成本。初期团队规模 5-8 人，核心成员来自计算机科学、软件工程、教育学等专业方向。')

    # 2.3
    add_heading_custom(doc, '三、公司战略', level=2)

    add_heading_custom(doc, '（一）产品与服务战略', level=3)
    add_body_text(doc, '以 AI 引擎为核心壁垒，围绕高校课程体系构建细粒度知识图谱和高质量课程向量库，确保答疑的准确性和课程相关性。产品采用"一横一纵"策略：横向覆盖更多高校和学科，纵向在每门课程上做到深度覆盖。')

    add_heading_custom(doc, '（二）市场战略', level=3)
    add_body_text(doc, '采用 B2B2C 模式切入：先与高校教师和院系建立合作关系，以免费试用撬动种子用户，再通过课程群裂变和学习打卡等社交机制实现 C 端规模化增长。初期聚焦 985/211 高校的理工科学生（约 1500 万目标用户），验证产品和商业模式后逐步拓展至全部本科和高职高专院校。')

    add_heading_custom(doc, '（三）客户服务与增值业务', level=3)
    add_body_text(doc, '基础 AI 答疑和学习路径规划免费开放，高级功能（无限答疑、深度诊断报告、个性化备考方案、课件智能解析等）通过 Pro 订阅收费。中长期拓展至考研、考公、考证等强付费意愿的备考场景，提供增值内容服务和 1v1 AI 辅导。')

    add_heading_custom(doc, '（四）技术战略', level=3)
    add_body_text(doc, '核心技术路线是基于开源大模型进行微调（Fine-tuning），结合自建课程知识库做 RAG 增强检索，在降低 API 调用成本的同时提升回答质量。构建"数据飞轮"：用户学习行为数据反哺模型优化，模型越精准，用户粘性越高，数据积累越多。')

    # 2.4
    add_heading_custom(doc, '四、核心技术', level=2)

    add_heading_custom(doc, '（一）大模型 + RAG 混合架构', level=3)
    add_body_text(doc, '底层接入 GPT-4o、Claude Opus 等主流大模型 API 作为基础推理引擎，同时基于 LangChain 框架搭建 RAG 管线。学生上传的课件、教材、笔记经向量化处理后存入 Milvus 向量数据库，查询时先检索相关知识点片段，再与大模型生成结果融合输出，有效解决大模型在具体课程内容上的"幻觉"问题。')

    add_heading_custom(doc, '（二）课程知识图谱', level=3)
    add_body_text(doc, '自建覆盖高等数学、线性代数、大学物理、数据结构、操作系统等 30+ 门核心课程的知识图谱。每个知识点标注前置依赖、难度层级、常见错误类型和关联考题。基于 Neo4j 图数据库存储和查询，支撑学习路径自动规划和薄弱环节诊断。')

    add_heading_custom(doc, '（三）学习路径推荐算法', level=3)
    add_body_text(doc, '结合知识图谱拓扑结构和学生历史答题数据，采用改进的贝叶斯知识追踪（BKT）模型评估每个知识点的掌握概率，动态生成个性化学习路径。推荐学习顺序遵循"先修知识点→目标知识点→高级拓展"的依赖关系，确保学习过程有逻辑、有节奏。')

    add_heading_custom(doc, '（四）技术壁垒', level=3)
    add_body_text(doc, '数据飞轮效应：用户规模增长带来学习行为数据积累，模型迭代加速，形成正向循环。课程知识图谱：构建成本高、维护难度大，非通用模型短期可复制。垂直场景优化：针对大学课程做了大量 Prompt Engineering 和 Fine-tuning，通用工具无法直接替代。')

    # 2.5
    add_heading_custom(doc, '五、价值主张', level=2)

    add_heading_custom(doc, '（一）对于学生', level=3)
    add_body_text(doc, '比辅导班便宜 90%，比自学高效 3 倍；7×24 小时即时答疑，不依赖老师的时间；每个学生都有自己的专属学习路径，告别"学不懂不知道从哪补"。')

    add_heading_custom(doc, '（二）对于高校', level=3)
    add_body_text(doc, '提升教学质量的数据化工具，精准定位班级知识薄弱点；降低挂科率和学业预警率；与教务系统对接，为教学改革提供数据支撑。')

    add_heading_custom(doc, '（三）对于投资者', level=3)
    add_body_text(doc, '进入 5000 亿在线教育赛道的 AI 入口；清晰的 Freemium + SaaS 双引擎商业模式；三年内有望覆盖千万级用户，具备独立上市或被巨头收购的退出路径。')

    # 2.6
    add_heading_custom(doc, '六、公司管理', level=2)

    add_body_text(doc, '核心团队（初期 5-8 人）：', indent=False)
    add_empty_line(doc)

    team_headers = ['角色', '职责', '背景要求']
    team_rows = [
        ['CEO/创始人', '公司战略、融资、对外合作', '计算机专业，有学生创业或竞赛经历'],
        ['CTO', '技术架构、AI 引擎研发', '软件工程方向，熟悉 LLM 和推荐系统'],
        ['CPO', '产品设计与用户体验', '有产品设计经验'],
        ['COO', '日常运营、高校渠道', '有校园活动组织经验'],
        ['教育顾问（兼职）', '课程内容把关、教育方法论', '教育学院教师或研究生'],
    ]
    build_table(doc, team_headers, team_rows)

    add_empty_line(doc)
    add_body_text(doc, '股权结构：创始人团队持股 80%，预留 20% 用于后续融资和员工期权池。')
    add_body_text(doc, '外部支持：合作律师事务所提供知识产权和公司法务支持；合作会计师事务所提供财务规范和审计服务；拟入驻学校创业孵化器或光谷创业园区。')
    add_body_text(doc, '董事会：初期由创始人团队 3 人 + 高校创业导师 1 人 + 外部行业顾问 1 人组成 5 人董事会，为公司战略决策提供指导。')

    # 2.7
    add_heading_custom(doc, '七、组织与协作关系', level=2)

    add_body_text(doc, '内部组织结构：CEO 下设 CTO、CPO、COO 三条线。CTO 负责 AI 研发、后端、前端；CPO 负责产品设计；COO 负责运营、市场、渠道。初期扁平化管理，随团队扩张逐步完善中层管理架构。')
    add_body_text(doc, '外部协作关系：与本校及周边高校的教务处、信息中心建立合作，推进产品试点；与高校教师合作开发课程知识图谱，确保内容权威性；使用阿里云/腾讯云算力资源，降低初期基础设施成本；与高校学生会、学习类社团合作进行校园推广。')

    # 2.8
    add_heading_custom(doc, '八、知识产权策略', level=2)
    add_body_text(doc, '软件著作权：对智学派 App、Web 平台、管理后台等软件系统进行著作权登记（3 项以上）。')
    add_body_text(doc, '商标保护：对"智学派""LearnGen"品牌名称及 Logo 在第 9 类（软件）、第 41 类（教育服务）、第 42 类（技术服务）进行商标注册。')
    add_body_text(doc, '专利布局：对学习路径推荐算法、知识图谱构建方法、RAG 课程问答系统三项核心技术申请发明专利（3 项）。')
    add_body_text(doc, '商业秘密：课程知识图谱数据、用户学习行为分析模型、核心 Prompt 模板等作为商业秘密保护。')
    add_body_text(doc, '开源合规：使用开源 LLM 和框架严格遵守相关许可协议。')

    # 2.9
    add_heading_custom(doc, '九、场地与设施', level=2)
    add_body_text(doc, '公司总部拟设在武汉光谷创业园区，初期租赁 80-120 平方米办公空间，满足 8-10 人团队办公需求。光谷地区拥有成熟的科创生态、便利的高校人才触达渠道和相对优惠的办公租金（约 60-80 元/㎡/月），是科技创业公司的理想起步地点。')
    add_body_text(doc, '核心 IT 设施采用云服务模式，无需自建机房：阿里云 GPU 实例用于模型微调和推理，OSS 对象存储 + 云数据库，CDN 加速保障全国用户访问体验。初期月度云服务成本预估 2-3 万元。')

    # 2.10
    add_heading_custom(doc, '十、风险分析', level=2)

    add_heading_custom(doc, '（一）技术风险：大模型幻觉问题', level=3)
    add_body_text(doc, 'LLM 在专业课程内容上可能生成似是而非的错误答案，影响用户信任。应对措施：RAG 检索增强 + 课程知识库约束生成范围 + 关键内容人工审核机制 + 答案附置信度标注。')

    add_heading_custom(doc, '（二）竞争风险：巨头入场', level=3)
    add_body_text(doc, '百度、字节跳动等大厂已在布局 AI+教育产品。应对措施：抢占细分赛道先发优势（专注大学课程而非 K12）；通过知识图谱和用户数据构建壁垒；高校渠道深度绑定。')

    add_heading_custom(doc, '（三）市场风险：用户付费意愿不足', level=3)
    add_body_text(doc, '大学生群体消费能力有限，C 端付费转化存在不确定性。应对措施：Freemium 模式降低体验门槛；高校合作版（B 端收入）平衡 C 端不确定性；优先切入考研、考公等高付费意愿场景。')

    add_heading_custom(doc, '（四）运营风险：用户增长不及预期', level=3)
    add_body_text(doc, '应对措施：以 B2B2C 策略保障初始用户基本盘；通过课程群裂变和学习打卡实现低成本获客；多轮融资保障资金链，适当延长烧钱周期。')

    add_heading_custom(doc, '（五）法律与合规风险', level=3)
    add_body_text(doc, '应对措施：严格遵守《网络安全法》《个人信息保护法》《生成式人工智能服务管理暂行办法》；用户数据脱敏处理；建立内容安全审核机制，防止有害内容生成。')

    doc.add_page_break()

    # ==================== CHAPTER 3: MARKET ANALYSIS ====================
    add_heading_custom(doc, '第三章　市场分析', level=1)
    add_empty_line(doc)

    add_heading_custom(doc, '一、市场介绍', level=2)
    add_body_text(doc, '中国在线教育市场是一个规模庞大且持续高速增长的赛道。据艾瑞咨询和前瞻产业研究院数据，2025 年中国在线教育市场规模已突破 5000 亿元人民币，预计 2028 年将超过 8000 亿元。其中 AI+教育细分领域增速尤为显著，年复合增长率超过 30%。')
    add_body_text(doc, '在政策层面，教育部持续推动教育数字化转型，《中国教育现代化 2035》明确提出"加快信息化时代教育变革"。国务院《新一代人工智能发展规划》将智能教育列为重点应用场景。2025 年《教育强国建设规划纲要》进一步强调利用人工智能促进个性化学习和教育公平。政策的持续加码为 AI+教育创业提供了有力的宏观支撑。')
    add_body_text(doc, '在技术层面，2024-2025 年大语言模型（LLM）技术实现质的飞跃，GPT-4o、Claude Opus、DeepSeek-R1 等先进模型在知识问答、逻辑推理、文本理解等能力上已接近甚至超越人类专家水平，为 AI 个性化教育从概念走向产品奠定了坚实的技术基础。')
    add_body_text(doc, '在市场痛点层面，中国高等教育长期面临三大矛盾：一是优质师资稀缺且分布不均，双一流高校和普通院校师资差距悬殊；二是传统辅导班和一对一家教价格昂贵（动辄数千元/学期），大部分学生无法负担；三是 MOOC 和录播课虽免费或低价，但缺乏互动和个性化指导，平均完课率不足 5%。这些痛点为 AI 驱动的个性化学习解决方案创造了巨大的市场空间。')

    add_heading_custom(doc, '二、目标市场', level=2)
    add_body_text(doc, '目标用户画像：在校大学生（含本科和高职高专），全国约 4000 万人。初期聚焦 985/211 及一本高校的理工科学生，约 1500 万人。年龄 18-24 岁，智能手机重度用户，对 AI 工具接受度高。有考试通过、绩点提升、考研备考等明确学习需求。')

    add_empty_line(doc)
    add_body_text(doc, '市场规模测算：', indent=False)
    add_empty_line(doc)

    mkt_headers = ['维度', '数据']
    mkt_rows = [
        ['中国高等教育在校生', '4000 万+'],
        ['理工科在校生', '约 1500 万'],
        ['目标渗透率（Year 3）', '20%（300 万）'],
        ['付费转化率', '15%（45 万）'],
        ['ARPU（年）', '299 元（Pro 订阅）'],
        ['C 端年收入潜力', '约 1.35 亿元'],
        ['B 端高校合作收入（100 所 × 70 万）', '约 7000 万元'],
        ['合计年收入潜力', '约 2 亿元'],
    ]
    build_table(doc, mkt_headers, mkt_rows)

    add_empty_line(doc)
    add_body_text(doc, '市场增长驱动因素：大模型推理成本持续下降，AI 产品边际成本不断降低；大学生对 AI 工具的使用习惯已在 2024-2025 年快速养成（ChatGPT/Kimi/豆包等普及）；高校教学改革对数字化工具的采购意愿和预算逐年增加。')

    add_heading_custom(doc, '三、顾客购买准则', level=2)
    add_body_text(doc, '大学生的教育产品消费决策具有以下特征：免费优先——愿意先试用免费版本，体验后再决定是否付费；口碑驱动——高度依赖同学推荐、课程群讨论和社交媒体种草（B 站、小红书等）；效果导向——付费的核心驱动力是"能帮我通过考试""能让我学懂这门课"；价格敏感——月付 30 元以内是可接受的心理价位，超过 50 元会大幅降低转化率；便利性——手机端操作方便、随时随地可用是刚需。')
    add_body_text(doc, '针对以上特征，智学派采用"免费试用 → 深度转化 → 社群裂变"的漏斗模型，以产品力驱动自然增长。')

    add_heading_custom(doc, '四、销售策略', level=2)
    add_body_text(doc, 'C 端销售：基础功能永久免费，每日 5 次 AI 答疑 + 基础学习路径生成；Pro 版 29.9 元/月（无限答疑 + 深度诊断 + 课件解析），年卡 199 元；新用户首月 Pro 体验价 9.9 元；3 人拼团享 8 折，刺激裂变传播。')
    add_body_text(doc, 'B 端销售：高校合作版 50-100 万/年，含全校学生账号 + 教学管理面板 + API 对接；与高校信息化建设预算对接，走学校采购流程；先从院系试点（10-20 万/年），验证效果后推广至全校。')
    add_body_text(doc, '销售渠道包括：应用商店（App Store、各大安卓市场）；高校课程群、学习交流群（课程群精准触达）；B 站/小红书/抖音 KOL 合作种草；校园代理/校园大使推广体系。')

    add_heading_custom(doc, '五、市场渗透与增长', level=2)
    add_body_text(doc, 'Year 1（种子期）：覆盖 50 所高校，累计注册用户 5 万人，Pro 订阅用户 3000-5000 人，B 端合作 5-10 所。重点打磨产品体验，验证 PMF（产品市场契合度）。')
    add_body_text(doc, 'Year 2（增长期）：覆盖 300 所高校，累计注册用户 50 万人，Pro 订阅用户 3-5 万人，B 端合作 50 所。启动校园大使体系和内容营销矩阵，实现规模化增长。')
    add_body_text(doc, 'Year 3（扩张期）：覆盖 1000 所高校，累计注册用户 300 万人，Pro 订阅用户 20-30 万人，B 端合作 100 所。启动考研/考公/考证等增值业务，探索第二增长曲线。')

    doc.add_page_break()

    # ==================== CHAPTER 4: COMPETITIVE ANALYSIS ====================
    add_heading_custom(doc, '第四章　竞争分析', level=1)
    add_empty_line(doc)

    add_body_text(doc, '当前 AI+教育赛道竞争可分为四个层级：')
    add_body_text(doc, '第一层——通用 AI 对话工具：代表产品包括 ChatGPT、Kimi、豆包、通义千问。优势在于品牌知名度高、基础能力强大、用户基数庞大；劣势在于不懂具体课程内容、产生幻觉、无学习路径规划、无知识诊断。这类产品是用户的"基础替代品"，但无法满足深度学习需求，类似于"计算器 vs 数学老师"的关系。')
    add_body_text(doc, '第二层——传统在线教育平台：代表产品包括超星学习通、智慧树、中国大学 MOOC。优势在于高校渠道深厚、课程资源丰富、已被学校强制使用；劣势在于产品体验老旧、无 AI 能力、学生被动使用而非主动使用、完课率极低。这类平台渠道强势但产品弱势，是智学派 B2B2C 谈判中的对标替代对象。')
    add_body_text(doc, '第三层——搜题/题库工具：代表产品包括作业帮、学小易、考途。优势在于题库积累深厚、用户习惯养成；劣势在于给学生的是"答案"而非"学会"、高校层面被抵制、缺乏 AI 个性化能力。这类产品在搜题场景有竞争，但智学派的核心价值是"帮你学会"而非"给你答案"。')
    add_body_text(doc, '第四层——新兴 AI+教育创业公司：代表产品包括学而思九章大模型、网易有道 AI、好未来 MathGPT。优势在于教育行业积累深、品牌信任度高；劣势在于多为 K12 方向，高校市场覆盖不足，大厂体制创新速度受限。这类公司是最直接的潜在竞争者，但目前主战场在 K12，与我们在高校赛道形成错位竞争。')

    add_empty_line(doc)
    add_body_text(doc, '我们的竞争优势：', indent=False)
    add_body_text(doc, '课程级知识图谱壁垒：自建 30+ 门核心课程的高细粒度知识图谱，非通用大模型训练数据可覆盖，复刻成本高。')
    add_body_text(doc, '先发优势：率先聚焦大学课程 AI 学习这一细分赛道，抢占用户心智。')
    add_body_text(doc, '数据飞轮：用户规模增长带来专属学习行为数据，模型越用越准，形成正向循环。')
    add_body_text(doc, '双引擎模式：C 端付费 + B 端高校合作，收入结构稳健，不至于因 C 端增长不及预期而崩盘。')
    add_body_text(doc, '团队灵活度：小团队决策快、迭代快，能与用户保持最近距离。')

    add_empty_line(doc)
    add_body_text(doc, '竞品详细对比：', indent=False)
    add_empty_line(doc)

    comp_headers = ['维度', '智学派', 'ChatGPT', '超星学习通', '作业帮', '学而思九章']
    comp_rows = [
        ['目标用户', '大学生', '全人群', '大学生', 'K12+大学', 'K12'],
        ['AI 个性化', '✓', '✗', '✗', '部分', '✓'],
        ['课程知识图谱', '✓', '✗', '✗', '✗', '✓（K12）'],
        ['学习路径规划', '✓', '✗', '✗', '✗', '部分'],
        ['知识点诊断', '✓', '✗', '✗', '✗', '✓'],
        ['免费使用', '✓', '✓', '✓', '✓', '部分'],
        ['高校渠道', 'B2B2C', '无', '强（学校采购）', '弱', '中'],
    ]
    build_table(doc, comp_headers, comp_rows)

    doc.add_page_break()

    # ==================== CHAPTER 5: PRODUCTS & SERVICES ====================
    add_heading_custom(doc, '第五章　产品与服务', level=1)
    add_empty_line(doc)

    add_heading_custom(doc, '一、核心产品模块', level=2)

    add_heading_custom(doc, '（一）个性化学习路径生成', level=3)
    add_body_text(doc, '学生输入当前课程名称和学习目标（如"期末考试 85+"或"高数上不挂科"），系统基于课程知识图谱自动生成个性化学习计划。学习路径以可视化知识树形式展示，标注每个知识点的建议学习时长、重要程度和前置依赖关系。学习进度实时追踪，已掌握和未掌握的知识点一目了然。')

    add_heading_custom(doc, '（二）AI 智能答疑', level=3)
    add_body_text(doc, '支持文字、语音、拍照三种方式提问。传统搜题工具只给答案，智学派 AI 给出的是"解题思路 + 分步讲解 + 同类题推荐"。系统基于 RAG 技术关联课程教材和课件，所有回答均标明知识来源，学生可一键跳转到相关知识点深入学习。对于复杂问题（如证明题、分析题），AI 提供逐步引导式解答而非直接给答案。')

    add_heading_custom(doc, '（三）知识点诊断与薄弱环节定位', level=3)
    add_body_text(doc, '学生完成 AI 生成的针对性练习题后，系统基于贝叶斯知识追踪模型评估每个知识点的掌握概率，生成个人知识薄弱点热力图。诊断结果精准到知识点层级，而非笼统的"第一章较弱"。针对薄弱知识点自动推送讲解视频、同类练习题和强化学习任务。')

    add_heading_custom(doc, '（四）学习数据看板', level=3)
    add_body_text(doc, '以可视化图表展示学习时长、每日答题数、正确率变化趋势、知识点覆盖率、学习进度完成百分比等关键指标。支持周报/月报自动生成，对比同班级/同校同学的匿名学习数据，给予适度的社交激励。')

    add_heading_custom(doc, '二、延展产品方向', level=2)
    add_body_text(doc, '考前冲刺训练营：期末考试前 2 周开设 AI 驱动的考前集训模式，包含高频考点串讲、模拟试卷 AI 批改、考前预测题推送。单独定价 19.9 元/次。')
    add_body_text(doc, '考研衔接 Pro：面向大三学生，提供考研公共课（数学、英语、政治）的 AI 学习路径规划 + 智能刷题 + 院校专业智能推荐。单独定价 39.9 元/月。')
    add_body_text(doc, '笔记社区：学生可分享 AI 辅助整理的学习笔记，优质笔记获得曝光和收益分成，形成 UGC 内容生态。')

    add_heading_custom(doc, '三、产品路线图', level=2)

    roadmap_headers = ['时间', '里程碑']
    roadmap_rows = [
        ['2026 Q3', 'MVP 上线：AI 答疑（文字）+ 高数/线代基础知识图谱 + 学习路径生成'],
        ['2026 Q4', '拍照提问、语音提问、课件上传解析，知识图谱扩展至 10 门课'],
        ['2027 Q1', '知识点诊断与薄弱环节定位上线，B 端高校管理平台上线'],
        ['2027 Q3', '知识图谱扩展至 30 门课，考研衔接 Pro 版上线'],
        ['2028 Q1', '笔记社区功能上线，启动国际版探索'],
    ]
    build_table(doc, roadmap_headers, roadmap_rows)

    doc.add_page_break()

    # ==================== CHAPTER 6: MARKETING & SALES ====================
    add_heading_custom(doc, '第六章　营销与销售', level=1)
    add_empty_line(doc)

    add_heading_custom(doc, '一、获客策略', level=2)
    add_body_text(doc, 'B2B2C 冷启动：与本校及周边合作高校的教务处或任课教师合作，以"AI 助力课程学习"的名义免费向学生开放试点。教师可在课堂上推荐学生使用，通过课程群发布邀请码，一个班级一个班级地积累种子用户。')
    add_body_text(doc, '课程群裂变：大学生学习行为高度依赖课程微信群/QQ 群。设计"邀请 3 人得 3 天 Pro 会员"裂变机制，利用课程群的自然传播效应实现低成本获客。核心逻辑是一个人在群里问了 AI 答得好，群友看到后自然就会下载。')
    add_body_text(doc, '校园大使计划：在目标高校招募校园大使（每校 3-5 人），负责校内推广和用户拉新。校园大使获得免费 Pro 会员 + 拉新佣金（每新用户 3-5 元）+ 实习证明。预计 Year 2 覆盖 200 所高校的校园大使网络。')
    add_body_text(doc, '内容营销与 KOL 合作：在 B 站、小红书投放 AI 学习相关短视频内容，如"用 AI 把高数从 60 学到 90""期末考试 AI 救命指南"等。与学习类 UP 主/博主合作产品测评和种草，目标单条视频 5-10 万播放量。')
    add_body_text(doc, '应用商店优化（ASO）：针对"大学高数""AI 答疑""期末考试""考研"等关键词做应用商店搜索优化。')

    add_heading_custom(doc, '二、用户留存与转化', level=2)
    add_body_text(doc, '新手引导：首日完成"学习 1 个知识点 + 问 1 个问题 + 做 3 道题"的 Aha Moment 引导，提升次日留存。')
    add_body_text(doc, '学习打卡：每日学习打卡获积分，连续 7 天打卡送 1 天 Pro 会员，21 天送 7 天。')
    add_body_text(doc, 'Push 推送：考试倒计时提醒、薄弱知识点复习提醒（非骚扰式，每天不超过 2 条）。')
    add_body_text(doc, '免费转付费：Pro 功能每日限免体验（如每 3 天可免费做 1 次知识点诊断），让用户养成使用习惯后自然付费。')

    add_heading_custom(doc, '三、品牌定位', level=2)
    add_body_text(doc, '口号："你的 AI 学长，24 小时在线"')
    add_body_text(doc, '品牌调性：年轻、科技感、陪伴式、懂学生。不搞"教育焦虑"营销，强调"学习本可以更轻松"。')

    doc.add_page_break()

    # ==================== CHAPTER 7: FINANCIAL PLAN ====================
    add_heading_custom(doc, '第七章　财务计划', level=1)
    add_empty_line(doc)

    add_heading_custom(doc, '一、初期融资需求', level=2)
    add_body_text(doc, '本次融资为 Pre-A 轮，融资目标 300 万元人民币，拟出让股权 10%-15%，投后估值 2000-3000 万元。')

    add_empty_line(doc)

    fund_headers = ['用途', '金额（万元）', '占比']
    fund_rows = [
        ['产品研发（AI 引擎、知识图谱、前后端）', '180', '60%'],
        ['市场推广（校园大使、KOL 投放、ASO）', '75', '25%'],
        ['运营储备（办公、云服务、法务财务）', '45', '15%'],
    ]
    build_table(doc, fund_headers, fund_rows)

    add_empty_line(doc)

    add_heading_custom(doc, '二、收入预测', level=2)

    rev_headers = ['项目', 'Year 1', 'Year 2', 'Year 3']
    rev_rows = [
        ['注册用户（万）', '5', '50', '300'],
        ['Pro 订阅用户（万）', '0.4', '4', '25'],
        ['C 端收入（万元）', '80', '800', '5,000'],
        ['B 端合作院校（所）', '8', '50', '100'],
        ['B 端收入（万元）', '240', '2,500', '7,000'],
        ['总收入（万元）', '320', '3,300', '12,000'],
    ]
    build_table(doc, rev_headers, rev_rows)

    add_empty_line(doc)

    add_heading_custom(doc, '三、成本与利润预测', level=2)

    cost_headers = ['项目', 'Year 1', 'Year 2', 'Year 3']
    cost_rows = [
        ['人力成本（万元）', '60', '180', '400'],
        ['云服务成本（万元）', '30', '150', '500'],
        ['市场推广（万元）', '75', '300', '800'],
        ['办公及其他（万元）', '35', '80', '200'],
        ['总成本（万元）', '200', '710', '1,900'],
        ['毛利（万元）', '120', '2,590', '10,100'],
        ['净利润（万元）', '-80', '-150', '约 0（盈亏平衡）'],
    ]
    build_table(doc, cost_headers, cost_rows)
    add_body_text(doc, '注：Year 3 目标为实现盈亏平衡，若市场增长超预期则有望实现盈利。', indent=False)

    add_empty_line(doc)

    add_heading_custom(doc, '四、关键财务假设', level=2)
    add_body_text(doc, '单个 Pro 用户月度 API 调用成本约 3-5 元（随模型降价逐年下降）。')
    add_body_text(doc, '获客成本（CAC）约 8-12 元/新用户。')
    add_body_text(doc, '用户生命周期价值（LTV）约 180 元。')
    add_body_text(doc, 'LTV/CAC 比率 > 15:1，处于健康水平。')
    add_body_text(doc, 'B 端合作均价 70 万元/校/年。')

    add_heading_custom(doc, '五、退出机制', level=2)
    add_body_text(doc, '并购退出：被在线教育头部企业（如网易有道、作业帮、好未来）或互联网大厂（字节跳动、百度）收购。估值参考同类 AI 教育公司并购案例，预期 3-5 年可实现 3-5 亿元级退出。')
    add_body_text(doc, '独立上市：若用户规模和收入达到相应门槛，可考虑科创板或港股上市。需进一步验证持续盈利能力。')
    add_body_text(doc, '后续融资：A 轮（Year 2）计划融资 800-1500 万元，B 轮（Year 3-4）计划融资 3000-5000 万元。')

    doc.add_page_break()

    # ==================== CHAPTER 8: APPENDIX ====================
    add_heading_custom(doc, '第八章　附　录', level=1)
    add_empty_line(doc)

    add_heading_custom(doc, '附录 A：课程知识图谱建设计划（首批 30 门课程）', level=2)
    add_body_text(doc, '数学类：高等数学（上/下）、线性代数、概率论与数理统计、离散数学。')
    add_body_text(doc, '计算机类：程序设计基础（C/C++）、数据结构、操作系统、计算机网络、数据库系统、计算机组成原理、算法设计与分析。')
    add_body_text(doc, '物理类：大学物理（上/下）。')
    add_body_text(doc, '电子类：电路分析、模拟电子技术、数字电子技术、信号与系统。')
    add_body_text(doc, '经管类：微观经济学、宏观经济学、管理学原理、会计学原理。')
    add_body_text(doc, '外语类：大学英语（4/6 级备考知识图谱）。')
    add_body_text(doc, '考研公共课：考研数学（一/二/三）、考研英语、考研政治。')

    add_empty_line(doc)

    add_heading_custom(doc, '附录 B：MVP 核心功能清单', level=2)

    mvp_headers = ['编号', '功能', '优先级', '预估工期']
    mvp_rows = [
        ['F01', 'AI 文字问答（基础）', 'P0', '4 周'],
        ['F02', '用户注册与登录', 'P0', '1 周'],
        ['F03', '学习路径生成（高数/线代）', 'P0', '3 周'],
        ['F04', '基础知识图谱入库', 'P0', '3 周'],
        ['F05', 'RAG 课件检索增强', 'P1', '3 周'],
        ['F06', '知识点诊断（BKT 模型）', 'P1', '4 周'],
        ['F07', '学习数据看板', 'P1', '2 周'],
        ['F08', '拍照提问（OCR）', 'P1', '3 周'],
        ['F09', '语音提问', 'P2', '2 周'],
        ['F10', 'B 端管理平台', 'P2', '4 周'],
    ]
    build_table(doc, mvp_headers, mvp_rows)

    add_empty_line(doc)

    add_heading_custom(doc, '附录 C：竞品详细对比表', level=2)
    add_body_text(doc, '（竞品详细对比表已在第四章中呈现，此处不再重复。）')

    add_empty_line(doc)

    add_heading_custom(doc, '附录 D：团队简历', level=2)
    add_body_text(doc, '（此处补充实际团队成员的具体简历、项目经历、获奖情况等。）')

    add_empty_line(doc)

    add_heading_custom(doc, '附录 E：相关法规与政策文件', level=2)
    add_body_text(doc, '《中华人民共和国网络安全法》')
    add_body_text(doc, '《中华人民共和国个人信息保护法》')
    add_body_text(doc, '《生成式人工智能服务管理暂行办法》（2023 年 8 月施行）')
    add_body_text(doc, '《中国教育现代化 2035》')
    add_body_text(doc, '《新一代人工智能发展规划》（国务院）')
    add_body_text(doc, '《教育强国建设规划纲要（2024-2035）》')

    # ==================== SAVE ====================
    output_path = r'D:\moniC\project\learn\商业计划书_智学派LearnGen.docx'
    doc.save(output_path)
    print(f'Done! Saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_docx()
