from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw


SRC = Path(r"D:\moniC\project\learn\00-inbox\docx-work\实习4-统计类滤波器+空域锐化-原稿.docx")
OUT = Path(r"D:\moniC\project\learn\00-inbox\docx-work\实习4-统计类滤波器+空域锐化-改写版-黑字6页内.docx")
MEDIA_DIR = Path(r"D:\moniC\project\learn\00-inbox\docx-work\media-extract")
ARTIFACT_DIR = Path(r"D:\moniC\project\learn\00-inbox\docx-work\generated-assets")


ACCENT = RGBColor(0, 0, 0)
TEXT = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)
LIGHT_FILL = "EDF3F9"


def set_run_font(run, name="宋体", size=Pt(11), bold=False, color=TEXT, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, label=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.1)
    r = p.add_run(text)
    set_run_font(
        r,
        name="等线" if label else "宋体",
        size=Pt(10.5),
        bold=label,
        color=ACCENT if label else TEXT,
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if label:
        shade_cell(cell, LIGHT_FILL)


def add_paragraph(doc, text="", *, style=None, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False,
                  before=0, after=6, line=1.35, font="宋体", size=11, bold=False, color=TEXT):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size=Pt(size), bold=bold, color=color)
    return p


def add_heading(doc, text, *, level=1):
    sizes = {1: 15, 2: 13}
    before = 10 if level == 1 else 6
    after = 4 if level == 1 else 3
    p = add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before=before,
        after=after,
        line=1.15,
        font="黑体",
        size=sizes[level],
        bold=True,
        color=ACCENT,
    )
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = add_paragraph(doc, before=0, after=2, line=1.25)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    r1 = p.add_run("• ")
    set_run_font(r1, name="等线", size=Pt(11), bold=True, color=ACCENT)
    r2 = p.add_run(text)
    set_run_font(r2, name="宋体", size=Pt(11), color=TEXT)
    return p


def add_code_block(doc, lines):
    p = add_paragraph(doc, before=1, after=6, line=1.1)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.keep_together = True
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=Pt(8.5), color=RGBColor(55, 55, 55))
        if idx != len(lines) - 1:
            r.add_break()
    return p


def add_caption(doc, text):
    p = add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=2,
        after=8,
        line=1.0,
        font="等线",
        size=9,
        color=MUTED,
    )
    return p


def add_picture(doc, path, *, width_inch):
    p = add_paragraph(doc, before=2, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inch))
    p.paragraph_format.keep_together = True
    return p


def add_image_grid(doc, image_paths, *, columns, width_inch):
    rows = (len(image_paths) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, path in enumerate(image_paths):
        cell = table.cell(i // columns, i % columns)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=2, after=2, line=1.0)
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inch))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table._cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for tag in ("top", "left", "bottom", "right"):
            elem = borders.find(qn(f"w:{tag}"))
            if elem is None:
                elem = OxmlElement(f"w:{tag}")
                borders.append(elem)
            elem.set(qn("w:val"), "nil")
    return table


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_table_borders(table, color="D9E2F0"):
    for cell in table._cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = borders.find(qn(f"w:{tag}"))
            if elem is None:
                elem = OxmlElement(f"w:{tag}")
                borders.append(elem)
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "8")
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), color)


def ensure_media_assets():
    ARTIFACT_DIR.mkdir(exist_ok=True)
    if not MEDIA_DIR.exists():
        raise FileNotFoundError(f"Missing media directory: {MEDIA_DIR}")
    for emf_path in MEDIA_DIR.glob("*.emf"):
        png_path = emf_path.with_suffix(".png")
        if png_path.exists():
            continue
        raise FileNotFoundError(f"Expected converted PNG for {emf_path.name}")


def build_gradient_collage():
    ensure_media_assets()
    selected = [
        (9, "原图像"),
        (10, "Roberts 梯度"),
        (13, "Prewitt 梯度"),
        (16, "Sobel 梯度"),
        (19, "Roberts 二值图"),
        (20, "Prewitt 二值图"),
        (21, "Sobel 二值图"),
    ]
    cards = []
    for idx, label in selected:
        img = Image.open(MEDIA_DIR / f"image{idx}.png").convert("RGB")
        img = img.resize((250, 250))
        card = Image.new("RGB", (290, 318), "white")
        draw = ImageDraw.Draw(card)
        draw.rectangle((0, 0, 289, 317), outline=(210, 220, 235), width=2)
        card.paste(img, (20, 16))
        draw.text((20, 276), label, fill=(34, 34, 34))
        cards.append(card)

    cols = 2
    rows = (len(cards) + cols - 1) // cols
    collage = Image.new("RGB", (cols * 310, rows * 336), "white")
    for i, card in enumerate(cards):
        x = (i % cols) * 310
        y = (i // cols) * 336
        collage.paste(card, (x, y))
    out = ARTIFACT_DIR / "gradient-collage.png"
    collage.save(out)
    return out


def build_simple_collage(indices, out_name, *, columns, thumb_size=(250, 250), card_padding=14, gap=18):
    ensure_media_assets()
    cards = []
    for idx in indices:
        img = Image.open(MEDIA_DIR / f"image{idx}.png").convert("RGB")
        img.thumbnail(thumb_size)
        card_w = thumb_size[0] + card_padding * 2
        card_h = thumb_size[1] + card_padding * 2
        card = Image.new("RGB", (card_w, card_h), "white")
        draw = ImageDraw.Draw(card)
        draw.rectangle((0, 0, card_w - 1, card_h - 1), outline=(215, 223, 236), width=2)
        x = (card_w - img.width) // 2
        y = (card_h - img.height) // 2
        card.paste(img, (x, y))
        cards.append(card)

    rows = (len(cards) + columns - 1) // columns
    card_w, card_h = cards[0].size
    sheet = Image.new(
        "RGB",
        (columns * card_w + (columns - 1) * gap, rows * card_h + (rows - 1) * gap),
        "white",
    )
    for i, card in enumerate(cards):
        x = (i % columns) * (card_w + gap)
        y = (i // columns) * (card_h + gap)
        sheet.paste(card, (x, y))
    out = ARTIFACT_DIR / out_name
    sheet.save(out)
    return out


def convert_to_jpg(src_path, out_name):
    img = Image.open(src_path).convert("RGB")
    out = ARTIFACT_DIR / out_name
    img.save(out, quality=95)
    return out


def build_document():
    doc = Document(str(SRC))
    ensure_media_assets()
    gradient_collage = build_gradient_collage()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    clear_body(doc)

    add_paragraph(
        doc,
        "南京信息工程大学",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=4,
        line=1.0,
        font="黑体",
        size=22,
        bold=True,
        color=TEXT,
    )
    add_paragraph(
        doc,
        "实验（实习）报告",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=8,
        line=1.0,
        font="黑体",
        size=19,
        bold=True,
        color=ACCENT,
    )
    add_paragraph(
        doc,
        "统计类滤波器与空域锐化",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=10,
        line=1.0,
        font="等线",
        size=14,
        bold=True,
        color=TEXT,
    )

    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.5), Cm(5.2), Cm(2.5), Cm(4.2)]
    rows = [
        ("实验名称", "统计类滤波器与空域锐化", "实验日期", "2026-04-29"),
        ("学院", "计算机学院、软件学院", "专业 / 年级", "计算机科学与技术 2024级"),
        ("班级", "3班", "姓名 / 学号", "周可名  202483290198"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value, width in zip(row.cells, values, widths):
            cell.width = width
            set_cell_text(cell, value, label=value in {"实验名称", "实验日期", "学院", "专业 / 年级", "班级", "姓名 / 学号"})
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
    add_paragraph(
        doc,
        "正文保留关键代码片段，完整实验过程通过运行截图与结果分析进行说明。",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=6,
        after=10,
        line=1.0,
        font="等线",
        size=9,
        color=MUTED,
    )

    add_heading(doc, "一、实验目的", level=1)
    add_bullet(doc, "理解统计类滤波器在抑制随机噪声、椒盐噪声时的作用机理及差异。")
    add_bullet(doc, "掌握 Roberts、Prewitt、Sobel、Laplacian 与 LoG 等空域锐化算子的基本实现方法。")
    add_bullet(doc, "结合实验结果分析不同滤波器在去噪、边缘增强和细节保持之间的权衡关系。")

    add_heading(doc, "二、实验内容与实现思路", level=1)
    add_paragraph(
        doc,
        "本次实验在 MATLAB 环境下完成，使用教材配套测试图像对统计类滤波与空域锐化算法进行验证。实验任务包括："
        "对含噪图像进行多种统计滤波比较；在高密度椒盐噪声场景下验证自适应中值滤波；"
        "分别实现一阶微分与二阶微分锐化方法，并通过 LoG 算子观察参数变化对滤波结果的影响。",
        first_indent=True,
        before=0,
        after=5,
        line=1.35,
    )
    add_bullet(doc, "统计类滤波器比较：算术均值、几何均值、中值和 Alpha 修正均值滤波。")
    add_bullet(doc, "自适应中值滤波：比较固定窗口中值滤波与窗口可变的自适应策略。")
    add_bullet(doc, "空域锐化：实现 Roberts、Prewitt、Sobel、Laplacian 和 LoG 算子，并比较其视觉效果。")

    add_heading(doc, "三、实验过程与结果分析", level=1)

    add_heading(doc, "3.1 Alpha 修正均值滤波器", level=2)
    add_paragraph(
        doc,
        "该部分以含椒盐噪声的电路板图像为对象，比较不同统计滤波器对噪声抑制与细节保持的影响。"
        "算术均值会对邻域像素直接求平均，平滑效果明显，但容易造成边缘模糊；几何均值在一定程度上减轻了亮度突变带来的影响；"
        "中值滤波和 Alpha 修正均值滤波更适合处理脉冲噪声，其中后者通过裁剪极端值后再求均值，兼顾了抗噪性与平滑性。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(doc, "关键代码片段：", before=0, after=2, line=1.0, font="等线", size=10.5, bold=True, color=ACCENT)
    add_code_block(
        doc,
        [
            "I = imread('Fig0512(b)(ckt-uniform-plus-saltpepr-prob-pt1).tif');",
            "avgKernel = fspecial('average', [3 3]);",
            "I_arith = imfilter(I, avgKernel);",
            "I_geo = uint8(exp(imfilter(log(double(I)+1), avgKernel)) - 1);",
            "I_median = medfilt2(I, [3 3]);",
            "I_alpha = alphaTrimmedMean(I, 3, 2);",
        ],
    )
    add_paragraph(
        doc,
        "从结果图可以看出，算术均值滤波虽然整体更平滑，但细小结构被明显抹平；几何均值在亮暗过渡区域略优于算术均值；"
        "中值滤波对椒盐噪声有较好的抑制作用，而 Alpha 修正均值滤波在去除异常噪点的同时保留了更多器件轮廓，综合效果最好。",
        first_indent=True,
        before=0,
        after=4,
        line=1.35,
    )
    add_image_grid(
        doc,
        [MEDIA_DIR / f"image{i}.png" for i in [1, 2, 3, 4, 5]],
        columns=2,
        width_inch=1.55,
    )
    add_caption(doc, "图 1 统计类滤波器输出结果对比")

    add_heading(doc, "3.2 自适应中值滤波器", level=2)
    add_paragraph(
        doc,
        "固定窗口中值滤波在中低密度椒盐噪声场景下效果较好，但当噪声密度较大时，窗口内部很可能被噪声像素主导，"
        "此时固定窗口给出的中值不再可靠。自适应中值滤波通过逐步扩大窗口，在更大的邻域中寻找满足条件的中值，从而提高对高密度脉冲噪声的处理能力。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(doc, "关键代码片段：", before=0, after=2, line=1.0, font="等线", size=10.5, bold=True, color=ACCENT)
    add_code_block(
        doc,
        [
            "I = imread('Fig0514(a)(ckt_saltpep_prob_pt25).tif');",
            "J_median = medfilt2(I);",
            "for each pixel, start with a 3x3 window;",
            "if median is still corrupted, enlarge window to 5x5, 7x7 ...;",
            "when a valid median is found, replace noisy pixel with that median;",
            "otherwise keep expanding until the maximum window size is reached;",
        ],
    )
    add_paragraph(
        doc,
        "实验结果表明，普通中值滤波能够明显降低椒盐噪声，但仍会留下部分孤立噪点，并造成局部纹理变钝。"
        "自适应中值滤波在高噪声区域表现更稳定，能更完整地恢复电路板条纹和芯片边界，不过其计算量也更大，执行时间相对更长。",
        first_indent=True,
        before=0,
        after=4,
        line=1.35,
    )
    add_image_grid(
        doc,
        [MEDIA_DIR / f"image{i}.png" for i in [6, 7, 8]],
        columns=1,
        width_inch=2.0,
    )
    add_caption(doc, "图 2 固定窗口中值滤波与自适应中值滤波结果")

    add_heading(doc, "3.3 一阶微分锐化算子", level=2)
    add_paragraph(
        doc,
        "Roberts、Prewitt 和 Sobel 算子都属于基于梯度的一阶微分锐化方法。"
        "其中 Roberts 模板最小，对细小边缘变化较敏感，但也更容易受到噪声干扰；"
        "Prewitt 采用 3×3 邻域估计水平和垂直变化；Sobel 在中心位置赋予更大权重，因此边缘响应通常更强、更稳定。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(doc, "关键代码片段：", before=0, after=2, line=1.0, font="等线", size=10.5, bold=True, color=ACCENT)
    add_code_block(
        doc,
        [
            "I = imread('rice.png');",
            "[Gx, Gy] = buildGradientKernel(method);",
            "gradientX = conv2(double(I), Gx, 'same');",
            "gradientY = conv2(double(I), Gy, 'same');",
            "Ig = uint8(sqrt(gradientX.^2 + gradientY.^2));",
            "use thresholding / edge highlight / background highlight for sharpened output;",
        ],
    )
    add_paragraph(
        doc,
        "从梯度图来看，Roberts 对斜向和局部细节的响应最直接，但噪声放大也更明显；Prewitt 输出较平滑；Sobel 的米粒边缘更连续、层次更清晰。"
        "在阈值化与高亮实验中，阈值越高，背景噪声越少，但部分弱边缘会被遗漏，因此阈值需要在边缘完整性与干净背景之间折中选择。",
        first_indent=True,
        before=0,
        after=4,
        line=1.35,
    )
    add_picture(doc, gradient_collage, width_inch=5.6)
    add_caption(doc, "图 3 一阶微分算子梯度图及不同锐化输出方式")

    add_heading(doc, "3.4 Laplacian 二阶微分锐化", level=2)
    add_paragraph(
        doc,
        "Laplacian 算子利用二阶导数突出灰度快速变化的位置，常用于增强边缘与细节。"
        "本实验分别观察了 Laplacian 响应图和“原图减去 Laplacian”两种使用方式：前者突出灰度变化本身，"
        "后者则把高频信息叠加回原图，从而得到视觉上更清晰的锐化结果。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(doc, "关键代码片段：", before=0, after=2, line=1.0, font="等线", size=10.5, bold=True, color=ACCENT)
    add_code_block(
        doc,
        [
            "I = im2double(imread('rice.png'));",
            "lapKernel = [0 1 0; 1 -4 1; 0 1 0];",
            "lap = imfilter(I, lapKernel, 'replicate', 'conv');",
            "sharpened = I - lap;",
        ],
    )
    add_paragraph(
        doc,
        "实验结果显示，Laplacian 响应图主要保留了米粒轮廓及灰度跃迁位置；"
        "将其与原图结合后，米粒边界更锋利，纹理层次更突出。与此同时，二阶微分也会同步放大噪声，"
        "因此实际应用中通常会先进行平滑处理，再执行 Laplacian 锐化。",
        first_indent=True,
        before=0,
        after=4,
        line=1.35,
    )
    add_picture(doc, MEDIA_DIR / "image25.png", width_inch=6.0)
    add_caption(doc, "图 4 Laplacian 响应与锐化结果")

    add_heading(doc, "3.5 LoG 算子系数可视化与滤波比较", level=2)
    add_paragraph(
        doc,
        "LoG（Laplacian of Gaussian）先利用高斯滤波抑制噪声，再通过二阶微分突出边缘。"
        "相较于直接使用 Laplacian，LoG 对噪声更稳健，但其输出效果与高斯标准差和滤波器尺寸密切相关。"
        "因此本实验分别展示不同参数组合下的 LoG 三维系数分布，以及作用于 Lena 灰度图后的输出结果。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(doc, "关键代码片段：", before=0, after=2, line=1.0, font="等线", size=10.5, bold=True, color=ACCENT)
    add_code_block(
        doc,
        [
            "filterSizes = [5, 9, 13];",
            "sigmas = [0.6, 1.0, 1.6];",
            "h = fspecial('log', filterSize, sigma);",
            "surf(X, Y, h);",
            "out = imfilter(I, h, 'replicate', 'conv');",
            "compare filtering results under different size / sigma settings;",
        ],
    )
    add_paragraph(
        doc,
        "从三维核函数可以看出，随着滤波器尺寸和标准差增大，LoG 曲面的中心凹陷与周边响应范围更宽，"
        "说明其对更大尺度结构更敏感。滤波结果中，小尺寸、小 sigma 更强调细小纹理，但噪声残留也更明显；"
        "较大的参数组合平滑作用更强，输出更稳定，不过部分细节会被削弱。这说明 LoG 参数选择需要服务于具体任务目标。",
        first_indent=True,
        before=0,
        after=4,
        line=1.35,
    )
    add_picture(doc, MEDIA_DIR / "image26.png", width_inch=5.85)
    add_caption(doc, "图 5 不同参数下的 LoG 算子三维可视化")
    add_picture(doc, MEDIA_DIR / "image27.png", width_inch=5.85)
    add_caption(doc, "图 6 LoG 算子对 Lena 灰度图的滤波结果")

    add_heading(doc, "四、实验总结与体会", level=1)
    add_paragraph(
        doc,
        "本次实验把“去噪”和“锐化”两类典型空域处理任务串联起来，使我更直观地理解了不同算法背后的适用条件。"
        "对于脉冲噪声，中值类方法明显优于简单均值滤波；而在锐化任务中，一阶微分更适合突出边缘方向信息，二阶微分更适合强调灰度突变。"
        "同一种算子在不同参数或阈值下的表现差异也十分明显，因此图像处理并不存在放之四海而皆准的“最佳算法”。",
        first_indent=True,
        before=1,
        after=4,
        line=1.35,
    )
    add_paragraph(
        doc,
        "通过本次编程实现，我进一步熟悉了 MATLAB 图像处理流程，也体会到实验报告不能只堆叠代码与截图，"
        "更重要的是把算法思路、参数选择依据和结果现象讲清楚。后续如果继续扩展实验，可以增加 PSNR、SSIM 或边缘检测精度等量化指标，"
        "使不同滤波器之间的比较更加客观。",
        first_indent=True,
        before=0,
        after=0,
        line=1.35,
    )

    doc.save(str(OUT))


if __name__ == "__main__":
    build_document()
