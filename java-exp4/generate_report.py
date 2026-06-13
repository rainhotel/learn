from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


STUDENT_NAME = "[请替换姓名]"
STUDENT_ID = "[请替换学号]"
CLASS_NAME = "[请替换班级]"
TEACHER_NAME = "[请替换教师]"
EXPERIMENT_DATE = "2026-06-11"
OUTPUT_FILE = "java-exp4-report.docx"


def set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(document: Document, text: str, color=None) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    if color is not None:
        run.font.color.rgb = color


def add_code_block(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def build_info_table(document: Document) -> None:
    table = document.add_table(rows=3, cols=4, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = table.rows
    rows[0].cells[0].text = "课程名称"
    rows[0].cells[1].text = "Java 程序设计"
    rows[0].cells[2].text = "实验日期"
    rows[0].cells[3].text = EXPERIMENT_DATE

    rows[1].cells[0].text = "实验题目"
    rows[1].cells[1].text = "6-1、6-4、11-1"
    rows[1].cells[2].text = "指导教师"
    rows[1].cells[3].text = TEACHER_NAME

    rows[2].cells[0].text = "姓名 / 学号"
    rows[2].cells[1].text = f"{STUDENT_NAME} / {STUDENT_ID}"
    rows[2].cells[2].text = "班级"
    rows[2].cells[3].text = CLASS_NAME


def main() -> None:
    document = Document()
    set_base_style(document)

    add_title(document, "Java 实验报告")
    add_title(document, "实验题目：6-1、6-4、11-1")
    document.add_paragraph()
    build_info_table(document)
    document.add_paragraph()

    add_heading(document, "一、实验目的")
    add_paragraph(document, "1. 掌握 HashSet 的基本特性和常用操作。")
    add_paragraph(document, "2. 掌握 Iterator 的遍历与安全删除方式。")
    add_paragraph(document, "3. 理解 JSP 页面中声明、脚本片段和表达式的用法。")
    add_paragraph(document, "4. 提升 Java 集合框架与 JSP 基础编程的综合应用能力。")

    add_heading(document, "二、实验环境")
    add_paragraph(document, "1. 操作系统：Windows")
    add_paragraph(document, "2. JDK：21")
    add_paragraph(document, "3. 开发方式：命令行编译运行 + 文本编辑")
    add_paragraph(document, "4. JSP 运行环境：Tomcat（本报告提供部署方法，本机未完成容器验证）")

    add_heading(document, "三、实验内容")
    add_paragraph(document, "1. 实验 6-1：使用 HashSet 完成集合的添加、删除、交集保留、交集删除、清空和 null 元素测试。")
    add_paragraph(document, "2. 实验 6-4：使用 Iterator 遍历 Vector，并删除长度大于 4 的字符串。")
    add_paragraph(document, "3. 实验 11-1：编写 JSP 页面，输入自然数并计算 1 到该数的连续和。")

    add_heading(document, "四、关键源码")
    add_paragraph(document, "1. HashSetTester.java")
    add_code_block(
        document,
        "Set<String> set = new HashSet<>(3);\n"
        "set.add(\"one\");\n"
        "set.add(\"two\");\n"
        "set.add(\"three\");\n"
        "set.retainAll(setToRetain);\n"
        "set.removeAll(setToRemove);\n"
        "set.clear();"
    )
    add_paragraph(document, "2. IteratorTester.java")
    add_code_block(
        document,
        "Iterator<String> nums = vector.iterator();\n"
        "while (nums.hasNext()) {\n"
        "    String aString = nums.next();\n"
        "    if (aString.length() > 4) {\n"
        "        nums.remove();\n"
        "    }\n"
        "}"
    )
    add_paragraph(document, "3. Ex11_1.jsp")
    add_code_block(
        document,
        "<%! int continuousSum(int n) { ... } %>\n"
        "<% String str = request.getParameter(\"number\"); %>\n"
        "<p><%= r %> 的连续和是 <%= continuousSum(r) %></p>"
    )

    add_heading(document, "五、运行结果")
    add_paragraph(document, "1. 实验 6-1 输出：")
    add_code_block(
        document,
        "The initial set: [one, two, three]\n"
        "The element 'three' is removed from the set: true\n"
        "The element 'three' is removed from the set once again: false\n"
        "The element 'three' is added to the set: true\n"
        "The element 'three' is added to the set once again: false\n"
        "The elements to retain: [one, two]\n"
        "The set after retaining: [one, two]\n"
        "The elements to remove: [two, three]\n"
        "The set after removing: [one]\n"
        "The set is empty after clearing: true\n"
        "The set now contains a 'null' element: true"
    )
    add_paragraph(document, "2. 实验 6-4 输出：")
    add_code_block(
        document,
        "The initial Vector is: [one, two, three, four, five, six, seven, eight, nine, ten]\n"
        "one\n"
        "two\n"
        "three\n"
        "four\n"
        "five\n"
        "six\n"
        "seven\n"
        "eight\n"
        "nine\n"
        "ten\n"
        "The Vector after iteration is: [one, two, four, five, six, nine, ten]"
    )
    add_paragraph(document, "3. 实验 11-1 预期页面结果：输入 10 后显示“10 的连续和是 55”。")
    add_paragraph(document, "如需提交截图，可在部署到 Tomcat 后补充浏览器运行截图。", RGBColor(128, 128, 128))

    add_heading(document, "六、实验分析")
    add_paragraph(document, "1. HashSet 的主要特点是元素唯一、查找与删除效率较高，但元素顺序不固定。")
    add_paragraph(document, "2. Iterator 适合在遍历过程中同步删除当前元素，避免并发修改异常。")
    add_paragraph(document, "3. JSP 可将 Java 逻辑和 HTML 页面组合，适合入门理解动态网页的处理流程。")

    add_heading(document, "七、实验总结")
    add_paragraph(document, "通过本次实验，我进一步掌握了 Java 集合框架中的 HashSet 与 Iterator 的使用方法，也理解了 JSP 页面接收参数、调用方法和输出结果的基本过程。三个实验分别对应集合操作、集合遍历和 Web 动态页面三个基础方向，为后续继续学习 ArrayList、Map、Servlet 和 JavaBean 奠定了基础。")

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    main()
