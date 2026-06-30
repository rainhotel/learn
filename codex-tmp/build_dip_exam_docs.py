from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\moniC\project\learn")
OUT_DIR = ROOT / "01-topics" / "digital-image-processing" / "exam-docx"


REVIEW_SECTIONS = [
    {
        "title": "一、数字图像处理基础",
        "lead": "这一部分通常出选择题和简答题，考你是否理解数字图像是如何从连续图像变成计算机矩阵的，以及影响图像质量的几个基本因素。",
        "points": [
            {
                "name": "图像采样",
                "body": "采样是把连续空间坐标离散化，把连续图像变成有限个像素点。采样间隔越小，空间分辨率越高，细节越丰富；采样过粗会出现细节丢失、锯齿和混叠。考试里看到“空间分辨率”通常就要联想到采样密度。",
                "exam": "常考问法：为什么同一场景用更高分辨率相机拍摄更清晰？答：采样点更多，空间细节保留更多。",
            },
            {
                "name": "图像量化",
                "body": "量化是把连续灰度值离散化为有限灰度级。例如8 bit灰度图有256个灰度级，范围通常为0到255。量化级数越多，灰度过渡越平滑；量化级数太少会出现伪轮廓。",
                "exam": "常考问法：256个灰度级需要多少bit？答：2^8=256，所以需要8 bit。",
            },
            {
                "name": "数字图像质量因素",
                "body": "重点记四个：空间分辨率、灰度级分辨率、对比度、清晰度。空间分辨率看像素密度；灰度级分辨率看灰度层次；对比度看最大灰度和最小灰度差异；清晰度常与边缘、细节和模糊程度有关。",
                "exam": "简答题可以按“定义-影响-改善方法”写：提高采样改善空间分辨率，增加量化位数改善灰度层次，灰度变换或直方图增强改善对比度，锐化增强边缘细节。",
            },
        ],
        "formulas": [
            ("灰度级与比特数", "L = 2^k；若L=256，则k=8"),
            ("未压缩数据量", "M x N x bit_depth；RGB真彩色8 bit/通道为M x N x 24 bit"),
        ],
    },
    {
        "title": "二、空域图像增强",
        "lead": "空域增强直接在像素矩阵上操作，是考试最容易出算法步骤、模板计算和简答分析的部分。重点是点运算、直方图处理、平滑滤波和锐化滤波。",
        "points": [
            {
                "name": "为什么做空域增强",
                "body": "目的是改善视觉质量或突出后续处理所需的信息，例如增强细节、拉伸对比度、去除噪声、突出边缘。空域增强的特点是直接操作f(x,y)，不需要先变换到频域。",
                "exam": "简答题可写：增强不是恢复真实原图，而是为了改善视觉效果或方便识别、分割、检测。",
            },
            {
                "name": "点运算与线性灰度变换",
                "body": "点运算的输出像素只依赖同位置输入像素，形式为s=T(r)。线性变换可进行亮度调整和对比度拉伸；反转s=L-1-r适合增强暗背景中的亮细节；对数变换扩展暗部、压缩亮部；幂律变换中gamma<1提亮，gamma>1压暗。",
                "exam": "常考判断：点运算不使用邻域，所以不能直接完成空间去噪。",
            },
            {
                "name": "直方图均衡化",
                "body": "直方图均衡化通过累计分布函数CDF重新映射灰度，使灰度动态范围更充分，从而增强对比度。它不保证输出直方图完全平坦，离散图像中还可能出现某些灰度级被合并或空缺。",
                "exam": "手算步骤：统计概率p(rk) -> 算CDF -> 乘L-1 -> 按题目规则取整 -> 建立映射。",
            },
            {
                "name": "直方图匹配/规定化",
                "body": "均衡化的目标是自动拉开灰度，而规定化是让图像直方图接近期望分布。它比均衡化更可控，适合希望图像呈现指定灰度风格的场景。",
                "exam": "问两者区别：均衡化目标近似均匀分布；规定化目标是指定分布。",
            },
            {
                "name": "平滑滤波器",
                "body": "均值滤波通过邻域平均降低随机噪声，但会模糊边缘；中值滤波是非线性排序滤波，特别适合椒盐噪声；最大值滤波可去椒噪声，最小值滤波可去盐噪声。",
                "exam": "常考作用：均值滤波平滑但模糊；中值滤波保护边缘能力较好；最大/最小滤波针对黑白脉冲噪声。",
            },
            {
                "name": "锐化技术",
                "body": "锐化利用灰度突变增强边缘和细节。一阶微分反映灰度变化强度，典型算子有Roberts、Prewitt、Sobel；二阶微分典型是Laplacian，对灰度突变更敏感，但也更容易放大噪声。",
                "exam": "Prewitt和Sobel区别：Sobel中心行/列权重为2，对噪声有更强平滑作用，边缘检测更稳。",
            },
        ],
        "formulas": [
            ("通用点运算", "s = T(r)"),
            ("反转", "s = L - 1 - r"),
            ("对数变换", "s = c log(1+r)"),
            ("幂律变换", "s = c r^gamma"),
            ("均值滤波", "g(x,y)=1/mn * sum f(s,t)"),
            ("Sobel近似幅值", "|G| ≈ |Gx| + |Gy|"),
        ],
    },
    {
        "title": "三、频域图像增强",
        "lead": "频域增强先把图像变换到频率域，再设计滤波器处理频率成分。考试重点是傅里叶变换意义、频谱特征、低通/高通/带通/带阻滤波器及其效果。",
        "points": [
            {
                "name": "频域概念与基本特征",
                "body": "空间域看像素位置，频域看灰度变化快慢。低频表示整体亮度和平滑区域，高频表示边缘、纹理、细节和噪声。频谱的幅度表示该频率成分强弱，相位决定结构位置，能量常集中在低频。",
                "exam": "常考：F(0,0)是直流分量，等于所有像素灰度和；fftshift后低频在中心。",
            },
            {
                "name": "傅里叶变换处理步骤",
                "body": "标准流程为：对图像做DFT得到F(u,v)，中心化频谱，构造滤波器H(u,v)，相乘得到G=HF，再反中心化并IDFT回空间域。频域滤波本质是乘法。",
                "exam": "算法题一定写清楚DFT、滤波器相乘、IDFT三个环节。",
            },
            {
                "name": "低通滤波器",
                "body": "低通保留低频、抑制高频，因此用于平滑和降噪，但会损失边缘细节。理想低通截止突然，容易振铃；巴特沃斯低通过渡可由阶数控制；高斯低通过渡最平滑，振铃最少。指数形、梯形低通也属于低通家族，核心仍是高频衰减。",
                "exam": "D0越小，图像越平滑越模糊；D0越大，细节保留更多但噪声也可能保留。",
            },
            {
                "name": "高通滤波器",
                "body": "高通保留高频、抑制低频，用于增强边缘和细节。理想高通同样可能振铃；巴特沃斯和高斯高通过渡更平滑。高通会增强噪声，所以常需要先平滑或使用高提升/高频强调保留背景信息。",
                "exam": "频谱滤波器图像中，中心暗、外围亮通常是高通；中心亮、外围暗通常是低通。",
            },
            {
                "name": "带通与带阻滤波器",
                "body": "带通保留某一频带，带阻抑制某一频带。周期噪声常在频谱中表现为成对对称亮点或亮带，可用陷波或带阻滤波抑制。",
                "exam": "看到规律条纹或频谱对称亮点，先想到周期噪声和陷波/带阻滤波。",
            },
        ],
        "formulas": [
            ("二维DFT", "F(u,v)=sum_x sum_y f(x,y) exp[-j2pi(ux/M+vy/N)]"),
            ("频域滤波", "G(u,v)=H(u,v)F(u,v)"),
            ("理想低通", "H=1 if D<=D0, else 0"),
            ("巴特沃斯低通", "H=1/[1+(D/D0)^(2n)]"),
            ("高斯低通", "H=exp[-D^2/(2D0^2)]"),
            ("高通互补", "H_hp = 1 - H_lp"),
        ],
    },
    {
        "title": "四、图像复原与噪声滤波",
        "lead": "图像复原强调根据退化模型恢复原图，比增强更客观。图片考点里同时列出了退化定义、逆滤波、维纳滤波、噪声类型和多类空间复原滤波器。",
        "points": [
            {
                "name": "图像退化的定义与性质",
                "body": "退化是成像过程中受到模糊、运动、大气湍流、噪声等影响，导致观测图像质量下降。复原要建立退化模型，再估计原图。空间域卷积在频域变为乘法，所以频域复原很常见。",
                "exam": "增强偏主观，复原偏模型；答题时要写退化模型和噪声项。",
            },
            {
                "name": "逆滤波法",
                "body": "若G=HF且无噪声，可用F_hat=G/H复原。但实际G=HF+N，逆滤波后F_hat=F+N/H，当H很小时噪声被严重放大。",
                "exam": "问逆滤波缺点：对噪声敏感，H接近0处不稳定。",
            },
            {
                "name": "维纳滤波法",
                "body": "维纳滤波在分母中加入噪声与图像功率谱比，综合退化函数和噪声统计，在最小均方误差意义下更稳定。噪声越强，抑制越明显。",
                "exam": "问为什么比逆滤波好：它考虑Sn/Sf，不是简单除以H。",
            },
            {
                "name": "常见噪声类型",
                "body": "高斯噪声呈钟形分布，常由传感器和电子系统产生；瑞利、伽马、指数、均匀噪声有各自概率模型；脉冲噪声也叫椒盐噪声，表现为随机黑点和白点。",
                "exam": "识别题：随机黑白点是椒盐噪声；直方图近似钟形是高斯噪声。",
            },
            {
                "name": "空间复原滤波器",
                "body": "算术均值适合平滑加性随机噪声但模糊边缘；几何均值平滑较温和；谐波均值适合盐噪声，不适合椒噪声；逆谐波均值Q>0适合椒噪声，Q<0适合盐噪声；中值滤波适合椒盐噪声；自适应滤波根据局部统计调整强度。",
                "exam": "最容易记反：Q>0去椒噪声，Q<0去盐噪声。",
            },
        ],
        "formulas": [
            ("退化模型", "g(x,y)=h(x,y)*f(x,y)+n(x,y)"),
            ("频域退化模型", "G(u,v)=H(u,v)F(u,v)+N(u,v)"),
            ("逆滤波", "F_hat=G/H=F+N/H"),
            ("维纳滤波", "F_hat=[H*/(|H|^2+Sn/Sf)]G"),
            ("逆谐波均值", "f_hat = sum g^(Q+1) / sum g^Q"),
        ],
    },
    {
        "title": "五、图像分割",
        "lead": "图像分割是本次图片里明确列出的重点，包含检测类分割、阈值分割和基于区域的分割。它既可能出选择题，也很适合出算法设计与分析题。",
        "points": [
            {
                "name": "点检测、线检测、边缘检测",
                "body": "点检测寻找孤立灰度突变点；线检测用方向模板寻找线状结构；边缘检测寻找区域边界，通常基于一阶或二阶导数。Prewitt和Sobel都是一阶边缘算子，Sobel权重更强调中心行/列，抗噪性更好。",
                "exam": "比较题：点、线、边缘检测的目标结构不同；边缘检测常用于分割前提取边界。",
            },
            {
                "name": "阈值处理",
                "body": "全局阈值整幅图一个T，适合光照均匀；自适应阈值根据局部统计变化，适合光照不均；Otsu法通过最大化类间方差自动选择阈值。",
                "exam": "Otsu不是找谷底，而是最大化sigma_B^2=P1P2(m1-m2)^2。",
            },
            {
                "name": "区域生长法",
                "body": "从种子点开始，把满足相似性准则并与区域连通的像素加入。三要素是种子点、生长规则、终止条件。常见规则是候选像素与当前区域均值差小于阈值。",
                "exam": "手算题注意：每加入一个像素后当前区域均值要更新。",
            },
            {
                "name": "区域分裂与合并",
                "body": "先判断区域是否满足一致性准则，不满足则分裂成子区域；相邻子区域满足相似性时合并。分裂关注区域内部是否均匀，合并关注相邻区域是否相似。",
                "exam": "分裂和合并的阈值含义不同，不能混用。",
            },
        ],
        "formulas": [
            ("二值阈值", "g=1 if f>T, else 0"),
            ("Otsu类间方差", "sigma_B^2=P1P2(m1-m2)^2"),
            ("区域完整性", "R1 union R2 union ... union Rn = R"),
            ("区域互斥", "Ri intersect Rj = empty, i != j"),
            ("区域生长规则示例", "|g(candidate)-mean(region)| < T"),
        ],
    },
    {
        "title": "六、彩色图像处理",
        "lead": "彩色图像处理的考点通常较概念化：彩色基础、颜色空间和为什么要在特定颜色空间中处理。选择题和简答题都可能出现。",
        "points": [
            {
                "name": "彩色基础",
                "body": "RGB是加色模型，面向显示设备；CMY/CMYK是减色模型，面向印刷；彩色图像在Matlab中通常是M x N x 3矩阵，每个通道表示一种颜色分量。",
                "exam": "数据量题：RGB真彩色8 bit/通道就是24 bit/像素。",
            },
            {
                "name": "颜色空间",
                "body": "RGB适合设备显示和通道处理；HSI/HSL更接近人眼感知，把色调、饱和度和亮度分离；YCbCr、HSV等也常用于目标检测、压缩或颜色分割。",
                "exam": "如果只想增强亮度而尽量不改变色调，优先转HSI处理I分量。",
            },
            {
                "name": "彩色处理技术",
                "body": "包括伪彩色增强、全彩色处理、彩色平滑锐化、颜色空间变换和基于颜色阈值的目标分割。伪彩色不是恢复真实颜色，而是把灰度映射为便于观察的颜色。",
                "exam": "颜色阈值法对光照变化敏感，答分析题时要写这一限制。",
            },
        ],
        "formulas": [
            ("CMY与RGB", "C=1-R, M=1-G, Y=1-B"),
            ("RGB数据量", "M x N x 3 bytes for uint8 true-color image"),
        ],
    },
]


SINGLE_CHOICE = [
    ("图像采样主要完成的是（ ）。", ["灰度值离散化", "空间坐标离散化", "颜色空间转换", "频域滤波"], "B", "采样对应空间坐标离散化，量化对应灰度值离散化。"),
    ("8 bit灰度图像的灰度级数通常为（ ）。", ["8", "16", "128", "256"], "D", "8 bit可以表示2^8=256个灰度级。"),
    ("直方图均衡化主要用于（ ）。", ["图像压缩", "增强对比度", "恢复退化函数", "彩色空间转换"], "B", "均衡化通过重新分布灰度增强对比度。"),
    ("下列滤波器中最适合去除椒盐噪声的是（ ）。", ["算术均值滤波", "中值滤波", "理想低通滤波", "高通滤波"], "B", "中值滤波对随机黑白点效果好，并能较好保护边缘。"),
    ("Sobel算子相比Prewitt算子的主要特点是（ ）。", ["不需要模板", "中心权重更大，抗噪性更好", "只能检测点", "属于频域滤波器"], "B", "Sobel模板中间行/列权重为2。"),
    ("频域中低频成分通常对应图像的（ ）。", ["整体亮度和平滑区域", "孤立噪声尖点", "边缘和纹理", "随机相位"], "A", "低频代表缓慢变化，高频代表快速变化。"),
    ("理想低通滤波器最典型的缺点是（ ）。", ["无法平滑图像", "频域截止突然导致振铃", "只能处理彩色图像", "会增强边缘噪声"], "B", "理想滤波器硬截止，空间域易出现振铃。"),
    ("高通滤波器的主要作用是（ ）。", ["平滑图像", "增强边缘和细节", "降低空间分辨率", "做灰度量化"], "B", "高通保留高频，高频包含边缘细节和噪声。"),
    ("图像复原中的频域退化模型是（ ）。", ["G=H+F+N", "G=HF+N", "G=F/H", "G=F-N"], "B", "空间卷积在频域变乘法，所以G=HF+N。"),
    ("逆滤波最怕的问题是（ ）。", ["灰度级太多", "H接近0时噪声被放大", "无法进行傅里叶变换", "只能处理二值图像"], "B", "F_hat=F+N/H，H小会放大噪声。"),
    ("Otsu法选择阈值的准则是（ ）。", ["最小化图像均值", "最大化类间方差", "最大化灰度级数", "固定取128"], "B", "Otsu也叫最大类间方差法。"),
    ("区域生长法的三要素不包括（ ）。", ["种子点", "生长规则", "终止条件", "傅里叶相位"], "D", "区域生长与傅里叶相位无关。"),
    ("RGB模型主要面向（ ）。", ["显示设备", "印刷油墨", "频域滤波", "噪声统计"], "A", "RGB是加色模型，显示设备常用。"),
    ("若只想增强彩色图像亮度而尽量保持色调，应优先处理（ ）。", ["RGB三个通道分别任意增强", "HSI中的I分量", "CMY中的C分量", "频谱相位"], "B", "HSI把亮度与色调、饱和度分离。"),
    ("频谱中出现关于中心对称的亮点，空间域常对应（ ）。", ["周期噪声", "量化不足", "采样过密", "灰度反转"], "A", "周期噪声在频域常表现为成对对称尖峰。"),
]


MULTIPLE_CHOICE = [
    ("影响数字图像质量的因素包括（ ）。", ["空间分辨率", "灰度级分辨率", "对比度", "清晰度"], "ABCD", "四项都是图片中列出的图像质量因素。"),
    ("下列属于频域滤波器类型的有（ ）。", ["低通滤波器", "高通滤波器", "带通滤波器", "带阻滤波器"], "ABCD", "四项都是常见频域滤波器。"),
    ("下列属于常见噪声模型的有（ ）。", ["高斯噪声", "瑞利噪声", "指数噪声", "脉冲噪声"], "ABCD", "图片中列出了高斯、瑞利、伽马、指数、均匀、脉冲噪声。"),
    ("图像分割中的阈值处理方法包括（ ）。", ["全局阈值", "自适应阈值", "Otsu法", "维纳滤波"], "ABC", "维纳滤波属于复原滤波，不是阈值分割方法。"),
    ("下列关于边缘检测算子的说法正确的有（ ）。", ["Roberts、Prewitt、Sobel属于一阶微分边缘算子", "Laplacian属于二阶微分算子", "Sobel通常比Prewitt抗噪性更好", "锐化可能放大噪声"], "ABCD", "四项均正确。"),
]


COMPREHENSIVE = [
    {
        "title": "综合题1：直方图均衡化手算（16分）",
        "question": [
            "某3 bit以下简化灰度图只使用4个灰度级r0=0,r1=1,r2=2,r3=3。各灰度级像素个数分别为n=[2,3,3,0]，总像素数为8。设L=4，按四舍五入规则计算均衡化映射s_k=round((L-1)CDF(r_k))。",
            "要求：1）写出各灰度级概率；2）写出CDF；3）写出映射关系；4）说明均衡化如何改善图像质量。",
        ],
        "answer": [
            "概率：p=[2/8,3/8,3/8,0]=[0.25,0.375,0.375,0]。",
            "CDF：C0=0.25，C1=0.625，C2=1.0，C3=1.0。",
            "乘L-1=3：0.75，1.875，3，3。按四舍五入得到s=[1,2,3,3]，即0->1，1->2，2->3，3->3。",
            "均衡化利用累计分布重新映射灰度，使灰度范围更充分，通常能扩大最大灰度与最小灰度差异，提高对比度和清晰感。离散图像中不保证输出直方图完全平坦。",
        ],
    },
    {
        "title": "综合题2：Otsu阈值分割（17分）",
        "question": [
            "一幅图像只有4个灰度级0、1、2、3，对应像素个数为[2,2,4,8]，总像素数16。设阈值k把灰度<=k分为C0，灰度>k分为C1。请分别计算k=0、1、2时的类间方差sigma_B^2=P0P1(m0-m1)^2，并确定最佳阈值。",
        ],
        "answer": [
            "概率p=[0.125,0.125,0.25,0.5]。",
            "k=0：P0=0.125，m0=0；P1=0.875，m1=(1*0.125+2*0.25+3*0.5)/0.875=2.4286；sigma_B^2≈0.645。",
            "k=1：P0=0.25，m0=(1*0.125)/0.25=0.5；P1=0.75，m1=(2*0.25+3*0.5)/0.75=2.6667；sigma_B^2≈0.880。",
            "k=2：P0=0.5，m0=(1*0.125+2*0.25)/0.5=1.25；P1=0.5，m1=3；sigma_B^2=0.7656。",
            "最大值出现在k=1，所以Otsu最佳阈值为1。答题关键词：最大类间方差。",
        ],
    },
    {
        "title": "综合题3：Sobel边缘检测与方法分析（17分）",
        "question": [
            "给定3x3邻域：第一行10,10,10；第二行10,50,90；第三行10,90,90。使用Sobel算子Gx=[-1 0 1; -2 0 2; -1 0 1]，Gy=[-1 -2 -1; 0 0 0; 1 2 1]。",
            "要求：1）计算Gx和Gy；2）计算梯度幅值sqrt(Gx^2+Gy^2)；3）若阈值T=300，判断中心像素是否为边缘；4）简述Prewitt与Sobel区别。",
        ],
        "answer": [
            "Gx=(-1*10+0*10+1*10)+(-2*10+0*50+2*90)+(-1*10+0*90+1*90)=0+160+80=240。",
            "Gy=(-1*10-2*10-1*10)+(0)+(1*10+2*90+1*90)=-40+280=240。",
            "梯度幅值=sqrt(240^2+240^2)=240sqrt(2)≈339.4。",
            "因为339.4>300，所以判断为边缘点。",
            "Prewitt模板权重较均匀；Sobel在中心行/列权重为2，兼具一定平滑作用，通常抗噪性更好。",
        ],
    },
]


SHORT_ANSWERS = [
    {
        "question": "简述图像退化模型，并比较逆滤波和维纳滤波。",
        "answer": "退化模型为空间域g=h*f+n，频域G=HF+N。逆滤波用F_hat=G/H，简单但当H接近0时会放大噪声。维纳滤波F_hat=[H*/(|H|^2+Sn/Sf)]G，考虑噪声功率谱和图像功率谱，在噪声存在时更稳定。",
    },
    {
        "question": "比较全局阈值、自适应阈值和Otsu法。",
        "answer": "全局阈值整幅图使用一个T，适合光照均匀、目标背景差异明显的图像。自适应阈值根据位置和邻域统计变化，适合光照不均。Otsu法自动遍历阈值并最大化类间方差，适合双峰或目标背景分离较明显的直方图。",
    },
    {
        "question": "均值滤波器、最大值滤波器、中值滤波器和最小值滤波器分别有什么作用？",
        "answer": "均值滤波器用于平滑随机噪声，但会模糊边缘。最大值滤波器用邻域最大值替换中心像素，常用于去除黑色椒噪声。最小值滤波器用邻域最小值替换中心像素，常用于去除白色盐噪声。中值滤波器取排序中值，适合椒盐噪声且较能保护边缘。",
    },
    {
        "question": "为什么彩色图像处理中常把RGB转换到HSI再处理？",
        "answer": "RGB三个通道耦合了颜色和亮度，直接分别增强容易偏色。HSI把色调H、饱和度S和亮度I分开，更符合人眼感知。若只想调整亮度或增强对比度，可只处理I分量，再转换回RGB，从而尽量保持原有颜色。",
    },
]


def set_east_asian_font(run, font="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "ReviewTitle" not in styles:
        title = styles.add_style("ReviewTitle", 1)
        title.font.name = "Calibri"
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        title.font.size = Pt(20)
        title.font.bold = True
        title.font.color.rgb = RGBColor.from_string("0B2545")
        title.paragraph_format.space_after = Pt(6)

    if "SubtitleLine" not in styles:
        subtitle = styles.add_style("SubtitleLine", 1)
        subtitle.font.name = "Calibri"
        subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        subtitle.font.size = Pt(10)
        subtitle.font.color.rgb = RGBColor.from_string("555555")
        subtitle.paragraph_format.space_after = Pt(12)

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_header_footer(doc, label):
    header = doc.sections[0].header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("数字图像处理考试复习")
    set_east_asian_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(label)
    set_east_asian_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("777777")


def add_para(doc, text, bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        run = p.add_run(bold_label)
        run.bold = True
        set_east_asian_font(run)
    run = p.add_run(text)
    set_east_asian_font(run)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    return p


def add_list(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(item)
        set_east_asian_font(run)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                set_east_asian_font(run)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_east_asian_font(run)
    return table


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(style="ReviewTitle")
    run = p.add_run(title)
    set_east_asian_font(run)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph(style="SubtitleLine")
    run = p.add_run(subtitle)
    set_east_asian_font(run)


def build_review_doc():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "针对性知识点复习")
    add_title(
        doc,
        "数字图像处理针对性知识点复习",
        "依据课堂图片考点整理：数字图像基础、空域增强、频域增强、图像复原、图像分割、彩色图像处理。",
    )

    doc.add_heading("考试结构与复习优先级", level=1)
    add_table(
        doc,
        ["题型", "分值", "复习策略"],
        [
            ("单选题", "15", "背概念、公式作用、方法适用场景，尤其注意易混概念。"),
            ("多选题", "15", "按类别记全：噪声类型、频域滤波器、阈值方法、边缘算子。"),
            ("综合题：算法应用与实现", "50", "会手算直方图均衡化、Otsu、模板滤波、频域滤波流程，并能写伪代码或Matlab思路。"),
            ("简答题：算法设计与分析", "20", "按“定义-步骤-优点-缺点-适用条件”组织答案。"),
        ],
        [1800, 1200, 6360],
    )
    add_para(
        doc,
        "复习时不要平均用力。综合题占50分，最值得投入的是算法步骤、手算流程和结果分析；选择题主要靠概念边界和方法分类拿稳分。"
    )

    for section in REVIEW_SECTIONS:
        doc.add_heading(section["title"], level=1)
        add_para(doc, section["lead"])
        for point in section["points"]:
            doc.add_heading(point["name"], level=2)
            add_para(doc, point["body"], "知识点：")
            add_para(doc, point["exam"], "考法：")
        doc.add_heading("本节必背公式/模型", level=2)
        add_table(doc, ["名称", "表达式/记忆点"], section["formulas"], [2700, 6660])

    doc.add_heading("七、算法实现模板", level=1)
    templates = [
        ("直方图均衡化", ["统计每个灰度级出现次数。", "除以总像素数得到概率。", "计算累计分布CDF。", "用s=round((L-1)CDF)得到映射。", "按映射替换原图灰度。"]),
        ("Otsu阈值", ["枚举所有候选阈值k。", "按k分成C0和C1。", "计算P0、P1、m0、m1。", "计算sigma_B^2=P0P1(m0-m1)^2。", "取sigma_B^2最大的k作为阈值。"]),
        ("空间模板滤波", ["确定邻域窗口和模板。", "模板与邻域对应位置相乘。", "全部求和。", "若模板有归一化系数再除以系数。", "对边界像素说明采用补零、复制或忽略。"]),
        ("频域滤波", ["DFT得到F(u,v)。", "fftshift把低频移到中心。", "构造H(u,v)。", "计算G=HF。", "ifftshift并IDFT得到输出图像。"]),
        ("区域生长", ["选种子点。", "设定相似性准则和连通方式。", "检查候选邻域像素。", "满足条件则加入区域并更新区域统计量。", "没有新像素可加入时停止。"]),
    ]
    for name, steps in templates:
        doc.add_heading(name, level=2)
        add_list(doc, steps, "List Number")

    doc.add_heading("八、考前速记", level=1)
    add_list(
        doc,
        [
            "采样管空间分辨率，量化管灰度级分辨率。",
            "点运算只看同位置像素；滤波看邻域。",
            "均衡化增强对比度，但不保证直方图完全平坦。",
            "均值滤波平滑但模糊；中值滤波适合椒盐噪声。",
            "Sobel比Prewitt多中心权重，抗噪性更好。",
            "低通平滑，高通锐化，带阻/陷波常用于周期噪声。",
            "逆滤波怕H接近0；维纳滤波考虑噪声统计。",
            "Otsu最大化类间方差，不是凭眼睛找谷底。",
            "区域生长要更新当前区域均值。",
            "只增强彩色图像亮度，优先处理HSI的I分量。",
        ],
    )
    return doc


def add_question_options(doc, stem, options):
    add_para(doc, stem)
    for idx, option in zip(["A", "B", "C", "D"], options):
        add_para(doc, f"{idx}. {option}")


def build_mock_exam_doc():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "模拟试题含答案")
    add_title(
        doc,
        "数字图像处理模拟试题",
        "题型按课堂图片结构设置：单选15分、多选15分、综合题50分、简答题20分。文末附答案与解析。",
    )

    doc.add_heading("一、单选题（共15题，每题1分，共15分）", level=1)
    for idx, (stem, options, _answer, _explain) in enumerate(SINGLE_CHOICE, 1):
        add_question_options(doc, f"{idx}. {stem}", options)

    doc.add_heading("二、多选题（共5题，每题3分，共15分）", level=1)
    for idx, (stem, options, _answer, _explain) in enumerate(MULTIPLE_CHOICE, 1):
        add_question_options(doc, f"{idx}. {stem}", options)

    doc.add_heading("三、综合题：算法应用与实现（共50分）", level=1)
    for item in COMPREHENSIVE:
        doc.add_heading(item["title"], level=2)
        add_list(doc, item["question"], "List Number")

    doc.add_heading("四、简答题：算法设计与分析（共4题，每题5分，共20分）", level=1)
    for idx, item in enumerate(SHORT_ANSWERS, 1):
        add_para(doc, f"{idx}. {item['question']}")

    doc.add_page_break()
    add_title(doc, "数字图像处理模拟试题答案", "建议先独立完成前半部分，再对照本答案订正。")

    doc.add_heading("一、单选题答案", level=1)
    add_table(
        doc,
        ["题号", "答案", "解析"],
        [(str(idx), answer, explain) for idx, (_stem, _options, answer, explain) in enumerate(SINGLE_CHOICE, 1)],
        [900, 900, 7560],
    )

    doc.add_heading("二、多选题答案", level=1)
    add_table(
        doc,
        ["题号", "答案", "解析"],
        [(str(idx), answer, explain) for idx, (_stem, _options, answer, explain) in enumerate(MULTIPLE_CHOICE, 1)],
        [900, 900, 7560],
    )

    doc.add_heading("三、综合题答案", level=1)
    for item in COMPREHENSIVE:
        doc.add_heading(item["title"], level=2)
        add_list(doc, item["answer"], "List Number")

    doc.add_heading("四、简答题参考答案", level=1)
    for idx, item in enumerate(SHORT_ANSWERS, 1):
        doc.add_heading(f"{idx}. {item['question']}", level=2)
        add_para(doc, item["answer"])

    doc.add_heading("评分提醒", level=1)
    add_list(
        doc,
        [
            "综合题如果公式正确但步骤不完整，会丢过程分；一定写出分组、概率、均值或模板乘加过程。",
            "简答题不要只写名词，按定义、步骤、优缺点、适用条件四层写更稳。",
            "多选题按类别记忆，少选和多选都容易失分，尤其注意维纳滤波不属于阈值分割方法。",
        ],
    )
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    review = build_review_doc()
    review_path = OUT_DIR / "dip-targeted-review-exam-points.docx"
    review.save(review_path)
    print(review_path)

    exam = build_mock_exam_doc()
    exam_path = OUT_DIR / "dip-mock-exam-with-answers.docx"
    exam.save(exam_path)
    print(exam_path)


if __name__ == "__main__":
    main()
