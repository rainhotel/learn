from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\moniC\project\learn")
OUT_DIR = ROOT / "01-topics" / "digital-image-processing" / "review-docx"


CHAPTERS = [
    {
        "file": "01-region-segmentation-review.docx",
        "title": "专题4 图像分割：基于区域的分割",
        "source": "01-专题4图像分割（基于区域的分割）",
        "overview": "本章把分割问题从“单个像素判阈值”推进到“区域整体是否相似”。考试重点是区域分割的五个条件、4/8连通、区域生长三要素、分裂合并法的判断规则，以及高铁轨道缺陷检测案例流程。",
        "key_points": [
            ("区域分割结果条件", "完整性、连通性、独立性、单一性、互斥性。所有子区域并集为原图区域；各区域连通且互不重叠；区域内部满足一致性，任意相邻区域合并后不应仍满足一致性。"),
            ("4连通与8连通", "4连通只看上、下、左、右；8连通还看四个对角邻域。8连通通常得到更大或相同的连通区域。"),
            ("区域生长三要素", "种子点选择、生长规则、终止条件。常见规则是候选像素与当前区域均值的差小于阈值。"),
            ("区域分裂合并", "先用一致性准则判断是否分裂；若区域不均匀就分裂为4个子区域；相邻子区域满足相似性规则时合并。"),
            ("工程案例", "高铁钢轨表面缺陷检测流程：图像采集、预处理、钢轨区域提取、缺陷区域分割、缺陷智能识别。"),
        ],
        "formulas": [
            ("完整性", "R1 ∪ R2 ∪ ... ∪ Rn = R"),
            ("互不重叠", "Ri ∩ Rj = ∅, i ≠ j"),
            ("区域一致性", "P(Ri)=TRUE, P(Ri ∪ Rj)=FALSE"),
            ("标准差准则", "std = sqrt(1/MN * ΣΣ[f(x,y)-m]^2)"),
            ("生长规则示例", "|g(candidate) - mean(region)| < T"),
        ],
        "exam": [
            "区域生长手算题：先确定种子点，再列出4邻域或8邻域候选点；每加入一个像素后，必须重新计算当前区域均值，再继续判断。",
            "分裂合并手算题：先算整块均值和标准差；若标准差超过阈值则分裂成4块；再分别算子块标准差；最后用相邻块均值差判断合并。",
            "问机器视觉检测钢轨缺陷的优缺点：优点是自动、实时、高效、高精度；缺点是难以检测钢轨内部缺陷。",
        ],
        "pitfalls": [
            "区域生长不是固定和种子点比较，PPT中的分析题是与当前区域均值比较。",
            "8连通不是更严格，而是邻接关系更多。",
            "区域分裂和合并的阈值含义不同：分裂看区域内部波动，合并看相邻区域均值相似。",
        ],
    },
    {
        "file": "02-threshold-segmentation-review.docx",
        "title": "专题4 图像分割：阈值分割",
        "source": "02-专题4图像分割（阈值分割）",
        "overview": "阈值分割用目标与背景灰度差异完成分割，是最常见的相似性分割方法。本章考点集中在单阈值/多阈值、全局/局部/自适应阈值、全局阈值迭代法、OTSU大津法以及噪声和光照对阈值的影响。",
        "key_points": [
            ("基本思想", "选一个或多个阈值，把像素灰度与阈值比较，将图像划分为目标和背景或多个类别。"),
            ("单阈值与多阈值", "直方图双峰明显时常用单阈值；多峰明显时可用多个阈值区分多个目标或背景层次。"),
            ("阈值类型", "全局阈值整幅图一个T；局部阈值依赖邻域性质；自适应阈值随位置和局部统计动态变化。"),
            ("迭代法", "先用初始T分成两类，求两类均值，再取均值中点作为新T，直到收敛。"),
            ("OTSU", "遍历所有灰度阈值，选择使类间方差最大的阈值；思想是类内差异小、类间差异大。"),
        ],
        "formulas": [
            ("二值阈值分割", "g(x,y)=1 if f(x,y)>T; g(x,y)=0 if f(x,y)<=T"),
            ("通用阈值", "T = T[x, y, p(x,y), f(x,y)]"),
            ("迭代阈值更新", "T_new = (μ1 + μ2)/2"),
            ("类间方差", "σB^2 = P1(m1-mG)^2 + P2(m2-mG)^2"),
            ("OTSU简化形式", "σB^2 = P1 P2 (m1-m2)^2"),
        ],
        "exam": [
            "迭代法计算：给初始T，分组G1/G2，算均值μ1/μ2，更新T=(μ1+μ2)/2，直到前后阈值差小于给定误差。",
            "OTSU计算：枚举阈值k，分别算P1、P2、m1、m2，再算P1P2(m1-m2)^2，取最大者。",
            "PPT主观题：对高斯噪声七边形图用OTSU时，噪声会让边缘和背景有误检，通常先平滑；对光照不均七边形，全局OTSU容易失败，应使用局部或自适应阈值。",
        ],
        "pitfalls": [
            "OTSU不是直接找直方图谷底，而是最大化类间方差。",
            "阈值太小会把背景误分成目标，阈值太大会漏掉目标内部。",
            "全局阈值不适合光照渐变或阴影明显的图像。",
        ],
    },
    {
        "file": "03-image-dehazing-review.docx",
        "title": "专题3 频域图像处理：前沿案例 图像去雾",
        "source": "03-专题3频域图像处理（前沿案例：图像去雾）",
        "overview": "图像去雾是基于物理成像模型的复原/增强案例。核心是大气散射模型、暗通道先验、大气光A估计、透射率t估计以及恢复公式。深度学习部分重点比较DehazeNet和AOD-Net。",
        "key_points": [
            ("去雾本质", "雾霾使对比度降低、画面模糊、颜色偏移和细节丢失。去雾不是简单锐化，而是根据雾的成像模型估计清晰图像。"),
            ("大气散射模型", "有雾图像由衰减后的真实场景光和大气光叠加组成。已知I，未知J、t、A，因此是病态问题。"),
            ("暗通道先验", "无雾自然图像局部窗口内，RGB至少有一个通道值很低，暗通道接近0。"),
            ("大气光估计", "计算暗通道，选暗通道最亮的前0.1%像素，再回到原图中取亮度最大的RGB作为A。"),
            ("透射率估计", "由暗通道先验推导t，实际加ω保留少量雾，使画面自然。"),
            ("深度学习方法", "DehazeNet用CNN估计透射率；AOD-Net把t和A重参数化为K，端到端输出J。"),
        ],
        "formulas": [
            ("大气散射模型", "I(x)=J(x)t(x)+A(1-t(x))"),
            ("暗通道", "J_dark(x)=min_c min_{y∈Ω(x)} J^c(y)"),
            ("透射率估计", "t~(x)=1-ω min_c min_{y∈Ω(x)} I^c(y)/A^c"),
            ("清晰图恢复", "J(x)=(I(x)-A)/max(t(x),t0)+A"),
            ("AOD-Net形式", "J(x)=K(x)I(x)-K(x)+b"),
        ],
        "exam": [
            "讨论题：没有配对无雾图的行车记录仪视频，可优先选DCP，因为它不依赖配对训练数据；若强调实时性，可用预训练轻量网络。",
            "计算题：给I、A、t，直接代入J=(I-A)/t+A；若t很小，要说明应使用max(t,t0)防止噪声放大。",
            "概念题：当t趋近0时，I≈A，表示极远处或雾最浓区域颜色接近大气光。",
        ],
        "pitfalls": [
            "A是大气光，不是清晰图像。",
            "暗通道先验在天空、白墙、水面、大面积亮区域可能失效。",
            "DCP效果好但计算慢，天空区域可能偏暗或有光晕。",
        ],
    },
    {
        "file": "04-image-restoration-review.docx",
        "title": "专题3 频域图像处理：图像复原",
        "source": "04-专题3频域图像处理（图像复原）",
        "overview": "图像复原强调根据退化模型反推原图，比图像增强更客观。考点包括退化模型、退化函数估计、运动模糊/大气湍流模型、逆滤波缺陷、维纳滤波公式和周期噪声频域处理。",
        "key_points": [
            ("复原与增强", "增强重视觉效果，主观性强；复原重退化模型和真实性，目标是接近原图。"),
            ("退化模型", "输入图像经过退化系统H并叠加噪声，得到退化图像。空间域卷积在频域中变为乘法。"),
            ("退化函数估计", "图像观察估计法、试验估计法、模型估计法。试验法可输入冲激图像，输出频谱除以冲激幅值A得到H。"),
            ("运动模糊", "匀速运动造成方向性模糊，退化函数与速度、方向和曝光时间有关。"),
            ("大气湍流", "用指数衰减模型描述，k越大，高频衰减越强，图像越模糊。"),
            ("逆滤波与维纳滤波", "逆滤波直接除以H，容易放大噪声；维纳滤波综合退化函数和噪声统计，最小均方误差意义下更稳定。"),
        ],
        "formulas": [
            ("空间退化模型", "g(x,y)=h(x,y)*f(x,y)+n(x,y)"),
            ("频域退化模型", "G(u,v)=H(u,v)F(u,v)+N(u,v)"),
            ("逆滤波", "F_hat(u,v)=G(u,v)/H(u,v)"),
            ("噪声放大", "F_hat=F+N/H"),
            ("维纳滤波", "F_hat=[H*/(|H|^2+Sn/Sf)]G"),
            ("大气湍流", "H(u,v)=exp[-k(u^2+v^2)^(5/6)]"),
        ],
        "exam": [
            "问为什么已知H也不能完全复原：因为有随机噪声N，逆滤波出现N/H；当H很小时噪声被严重放大。",
            "反滤波步骤：对g做DFT得到G，计算H，用G/H估计F，再IDFT得到复原图。",
            "维纳滤波极限：噪声功率趋近0时接近逆滤波；原图功率趋近0时说明没有有用信息可恢复。",
        ],
        "pitfalls": [
            "复原不是单纯让图像更好看，而是依据退化过程反推。",
            "直接全频逆滤波效果通常差，需要限制频率范围或使用维纳滤波。",
            "周期噪声在频谱中常表现为成对亮点，常用陷波或带阻滤波处理。",
        ],
    },
    {
        "file": "05-high-pass-filters-review.docx",
        "title": "专题3 频域图像处理：高通、带通与带阻滤波器",
        "source": "05-专题3频域图像处理（高通滤波器）",
        "overview": "高通滤波用于强调边缘和细节，带阻/陷波用于抑制周期噪声。本章考点包括理想/巴特沃斯/高斯高通、高提升滤波、高频强调、同态滤波思想、带阻与带通互补关系。",
        "key_points": [
            ("高通滤波", "保留高频、抑制低频，图像边缘和细节被增强，但噪声也可能增强。"),
            ("高通类型", "理想高通截止突然，可能振铃；巴特沃斯高通过渡可调；高斯高通过渡最平滑。"),
            ("高提升滤波", "在高通结果中保留一部分原图低频，使图像既锐化又保留背景。"),
            ("高频强调", "H_hfe=a+bH_hp，a保留低频，b增强高频。"),
            ("带阻与带通", "带阻去除某一频带，带通保留某一频带，二者互补；周期噪声常用带阻或陷波。"),
            ("同态滤波", "通过对数变换把照度和反射乘法模型变成加法，在频域压低照度低频、增强反射高频。"),
        ],
        "formulas": [
            ("频域滤波", "G(u,v)=H(u,v)F(u,v)"),
            ("高通与低通互补", "H_hp(u,v)=1-H_lp(u,v)"),
            ("高提升滤波", "H_hb=(A-1)+H_hp"),
            ("高频强调", "H_hfe=a+bH_hp"),
            ("带通互补", "H_bp=1-H_br"),
            ("同态滤波常见形式", "H=(γH-γL)(1-exp[-cD^2/D0^2])+γL"),
        ],
        "exam": [
            "判断滤波器类型：中心为暗、外围为亮的是高通；中心为亮、外围为暗的是低通。",
            "问A=1的高提升滤波是什么：退化为普通高通滤波。",
            "周期噪声题：频谱有规则对称亮点，用带阻、带通观察或陷波滤波去除。",
        ],
        "pitfalls": [
            "高通不是去噪工具，高频噪声会被保留或增强。",
            "高提升和高频强调都不是完全删除低频，而是保留一定背景信息。",
            "理想高通/低通因为频域突变，空间域更容易产生振铃。",
        ],
    },
    {
        "file": "06-low-pass-filters-review.docx",
        "title": "专题3 频域图像处理：低通滤波器",
        "source": "06-专题3频域图像处理（低通滤波器）",
        "overview": "低通滤波用于频域平滑和降噪，核心是保留低频、抑制高频。考点包括频域滤波流程、理想/巴特沃斯/高斯低通、截止频率D0影响和振铃现象。",
        "key_points": [
            ("低频与高频", "低频代表背景、亮度和缓慢变化；高频代表边缘、细节和噪声。"),
            ("低通作用", "保留低频使图像平滑，去除高频噪声，但会牺牲边缘细节。"),
            ("理想低通", "D<=D0时H=1，否则H=0；频域边界突然，空间域易产生振铃。"),
            ("巴特沃斯低通", "通过阶数n控制过渡陡峭程度；n越大越接近理想低通。"),
            ("高斯低通", "过渡平滑，没有突然截断，振铃最少。"),
            ("D0影响", "D0越小，保留频率越少，图像越模糊；D0越大，保留细节越多。"),
        ],
        "formulas": [
            ("理想低通", "H=1 if D(u,v)<=D0; H=0 if D(u,v)>D0"),
            ("巴特沃斯低通", "H=1/[1+(D/D0)^(2n)]"),
            ("高斯低通", "H=exp[-D^2/(2D0^2)]"),
            ("滤波输出", "G(u,v)=H(u,v)F(u,v)"),
        ],
        "exam": [
            "选择去噪滤波器：随机高频噪声通常选低通或平滑处理。",
            "问理想低通缺点：频域截断过陡，空间域出现振铃。",
            "判断D0变化：D0变大，图像更清晰但噪声可能保留；D0变小，图像更模糊。",
        ],
        "pitfalls": [
            "低通不是增强边缘，而是抑制边缘和噪声等高频。",
            "高斯低通不是截止型硬阈值，而是连续衰减。",
            "频谱图经过fftshift后中心才是低频。",
        ],
    },
    {
        "file": "07-fourier-transform-review.docx",
        "title": "专题3 频域图像处理：傅里叶变换",
        "source": "07-专题3频域图像处理（傅里叶变换）",
        "overview": "傅里叶变换是频域图像处理的基础。本章要掌握2D DFT/IDFT、频谱意义、低频高频对应关系、F(0,0)与图像灰度和、频谱中心化和频域滤波基本流程。",
        "key_points": [
            ("空间域与频域", "空间域关注像素位置和灰度；频域关注不同频率成分的强弱。"),
            ("低频与高频", "低频表示平滑区域和整体亮度，高频表示边缘、纹理、细节和噪声。"),
            ("直流分量", "F(0,0)等于所有像素灰度和，图像均值为F(0,0)/(MN)。"),
            ("频谱中心化", "fftshift把低频移动到图像中心，便于观察和设计滤波器。"),
            ("频域滤波流程", "原图DFT得到F，乘以滤波器H得到G，再IDFT回空间域。"),
        ],
        "formulas": [
            ("二维DFT", "F(u,v)=ΣΣ f(x,y)e^{-j2π(ux/M+vy/N)}"),
            ("二维IDFT", "f(x,y)=1/MN ΣΣ F(u,v)e^{j2π(ux/M+vy/N)}"),
            ("直流分量", "F(0,0)=ΣΣ f(x,y)"),
            ("图像均值", "mean=F(0,0)/(MN)"),
            ("频域滤波", "G(u,v)=H(u,v)F(u,v)"),
        ],
        "exam": [
            "给2×2或3×3图像求F(0,0)：直接把所有像素求和。",
            "问低频/高频意义：低频对应平滑和亮度，高频对应边缘细节噪声。",
            "问fftshift作用：把低频移动到频谱中心，而不是改变图像本身内容。",
        ],
        "pitfalls": [
            "未shift时低频在角落，shift后低频在中心。",
            "频域滤波是乘法，不是把H和F相加。",
            "高频不总是有用细节，很多噪声也是高频。",
        ],
    },
    {
        "file": "08-color-image-processing-review.docx",
        "title": "专题2 空域图像处理：彩色图像处理",
        "source": "08-专题2空域图像处理（彩色图像处理）",
        "overview": "彩色图像处理考点包括RGB、CMY/CMYK、HSI模型、伪彩色增强、彩色图像平滑锐化以及基于颜色规则的目标检测。重点是理解不同颜色模型适合处理的问题。",
        "key_points": [
            ("RGB模型", "用红、绿、蓝三个分量叠加表示颜色，数字图像常用M×N×3数组存储。"),
            ("CMY/CMYK模型", "青、品红、黄是RGB的补色；印刷常加入黑色K提高表现和节省油墨。"),
            ("HSI模型", "H表示色调，S表示饱和度，I表示亮度，更贴近人类对颜色的感知。"),
            ("彩色增强", "可以在RGB三通道分别处理，也可以转到HSI只处理I分量以尽量保持色彩。"),
            ("伪彩色", "把灰度范围映射为不同颜色，提高人眼区分能力。"),
            ("肤色检测", "可用RGB阈值规则检测皮肤区域，但受光照和人种等因素影响。"),
        ],
        "formulas": [
            ("CMY转换", "C=1-R, M=1-G, Y=1-B"),
            ("RGB数据量", "M×N×3 bytes for uint8 RGB image"),
            ("肤色规则示例", "R>95, G>40, B>20, R>G, R>B, max-min>15, |R-G|>15"),
        ],
        "exam": [
            "问只改变亮度尽量不改变颜色：转到HSI空间处理I分量。",
            "问RGB与HSI区别：RGB面向设备显示，HSI更符合人眼感知。",
            "给RGB像素判断是否满足肤色条件：逐项代入规则，所有条件同时成立才判为肤色。",
        ],
        "pitfalls": [
            "RGB三通道分别增强可能改变颜色平衡。",
            "伪彩色不是恢复真实颜色，而是把灰度映射为便于观察的颜色。",
            "颜色阈值法对光照变化敏感。",
        ],
    },
    {
        "file": "09-spatial-sharpening-review.docx",
        "title": "专题2 空域图像处理：空域锐化",
        "source": "09-专题2空域图像处理（空域锐化）",
        "overview": "空域锐化通过一阶或二阶导数增强灰度突变，突出边缘和细节。考试重点是梯度、Roberts/Prewitt/Sobel算子、拉普拉斯算子、锐化公式以及噪声放大风险。",
        "key_points": [
            ("锐化本质", "增强灰度快速变化区域，使边缘、线条和细节更明显。"),
            ("一阶导数", "梯度用于检测边缘强度和方向，常用Roberts、Prewitt、Sobel算子。"),
            ("Sobel特点", "中间行或列权重为2，具有一定平滑效果，比Prewitt更抗噪。"),
            ("二阶导数", "拉普拉斯算子各向同性，响应灰度突变，常用于边缘增强。"),
            ("锐化风险", "噪声也常是高频，锐化会增强噪声，因此噪声图像通常先平滑再锐化。"),
        ],
        "formulas": [
            ("梯度幅值近似", "|G|=|Gx|+|Gy|"),
            ("梯度幅值精确", "|G|=sqrt(Gx^2+Gy^2)"),
            ("拉普拉斯", "∇^2 f = ∂^2f/∂x^2 + ∂^2f/∂y^2"),
            ("锐化", "g=f+c∇^2f 或 g=f-c∇^2f，取决于模板符号"),
            ("Sobel Gx", "[-1 0 1; -2 0 2; -1 0 1]"),
        ],
        "exam": [
            "Sobel手算：模板与3×3邻域对应相乘并求和，分别求Gx和Gy，再算幅值。",
            "选择题：Sobel比Prewitt抗噪更好，因为中间权重为2。",
            "判断题：锐化能增强边缘，但也会增强噪声。",
        ],
        "pitfalls": [
            "拉普拉斯模板符号不同，最后加减号也会不同。",
            "高通/锐化不等于去噪。",
            "一阶导数主要用于边缘强度，二阶导数常用于过零点和边缘定位。",
        ],
    },
    {
        "file": "10-spatial-smoothing-review.docx",
        "title": "专题2 空域图像处理：空域平滑",
        "source": "10-专题2空域图像处理（空域平滑）",
        "overview": "空域平滑主要用于去噪和模糊，考点包括算术均值、几何均值、谐波均值、逆谐波均值、阿尔法修剪均值和中值滤波，尤其是不同噪声对应的滤波器选择。",
        "key_points": [
            ("算术均值", "邻域平均，适合随机噪声，但会模糊边缘。"),
            ("几何均值", "对乘性形式更友好，平滑同时保留细节较好。"),
            ("谐波均值", "适合盐噪声，不适合椒噪声。"),
            ("逆谐波均值", "通过参数Q控制去除椒噪声或盐噪声。"),
            ("中值滤波", "非线性滤波，排序取中位数，特别适合椒盐噪声。"),
            ("阿尔法修剪均值", "先去掉若干最大值和最小值，再平均，兼顾均值和中值思想。"),
        ],
        "formulas": [
            ("算术均值", "g=1/mn Σ f(s,t)"),
            ("几何均值", "g=(Π f(s,t))^(1/mn)"),
            ("谐波均值", "g=mn / Σ[1/f(s,t)]"),
            ("逆谐波均值", "g=Σ f^(Q+1) / Σ f^Q"),
            ("逆谐波特例", "Q=0为算术均值，Q=-1为谐波均值"),
        ],
        "exam": [
            "中值滤波手算：把邻域像素排序，奇数个取中间值；PPT例子排序后中值为205。",
            "逆谐波判断：Q>0去椒噪声，Q<0去盐噪声。",
            "问均值滤波缺点：平滑噪声的同时会模糊边缘和细节。",
        ],
        "pitfalls": [
            "中值滤波不是线性滤波。",
            "椒噪声是黑点，盐噪声是白点，不要把Q符号记反。",
            "几何/谐波均值在存在0值时要小心计算。",
        ],
    },
    {
        "file": "11-spatial-filtering-review.docx",
        "title": "专题2 空域图像处理：空域滤波",
        "source": "11-专题2空域图像处理（空域滤波）",
        "overview": "空域滤波是后续平滑和锐化的基础。核心考点包括模板/核/窗口、相关与卷积、边界处理、线性滤波加权求和以及模板手算。",
        "key_points": [
            ("空域滤波定义", "在图像空间域中用邻域模板对每个像素进行局部运算。"),
            ("模板运算", "模板覆盖邻域，权值与像素对应相乘后求和，得到中心像素输出。"),
            ("相关与卷积", "相关不翻转模板；卷积要把模板旋转180°。若模板对称，两者结果相同。"),
            ("边界处理", "常见有零填充、复制边界、镜像扩展、忽略边界等。"),
            ("滤波用途", "平滑滤波用于降噪和模糊，锐化滤波用于边缘和细节增强。"),
        ],
        "formulas": [
            ("线性空域滤波", "g(x,y)=ΣΣ w(s,t)f(x+s,y+t)"),
            ("均值模板", "1/9 * ones(3,3)"),
            ("加权均值模板", "1/16 * [1 2 1; 2 4 2; 1 2 1]"),
        ],
        "exam": [
            "模板手算：对应位置相乘，全部求和，若模板有归一化系数再相除。",
            "判断相关/卷积：题目如果强调卷积且模板不对称，要先旋转模板。",
            "PPT均值滤波填空：按3×3邻域平均计算中心值，注意只对指定窗口内像素求和。",
        ],
        "pitfalls": [
            "不要把卷积和相关无条件等同。",
            "边界处理方法不同，图像边缘输出会不同。",
            "滤波模板权值和不一定为1；锐化模板常常权值和为0或1。",
        ],
    },
    {
        "file": "12-pixel-relations-noise-review.docx",
        "title": "专题2 空域图像处理：像素间联系与随机噪声估计",
        "source": "12-专题2空域图像处理（像素间联系+随机噪声估计）",
        "overview": "本章连接图像几何基础和噪声建模。考点包括4邻域、8邻域、m邻接、连通性、距离度量，以及常见随机噪声的概率分布和识别。",
        "key_points": [
            ("邻域", "4邻域包括上下左右；对角邻域包括四个对角；8邻域为二者合并。"),
            ("邻接", "4邻接只允许4邻域相邻；8邻接允许8邻域；m邻接用于避免8邻接引起的连接歧义。"),
            ("连通路径", "若像素序列中相邻像素满足指定邻接关系，则形成路径；两个像素之间存在路径即连通。"),
            ("距离度量", "欧氏距离、D4城市街区距离、D8棋盘距离。"),
            ("噪声识别", "根据噪声分布、直方图形状和图像表现判断高斯、瑞利、伽马、指数、均匀、椒盐等噪声。"),
        ],
        "formulas": [
            ("欧氏距离", "De=sqrt((x-s)^2+(y-t)^2)"),
            ("D4距离", "D4=|x-s|+|y-t|"),
            ("D8距离", "D8=max(|x-s|,|y-t|)"),
            ("高斯噪声", "p(z)=1/(sqrt(2π)σ) exp[-(z-μ)^2/(2σ^2)]"),
        ],
        "exam": [
            "距离计算题：两点相差(3,3)时，De=3sqrt(2)，D4=6，D8=3。",
            "m邻接判断：若两个对角像素的共同4邻域中存在同集合像素，则不能认为m邻接。",
            "噪声题：椒盐噪声表现为随机黑白点；高斯噪声直方图近似钟形。",
        ],
        "pitfalls": [
            "8邻接可能把只在角上接触的区域连起来，m邻接用于消除这种歧义。",
            "D4和D8不是欧氏距离的近似写法，而是不同的离散距离定义。",
            "椒噪声是低灰度黑点，盐噪声是高灰度白点。",
        ],
    },
    {
        "file": "13-histogram-processing-review.docx",
        "title": "专题2 空域图像处理：直方图处理",
        "source": "13-专题2空域图像处理（直方图处理）",
        "overview": "直方图描述灰度分布，是图像增强和阈值分割的重要工具。考点包括直方图统计、归一化直方图、直方图均衡化、直方图规定化，以及算术/逻辑运算中图像平均降噪。",
        "key_points": [
            ("直方图", "统计各灰度级出现次数；归一化后表示每个灰度级出现的概率。"),
            ("均衡化思想", "通过累积分布函数进行灰度映射，使灰度范围展开，增强全局对比度。"),
            ("均衡化步骤", "统计直方图，归一化，计算CDF，乘以L-1并取整，完成灰度映射。"),
            ("图像平均降噪", "多幅独立同分布噪声图像求平均，噪声均值趋近0，噪声标准差约缩小为1/sqrt(M)。"),
            ("算术/逻辑运算", "加减乘除可用于增强、背景校正、差分检测；逻辑运算常用于二值图像掩膜处理。"),
        ],
        "formulas": [
            ("归一化直方图", "p(rk)=nk/n"),
            ("均衡化映射", "sk=(L-1)Σ_{j=0}^k p(rj)"),
            ("平均降噪", "σ_avg=σ/sqrt(M)"),
        ],
        "exam": [
            "均衡化手算：先算累计概率CDF，再乘L-1，最后按题目要求取整。",
            "图像平均题：若平均M幅独立噪声图像，噪声标准差变为原来的1/sqrt(M)。",
            "判断题：直方图均衡化不保证输出直方图完全平坦。",
        ],
        "pitfalls": [
            "均衡化可能增强噪声，不一定使主观效果更好。",
            "直方图只反映灰度分布，不包含空间位置信息。",
            "减法运算常用于变化检测，但可能出现负值，需要截断或归一化。",
        ],
    },
    {
        "file": "14-gray-transform-review.docx",
        "title": "专题2 空域图像处理：灰度变换",
        "source": "14-专题2空域图像处理（灰度变换）",
        "overview": "灰度变换是最直接的点运算增强方法。本章考点包括线性变换、图像反转、对数变换、幂律/伽马变换、分段线性变换和灰度级切片。",
        "key_points": [
            ("点运算", "输出像素只依赖对应输入像素，不考虑邻域。形式为s=T(r)。"),
            ("图像反转", "适合增强暗背景中的亮细节，如X光片中的白色结构。"),
            ("对数变换", "扩展低灰度、压缩高灰度，适合增强暗部细节或压缩动态范围。"),
            ("幂律变换", "γ<1使图像变亮，γ>1使图像变暗，可用于显示设备伽马校正。"),
            ("分段线性变换", "通过不同灰度区间不同斜率实现对比度拉伸或灰度切片。"),
        ],
        "formulas": [
            ("通用灰度变换", "s=T(r)"),
            ("反转", "s=L-1-r"),
            ("对数变换", "s=c log(1+r)"),
            ("幂律变换", "s=c r^γ"),
            ("归一化对数常数", "c=(L-1)/log(1+r_max)"),
        ],
        "exam": [
            "判断曝光过度/不足：曝光过度用γ>1压暗；曝光不足用γ<1提亮。",
            "对数变换计算：若8位图最大值映射到255，则c=255/log(256)。",
            "灰度切片：突出某个灰度范围，可保持背景或将背景压到固定值。",
        ],
        "pitfalls": [
            "γ<1和γ>1效果最容易记反。",
            "对数变换不是单纯整体变亮，而是暗部扩展、亮部压缩。",
            "点运算不会利用邻域信息，不能直接完成空间去噪。",
        ],
    },
    {
        "file": "15-digital-image-fundamentals-review.docx",
        "title": "专题1 数字图像基础",
        "source": "15-专题1数字图像基础",
        "overview": "基础章关注数字图像如何在计算机中表示。考点包括像素、二维矩阵、二值/灰度/RGB图像、数据类型、数据量计算、位平面、LSB水印和Matlab读取图像数据。",
        "key_points": [
            ("像素", "数字图像由像素按矩阵紧密排列组成。每个像素有位置坐标和像素值两个基本属性。"),
            ("二维矩阵表示", "灰度图像可表示为M×N矩阵F，矩阵元素fij对应像素灰度值。"),
            ("二值图像", "每个像素用1个二进制位表示，通常1代表白色、0代表黑色；Matlab中常为logical。"),
            ("灰度图像", "常用0到255的整数表示，0为黑、255为白，中间为不同深浅灰色；Matlab中常为uint8。"),
            ("RGB彩色图像", "每个像素包含R/G/B三个分量，每个分量通常8位，可表示256^3=16777216种颜色。"),
            ("位平面", "8位灰度图可分成8个位平面，高位平面承载主要视觉信息，低位平面更像细节和噪声，可用于水印。"),
        ],
        "formulas": [
            ("灰度图数据量", "M×N×bits_per_pixel / 8 bytes"),
            ("RGB真彩色数据量", "M×N×3 bytes for 8-bit RGB"),
            ("颜色数量", "256×256×256=16777216"),
            ("位平面1提取", "bitand(f,1,'uint8')"),
            ("位平面3-8组合", "bitand(f,252,'uint8')"),
        ],
        "exam": [
            "灰度级0-255需要多少比特：256=2^8，所以需要8 bit。",
            "1024×1024真彩色图像数据量：1024×1024×24 bit = 1024×1024×3 byte = 3 MB。",
            "提取位平面3-8组合：二进制掩码11111100为252，所以Matlab代码为bitand(f,252,'uint8')。",
        ],
        "pitfalls": [
            "灰度级数256不是需要256 bit，而是8 bit。",
            "RGB图像在Matlab中尺寸通常是M×N×3。",
            "低位平面改动对视觉影响小，但不能说完全无影响。",
        ],
    },
    {
        "file": "16-matlab-image-programming-review.docx",
        "title": "Matlab 图像编程基础",
        "source": "16-Matlab 图像编程基础",
        "overview": "本章服务于上机和作业，重点不是图像理论，而是Matlab矩阵、控制流、函数、输入输出和基本图像处理命令。考试或作业常要求写出命令、解释矩阵索引和完成简单图像操作。",
        "key_points": [
            ("矩阵思想", "Matlab以矩阵为核心，图像本质也是矩阵或三维数组。"),
            ("基础命令", "clc清命令窗口，clear清变量，close all关闭图窗，who/whos查看工作区变量，save/load保存或加载变量。"),
            ("矩阵创建", "冒号表达式、linspace、zeros、ones、reshape等用于生成和改变矩阵。"),
            ("矩阵索引", "A(i,j)访问第i行第j列；A(:,j)取第j列；A(i,:)取第i行；end表示末尾索引。"),
            ("控制流", "if判断、while循环、for循环用于实现算法流程。"),
            ("图像操作", "imread读取图像，imshow显示，imhist显示直方图，histeq均衡化，imwrite保存。"),
        ],
        "formulas": [
            ("for循环", "for i=1:n; statements; end"),
            ("while循环", "while condition; statements; end"),
            ("if语句", "if condition; statements; end"),
            ("读图显示", "I=imread('file.tif'); imshow(I)"),
            ("直方图均衡化", "I2=histeq(I); imhist(I2)"),
        ],
        "exam": [
            "矩阵题：已知y=[1 2 3;4 5 6]，y(1,:)为第一行，y(:,2)为第二列，y(2:end,2:end)取右下子矩阵。",
            "表达式题：计算矩阵t的z=1/2*exp(2t)*log(t+sqrt(1+t.^2))时，矩阵逐元素运算要用点乘、点除、点幂。",
            "作业题：分别读入二值、灰度、彩色图像，用whos比较logical、uint8、M×N与M×N×3的差异。",
        ],
        "pitfalls": [
            "矩阵逐元素运算要用.*, ./, .^，不能误用矩阵乘法。",
            "Matlab索引从1开始，不是从0开始。",
            "命令后加分号不会显示计算结果，但仍会执行。",
        ],
    },
    {
        "file": "17-course-overview-matlab-basics-review.docx",
        "title": "课程概述与 Matlab 图像处理基础",
        "source": "17-课程概述+Matlab图像处理基础",
        "overview": "本章是课程导入，解释为什么要处理图像、什么是数字图像处理、AI时代为什么仍要学传统DIP，并介绍课程结构、考核方式与Matlab环境。复习时重点是建立课程地图和基础工具意识。",
        "key_points": [
            ("为什么处理图像", "图像处理服务于图像获取、显示打印、存储传输、增强复原、分析识别、特效和变形等任务。"),
            ("DIP流程", "从问题出发，经过图像获取、增强、复原、形态学处理、分割、目标检测、特征表示与描述、图像压缩等环节。"),
            ("传统DIP与AI", "传统方法提供可解释的基础原理和工程工具；AI不是唯一工具，传统方法仍适合数据少、资源受限、需解释验证的场景。"),
            ("课程内容", "专题包括数字图像基础、空域图像处理、频域图像处理、图像分割、特征表示与深度学习等。"),
            ("考核方式", "总评=期末笔试60%+平时成绩40%；平时包括作业、雨课堂/学习通、上机实习和期中考试等。"),
            ("Matlab基础", "Matlab是矩阵实验室，适合图像矩阵计算、绘图、算法验证和上机实践。"),
        ],
        "formulas": [
            ("课程主线", "处理图像 -> 分析图像 -> 理解图像"),
            ("总评比例", "期末60% + 平时40%"),
            ("Matlab变量赋值", "variable_name = expression;"),
            ("常用函数形式", "result = function_name(input);"),
        ],
        "exam": [
            "雨课堂选择题：总评成绩中期末笔试60%，平时成绩40%，答案选“期末60%，平时40%”。",
            "简答题：AI时代为什么仍要学传统DIP？答可解释、数据少成本低、资源受限可落地、是AI视觉算法的基础零部件。",
            "命令题：randn生成正态随机数，可用help randn查询；生成均值3、方差1的500个数据可写为x=randn(1,500)+3。",
        ],
        "pitfalls": [
            "不要把数字图像处理等同于Photoshop修图；本课程强调算法、模型和工程流程。",
            "AI方法不是所有场景最优，传统方法在可解释性和成本方面仍重要。",
            "Matlab中的逻辑运算符包括&, |, xor, ~；关系运算结果为0或1。",
        ],
    },
]


def set_east_asian_font(run, font_name="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph(style="FormulaBlock")
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    return p


def add_list(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(item)
        set_east_asian_font(run)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    title = styles.add_style("ReviewTitle", 1)
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.15

    subtitle = styles.add_style("SubtitleLine", 1)
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(10)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)

    formula = styles.add_style("FormulaBlock", 1)
    formula.font.name = "Consolas"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    formula.font.size = Pt(9.5)
    formula.paragraph_format.left_indent = Inches(0.18)
    formula.paragraph_format.right_indent = Inches(0.05)
    formula.paragraph_format.space_before = Pt(2)
    formula.paragraph_format.space_after = Pt(5)
    formula.paragraph_format.line_spacing = 1.15

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_header_footer(doc, chapter_title):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("数字图像处理复习资料")
    set_east_asian_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(chapter_title)
    set_east_asian_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("777777")


def add_key_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    hdr = table.rows[0].cells
    hdr[0].text = "考点"
    hdr[1].text = "复习要点"
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.runs[0].bold = True
            set_east_asian_font(p.runs[0])
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            for p in cell.paragraphs:
                if p.runs:
                    set_east_asian_font(p.runs[0])
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
    return table


def add_formula_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    hdr = table.rows[0].cells
    hdr[0].text = "公式/模型"
    hdr[1].text = "表达式"
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.runs[0].bold = True
            set_east_asian_font(p.runs[0])
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for p in cells[0].paragraphs:
            if p.runs:
                set_east_asian_font(p.runs[0])
        for p in cells[1].paragraphs:
            for run in p.runs:
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
    return table


WORKED_EXAMPLES = {
    "01-region-segmentation-review.docx": [
        {
            "title": "区域生长：每加入一个像素都要更新均值",
            "question": "种子灰度为100，已加入像素102、104，阈值T=5。候选点灰度为109，是否能加入区域？",
            "steps": [
                "先算当前区域均值：(100+102+104)/3=102。",
                "把候选点与当前区域均值比较，而不是只和最初种子点比较：|109-102|=7。",
                "因为7>T，所以候选点不能加入。本题考的就是生长规则和终止条件。"
            ],
            "answer": "不能加入。答题时写清楚“当前区域均值”四个字，避免被扣步骤分。",
        },
        {
            "title": "分裂合并：先分裂再合并",
            "question": "一个区域内部灰度标准差大于阈值，分成4块后，相邻两块均值差小于合并阈值，应如何处理？",
            "steps": [
                "判断整块是否均匀：标准差大于阈值，说明不满足一致性，先分裂。",
                "分别检查子块：每块满足一致性后，才进入合并判断。",
                "相邻子块均值差小于阈值，说明合并后仍相似，可以合并。"
            ],
            "answer": "先分裂为4个子区域，再把满足相似性条件的相邻子区域合并。",
        },
    ],
    "02-threshold-segmentation-review.docx": [
        {
            "title": "迭代阈值：分组-求均值-更新",
            "question": "灰度集合为{10,12,14,200,210,220}，初始阈值T0=100，求一次迭代后的阈值。",
            "steps": [
                "按T0分组：G1={10,12,14}，G2={200,210,220}。",
                "求两组均值：μ1=12，μ2=210。",
                "更新阈值：T1=(μ1+μ2)/2=(12+210)/2=111。"
            ],
            "answer": "一次迭代后阈值为111。若再迭代分组不变，则阈值收敛。",
        },
        {
            "title": "OTSU：不是找谷底，是最大化类间方差",
            "question": "若阈值k1得到P1P2(m1-m2)^2=1200，阈值k2得到1500，应选哪个阈值？",
            "steps": [
                "OTSU目标是让目标和背景分得最开，即类间方差最大。",
                "比较各候选阈值的类间方差。",
                "1500大于1200，所以选择k2。"
            ],
            "answer": "选择k2。写理由时用“类间方差最大”而不是“看起来谷底更深”。",
        },
    ],
    "03-image-dehazing-review.docx": [
        {
            "title": "大气散射模型恢复清晰像素",
            "question": "单通道下I=180，A=220，t=0.4，求恢复后的J。",
            "steps": [
                "使用恢复式J=(I-A)/t+A。",
                "代入：J=(180-220)/0.4+220。",
                "计算：(-40)/0.4=-100，J=120。"
            ],
            "answer": "J=120。若t很小，实际算法要用max(t,t0)防止噪声被放大。",
        },
        {
            "title": "暗通道先验的解题流程",
            "question": "简述暗通道先验去雾如何估计A和t。",
            "steps": [
                "对每个局部窗口取RGB三通道最小值，得到暗通道图。",
                "在暗通道中找最亮的一小部分像素，回到原图寻找亮度最大的像素估计大气光A。",
                "用t(x)=1-ω min_c min_y I^c(y)/A^c估计透射率，再用恢复公式求J。"
            ],
            "answer": "先暗通道，再估计A，再估计t，最后恢复J。这四步是主观题的标准骨架。",
        },
    ],
    "04-image-restoration-review.docx": [
        {
            "title": "逆滤波为什么怕噪声",
            "question": "频域中H=0.1，噪声项N=2，逆滤波后噪声会变成多大？",
            "steps": [
                "逆滤波估计F_hat=G/H=F+N/H。",
                "噪声项被除以H：N/H=2/0.1=20。",
                "H很小时，噪声会被严重放大。"
            ],
            "answer": "噪声变成20。这就是逆滤波在H接近0处不稳定的原因。",
        },
        {
            "title": "维纳滤波答题模板",
            "question": "为什么维纳滤波通常比直接逆滤波稳？",
            "steps": [
                "先写退化模型G=HF+N。",
                "指出逆滤波只除以H，没有考虑噪声统计。",
                "维纳滤波在分母加入Sn/Sf，噪声越强，抑制越明显，目标是最小均方误差。"
            ],
            "answer": "维纳滤波把退化函数和噪声/图像功率谱一起考虑，所以比直接G/H更稳。",
        },
    ],
    "05-high-pass-filters-review.docx": [
        {
            "title": "高提升滤波：A=1时退化为高通",
            "question": "高提升滤波H_hb=(A-1)+H_hp，当A=1时是什么？",
            "steps": [
                "把A=1代入公式。",
                "得到H_hb=(1-1)+H_hp=H_hp。",
                "所以只剩普通高通滤波，没有额外保留原图低频。"
            ],
            "answer": "A=1时为普通高通滤波；A>1时既保留背景又增强高频。",
        },
        {
            "title": "周期噪声的频域处理",
            "question": "频谱图中出现关于中心对称的亮点，应考虑什么滤波器？",
            "steps": [
                "周期噪声在频域中常表现为离散、成对、对称的尖峰。",
                "如果只想去掉这些尖峰，优先使用陷波滤波器。",
                "如果噪声占据一圈频带，可使用带阻滤波器。"
            ],
            "answer": "优先答陷波滤波；若题目说是一段频带，则答带阻滤波。",
        },
    ],
    "06-low-pass-filters-review.docx": [
        {
            "title": "巴特沃斯低通手算",
            "question": "巴特沃斯低通H=1/[1+(D/D0)^(2n)]，若D=D0，n=2，H是多少？",
            "steps": [
                "代入D/D0=1。",
                "计算1^(2n)=1。",
                "H=1/(1+1)=0.5。"
            ],
            "answer": "H=0.5。截止频率处的响应是二分之一，是常见考点。",
        },
        {
            "title": "D0变化判断",
            "question": "低通滤波中D0变小，图像会怎样？",
            "steps": [
                "D0是保留低频范围的半径。",
                "D0变小，保留的频率范围变窄，高频细节被更多抑制。",
                "结果是噪声减少，但边缘和纹理也更模糊。"
            ],
            "answer": "图像更平滑、更模糊，噪声可能减少，细节损失更明显。",
        },
    ],
    "07-fourier-transform-review.docx": [
        {
            "title": "F(0,0)就是灰度和",
            "question": "2x2图像[[1,2],[3,4]]的F(0,0)和均值是多少？",
            "steps": [
                "DFT直流分量F(0,0)=所有像素求和。",
                "1+2+3+4=10。",
                "均值=F(0,0)/(MN)=10/4=2.5。"
            ],
            "answer": "F(0,0)=10，图像均值为2.5。",
        },
        {
            "title": "频域滤波完整流程",
            "question": "写出用频域低通滤波平滑图像的步骤。",
            "steps": [
                "对原图f(x,y)做DFT得到F(u,v)，通常再fftshift中心化。",
                "构造低通滤波器H(u,v)，让中心低频通过，远离中心的高频衰减。",
                "相乘G=HF，再ifftshift并IDFT，取实部得到输出图像。"
            ],
            "answer": "DFT - 构造H - 相乘 - IDFT，是频域滤波题的固定答题顺序。",
        },
    ],
    "08-color-image-processing-review.docx": [
        {
            "title": "只增强亮度时为什么转HSI",
            "question": "如果想让彩色图像变亮，但尽量不改变色调，应怎么做？",
            "steps": [
                "RGB三通道分别增强可能改变颜色比例，导致偏色。",
                "HSI把色调H、饱和度S、亮度I分开。",
                "只处理I分量，再转换回RGB，可以更好保持颜色。"
            ],
            "answer": "转到HSI，处理I分量，再转回RGB。",
        },
        {
            "title": "CMY与RGB互补关系",
            "question": "归一化RGB为R=0.2，G=0.7，B=0.4，求CMY。",
            "steps": [
                "使用C=1-R，M=1-G，Y=1-B。",
                "C=0.8，M=0.3，Y=0.6。",
                "注意CMY常用于印刷减色模型，RGB常用于显示加色模型。"
            ],
            "answer": "CMY=(0.8,0.3,0.6)。",
        },
    ],
    "09-spatial-sharpening-review.docx": [
        {
            "title": "Sobel模板手算",
            "question": "3x3邻域为[[10,10,10],[10,50,90],[10,90,90]]，用Sobel求Gx。",
            "steps": [
                "取Gx模板[-1 0 1; -2 0 2; -1 0 1]。",
                "对应相乘求和：(-10+10)+(-20+180)+(-10+90)=240。",
                "若题目还要求边缘强度，再求Gy并计算sqrt(Gx^2+Gy^2)或近似|Gx|+|Gy|。"
            ],
            "answer": "Gx=240。模板题关键是位置对应，不要把模板写反。",
        },
        {
            "title": "拉普拉斯锐化符号判断",
            "question": "为什么有的书写g=f-∇²f，有的写g=f+∇²f？",
            "steps": [
                "拉普拉斯模板有两套符号：中心为负或中心为正。",
                "若模板中心为负，锐化时通常用f-∇²f。",
                "若模板中心为正，锐化时通常用f+∇²f。"
            ],
            "answer": "符号要和模板配套；考试看模板中心符号，不要死背加或减。",
        },
    ],
    "10-spatial-smoothing-review.docx": [
        {
            "title": "中值滤波排序法",
            "question": "邻域像素为0,201,203,204,205,206,207,208,255，中值滤波输出是多少？",
            "steps": [
                "中值滤波先排序，本题已经按从小到大排列。",
                "9个数取第5个。",
                "第5个是205。"
            ],
            "answer": "输出205。它能去除孤立的0和255，所以适合椒盐噪声。",
        },
        {
            "title": "逆谐波均值Q的符号",
            "question": "逆谐波均值滤波中，Q>0和Q<0分别适合处理什么噪声？",
            "steps": [
                "记住盐噪声是白点，高灰度；椒噪声是黑点，低灰度。",
                "Q>0压制低灰度异常点，所以适合椒噪声。",
                "Q<0压制高灰度异常点，所以适合盐噪声。"
            ],
            "answer": "Q>0去椒噪声，Q<0去盐噪声。",
        },
    ],
    "11-spatial-filtering-review.docx": [
        {
            "title": "3x3均值滤波",
            "question": "3x3邻域像素和为450，使用3x3算术均值模板，中心输出是多少？",
            "steps": [
                "3x3均值模板每个权重为1/9。",
                "输出=邻域总和/9。",
                "450/9=50。"
            ],
            "answer": "中心输出为50。若模板带其他系数，要先乘权重再求和。",
        },
        {
            "title": "相关与卷积",
            "question": "模板不对称时，相关和卷积的手算差异是什么？",
            "steps": [
                "相关是模板直接覆盖邻域，对应相乘求和。",
                "卷积要先把模板旋转180度，再对应相乘求和。",
                "若模板中心对称，两者结果可能一样；不对称时不能混用。"
            ],
            "answer": "卷积先旋转模板，相关不旋转。",
        },
    ],
    "12-pixel-relations-noise-review.docx": [
        {
            "title": "三种距离度量",
            "question": "两点坐标差为(3,3)，求欧氏距离、D4距离、D8距离。",
            "steps": [
                "欧氏距离De=sqrt(3^2+3^2)=3sqrt(2)。",
                "D4距离=|3|+|3|=6。",
                "D8距离=max(|3|,|3|)=3。"
            ],
            "answer": "De=3sqrt(2)，D4=6，D8=3。",
        },
        {
            "title": "m邻接消除8邻接歧义",
            "question": "两个对角像素都属于同一集合，但它们共同4邻域中还有同集合像素，是否m邻接？",
            "steps": [
                "m邻接允许4邻接直接相连。",
                "对角相连时，需要检查共同4邻域。",
                "若共同4邻域中存在同集合像素，则这两个对角像素不算m邻接。"
            ],
            "answer": "不算m邻接。m邻接就是为了避免8邻接造成路径歧义。",
        },
    ],
    "13-histogram-processing-review.docx": [
        {
            "title": "直方图均衡化手算",
            "question": "L=4，灰度0、1、2的概率分别为0.25、0.25、0.5，求均衡化映射。",
            "steps": [
                "计算CDF：T(0)=0.25，T(1)=0.5，T(2)=1。",
                "乘L-1=3：0.75、1.5、3。",
                "按题目要求四舍五入可得1、2、3；若要求向下取整则为0、1、3。"
            ],
            "answer": "常见四舍五入映射为0->1，1->2，2->3。一定说明取整规则。",
        },
        {
            "title": "图像平均降噪",
            "question": "独立同分布噪声图像平均M=16幅后，噪声标准差变为原来的多少？",
            "steps": [
                "图像平均后噪声方差变为原来的1/M。",
                "标准差是方差开方，所以变为1/sqrt(M)。",
                "sqrt(16)=4，因此标准差变为原来的1/4。"
            ],
            "answer": "噪声标准差变为原来的1/4。",
        },
    ],
    "14-gray-transform-review.docx": [
        {
            "title": "对数变换常数",
            "question": "8位图像使用s=c log(1+r)，希望r=255映射到s=255，c是多少？",
            "steps": [
                "把最大输入代入最大输出：255=c log(1+255)。",
                "得到c=255/log(256)。",
                "若题目指定log底数，按指定底数；没指定时通常按自然对数理解。"
            ],
            "answer": "c=255/log(256)。",
        },
        {
            "title": "伽马变换亮暗判断",
            "question": "归一化灰度r=0.25，比较γ=0.5和γ=2的输出。",
            "steps": [
                "γ=0.5时，s=sqrt(0.25)=0.5，输出变大，图像变亮。",
                "γ=2时，s=0.25^2=0.0625，输出变小，图像变暗。",
                "所以γ<1提亮暗部，γ>1压暗图像。"
            ],
            "answer": "γ=0.5输出0.5；γ=2输出0.0625。",
        },
    ],
    "15-digital-image-fundamentals-review.docx": [
        {
            "title": "灰度级与比特数",
            "question": "256个灰度级需要多少bit表示一个像素？",
            "steps": [
                "n bit可以表示2^n个不同状态。",
                "256=2^8。",
                "所以每个像素需要8 bit。"
            ],
            "answer": "需要8 bit，不是256 bit。",
        },
        {
            "title": "真彩色图像数据量",
            "question": "1024x1024真彩色RGB图像，每通道8 bit，未压缩数据量是多少？",
            "steps": [
                "RGB三个通道，每像素24 bit=3 byte。",
                "像素数为1024x1024。",
                "数据量=1024x1024x3 byte=3 MB。"
            ],
            "answer": "未压缩数据量约为3 MB。",
        },
    ],
    "16-matlab-image-programming-review.docx": [
        {
            "title": "矩阵索引",
            "question": "y=[1 2 3;4 5 6]，y(:,2)和y(2,end)分别是多少？",
            "steps": [
                "Matlab索引从1开始。",
                "y(:,2)表示所有行的第2列，所以是[2;5]。",
                "y(2,end)表示第2行最后一列，所以是6。"
            ],
            "answer": "y(:,2)=[2;5]，y(2,end)=6。",
        },
        {
            "title": "逐元素运算",
            "question": "对矩阵t逐元素平方并逐元素相乘，应使用什么符号？",
            "steps": [
                "矩阵乘法用*和^，要求线性代数维度匹配。",
                "逐元素乘法用.*，逐元素幂用.^。",
                "所以逐元素平方写t.^2，逐元素相乘写a.*b。"
            ],
            "answer": "使用.*、./、.^这一组带点运算符。",
        },
    ],
    "17-course-overview-matlab-basics-review.docx": [
        {
            "title": "课程考核比例",
            "question": "雨课堂选择题问总评成绩构成，期末和平时比例是多少？",
            "steps": [
                "PPT给出的课程考核为期末笔试60%，平时成绩40%。",
                "平时成绩一般来自课堂、作业、实验或雨课堂活动。",
                "选择题直接选“期末60%，平时40%”。"
            ],
            "answer": "期末60%，平时40%。",
        },
        {
            "title": "randn生成指定均值数据",
            "question": "如何生成均值约为3、方差约为1的500个正态随机数？",
            "steps": [
                "randn默认生成均值0、方差1的标准正态随机数。",
                "要把均值平移到3，只需要加3。",
                "Matlab代码：x=randn(1,500)+3。"
            ],
            "answer": "x=randn(1,500)+3。",
        },
    ],
}


def add_labeled_para(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    set_east_asian_font(run)
    run = p.add_run(text)
    set_east_asian_font(run)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_knowledge_explainer(doc, chapter):
    for idx, (key, value) in enumerate(chapter["key_points"], 1):
        doc.add_heading(f"{idx}. {key}", level=2)
        add_labeled_para(doc, "核心理解：", value)
        add_labeled_para(
            doc,
            "考试抓手：",
            "先判断题目是在考概念辨析、公式计算还是流程复述；再把本考点对应的关键词写出来。"
            "如果题目给了图像、矩阵或频谱，不要急着套结论，先说明你观察到的目标、背景、噪声、频率或邻域关系。"
        )
        add_labeled_para(
            doc,
            "答题表达：",
            f"涉及“{key}”时，建议用“定义/公式 - 作用 - 适用条件 - 易错限制”的顺序作答。"
        )


def add_formula_usage(doc, chapter):
    add_para(
        doc,
        "公式题不要只背表达式。真正拿分的顺序是：先写变量含义，再判断适用条件，最后代入计算。"
        "遇到需要取整、归一化、截断或阈值比较的题目，要把处理规则单独写出来。"
    )
    for name, expr in chapter["formulas"]:
        add_labeled_para(doc, f"{name}：", f"{expr}。使用时先确认输入范围、变量单位和题目是否要求归一化。")


def add_worked_examples(doc, chapter):
    examples = WORKED_EXAMPLES.get(chapter["file"], [])
    for idx, example in enumerate(examples, 1):
        doc.add_heading(f"例题{idx}：{example['title']}", level=2)
        add_labeled_para(doc, "题目：", example["question"])
        add_labeled_para(doc, "解题步骤：", "")
        add_list(doc, example["steps"], "List Number")
        add_labeled_para(doc, "标准答案：", example["answer"])

    doc.add_heading("PPT常考题型怎么写", level=2)
    for idx, item in enumerate(chapter["exam"], 1):
        title = item.split("：", 1)[0] if "：" in item else f"题型{idx}"
        doc.add_heading(f"题型{idx}：{title}", level=3)
        add_labeled_para(doc, "题目特征：", item)
        add_list(
            doc,
            [
                "先写本题考查的概念或公式名称，避免答案一上来只有计算。",
                "按PPT中的处理流程列步骤：输入是什么、判断条件是什么、输出是什么。",
                "最后补一句限制条件或易错点，这通常是简答题和分析题的得分点。"
            ],
            "List Number",
        )


def add_exam_qa(doc, chapter):
    qa_rows = []
    for key, value in chapter["key_points"][:3]:
        qa_rows.append((f"问：{key}常怎么考？", f"答：先写核心定义：{value} 再结合题目给出的图像、矩阵或应用场景判断适用条件。"))
    for pitfall in chapter["pitfalls"][:3]:
        qa_rows.append(("问：这章最容易错在哪里？", f"答：{pitfall}"))

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3300, 6060])
    hdr = table.rows[0].cells
    hdr[0].text = "易考问法"
    hdr[1].text = "标准答法"
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                set_east_asian_font(p.runs[0])
    for q, a in qa_rows:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = a
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_east_asian_font(run)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15


def build_doc(chapter):
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, chapter["title"])

    p = doc.add_paragraph(style="ReviewTitle")
    run = p.add_run(chapter["title"])
    set_east_asian_font(run)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph(style="SubtitleLine")
    run = p.add_run(f"来源 PPT：{chapter['source']} | 用途：考前复习、公式回看、题型训练")
    set_east_asian_font(run)

    doc.add_heading("一、章节定位", level=1)
    add_para(doc, chapter["overview"])

    doc.add_heading("二、PPT考点地图", level=1)
    add_key_table(doc, chapter["key_points"])

    doc.add_heading("三、知识点精讲", level=1)
    add_knowledge_explainer(doc, chapter)

    doc.add_heading("四、必背公式与模型", level=1)
    add_formula_table(doc, chapter["formulas"])
    add_formula_usage(doc, chapter)

    doc.add_heading("五、手把手例题与题型", level=1)
    add_worked_examples(doc, chapter)

    doc.add_heading("六、PPT问题与易考问答", level=1)
    add_exam_qa(doc, chapter)

    doc.add_heading("七、易错点", level=1)
    add_list(doc, chapter["pitfalls"], "List Bullet")

    doc.add_heading("八、考前速记", level=1)
    quick = []
    for key, value in chapter["key_points"][:4]:
        quick.append(f"{key}：{value.split('。')[0]}。")
    quick.extend(chapter["pitfalls"][:2])
    add_list(doc, quick, "List Bullet")

    return doc


def write_readme():
    readme = OUT_DIR.parent / "README.md"
    if not readme.exists():
        readme.write_text(
            "# Digital Image Processing\n\n"
            "本主题用于整理数字图像处理课程的复习资料、考点、公式和标准题型。\n\n"
            "当前重点资料：`review-docx/` 中的逐 PPT 章节复习文档。\n",
            encoding="utf-8",
        )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_readme()
    for chapter in CHAPTERS:
        doc = build_doc(chapter)
        out_path = OUT_DIR / chapter["file"]
        doc.save(out_path)
        print(out_path)


if __name__ == "__main__":
    main()
