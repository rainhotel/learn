from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\moniC\project\learn")
OUT = ROOT / "01-topics" / "408-computer-organization" / "408-central-processing-unit-guide.docx"


CONTENT = r"""
## 0. 使用说明：这份手册怎么用

这份文档按 408 统考公开真题的常见命题形态，以及王道课程中“先主线、再题型、再易错点”的讲法整理。它不是历年真题题干的复刻，也不逐字搬运课程讲义；重点是把中央处理器这一章转成可执行的做题方法。

本章最重要的一句话是：**CPU 题的核心不是背部件名，而是追踪“一条指令在数据通路中怎样流动，控制器在每个节拍发出什么信号”。**

建议学习顺序：

1. 先读“命题地图”，知道 408 爱怎么考。
2. 再读“指令执行主线”，把 PC、IR、MAR、MDR、ALU、寄存器组连起来。
3. 然后按题型刷：微操作题、数据通路题、微程序题、流水线题、中断题。
4. 每做错一道题，都回到本手册对应题型的“判别句”和“答题模板”。

> 资料边界：文中例题均为自拟改编题，用来模拟 408 常见问法；涉及王道课程的部分只提炼方法，不复制课程原文。

## 1. 408 命题地图：中央处理器章到底考什么

| 考点群 | 历年真题常见问法 | 王道课抓手 | 标准输出 |
|---|---|---|---|
| CPU 功能与寄存器 | 问 PC、IR、MAR、MDR、PSW 的作用和变化时刻 | 把寄存器分成“地址、数据、指令、状态、通用”五类 | 判断题、选择题、传送语句 |
| 指令周期与微操作 | 给一条指令，问取指、间址、执行、中断周期的微操作 | 先写公共取指，再按寻址方式和指令功能补执行 | T0/T1/T2 节拍序列 |
| 数据通路 | 给数据通路图，问某步需要哪些控制信号，或哪些传送不能并行 | 看“谁上总线、谁接收、ALU 做什么、存储器读写什么” | 控制信号集合 |
| 硬布线控制器 | 问输入输出、速度、可修改性、适用指令系统 | 控制信号是组合逻辑函数 | 概念辨析、表达式 |
| 微程序控制器 | 问微命令、微指令、微程序、控制存储器、微地址 | 机器指令对应微程序，微指令发出微命令 | 概念、容量、字段位数 |
| 微指令编码 | 给互斥微命令组，求操作控制字段长度 | 直接编码长且快，字段编码短但要译码 | ceil(log2(n)) 或 ceil(log2(n+1)) |
| CPU 性能 | 给指令数、CPI、主频，求执行时间、MIPS、CPI 加权平均 | 单位先统一，公式只用两条 | 时间、频率、CPI |
| 流水线基本性能 | 给 k 段、n 条、段延迟，求总时间、吞吐率、加速比 | 先定时钟周期，再算装入、稳定、排空 | (k+n-1)Δt |
| 流水线冒险 | 判断结构/数据/控制冒险，问插入几个气泡 | RAW 真相关最常考；load-use 特别高频 | 时空图、暂停周期 |
| 中断与异常 | 问响应时机、保存断点、现场保护、中断向量 | 一条指令结束后响应，先保存后转入口 | 流程排序、概念辨析 |
| 综合题 | 把指令执行、数据通路、流水线或 CPI 混在一起 | 先分类，再画表，不要脑算 | 分步列式和时序表 |

408 中央处理器章的高频难点集中在三个地方：

- **第一高频：流水线性能和冒险。** 这类题计算性强，容易和指令系统、存储系统综合。
- **第二高频：微操作与数据通路。** 题面经常给图，答案看起来像“控制信号清单”。
- **第三高频：微程序控制。** 选择题考概念，计算题考字段编码和控制存储器容量。

## 2. 章节主线：一条指令在 CPU 里的生命线

中央处理器章可以用下面这条链串起来：

```text
程序中的机器指令
→ PC 给出指令地址
→ 主存读出指令到 IR
→ 控制器译码
→ 产生控制信号
→ 数据通路完成取数、运算、访存、写回
→ PC 指向下一条指令或转移目标
→ 若有中断/异常，保存断点并转入处理程序
```

也可以记成四个问题：

- **当前指令在哪？** PC、MAR、主存、MDR、IR。
- **当前指令要干什么？** 操作码译码、寻址方式、控制器。
- **数据怎样流动？** 寄存器、总线、ALU、移位器、存储器。
- **下一步去哪？** PC+1、转移目标、中断向量。

| 部件 | 本质作用 | 真题常见陷阱 |
|---|---|---|
| 运算器 | 对数据做算术/逻辑/移位处理 | ALU 只加工数据，不负责“决定整条指令怎么执行” |
| 控制器 | 取指、译码、产生控制信号、控制时序 | 控制器不是主存；控制存储器属于微程序控制器内部 |
| 寄存器组 | 临时保存地址、数据、指令、状态 | PC 保存“下一条指令地址”，IR 保存“当前指令” |
| 内部总线 | CPU 内部信息传送通路 | 单总线结构同一时刻通常只能有一个部件向总线输出 |
| ALU | 执行加、减、与、或、比较等 | 地址计算也常由 ALU 完成 |
| PSW/标志寄存器 | 保存条件码、中断允许位、处理机状态等 | 转移、中断、异常题经常要看状态位 |

## 3. 必背基础：CPU 功能与常见寄存器

### 3.1 CPU 的四类功能

| 功能 | 含义 | 典型题眼 |
|---|---|---|
| 指令控制 | 控制程序中指令的执行顺序 | PC 自增、转移、调用、返回 |
| 操作控制 | 发出微命令，让各部件执行规定动作 | 控制信号、微操作、数据通路 |
| 时间控制 | 决定每个操作在什么节拍发生 | T0/T1/T2、机器周期、时钟周期 |
| 数据加工 | 对数据进行算术逻辑运算 | ALU、标志位、溢出 |

### 3.2 高频寄存器速查

| 寄存器 | 中文 | 保存内容 | 做题关键词 |
|---|---|---|---|
| PC | 程序计数器 | 下一条将要取出的指令地址 | 取指前给 MAR，取指后自增，转移时改写 |
| IR | 指令寄存器 | 当前正在执行的机器指令 | 译码依据，通常由 MDR 写入 |
| MAR | 存储器地址寄存器 | 本次访问主存的地址 | 地址送主存地址线 |
| MDR/MBR | 存储器数据寄存器 | 读出或写入主存的数据 | 连接 CPU 和主存数据线 |
| ACC | 累加器 | 运算操作数或结果 | 早期累加器结构常见 |
| GPR | 通用寄存器组 | 程序可见的操作数/地址/结果 | R0、R1、R2 等 |
| PSW | 程序状态字 | 条件码、中断允许、CPU 状态 | 中断、异常、条件转移 |
| SP | 栈指针 | 栈顶地址 | 调用返回、中断现场保护 |

记忆口诀：**PC 管下一条，IR 管当前条；MAR 管地址，MDR 管数据；PSW 管状态。**

### 3.3 题型 1：寄存器与 CPU 功能辨析题

真题常见问法：

- “下列寄存器中，取指时内容送入地址总线的是哪一个？”
- “CPU 响应中断时保存的断点通常是什么？”
- “指令寄存器的内容来自哪里？”
- “PC 在什么情况下不是简单加 1？”

解法：

第 1 步，先判断题目问的是地址、数据、指令还是状态。

第 2 步，把寄存器按功能归类：

```text
地址类：PC、MAR、SP、基址/变址寄存器
数据类：MDR、ACC、GPR
指令类：IR
状态类：PSW/FLAGS
```

第 3 步，遇到“下一条”“断点”“返回地址”，优先想到 PC；遇到“当前指令”“译码”，优先想到 IR；遇到“主存地址线”，优先想到 MAR；遇到“主存数据线”，优先想到 MDR。

答题写法：

```text
本题考查寄存器功能。取指时由 PC 给出指令地址，先送入 MAR，再访问主存；
主存读出的指令字先进入 MDR，再写入 IR。因此与地址总线直接对应的是 MAR，
而 PC 保存的是下一条指令地址。
```

改编例题：

> CPU 正在执行一条条件转移指令，若转移条件成立，最可能被改写的是哪个寄存器？

标准解：

条件转移成立意味着下一条指令地址不再是顺序地址，因此要改写 PC。IR 保存当前指令，不表示下一条地址；MAR 只保存某次访存地址，不稳定保存程序执行流。

易错点：

- 不要说“PC 保存当前指令地址”。更准确地说，PC 通常保存**下一条将要取的指令地址**。在取指过程中，PC 可能先送出当前取指地址，然后自增到下一条。
- 中断断点通常是 PC 内容，也就是被中断程序下一条要执行的指令地址。

## 4. 指令周期、机器周期、时钟周期

### 4.1 三个周期的层级关系

```text
指令周期 >= 机器周期 >= 时钟周期
```

| 名称 | 含义 | 例子 |
|---|---|---|
| 时钟周期 | CPU 最小节拍单位 | 一个 T 状态 |
| 机器周期 | 完成一个基本 CPU 操作的时间 | 一次取指、一次读主存、一次写主存 |
| 指令周期 | 执行一条完整机器指令所需时间 | 取指 + 间址 + 执行 + 中断检查 |

典型指令周期组成：

```text
取指周期 → 间址周期（可选）→ 执行周期 → 中断周期（可选）
```

不是每条指令都有间址周期，也不是每条指令结束后都进入中断周期。

### 4.2 公共取指微操作

408 中非常喜欢考“取指阶段哪些寄存器发生了什么变化”。公共取指过程一般写成：

```text
T0: PC → MAR
T1: M[MAR] → MDR, PC + 1 → PC
T2: MDR → IR
T3: OP(IR) → CU，Ad(IR) → 地址形成逻辑
```

不同教材会把 PC+1 放在 T0、T1 或 T2，这取决于数据通路和控制时序。考试作答时按题目给出的数据通路和节拍约束写；若题目没有限定，写清楚“取指后 PC 自增”即可。

### 4.3 间址周期

如果指令采用间接寻址，指令地址字段给出的不是操作数地址，而是“操作数地址所在单元的地址”。因此需要多访问一次主存：

```text
T0: Ad(IR) → MAR
T1: M[MAR] → MDR
T2: MDR → 有效地址寄存器 / IR 地址字段
```

一句话：**间址周期取的是有效地址，不是最终操作数。**

### 4.4 执行周期按指令类型分类

| 指令类型 | 执行阶段核心动作 | 微操作写法 |
|---|---|---|
| 寄存器-寄存器运算 | 从寄存器取两个操作数，经 ALU 运算，写回寄存器 | R1, R2 → ALU；ALUout → R1 |
| 取数 LOAD | 形成有效地址，读主存，写入寄存器 | EA → MAR；M[MAR] → MDR；MDR → R |
| 存数 STORE | 形成有效地址，把寄存器内容写入主存 | EA → MAR；R → MDR；MDR → M[MAR] |
| 无条件转移 JMP | 目标地址写入 PC | Target → PC |
| 条件转移 BEQ/BNE | 判断标志或比较结果，满足则改 PC | if Cond then Target → PC |
| 调用 CALL | 保存返回地址，PC 改为子程序入口 | PC → Stack；Target → PC |
| 返回 RET | 从栈取返回地址写 PC | StackTop → PC |

### 4.5 题型 2：指令周期/微操作序列题

真题常见问法：

- “写出取指周期的微操作序列。”
- “某指令采用间接寻址，执行前后 MAR/MDR/PC/IR 如何变化？”
- “给出单总线数据通路，完成某指令至少需要几个节拍？”
- “下列微操作中，哪些可以安排在同一节拍？”

解法：

第 1 步，先写公共取指：

```text
PC → MAR
M[MAR] → MDR
MDR → IR
PC + 指令长度 → PC
```

第 2 步，看寻址方式。若是间接寻址，要补“取有效地址”；若是立即寻址，不要再去主存取操作数。

第 3 步，看指令功能。LOAD、STORE、ALU、JMP、CALL、RET 的执行阶段不同。

第 4 步，看数据通路限制。单总线同一节拍通常只能有一个源部件向总线输出；两个寄存器同时接收同一总线内容有时可以并行。

答题模板：

```text
取指周期：
T0: PCout, MARin
T1: Read, MDRin, PC + 1 → PC
T2: MDRout, IRin

执行周期：
T3: 按寻址方式形成 EA
T4: 按指令功能进行读/写/运算/转移
T5: 结果写回或 PC 修改
```

改编例题：

> 设指令 `LOAD R1, (A)` 采用间接寻址，A 为指令地址字段。写出从取指到把操作数装入 R1 的主要微操作。

标准解：

```text
取指：
T0: PC → MAR
T1: M[MAR] → MDR, PC + 1 → PC
T2: MDR → IR

间址：
T3: Ad(IR) → MAR
T4: M[MAR] → MDR
T5: MDR → MAR        ; MDR 中是有效地址 EA

执行：
T6: M[MAR] → MDR     ; 取操作数
T7: MDR → R1
```

扣分点：

- 把间址周期取出的“有效地址”误当成“操作数”。
- STORE 指令写反方向，把 `M[MAR] → MDR` 写成存数。
- 忘记 PC 自增，或把 PC 自增写成修改 IR。

## 5. 数据通路：读图题的核心

### 5.1 数据通路要问的三个问题

任何数据通路题都可以拆成三问：

1. 数据从哪里来？即哪个寄存器或存储器输出。
2. 数据经过哪里？即总线、MUX、ALU、移位器、MDR 等。
3. 数据到哪里去？即哪个寄存器、PC、MAR、MDR 或主存写入。

### 5.2 单总线数据通路的典型限制

单总线结构中，多个部件共享一条内部总线，常见规则是：

- 一个节拍内一般只能有一个部件输出到总线。
- 一个节拍内可以有多个部件从总线接收同一数据，前提是硬件允许同时写入。
- ALU 若需要两个操作数，常要先把一个操作数送入暂存寄存器 Y，再让另一个操作数上总线。
- ALU 输出常先进入 Z，再由 Z 写回寄存器或 PC。

典型单总线加法：

```text
T0: R1out, Yin
T1: R2out, ALU=Add, Zin
T2: Zout, R3in
```

这说明 `R1 + R2 → R3` 在单总线上通常不能一个节拍完成。

### 5.3 多总线数据通路

多总线结构能同时读多个寄存器并送入 ALU，例如常见寄存器堆有两个读端口、一个写端口：

```text
R[rs] → ALU.A
R[rt] → ALU.B
ALUout → R[rd]
```

因此多总线结构更适合流水线和 RISC 数据通路。408 遇到 MIPS 风格图时，要特别关注 `RegDst`、`ALUSrc`、`MemRead`、`MemWrite`、`MemtoReg`、`RegWrite`、`Branch`、`PCSrc` 等控制信号。

### 5.4 题型 3：数据通路控制信号题

真题常见问法：

- “完成 `R1 ← R2 + R3` 需要哪些控制信号？”
- “在给定数据通路中，某指令执行阶段 ALU 的两个输入分别来自哪里？”
- “下列控制信号组合是否可能在同一时钟周期有效？”
- “单总线结构下，完成某寄存器传送至少需要几个节拍？”

解法：

第 1 步，先把目标写成寄存器传送语句：

```text
目的部件 ← 源数据经过的路径
```

第 2 步，对照数据通路图，把路径拆成“输出信号 + 选择信号 + 运算信号 + 写入信号”。

第 3 步，检查总线冲突。若两个源都想同时上同一条总线，就必须拆节拍。

第 4 步，检查写入时机。寄存器写入、MDR 写入、主存写入通常有不同控制信号。

答题写法：

```text
本步需要 R2 和 R3 经 ALU 相加并写入 R1。
若是单总线结构，可安排：
T0: R2out, Yin
T1: R3out, ALU=Add, Zin
T2: Zout, R1in
因此控制信号分别是 R2out/Yin、R3out/Add/Zin、Zout/R1in。
```

改编例题：

> 单总线 CPU 中，ALU 的一个输入来自 Y 寄存器，另一个输入来自内部总线，ALU 输出写入 Z。完成 `PC ← PC + 4` 如何安排？

标准解：

```text
T0: PCout, Yin
T1: 常数4out, ALU=Add, Zin
T2: Zout, PCin
```

若数据通路支持 PC 专用自增器，则也可能在取指阶段直接 `PC + 4 → PC`，但这必须由题图说明。

易错点：

- 看到 `PC+4` 就默认一定要 ALU，实际可能有专用加法器。
- 忘记暂存寄存器 Y/Z，直接写成 `PCout, 4out, PCin`，这在单总线结构中通常不合法。

- 把控制信号写成自然语言而没有指出“输出/输入/ALU 操作/读写存储器”。

### 5.5 MIPS 风格数据通路控制信号

408 真题和王道题中经常出现 MIPS 风格单周期/多周期数据通路。不要死背图，要背控制信号的含义。

| 控制信号 | 作用 | 取 0 常见含义 | 取 1 常见含义 |
|---|---|---|---|
| RegDst | 选择写寄存器号 | rt 字段 | rd 字段 |
| ALUSrc | 选择 ALU 第二操作数 | 寄存器 rt | 立即数扩展 |
| MemtoReg | 选择写回寄存器的数据 | ALU 结果 | 主存读出数据 |
| RegWrite | 是否写寄存器 | 不写 | 写 |
| MemRead | 是否读主存 | 不读 | 读 |
| MemWrite | 是否写主存 | 不写 | 写 |
| Branch | 是否为条件分支 | 非分支 | 分支 |
| ALUOp | ALU 操作类别 | 加法/由功能码决定/比较 | 题图规定 |
| PCSrc | 下一 PC 来源 | PC+4 | 分支/跳转目标 |

典型指令控制信号可以这样记：

| 指令 | RegDst | ALUSrc | MemtoReg | RegWrite | MemRead | MemWrite | Branch | ALU 主要动作 |
|---|---|---|---|---|---|---|---|---|
| R-type | 1 | 0 | 0 | 1 | 0 | 0 | 0 | 由 funct 决定 |
| lw | 0 | 1 | 1 | 1 | 1 | 0 | 0 | 基址 + 偏移 |
| sw | X | 1 | X | 0 | 0 | 1 | 0 | 基址 + 偏移 |
| beq | X | 0 | X | 0 | 0 | 0 | 1 | 减法/比较 |
| j | X | X | X | 0 | 0 | 0 | 0 | PC 取跳转目标 |

`X` 表示无关项，因为该指令不写寄存器或不使用对应数据选择器的输出。考试若要求完整填表，X 可以写作“don’t care”；若选项没有 X，按不会影响结果的任一值处理，但要遵守题目说明。

### 5.6 题型 3-扩展：MIPS 控制信号填表题

真题常见问法：

- “给出单周期数据通路，填写 lw/sw/R 型/beq 的控制信号。”
- “某控制信号错误置 0/1，会影响哪些指令？”
- “执行某条指令时，ALU 的输入分别来自哪里？”
- “MemtoReg、RegDst、ALUSrc 的选择含义是什么？”

解法：

第 1 步，先问该指令是否写寄存器。写则 `RegWrite=1`，不写则 `RegWrite=0`。

第 2 步，问 ALU 第二操作数来自寄存器还是立即数。R 型和 beq 常来自寄存器，lw/sw 地址计算来自立即数偏移。

第 3 步，问是否访问数据存储器。lw 读，sw 写，R 型和 beq 不访问数据存储器。

第 4 步，问写回数据来自 ALU 还是存储器。R 型写 ALU 结果，lw 写主存读出数据。

答题模板：

```text
该指令为 lw，需要用寄存器基址和符号扩展立即数经 ALU 相加形成地址，
读数据存储器，并把读出的数据写回 rt。
因此 RegDst=0, ALUSrc=1, MemtoReg=1, RegWrite=1, MemRead=1, MemWrite=0, Branch=0。
```

改编例题：

> 若单周期 MIPS 中 `MemtoReg` 控制信号恒为 0，哪些指令最直接出错？

标准解：

`MemtoReg=0` 表示写回寄存器的数据来自 ALU 结果。R 型指令本来就写 ALU 结果，通常不受影响；`lw` 需要把数据存储器读出的数据写回寄存器，若恒为 0，会错误写入地址计算的 ALU 结果。因此最直接出错的是 `lw`。

易错点：

- sw 不写寄存器，因此 RegDst 和 MemtoReg 对 sw 无关。
- beq 需要 ALU 做比较，但不写寄存器，也不读写数据存储器。
- lw/sw 的 ALU 动作通常都是“基址 + 偏移”，不是把主存数据送 ALU。

## 6. 控制器：硬布线与微程序

### 6.1 控制器的输入与输出

控制器的任务是根据指令和时序产生控制信号。抽象表达式是：

```text
控制信号 = f(操作码, 时序信号, 状态标志, 外部请求)
```

| 输入 | 例子 | 作用 |
|---|---|---|
| 操作码 | ADD、LOAD、STORE、JMP | 决定指令类型 |
| 时序信号 | T0、T1、T2 | 决定当前执行到哪一步 |
| 状态标志 | ZF、CF、OF、SF | 条件转移、异常判断 |
| 外部请求 | 中断请求、总线请求 | 决定是否响应外部事件 |

输出就是各种控制信号，例如 `PCout`、`MARin`、`Read`、`ALUop`、`RegWrite`、`MemWrite`。

### 6.2 硬布线控制器

硬布线控制器用组合逻辑和时序逻辑直接产生控制信号。

| 维度 | 硬布线控制 |
|---|---|
| 本质 | 逻辑电路 |
| 优点 | 速度快 |
| 缺点 | 设计复杂，修改困难 |
| 适合 | 指令格式规整、控制较简单的 RISC |
| 真题陷阱 | “硬布线不能实现复杂指令”是过度绝对；应说复杂时设计和修改成本高 |

硬布线题常见判断：

```text
若某控制信号只在 ADD 指令执行阶段 T3 且无异常时有效，
可以写成：C = ADD · T3 · ¬Exception
```

### 6.3 微程序控制器

微程序控制器把控制信号组织成“微指令”，再由一串微指令完成一条机器指令。

层级关系必须背熟：

```text
机器指令
→ 微程序
→ 微指令
→ 微命令
→ 微操作
```

| 概念 | 含义 | 常见混淆 |
|---|---|---|
| 微操作 | CPU 内部最基本操作，如 PC → MAR | 它是动作，不是编码 |
| 微命令 | 控制微操作的控制命令 | 可理解为控制信号 |
| 微指令 | 一组微命令加下一地址控制信息 | 存在控制存储器中 |
| 微程序 | 完成一条机器指令的一段微指令序列 | 不是用户程序 |
| 控制存储器 CM | 存放微程序 | 不是主存 |
| 微地址 | 微指令在 CM 中的地址 | 不是主存地址 |
| 微程序计数器 μPC/CMAR | 指向下一条微指令 | 类似微层次的 PC |

### 6.4 题型 4：硬布线与微程序概念辨析题

真题常见问法：

- “微程序控制器中，控制存储器存放什么？”
- “微指令和机器指令的关系是什么？”
- “硬布线控制器和微程序控制器哪个速度更快？”
- “RISC/CISC 更适合哪种控制方式？”

解法：

第 1 步，看到“控制存储器”就回答“微程序/微指令”，不要回答“用户程序/机器指令”。

第 2 步，比较速度与灵活性：

```text
硬布线：快，不易改。
微程序：较慢，易改，规整。
```

第 3 步，和指令系统联系：

```text
RISC：指令简单、格式规整，常用硬布线。
CISC：指令复杂、寻址方式多，常用微程序。
```

答题模板：

```text
硬布线控制器直接由逻辑电路根据操作码、节拍和状态产生控制信号，速度快但修改困难；
微程序控制器把控制信号编码为微指令，存入控制存储器，设计规整且便于修改，但需要取微指令，速度通常较慢。
```

易错点：

- 不要把“微程序”说成“比机器语言更底层、由程序员编写的程序”。微程序通常对程序员不可见。
- 控制存储器常用 ROM 或只读/可写控制存储器实现，不等同于 Cache 或主存。

## 7. 微指令格式与控制存储器计算

### 7.1 微指令的两个字段

```text
微指令 = 操作控制字段 + 顺序控制字段
```

| 字段 | 作用 | 常考内容 |
|---|---|---|
| 操作控制字段 | 本微周期发出哪些微命令 | 直接编码、字段编码、字段间接编码 |
| 顺序控制字段 | 下一条微指令地址如何产生 | 顺序、转移、判别、操作码映射 |

### 7.2 微命令编码方式

| 编码方式 | 做法 | 优点 | 缺点 | 适用 |
|---|---|---|---|---|
| 直接编码 | 每位对应一个微命令 | 不需译码，速度快 | 微指令很长 | 控制信号不多或追求速度 |
| 字段直接编码 | 互斥微命令放同一字段编码 | 缩短微指令 | 需要译码 | 常考计算 |
| 字段间接编码 | 某字段含义由其他字段解释 | 更短 | 译码慢且复杂 | 概念题为主 |

互斥微命令的意思是：同一时刻不可能同时有效。例如同一总线的多个输出信号通常互斥。

### 7.3 位数计算规则

如果某字段有 n 个互斥微命令：

```text
若字段每个编码都表示一个有效微命令：位数 = ceil(log2 n)
若需要一个“不发命令/空操作”编码：位数 = ceil(log2(n + 1))
```

控制存储器容量：

```text
CM 容量 = 微指令条数 × 微指令字长
微地址位数 = ceil(log2(控制存储器可寻址微指令数))
```

若题目给出“控制存储器有 1024 个微指令单元，每条微指令 40 位”：

```text
容量 = 1024 × 40 bit = 40960 bit = 5 KB
微地址位数 = log2 1024 = 10 位
```

注意 bit 和 Byte 的换算：`1 B = 8 bit`。

### 7.4 后继微地址形成方式

| 方式 | 含义 | 常见题眼 |
|---|---|---|
| 顺序方式 | μPC + 1 | 普通微指令下一条 |
| 断定方式 | 按条件选择下一微地址 | 根据标志位/状态位转移 |
| 操作码映射 | 由机器指令操作码产生微程序入口 | 取指后按指令分派 |
| 直接给出 | 微指令中含下一微地址字段 | 字段位数等于微地址位数 |

### 7.5 题型 5：微指令编码与控制存储器容量题

真题常见问法：

- “共有 m 个微命令，若采用直接编码，操作控制字段至少多少位？”
- “若将互斥微命令分为若干组，每组字段直接编码，微指令操作控制字段多少位？”
- “控制存储器容量是多少？”
- “微地址字段至少需要多少位？”
- “某字段有 n 种转移方式，至少需要多少位编码？”

解法：

第 1 步，先分清求的是“操作控制字段”“顺序控制字段”还是“整条微指令字长”。

第 2 步，若直接编码：一个微命令一位。

第 3 步，若字段编码：每组互斥命令单独算位数，最后相加。是否加 1 取决于题目是否允许该组本周期不发任何命令。

第 4 步，顺序控制部分通常包括：

```text
下一地址字段 + 判别/转移方式字段
```

第 5 步，控制存储器容量用“微指令条数 × 微指令字长”，不要把主存容量带进来。

答题模板：

```text
第 1 组有 a 个互斥微命令，需 ceil(log2(a+1)) 位；
第 2 组有 b 个互斥微命令，需 ceil(log2(b+1)) 位；
...
操作控制字段共为这些字段位数之和。
若 CM 有 2^k 条微指令，则下一微地址字段需 k 位。
微指令字长 = 操作控制字段 + 顺序控制字段。
CM 容量 = 微指令条数 × 微指令字长。
```

改编例题：

> 某微程序控制器有 24 个微命令。若采用直接编码，操作控制字段多少位？若分成 4 组互斥微命令，组内微命令数分别为 5、6、3、4，且每组都需要一个“无操作”状态，操作控制字段至少多少位？

标准解：

```text
直接编码：24 位。

字段编码：
第 1 组：ceil(log2(5+1)) = 3 位
第 2 组：ceil(log2(6+1)) = 3 位
第 3 组：ceil(log2(3+1)) = 2 位
第 4 组：ceil(log2(4+1)) = 3 位
合计：3 + 3 + 2 + 3 = 11 位
```

易错点：

- 忘记“无操作”编码。
- 把 `ceil(log2 5)` 算成 2；实际 2 位只能表示 4 种状态。
- 控制存储器容量单位写错，bit 和 B 混用。
- 直接编码不是“二进制编码所有微命令”，而是“一位一个微命令”。

## 8. CPU 性能计算：CPI、主频、MIPS

### 8.1 核心公式

```text
CPU 执行时间 = 指令条数 × CPI × 时钟周期
CPU 执行时间 = 指令条数 × CPI / 主频
总时钟周期数 = 指令条数 × CPI
CPI = 总时钟周期数 / 指令条数
```

加权平均 CPI：

```text
CPI_avg = Σ(某类指令比例 × 该类指令 CPI)
```

MIPS：

```text
MIPS = 指令条数 / (执行时间 × 10^6)
     = 主频 / (CPI × 10^6)
```

注意：MIPS 只适合粗略比较同一类指令系统，跨 ISA 比较可能误导。

### 8.2 题型 6：CPU 执行时间与 CPI 计算题

真题常见问法：

- “给出主频、指令条数、CPI，求程序执行时间。”
- “给出各类指令比例和 CPI，求平均 CPI。”
- “某优化使某类指令 CPI 降低，求加速比。”
- “给出程序执行时间和主频，反求总时钟周期或 CPI。”

解法：

第 1 步，统一单位：

```text
主频 GHz → Hz
时间 ms/us/ns → s
周期 ns → s
```

第 2 步，先算平均 CPI，再算总周期，再算时间。

第 3 步，优化题用：

```text
加速比 = 优化前时间 / 优化后时间
```

若只优化某部分，可用 Amdahl 思想：

```text
T_new = T_old × [(1 - f) + f / S]
```

答题模板：

```text
CPI_avg = p1*CPI1 + p2*CPI2 + ...
Cycles = IC * CPI_avg
Time = Cycles / f
```

改编例题：

> 某程序有 1.0×10^8 条指令。A 类占 50%，CPI=1；B 类占 30%，CPI=2；C 类占 20%，CPI=4。处理器主频 2 GHz，求执行时间。

标准解：

```text
CPI_avg = 0.5×1 + 0.3×2 + 0.2×4 = 1.9
Cycles = 1.0×10^8 × 1.9 = 1.9×10^8
Time = 1.9×10^8 / (2×10^9) = 0.095 s
```

易错点：

- 把主频 2 GHz 当成周期 2 ns。实际周期为 0.5 ns。
- 先把每类 CPI 相加再除 3，这是错误的，必须按指令比例加权。
- MIPS 越大不一定表示真实性能越好，尤其跨指令系统时。

## 9. 指令流水线基础

### 9.1 流水线提高的是吞吐率

经典五段流水线：

```text
IF  取指
ID  译码 / 读寄存器
EX  执行 / 地址计算
MEM 访存
WB  写回
```

流水线的关键结论：

- 单条指令的完成时间通常不会变短。
- 多条指令连续执行时，平均每条指令完成时间变短。
- 理想情况下，装满流水线后每个时钟周期完成一条指令。
- 流水线时钟周期由最慢流水段决定，还要加流水寄存器开销。

### 9.2 基本公式

若 k 段流水线，每段时间相等为 Δt，执行 n 条指令：

```text
流水线时间 T_pipe = (k + n - 1) × Δt
非流水时间 T_non = n × k × Δt
加速比 S = T_non / T_pipe = nk / (k+n-1)
吞吐率 TP = n / T_pipe
最大吞吐率 TP_max = 1 / Δt
效率 E = S / k
```

当 n 远大于 k 时：

```text
S ≈ k
TP ≈ 1/Δt
```

若各段延迟不等：

```text
流水线时钟周期 = max(各段组合逻辑延迟) + 流水寄存器延迟
非流水单条指令时间 ≈ 各段延迟之和
```

### 9.3 题型 7：理想流水线性能计算题

真题常见问法：

- “k 段流水线执行 n 条指令需要多少时间？”
- “各段延迟不同，流水线时钟周期是多少？”
- “求加速比、吞吐率、效率。”
- “流水寄存器开销是否计入时钟周期？”

解法：

第 1 步，确定 k 和 n。

第 2 步，确定 Δt。若段延迟相等，直接用给定段时间；若不等，取最大段延迟并加寄存器开销。

第 3 步，使用：

```text
T_pipe = (k+n-1)Δt
```

第 4 步，若问加速比，要先明确非流水时间。等长段常用 `n*k*Δt`；不等长段常用 `n*Σti`。

答题模板：

```text
流水线共有 k 段，执行 n 条指令。
时钟周期 Δt = max(t1,t2,...,tk) + treg。
总时间 = (k+n-1)Δt。
非流水时间 = nΣti。
加速比 = 非流水时间 / 流水线时间。
```

改编例题：

> 一条指令分 5 段，段延迟分别为 80ns、90ns、70ns、100ns、60ns，流水寄存器延迟 10ns。执行 100 条指令，求流水线时间。

标准解：

```text
Δt = max(80,90,70,100,60) + 10 = 110ns
T = (5 + 100 - 1) × 110ns = 104 × 110ns = 11440ns
```

若求非流水时间：

```text
T_non = 100 × (80+90+70+100+60)ns = 40000ns
S = 40000 / 11440 ≈ 3.50
```

易错点：

- 把流水线总时间写成 `n × Δt`，漏掉装入和排空。
- 段延迟不等时，把 Δt 取平均值。
- 忘记流水寄存器开销。
- 以为流水线一定能达到 k 倍加速，实际受 n、段不均衡、冒险影响。

## 10. 流水线冒险：结构、数据、控制

### 10.1 三类冒险

| 冒险类型 | 本质 | 典型例子 | 解决方法 |
|---|---|---|---|
| 结构冒险 | 硬件资源冲突 | 同一周期 IF 和 MEM 都要访问统一存储器 | 增加资源、暂停、指令/数据分离 |
| 数据冒险 | 指令之间存在数据依赖 | 后一条读前一条尚未写回的寄存器 | 转发、暂停、编译调度、重命名 |
| 控制冒险 | 下一条 PC 不确定 | 分支、跳转、中断、异常 | 暂停、预测、延迟分支、提前判定 |

### 10.2 数据相关类型

| 类型 | 名称 | 含义 | 408 重点 |
|---|---|---|---|
| RAW | 写后读 | 后指令要读前指令将写的结果 | 真相关，最常考 |
| WAR | 读后写 | 后指令写前指令要读的寄存器 | 乱序流水中更常见 |
| WAW | 写后写 | 两条指令写同一寄存器，顺序不能乱 | 乱序/多发射更常见 |

经典五段顺序流水线中，最典型的是 RAW。WAR/WAW 在简单顺序五段流水中通常不突出，因为读寄存器在 ID，写回在 WB，指令按序流动。

### 10.3 转发与暂停的常用规则

以五段流水线 `IF-ID-EX-MEM-WB` 为例：

| 情况 | 无转发 | 有转发 |
|---|---|---|
| ALU 指令结果被下一条 ALU 指令使用 | 常需暂停到写回后再读 | 通常不暂停，EX/MEM 或 MEM/WB 转发到 EX |
| LOAD 结果被下一条指令立即使用 | 需暂停较多周期 | 通常仍需 1 个气泡，因为数据到 MEM 末才可用 |
| 分支依赖前一条结果 | 看分支判定阶段和转发路径 | 可能需要暂停 |
| STORE 使用前一条结果作为写入数据 | 看写数据进入 MEM 的时刻 | 有些数据通路可转发到 MEM |

具体停几个周期必须看题目给定假设，尤其是：

- 寄存器堆是否“前半周期写、后半周期读”。
- 是否有 EX/MEM、MEM/WB 到 ALU 输入的转发。
- 分支在 ID 还是 EX 阶段决定。
- 访存是否统一存储器。

### 10.4 题型 8：流水线时空图与数据冒险题

真题常见问法：

- “下列指令序列在五段流水线中是否发生数据相关？”
- “有/无转发时需要插入几个 nop？”
- “画出流水线执行时空图。”
- “执行该指令序列共需多少周期？”

解法：

第 1 步，标出每条指令的读写寄存器：

```text
I1: 写 R1
I2: 读 R1
```

第 2 步，判断相关类型。后一条读前一条写，就是 RAW。

第 3 步，确定结果何时产生、消费者何时需要：

```text
ALU 结果：EX 末产生
LOAD 数据：MEM 末产生
ALU 操作数：EX 初需要
分支比较数：ID 或 EX 需要，按题设
```

第 4 步，若生产时间晚于消费时间，就插入气泡，直到消费时数据可用。

第 5 步，画时空图。先画理想流水，再在冲突处停住相关指令，后续顺延。

答题模板：

```text
I1 写 R1，I2 读 R1，存在 RAW 相关。
若无转发，I2 必须等 I1 WB 后才能在 ID 读到 R1，因此插入若干气泡。
若有转发，ALU-ALU 相关可由转发解决；但 load-use 因数据在 MEM 末才产生，通常仍需 1 个气泡。
```

改编例题 1：ALU-ALU 相关

```asm
I1: ADD R1, R2, R3
I2: SUB R4, R1, R5
```

标准解：

`I1` 写 R1，`I2` 读 R1，存在 RAW 相关。若有 EX/MEM 到 ALU 输入的转发，`I1` 的 ALU 结果在 EX 末产生，可在下一周期转发给 `I2` 的 EX 阶段，所以通常不需要停顿。若无转发，则 `I2` 要等 `I1` WB 后才能读到 R1，需要插入气泡，周期数按题设寄存器读写时序确定。

改编例题 2：load-use 相关

```asm
I1: LW  R1, 0(R2)
I2: ADD R3, R1, R4
```

标准解：

`LW` 的数据在 MEM 末才从存储器得到，而 `ADD` 在 EX 初就需要 R1。即使有转发，紧随其后的使用也通常需要插入 1 个气泡：

```text
Cycle: 1   2   3   4   5   6   7
I1:    IF  ID  EX  MEM WB
I2:        IF  ID  St  EX  MEM WB
```

这里 `St` 表示暂停/气泡。实际表格可能把 IF/ID 冻结、EX 插入 bubble，写法按题目要求即可。

易错点：

- 只看到寄存器名相同就判定相关。必须看读写方向。
- 把 WAR/WAW 当成五段顺序流水的高频问题。
- 忘记 load-use 即使有转发也常要停 1 拍。

- 画时空图时只停当前指令，忘了后续指令也要顺延。

### 10.5 无转发流水线的精确写法

很多 408 题会特别说明“无数据转发”或“寄存器堆前半周期写、后半周期读”。这时不要套“ALU-ALU 不停顿”的结论，而要按读写阶段排时间。

常用假设：五段流水线中，寄存器在 ID 阶段读，在 WB 阶段写。

若没有转发，消费者必须等生产者 WB 后才能在 ID 读到正确值：

```text
生产者：IF ID EX MEM WB
消费者：   IF ID EX MEM WB   ; 理想情况，但 ID 太早
```

如果寄存器堆支持同一周期“前半周期写、后半周期读”，消费者的 ID 可以和生产者的 WB 放在同一周期；否则消费者 ID 还要再晚一个周期。题目若没明说，按题设或教材默认处理，答题时写出假设。

改编例题：

```asm
I1: ADD R1, R2, R3
I2: SUB R4, R1, R5
```

无转发、且寄存器堆可同周期先写后读时：

```text
Cycle: 1   2   3   4   5   6   7   8
I1:    IF  ID  EX  MEM WB
I2:        IF  St  St  ID  EX  MEM WB
```

解释：`I2` 原本第 3 周期进入 ID 读 R1，但 `I1` 第 5 周期才 WB。若同周期先写后读，则 `I2` 的 ID 可以安排在第 5 周期，所以第 3、4 周期停顿，共 2 个气泡。

答题模板：

```text
无转发时，结果只有到生产者 WB 后才对寄存器堆可见。
消费者在 ID 阶段读寄存器，因此必须把消费者 ID 推迟到生产者 WB 同周期或之后。
根据题设是否允许同周期先写后读，确定停顿周期数。
```

易错点：

- 把“有转发”的结论误用于“无转发”题。
- 没看题设中的“寄存器堆先写后读”条件。
- 只数插入的 nop，不重新计算总周期。

## 11. 控制冒险、分支预测与延迟分支

### 11.1 控制冒险的本质

控制冒险来自 PC 不确定：

```text
顺序执行：PC ← PC + 4
分支成立：PC ← branch target
跳转：PC ← jump target
中断/异常：PC ← handler entry
```

在分支结果确定前，流水线可能已经取入了若干条后继指令。如果分支成立，这些指令要取消或作废。

### 11.2 常见处理方法

| 方法 | 思路 | 代价 |
|---|---|---|
| 暂停取指 | 等分支结果出来再取指 | 简单但慢 |
| 静态预测 | 总预测不跳/总预测跳/按方向预测 | 硬件简单 |
| 动态预测 | 根据历史行为预测 | 硬件复杂，准确率高 |
| 延迟分支 | 分支后若干条指令总会执行 | 依赖编译器调度 |
| 提前判定 | 在 ID 阶段比较并形成目标地址 | 缩短分支惩罚 |

### 11.3 题型 9：分支与控制冒险题

真题常见问法：

- “分支在第几段确定时，会造成几个周期损失？”
- “若预测失败，需要冲刷流水线中几条指令？”
- “采用延迟分支后，哪些指令可放入延迟槽？”
- “条件转移指令修改的是哪个寄存器？”

解法：

第 1 步，确定分支结果在哪个阶段产生。若在 EX 末确定，IF、ID 中可能已有错误路径指令。

第 2 步，确定预测策略。预测正确不冲刷；预测失败要冲刷错误取入的指令。

第 3 步，计算惩罚周期：

```text
分支惩罚 ≈ 从取分支后到确定正确 PC 之间已经错误取入的周期数
```

第 4 步，延迟分支题要判断候选指令是否与分支结果无关，且移动后不改变程序语义。

答题模板：

```text
该分支在 EX 阶段末确定。若预测失败，在分支之后已经进入 IF/ID 的指令属于错误路径，
需要冲刷，因此损失为题设给出的分支惩罚周期。若分支频率为 p，预测失败率为 q，
额外 CPI = p × q × penalty。
```

改编例题：

> 某五段流水线分支在 EX 末确定，采用“预测不跳转”。分支指令占 20%，其中 60% 实际跳转。预测失败代价为 2 个周期，理想 CPI=1，求平均 CPI。

标准解：

```text
预测失败率 = 实际跳转比例 = 60%
额外 CPI = 分支比例 × 失败率 × 失败代价
        = 0.20 × 0.60 × 2 = 0.24
CPI = 1 + 0.24 = 1.24
```

易错点：

- 把“分支占比”和“跳转比例”混为一个数。
- 忘记只有预测失败才有冲刷惩罚。
- 延迟槽中的指令必须无论分支是否成立都可安全执行。

## 12. 中断与异常

### 12.1 基本概念

中断和异常都是改变正常控制流的机制。

| 类型 | 来源 | 例子 | 常见特点 |
|---|---|---|---|
| 外中断 | CPU 外部 | I/O 完成、时钟中断 | 通常异步 |
| 内异常 | CPU 执行指令内部 | 除零、溢出、非法指令、缺页 | 通常同步 |
| 软件中断/陷入 | 指令主动触发 | 系统调用 | 同步、可预期 |

408 中常用宽泛说法：中断多来自外部，异常多来自内部。更精确地说，外中断通常异步，异常通常与当前指令执行同步。

### 12.2 中断响应流程

```text
中断请求
→ 中断判优
→ 检查中断允许
→ 当前指令执行结束
→ 保存断点
→ 关中断或改变优先级
→ 形成中断服务程序入口地址
→ 转入中断服务程序
→ 保存现场
→ 执行服务
→ 恢复现场
→ 中断返回
```

关键点：

- 断点通常是 PC 内容，即下一条将执行的指令地址。
- 现场通常包括通用寄存器、PSW 等，保证返回后程序可继续。
- 中断向量通常保存中断服务程序入口地址，或保存形成入口地址所需的信息。
- CPU 通常在一条指令执行结束后响应可屏蔽中断。
- 异常可能要求重新执行当前指令或终止程序，具体看异常类型。

### 12.3 题型 10：中断/异常流程题

真题常见问法：

- “CPU 响应中断时首先保存什么？”
- “中断向量的作用是什么？”
- “中断响应发生在指令执行的哪个时刻？”
- “中断服务程序返回时需要恢复哪些内容？”
- “异常和中断有什么区别？”

解法：

第 1 步，判断来源：外部事件多为中断，当前指令执行导致多为异常。

第 2 步，判断响应时机：普通可屏蔽外中断通常在一条指令结束后响应。

第 3 步，回答保存内容：

```text
断点：PC
现场：通用寄存器、PSW、必要的控制寄存器
```

第 4 步，回答入口形成：

```text
中断类型号 → 中断向量表 → 服务程序入口地址
```

答题模板：

```text
CPU 在当前指令执行结束后检测到可响应中断，先保存断点 PC 和必要状态，
再根据中断类型号查中断向量表得到服务程序入口地址，将该地址送入 PC。
服务程序开始后保存现场，处理完毕后恢复现场并执行中断返回指令。
```

改编例题：

> 为什么保存断点和保存现场不是同一个概念？

标准解：

断点是程序控制流返回的位置，通常体现为 PC；现场是被中断程序继续运行所需的处理器状态，如通用寄存器、PSW 等。只保存 PC 可以回到原程序位置，但若寄存器内容被中断服务程序破坏，原程序仍不能正确继续。因此二者功能不同。

易错点：

- 把“保存断点”写成“保存当前 IR”。IR 不是返回地址。
- 把中断向量说成“中断服务程序本身”。中断向量通常是入口地址或入口信息。

- 以为所有中断都能打断机器指令中间。普通可屏蔽中断一般在指令结束后响应。

### 12.4 多级中断、优先级和屏蔽字

多级中断题的关键是区分两个概念：**响应优先级**和**处理优先级**。

| 概念 | 含义 | 题目怎么考 |
|---|---|---|
| 响应优先级 | 多个中断同时请求时，CPU 先响应谁 | 判优电路、优先级编码器 |
| 处理优先级 | 某中断服务程序执行期间，允许哪些更高级中断打断它 | 中断屏蔽字、嵌套中断 |
| 中断屏蔽字 | 控制某些中断源是否被屏蔽 | 给屏蔽矩阵判断嵌套顺序 |
| 中断嵌套 | 高优先级中断打断低优先级中断服务程序 | 现场保护要成栈保存 |

解法：

第 1 步，列出所有中断请求及其优先级。

第 2 步，看当前是否开中断，以及对应中断源是否被屏蔽。

第 3 步，若多个请求可响应，按响应优先级选择一个。

第 4 步，进入中断服务程序后，根据屏蔽字判断是否允许更高优先级中断嵌套。

第 5 步，返回顺序通常后进先出：最后嵌套进入的中断最先返回。

答题模板：

```text
同时发生多个中断请求时，CPU 先检查中断允许位和屏蔽字，排除被屏蔽的请求；
对可响应请求按响应优先级判优。进入某中断服务程序后，若开启中断且屏蔽字允许，
更高处理优先级的中断可以嵌套打断当前服务程序。现场和断点按栈式顺序保存和恢复。
```

改编例题：

> A、B、C 三个中断源处理优先级为 A > B > C。CPU 正在执行 C 的服务程序，若此时 A 和 B 同时请求，且屏蔽字允许 A、B 打断 C，则响应顺序如何？返回顺序如何？

标准解：

A 和 B 都允许打断 C，且 A 优先级高于 B，所以先响应 A。A 服务完成返回后，若 B 请求仍在且允许响应，再处理 B，最后回到 C。返回顺序是 A 返回到 C 的嵌套现场或中断调度点，B 返回后再继续 C，最终 C 返回主程序。若题目采用严格栈式嵌套，可写成进入顺序 C→A→B 或 C→A 后再 B，返回顺序反向，具体取决于 A 服务期间是否允许 B 嵌套。

易错点：

- 把响应优先级和处理优先级混为一谈。
- 忘记屏蔽字会改变“能否被响应”，不是只看优先级数字。
- 嵌套中断需要多次保存断点和现场，返回顺序不是按优先级从高到低固定返回。

## 13. RISC、CISC 与流水线友好性

408 可能把 CPU 控制方式、指令系统和流水线综合起来考。

| 维度 | RISC | CISC |
|---|---|---|
| 指令数量 | 少 | 多 |
| 指令格式 | 规整、长度较固定 | 复杂、长度可变 |
| 寻址方式 | 少 | 多 |
| 访存方式 | Load/Store 结构常见 | 运算指令可能直接访存 |
| 控制器 | 常用硬布线 | 常用微程序 |
| 流水线 | 更友好 | 较难规整流水 |

做题判断：

- “指令长度固定、格式规整、Load/Store、寄存器多、硬布线、流水线友好”通常指 RISC。
- “指令复杂、寻址方式多、微程序控制、单条指令功能强”通常指 CISC。

易错点：

- 不要说 RISC 一定比 CISC 快。实际性能受实现、编译器、缓存、流水线等影响。
- 不要说 CISC 不能流水线。现代处理器可把复杂指令译成内部微操作再流水执行。

## 14. 综合题解法：看到大题先分层

中央处理器综合题常把多个知识点拼起来。建议按四层拆：

### 14.1 第一层：题目在问哪类输出

| 输出类型 | 你应该写什么 |
|---|---|
| 微操作序列 | T0/T1/T2 + 寄存器传送语句 |
| 控制信号 | 源输出、目的写入、ALUop、MemRead/MemWrite、MUX 选择 |
| 性能时间 | 公式、单位、代入、结果 |
| 流水线图 | 指令为行、周期为列、阶段填入格子 |
| 冒险说明 | 相关类型、是否能转发、插入几个气泡 |
| 中断过程 | 断点、现场、向量、入口、返回 |

### 14.2 第二层：先画三个小图

写在草稿纸上即可：

```text
指令流：PC → IR → 译码 → 执行 → PC
数据流：源寄存器/主存 → ALU/MDR → 目的寄存器/主存
时间流：IF → ID → EX → MEM → WB
```

### 14.3 第三层：避免“脑内并行”

大题最容易错在把多个动作想当然放到同一拍。判断能否同拍，看三个约束：

- 资源是否冲突：同一总线、同一存储器端口、同一 ALU。
- 数据是否已经产生：上一阶段末产生，下一阶段初需要。
- 控制是否确定：分支、异常、中断是否已经给出正确 PC。

### 14.4 题型 11：CPU 综合大题的书写模板

答题模板：

```text
（1）指令功能分析：
该指令需要先取指，再形成有效地址/读取寄存器，然后执行运算/访存/转移。

（2）取指阶段：
T0: ...
T1: ...
T2: ...

（3）执行阶段：
按题图数据通路，列出每个节拍控制信号。

（4）若涉及流水线：
列出每条指令读写寄存器，判断 RAW/WAR/WAW。
根据是否有转发和分支判定阶段插入气泡。
最后用“周期数 = 理想周期 + 停顿/冲刷周期”计算。

（5）若涉及性能：
先算 CPI 或 Δt，再算执行时间/加速比。
```

改编综合例题：

> 五段流水线 IF-ID-EX-MEM-WB，支持 ALU 结果转发，但 load-use 需停 1 拍。执行：
> `LW R1, 0(R2)`；`ADD R3, R1, R4`；`SUB R5, R3, R6`。问数据相关和停顿。

标准解：

```text
I1 写 R1，I2 读 R1：RAW，且是 load-use，需停 1 拍。
I2 写 R3，I3 读 R3：RAW，但为 ALU-ALU 相关，有转发可解决，不停顿。
因此总停顿 1 个周期。
若无其他冒险，3 条指令五段流水理想周期为 k+n-1=5+3-1=7，
加入 1 个停顿后共 8 个周期。
```

## 15. 每种题目的“判别句”速查

| 题型 | 看到什么关键词 | 第一反应 | 核心公式/模板 |
|---|---|---|---|
| 寄存器辨析 | PC、IR、MAR、MDR、PSW | 先分地址/数据/指令/状态 | PC 下一条，IR 当前条 |
| 指令周期 | 取指、间址、执行、中断 | 先写公共取指 | PC→MAR；M→MDR；MDR→IR |
| 微操作并行 | 同一节拍、单总线 | 查资源冲突 | 一源上总线，多目的可接收 |
| 数据通路控制 | 控制信号、MUX、ALUop | 写源-路-目的 | output + select + ALUop + input |
| 硬布线 | 组合逻辑、节拍、操作码 | 快但难改 | C=f(op,T,flag) |
| 微程序 | 控制存储器、微指令 | 机器指令对应微程序 | 微指令=控制字段+顺序字段 |
| 微指令编码 | 互斥微命令、字段 | 每组单独算 | ceil(log2(n+1)) |
| CPU 时间 | CPI、主频、指令数 | 单位统一 | T=IC×CPI/f |
| 理想流水线 | k 段 n 条 | 装入+排空 | T=(k+n-1)Δt |
| 不等段流水线 | 各段延迟、寄存器开销 | 取最大段 | Δt=max(ti)+treg |
| 数据冒险 | 前写后读 | RAW 真相关 | 看产生时刻和需要时刻 |
| load-use | LW 后立即用 | 即使转发也常停 1 拍 | 插入 bubble |
| 控制冒险 | 分支、预测失败 | 看分支确定阶段 | 额外 CPI=p×q×penalty |
| 中断异常 | 断点、现场、向量 | PC+PSW | 保存断点→入口→服务→返回 |

## 16. 高频易错清单

### 16.1 概念类易错

- PC 不是 IR；PC 管下一条地址，IR 管当前指令。
- MAR 不是 MDR；MAR 是地址，MDR 是数据。
- 控制存储器不是主存；微程序不是用户程序。
- 微指令不是机器指令；微命令不是微操作。
- 硬布线不是“不能改”，而是修改困难。
- 流水线提高吞吐率，不保证单条指令延迟变短。
- 中断向量不是中断服务程序本身，而通常是入口地址或入口信息。

### 16.2 计算类易错

- `GHz` 和 `ns` 换算错：1GHz 对应 1ns，2GHz 对应 0.5ns。
- CPI 要按指令比例加权，不能简单平均。
- 流水线总时间要有 `k+n-1`，不是只算 n。
- 不等长流水段取最大段延迟，不取平均。
- 微指令字段编码常要给“无操作”留编码。
- 控制存储器容量用 bit 还是 Byte 要看题目要求。
- 预测失败惩罚只乘失败率，不乘预测正确的部分。

### 16.3 书写类易错

- 微操作题只写“取指、译码、执行”，没有写具体寄存器传送。
- 控制信号题只写自然语言，没有列出 `out/in/Read/Write/ALUop`。
- 流水线题没有画表，导致漏停顿或漏顺延。
- 中断流程题把保存断点和保存现场混成一句。
- 大题没有单位，最后结果不知道是 ns、周期还是秒。

## 17. 一页公式速查

```text
CPU 执行时间 = 指令条数 × CPI × 时钟周期
CPU 执行时间 = 指令条数 × CPI / 主频
CPI_avg = Σ(指令比例 × 该类 CPI)
MIPS = 主频 / (CPI × 10^6)

理想流水线时间 = (k+n-1)Δt
非流水时间（等长段） = n×k×Δt
非流水时间（不等长段） = n×Σti
流水线时钟周期 = max(ti) + 流水寄存器开销
加速比 = T_non / T_pipe
吞吐率 = n / T_pipe
最大吞吐率 = 1 / Δt
效率 = 加速比 / 段数

字段编码位数 = ceil(log2(互斥命令数 + 是否需要空操作))
微地址位数 = ceil(log2(控制存储器微指令条数))
控制存储器容量 = 微指令条数 × 微指令字长

分支额外 CPI = 分支比例 × 预测失败率 × 失败惩罚周期
总周期数 = 理想周期数 + 停顿周期 + 冲刷周期
```

## 18. 刷题复盘表

每刷完一组王道章节题或 408 真题，可以按下表复盘：

| 复盘项 | 要记录什么 |
|---|---|
| 错题属于哪类 | 概念、微操作、数据通路、微程序、流水线、中断 |
| 错因 | 概念混淆、公式错、单位错、图没画、题设没看清 |
| 正确第一步 | 例如“先列读写寄存器”“先算 Δt”“先写公共取指” |
| 可复用模板 | 把本题答案压缩成 2-3 行模板 |
| 下次检查点 | 做同类题前先看哪条易错清单 |

建议每章至少整理 10 道代表题：流水线 4 道，微程序 2 道，数据通路 2 道，中断/异常 1 道，CPI 性能 1 道。整理时不要抄题干，重点写“题眼 + 解法 + 易错原因”。

## 19. 最后总结

中央处理器章的本质是“控制 + 数据通路 + 时间”。408 不会只考你记不记得某个寄存器名，而是考你能不能在一个具体题面中回答：

- 数据从哪里来，经过哪里，到哪里去？
- 控制器在当前节拍要发出哪些信号？
- 这一拍能否并行，有没有资源冲突？
- 前后指令有没有数据或控制相关？
- 用什么公式把周期数、CPI、主频和时间连起来？

把这五个问题练熟，中央处理器章就会从“概念很多”变成“题型很固定”。
"""


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_run_font(run, *, latin="Calibri", east="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    widths = [dxa(w) for w in widths_in]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = tbl_borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "D0D7DE")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        row.height = None
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.12


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_runs_with_bold(paragraph, text: str, *, size=None, color=None):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_body(doc, text: str):
    p = doc.add_paragraph()
    add_runs_with_bold(p, text)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    add_runs_with_bold(p, text)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    add_runs_with_bold(p, text)
    return p


def add_code_block(doc, code: str):
    p = doc.add_paragraph(style="CodeBlock")
    set_paragraph_shading(p, "F7F9FC")
    lines = code.rstrip("\n").splitlines()
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, latin="Consolas", east="Microsoft YaHei", size=9.2, color="17202A")
        if i != len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    add_runs_with_bold(p, text)
    for run in p.runs:
        set_run_font(run, size=10.5, color="0B2545")
    doc.add_paragraph()


def parse_table(lines):
    cleaned = []
    for line in lines:
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        cleaned.append([c.strip() for c in line.split("|")])
    if len(cleaned) >= 2 and all(re.fullmatch(r":?-{3,}:?", c or "---") for c in cleaned[1]):
        return cleaned[0], cleaned[2:]
    return cleaned[0], cleaned[1:]


def widths_for(cols: int):
    patterns = {
        2: [1.65, 4.85],
        3: [1.45, 2.45, 2.60],
        4: [1.35, 1.65, 2.35, 1.15],
        5: [1.15, 1.20, 1.45, 1.45, 1.25],
    }
    return patterns.get(cols, [6.5 / cols] * cols)


def add_table_md(doc, table_lines):
    headers, rows = parse_table(table_lines)
    cols = len(headers)
    table = doc.add_table(rows=1, cols=cols)
    set_table_geometry(table, widths_for(cols))
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cols >= 4 else WD_ALIGN_PARAGRAPH.LEFT
        add_runs_with_bold(p, header, size=9.5)
        for run in p.runs:
            set_run_font(run, size=9.5, bold=True, color="0B2545")
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values[:cols]):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            if idx == 0 and cols >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_bold(p, value, size=9.2)
            for run in p.runs:
                set_run_font(run, size=9.2, color="111827")
    set_table_geometry(table, widths_for(cols))
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "CodeBlock" not in [s.name for s in styles]:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9.2)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05

    header = section.header
    clear_paragraph(header.paragraphs[0])
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("408 计算机组成原理｜中央处理器")
    set_run_font(run, size=9, color="6B7280")

    footer = section.footer
    clear_paragraph(footer.paragraphs[0])
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("408 CPU 复习手册  |  第 ")
    set_run_font(run, size=9, color="6B7280")
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color="6B7280")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("408 计算机组成原理")
    set_run_font(r, size=12, bold=True, color="1F4D78")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("中央处理器章节系统复习与题型解法")
    set_run_font(r, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("结合 408 历年真题命题形态与王道课程主线整理")
    set_run_font(r, size=12, color="4B5563")

    add_callout(
        doc,
        "定位：这是一份面向考研 408 的 CPU 章节复习讲义。重点不是堆概念，而是把每类题转成可复用的作答流程。"
    )

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1.65, 4.85])
    entries = [
        ("适用范围", "408 计算机组成原理：中央处理器、数据通路、控制器、流水线、中断异常"),
        ("使用场景", "第一轮建框架、王道章节题订正、真题二刷、考前公式回看"),
        ("内容原则", "不复刻真题原文；按公开命题形态提炼考点、题型、写法和解法"),
        ("生成日期", "2026-06-20"),
    ]
    for i, (k, v) in enumerate(entries):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        set_cell_shading(meta.cell(i, 0), "E8EEF5")
        for cell in meta.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(cell == meta.cell(i, 0)), color="111827")
    set_table_geometry(meta, [1.65, 4.85])

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("快速抓手")
    set_run_font(r, size=13, bold=True, color="2E74B5")
    for item in [
        "先用“数据从哪来、到哪去、何时有效”定位数据通路题。",
        "流水线题先列读写寄存器，再画时空表，不靠脑算。",
        "微程序题先分清微命令、微指令、微程序、控制存储器。",
        "中断题先分断点和现场，再说中断向量和返回。",
    ]:
        add_bullet(doc, item)


def render_markdown(doc, content: str):
    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "[[PAGEBREAK]]":
            doc.add_page_break()
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_md(doc, table_lines)
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_callout(doc, " ".join(quote_lines))
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            add_number(doc, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            add_body(doc, stripped)
        i += 1


def audit_document(doc_path: Path):
    # Lightweight structural audit for sections/tables before visual rendering.
    doc = Document(doc_path)
    assert len(doc.paragraphs) > 80, "Document looks too short."
    assert len(doc.tables) >= 12, "Expected summary tables were not generated."
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                assert cell.text.strip(), "Empty table cell detected."


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    render_markdown(doc, CONTENT)
    doc.save(OUT)
    audit_document(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
