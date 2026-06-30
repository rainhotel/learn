from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\moniC\project\learn")
OUT_DIR = ROOT / "01-topics" / "digital-image-processing" / "exam-docx"
OUT_FILE = OUT_DIR / "dip-algorithm-special-review-with-answers.docx"


FORMULA_ROWS = [
    ("采样与量化", "采样决定空间分辨率；量化决定灰度级分辨率；L=2^k"),
    ("数据量", "M x N x bit_depth；RGB uint8为M x N x 3 byte"),
    ("灰度反转", "s=L-1-r"),
    ("对数变换", "s=c log(1+r)，扩展暗部、压缩亮部"),
    ("幂律变换", "s=c r^gamma；gamma<1变亮，gamma>1变暗"),
    ("直方图均衡化", "s_k=round((L-1)CDF(r_k))，按题目取整规则执行"),
    ("均值滤波", "g=1/mn sum f(s,t)，平滑噪声但模糊边缘"),
    ("中值滤波", "邻域排序取中间值，适合椒盐噪声"),
    ("逆谐波均值", "f_hat=sum g^(Q+1)/sum g^Q；Q>0去椒，Q<0去盐"),
    ("Sobel幅值", "|G|=sqrt(Gx^2+Gy^2)，也常用|Gx|+|Gy|近似"),
    ("二维傅里叶变换/DFT", "F(u,v)=sum_x sum_y f(x,y)e^{-j2pi(ux/M+vy/N)}"),
    ("频域滤波", "G(u,v)=H(u,v)F(u,v)"),
    ("低通/高通互补", "H_hp=1-H_lp；低通平滑，高通锐化"),
    ("巴特沃斯低通", "H=1/[1+(D/D0)^(2n)]"),
    ("退化模型", "g=h*f+n；G=HF+N"),
    ("逆滤波", "F_hat=G/H=F+N/H，H小时放大噪声"),
    ("维纳滤波", "F_hat=[H*/(|H|^2+Sn/Sf)]G"),
    ("Otsu", "sigma_B^2=P0P1(m0-m1)^2，取最大者"),
    ("区域生长", "|candidate-mean(region)|<T，加入后更新均值"),
    ("大气散射去雾", "I=Jt+A(1-t)，J=(I-A)/max(t,t0)+A"),
]


ALGORITHM_TEMPLATES = [
    (
        "直方图均衡化",
        [
            "统计每个灰度级像素数n_k，除以总像素数得到概率p(r_k)。",
            "逐级累加得到CDF(r_k)。",
            "代入s_k=(L-1)CDF(r_k)，按题目要求四舍五入或向下取整。",
            "写出灰度映射关系，并说明对比度为什么被拉伸。",
        ],
    ),
    (
        "Otsu最大类间方差",
        [
            "枚举候选阈值k，把灰度<=k分成C0，把灰度>k分成C1。",
            "分别计算P0、P1、m0、m1。",
            "计算sigma_B^2=P0P1(m0-m1)^2。",
            "比较所有sigma_B^2，最大值对应的k就是阈值。",
        ],
    ),
    (
        "空间模板滤波",
        [
            "确定模板和3x3或指定邻域。",
            "模板与邻域对应位置相乘，再全部求和。",
            "若模板带1/9、1/16等归一化系数，最后再相除。",
            "边界处理要说明补零、复制边界或只处理有效区域。",
        ],
    ),
    (
        "Sobel/Prewitt边缘检测",
        [
            "分别用Gx和Gy模板做对应乘加。",
            "计算梯度幅值sqrt(Gx^2+Gy^2)或近似|Gx|+|Gy|。",
            "与阈值T比较，超过T判为边缘。",
            "分析题补充：Sobel中心权重为2，抗噪性通常优于Prewitt。",
        ],
    ),
    (
        "频域滤波",
        [
            "对f(x,y)做DFT得到F(u,v)，必要时fftshift把低频移到中心。",
            "根据任务构造H(u,v)：低通平滑，高通锐化，带阻/陷波去周期噪声。",
            "计算G=HF。",
            "做ifftshift和IDFT回到空间域，取实部并归一化显示。",
        ],
    ),
    (
        "图像复原",
        [
            "先写退化模型G=HF+N。",
            "若噪声可忽略，可讨论逆滤波F_hat=G/H。",
            "若噪声明显，应优先写维纳滤波，说明它考虑Sn/Sf。",
            "分析H接近0会放大噪声，必要时限制频率范围或加入正则项。",
        ],
    ),
    (
        "区域分割",
        [
            "区域生长题：写种子点、生长规则、连通方式、终止条件。",
            "每加入一个像素后更新当前区域均值或区域统计量。",
            "分裂合并题：分裂看区域内部是否均匀，合并看相邻区域是否相似。",
            "输出时说明区域完整、连通、互斥、一致。",
        ],
    ),
]


QUESTIONS = [
    {
        "title": "题1 数字图像基础：采样、量化与数据量",
        "score": "8分",
        "test_points": "采样、量化、空间分辨率、灰度级分辨率、RGB数据量",
        "question": [
            "某灰度图像大小为1024 x 1024，灰度级为256级。另有一幅同尺寸RGB真彩色图像，每个通道8 bit。",
            "1）灰度图像每个像素需要多少bit？",
            "2）灰度图像未压缩数据量是多少MB？",
            "3）RGB图像未压缩数据量是多少MB？",
            "4）说明采样和量化分别影响什么图像质量因素。",
        ],
        "answer": [
            "256=2^8，所以灰度图像每个像素需要8 bit。",
            "灰度图像数据量=1024 x 1024 x 8 bit=1024 x 1024 byte=1 MB。",
            "RGB图像每像素3通道 x 8 bit=24 bit=3 byte，所以数据量=1024 x 1024 x 3 byte=3 MB。",
            "采样影响空间分辨率，采样点越多，空间细节越丰富；量化影响灰度级分辨率，灰度级越多，灰度过渡越平滑。"
        ],
        "tips": "这类题分值不高但很稳，关键是别把256灰度级误写成256 bit。",
    },
    {
        "title": "题2 灰度变换：反转、对数和伽马",
        "score": "10分",
        "test_points": "点运算、反转、对数变换、幂律变换、增强效果判断",
        "question": [
            "8 bit灰度图像中某像素r=64。设反转变换s=255-r；对数变换s=c log(1+r)，要求r=255时s=255；归一化幂律变换s=r^gamma，其中r已经归一化到[0,1]。",
            "1）反转结果是多少？",
            "2）对数变换中c是多少？",
            "3）若归一化r=0.25，分别计算gamma=0.5和gamma=2时的输出，并说明亮暗变化。",
        ],
        "answer": [
            "反转：s=255-64=191。",
            "对数变换最大值约束：255=c log(1+255)，所以c=255/log(256)。",
            "gamma=0.5时，s=sqrt(0.25)=0.5，输出变大，图像变亮；gamma=2时，s=0.25^2=0.0625，输出变小，图像变暗。",
            "总结：对数变换扩展低灰度、压缩高灰度；gamma<1提亮，gamma>1压暗。"
        ],
        "tips": "灰度变换是点运算，输出只依赖同位置输入像素，不利用邻域。",
    },
    {
        "title": "题3 直方图均衡化完整手算",
        "score": "14分",
        "test_points": "概率、CDF、灰度映射、对比度增强",
        "question": [
            "一幅图像只有4个灰度级0、1、2、3，总像素数为16，各灰度级像素个数为n=[4,4,6,2]。设L=4，使用s_k=round((L-1)CDF(r_k))。",
            "求：1）各灰度级概率；2）CDF；3）均衡化映射；4）说明均衡化如何改善图像质量。",
        ],
        "answer": [
            "概率p=[4/16,4/16,6/16,2/16]=[0.25,0.25,0.375,0.125]。",
            "CDF=[0.25,0.50,0.875,1.00]。",
            "乘L-1=3得到[0.75,1.50,2.625,3.00]，四舍五入得到[1,2,3,3]。映射为0->1，1->2，2->3，3->3。",
            "均衡化通过累计分布把频繁灰度重新分配到较宽灰度范围，通常能提高最大灰度与最小灰度差异，增强对比度。但离散图像中不保证输出直方图完全平坦。"
        ],
        "tips": "一定写清楚取整规则。题目没说时，要说明自己采用四舍五入或向下取整。",
    },
    {
        "title": "题4 直方图规定化/匹配",
        "score": "10分",
        "test_points": "均衡化与规定化区别、目标直方图、映射选择",
        "question": [
            "某图像均衡化后灰度0、1、2、3对应的累计值分别为S=[0.20,0.45,0.80,1.00]。目标图像的累计分布为G=[0.10,0.50,0.75,1.00]。",
            "按“找目标CDF中与S最接近的灰度级”原则，写出源灰度0、1、2、3的规定化映射，并说明直方图均衡化和规定化的区别。",
        ],
        "answer": [
            "S0=0.20，与G0=0.10差0.10，与G1=0.50差0.30，所以0->0。",
            "S1=0.45，与G1=0.50最接近，所以1->1。",
            "S2=0.80，与G2=0.75差0.05，与G3=1.00差0.20，所以2->2。",
            "S3=1.00，与G3=1.00完全一致，所以3->3。",
            "映射为0->0，1->1，2->2，3->3。均衡化目标是尽量扩展灰度动态范围；规定化目标是让输出直方图接近指定分布，更可控。"
        ],
        "tips": "匹配题不一定改变所有灰度，关键看CDF距离最小。",
    },
    {
        "title": "题5 图像平均降噪与算术运算",
        "score": "8分",
        "test_points": "图像平均、噪声方差、算术运算增强",
        "question": [
            "同一静止场景拍摄M幅含独立零均值噪声的图像并求平均。若单幅图像噪声标准差为20，M=16，平均后噪声标准差是多少？图像平均为什么能改善图像质量？",
        ],
        "answer": [
            "独立噪声平均后方差变为原来的1/M，标准差变为原来的1/sqrt(M)。",
            "M=16，所以标准差=20/sqrt(16)=20/4=5。",
            "因为真实场景信号在多幅图中基本相同，平均后仍保留；独立随机噪声正负抵消，方差降低，所以图像更平滑、更清晰。"
        ],
        "tips": "问标准差是1/sqrt(M)，问方差才是1/M。",
    },
    {
        "title": "题6 空域平滑：均值、中值、最大值和最小值滤波",
        "score": "12分",
        "test_points": "空间滤波、椒盐噪声、滤波器选择",
        "question": [
            "某3x3邻域像素为：0, 98, 100, 101, 102, 103, 104, 105, 255。",
            "1）算术均值滤波输出是多少？",
            "2）中值滤波输出是多少？",
            "3）最大值滤波和最小值滤波分别适合去除什么噪声？",
            "4）若图像有椒盐噪声，优先选哪种滤波器？",
        ],
        "answer": [
            "均值=(0+98+100+101+102+103+104+105+255)/9=968/9≈107.56。",
            "排序后第5个数是102，所以中值滤波输出102。",
            "最大值滤波用邻域最大值替代中心，能去除黑色椒噪声；最小值滤波能去除白色盐噪声。",
            "椒盐噪声优先选中值滤波，因为它能去除孤立极大/极小值，同时较好保护边缘。"
        ],
        "tips": "最大值去椒、最小值去盐，中值去椒盐，别把盐和椒的颜色记反。",
    },
    {
        "title": "题7 逆谐波均值滤波器",
        "score": "10分",
        "test_points": "谐波/逆谐波均值、Q符号、盐椒噪声",
        "question": [
            "逆谐波均值滤波器为f_hat=sum g^(Q+1)/sum g^Q。",
            "1）Q>0适合去除盐噪声还是椒噪声？Q<0呢？",
            "2）为什么不能同时用同一个Q去除盐噪声和椒噪声？",
        ],
        "answer": [
            "Q>0适合去除椒噪声，也就是低灰度黑点；Q<0适合去除盐噪声，也就是高灰度白点。",
            "Q的正负决定滤波器对高灰度或低灰度异常值的抑制方向。若Q选反，会强化相反类型的脉冲噪声；同一个Q通常不能同时有效去除盐和椒两类极端噪声。",
            "若同时存在明显盐和椒噪声，通常优先使用中值滤波。"
        ],
        "tips": "这是选择题和简答题都爱考的符号题：Q>0去椒，Q<0去盐。",
    },
    {
        "title": "题8 Sobel边缘检测手算",
        "score": "14分",
        "test_points": "一阶微分、Sobel算子、阈值边缘判定",
        "question": [
            "给定3x3邻域：第一行10,10,10；第二行10,50,90；第三行10,90,90。Sobel模板为Gx=[-1 0 1; -2 0 2; -1 0 1]，Gy=[-1 -2 -1; 0 0 0; 1 2 1]。",
            "1）计算Gx和Gy；2）计算sqrt(Gx^2+Gy^2)；3）若阈值T=300，中心像素是否为边缘？4）说明Prewitt和Sobel区别。",
        ],
        "answer": [
            "Gx=(-10+10)+(-20+180)+(-10+90)=240。",
            "Gy=(-10-20-10)+(10+180+90)=240。",
            "梯度幅值=sqrt(240^2+240^2)=240sqrt(2)≈339.4。",
            "339.4>300，所以判定为边缘点。",
            "Prewitt模板权重较均匀；Sobel在中心行或中心列权重为2，相当于增加平滑作用，通常抗噪性更好。"
        ],
        "tips": "模板题要按位置相乘，不要只背结论。Sobel、Prewitt都是一阶微分边缘检测。",
    },
    {
        "title": "题9 Laplacian锐化与模板符号",
        "score": "10分",
        "test_points": "二阶微分、拉普拉斯锐化、符号配套",
        "question": [
            "某点原灰度f=100，使用中心为-4的四邻域Laplacian模板[0 1 0; 1 -4 1; 0 1 0]得到响应L=-20。若采用g=f-L进行锐化，输出是多少？为什么有些教材写g=f+L？",
        ],
        "answer": [
            "按题目公式g=f-L=100-(-20)=120。",
            "Laplacian模板有两套符号：中心为负时常写g=f-L；中心为正时常写g=f+L。关键不是死背加号或减号，而是让边缘细节被增强，并与模板符号配套。",
            "二阶微分对噪声敏感，实际使用时常先平滑或注意噪声放大。"
        ],
        "tips": "看到Laplacian先看模板中心符号，再决定加减。",
    },
    {
        "title": "题10 傅里叶变换与频域基础：F(0,0)",
        "score": "10分",
        "test_points": "DFT、直流分量、低频高频、频谱中心化",
        "question": [
            "给定2x2图像[[1,2],[3,4]]。",
            "1）求F(0,0)；2）图像均值是多少？3）说明低频和高频分别对应什么图像内容；4）fftshift的作用是什么？",
        ],
        "answer": [
            "F(0,0)=所有像素灰度和=1+2+3+4=10。",
            "均值=F(0,0)/(MN)=10/4=2.5。",
            "低频对应整体亮度、背景和平滑区域；高频对应边缘、纹理、细节和噪声。",
            "fftshift把频谱低频成分移动到中心，便于观察和设计滤波器；它不改变图像本身信息。"
        ],
        "tips": "频谱未中心化时低频在角落，中心化后低频在中心。",
    },
    {
        "title": "题11 低通滤波器设计与比较",
        "score": "12分",
        "test_points": "理想低通、巴特沃斯低通、高斯低通、振铃、D0",
        "question": [
            "巴特沃斯低通H=1/[1+(D/D0)^(2n)]。",
            "1）当D=D0，n=2时H是多少？",
            "2）D0变小对图像有什么影响？",
            "3）比较理想低通、巴特沃斯低通和高斯低通的过渡与振铃特点。",
        ],
        "answer": [
            "D=D0时，D/D0=1，所以H=1/(1+1^(2n))=1/2=0.5。",
            "D0变小，保留的低频范围变窄，高频被更多抑制，图像更平滑更模糊，噪声可能减少但边缘细节损失更明显。",
            "理想低通截止最突然，最容易产生振铃；巴特沃斯低通过渡由阶数n控制，n越大越接近理想低通；高斯低通过渡最平滑，振铃最少。"
        ],
        "tips": "低通用于平滑去噪，不用于增强边缘。",
    },
    {
        "title": "题12 高通、高提升、带阻与周期噪声",
        "score": "14分",
        "test_points": "高通锐化、高提升滤波、带阻/陷波、周期噪声",
        "question": [
            "1）高通滤波器的作用是什么？为什么会放大噪声？",
            "2）高提升滤波H_hb=(A-1)+H_hp，当A=1时是什么？A>1有什么意义？",
            "3）若图像空间域出现规律条纹，频谱中有关于中心对称的亮点，应采用什么滤波器？写出处理流程。",
        ],
        "answer": [
            "高通保留高频、抑制低频，用于增强边缘和细节。噪声常含高频成分，所以高通也可能把噪声一起增强。",
            "A=1时，H_hb=H_hp，退化为普通高通滤波。A>1时保留一部分原图背景低频，同时增强高频，图像不会只剩边缘轮廓。",
            "规律条纹通常是周期噪声，频谱中表现为对称亮点。应使用陷波滤波器抑制具体亮点；若噪声占一圈频带，可用带阻滤波器。",
            "流程：DFT并中心化 -> 找到噪声亮点位置 -> 构造陷波/带阻H -> G=HF -> 反变换得到去噪图像。"
        ],
        "tips": "看到“周期条纹、对称亮点”，第一反应是陷波或带阻。",
    },
    {
        "title": "题13 图像复原：逆滤波与维纳滤波",
        "score": "14分",
        "test_points": "退化模型、逆滤波噪声放大、维纳滤波",
        "question": [
            "频域退化模型为G=HF+N。某频率处H=0.1，噪声N=2。",
            "1）逆滤波后噪声项变为多少？",
            "2）为什么已知H也不一定能完全复原？",
            "3）写出维纳滤波基本形式，并说明它为什么比逆滤波稳定。",
        ],
        "answer": [
            "逆滤波F_hat=G/H=F+N/H，噪声项N/H=2/0.1=20，噪声被放大10倍。",
            "因为观测图像含噪声N，当H很小时，N/H会被严重放大；此外H可能存在零点或接近零的频率，导致复原不稳定。",
            "维纳滤波：F_hat=[H*/(|H|^2+Sn/Sf)]G。其中Sn/Sf表示噪声与图像功率谱比。",
            "维纳滤波在分母中考虑噪声统计，噪声越强抑制越明显，在最小均方误差意义下通常比直接G/H更稳定。"
        ],
        "tips": "复原题一定先写G=HF+N，再分析N/H。",
    },
    {
        "title": "题14 噪声类型与复原滤波器选择",
        "score": "12分",
        "test_points": "高斯、瑞利、伽马、指数、均匀、椒盐噪声；滤波器适用条件",
        "question": [
            "请根据现象选择合适方法，并说明理由：",
            "1）图像中出现随机黑白亮暗点。",
            "2）图像叠加近似钟形分布的加性噪声。",
            "3）图像中有黑色椒噪声为主。",
            "4）图像中有白色盐噪声为主。",
        ],
        "answer": [
            "随机黑白亮暗点是椒盐噪声，优先用中值滤波。",
            "近似钟形分布通常是高斯噪声，可用算术均值滤波、低通滤波或自适应局部降噪方法，但要注意边缘模糊。",
            "黑色椒噪声为低灰度异常点，可用最大值滤波或逆谐波均值滤波Q>0。",
            "白色盐噪声为高灰度异常点，可用最小值滤波或逆谐波均值滤波Q<0。"
        ],
        "tips": "常见噪声名字要能识别，但考试更常考“现象-方法-理由”。",
    },
    {
        "title": "题15 Otsu阈值分割完整计算",
        "score": "16分",
        "test_points": "阈值分割、全局阈值、最大类间方差",
        "question": [
            "一幅图像只有灰度级0、1、2、3，对应像素个数为[2,2,4,8]，总数16。阈值k把灰度<=k分为C0，灰度>k分为C1。计算k=0、1、2时的sigma_B^2=P0P1(m0-m1)^2，并确定最佳阈值。",
        ],
        "answer": [
            "概率p=[0.125,0.125,0.25,0.5]。",
            "k=0：P0=0.125，m0=0；P1=0.875，m1=(1*0.125+2*0.25+3*0.5)/0.875=2.4286；sigma_B^2≈0.125*0.875*(2.4286)^2≈0.645。",
            "k=1：P0=0.25，m0=(1*0.125)/0.25=0.5；P1=0.75，m1=(2*0.25+3*0.5)/0.75=2.6667；sigma_B^2≈0.25*0.75*(2.1667)^2≈0.880。",
            "k=2：P0=0.5，m0=(1*0.125+2*0.25)/0.5=1.25；P1=0.5，m1=3；sigma_B^2=0.5*0.5*(1.75)^2=0.7656。",
            "最大类间方差出现在k=1，所以最佳阈值为1。"
        ],
        "tips": "Otsu不是凭眼睛找谷底，而是最大化类间方差。",
    },
    {
        "title": "题16 迭代全局阈值",
        "score": "10分",
        "test_points": "全局阈值、迭代法、均值更新",
        "question": [
            "灰度集合为{10,12,14,200,210,220}，初始阈值T0=100。按迭代阈值法求收敛阈值。",
        ],
        "answer": [
            "按T0=100分组：G1={10,12,14}，G2={200,210,220}。",
            "均值：mu1=12，mu2=210。",
            "更新T1=(mu1+mu2)/2=(12+210)/2=111。",
            "再用T=111分组，分组仍不变，阈值仍为111，所以收敛阈值为111。"
        ],
        "tips": "迭代阈值题步骤固定：分组、求均值、更新、判断收敛。",
    },
    {
        "title": "题17 区域生长算法",
        "score": "12分",
        "test_points": "区域生长、种子点、生长准则、当前区域均值",
        "question": [
            "种子点灰度为100，当前区域已包含像素100、102、104，阈值T=5。候选邻接像素灰度依次为106、109、101。假设每次只要满足|candidate-mean(region)|<T就加入，且加入后更新均值。判断哪些像素能加入。",
        ],
        "answer": [
            "初始区域均值=(100+102+104)/3=102。",
            "候选106：|106-102|=4<5，可以加入。新均值=(100+102+104+106)/4=103。",
            "候选109：|109-103|=6>5，不能加入。",
            "候选101：当前区域仍为{100,102,104,106}，均值103；|101-103|=2<5，可以加入。",
            "最终能加入的候选像素为106和101，109不能加入。"
        ],
        "tips": "区域生长最容易错在一直和种子点比较。PPT和考试更常用当前区域均值。",
    },
    {
        "title": "题18 区域分裂与合并",
        "score": "10分",
        "test_points": "区域一致性、分裂合并、标准差和均值差",
        "question": [
            "某区域R的灰度标准差大于分裂阈值，因此分裂为R1、R2、R3、R4。分裂后R1和R2内部标准差都小于阈值，且二者均值差小于合并阈值。请说明处理流程和理由。",
        ],
        "answer": [
            "整块R标准差大于分裂阈值，说明区域内部不均匀，不满足一致性，应先分裂为子区域。",
            "R1和R2分裂后内部标准差小于阈值，说明各自内部一致。",
            "R1和R2均值差小于合并阈值，说明相邻区域相似，可以合并。",
            "分裂看区域内部波动，合并看相邻区域之间的相似性。最终目标是得到连通、互斥、内部一致且相邻区域不应再合并的分割结果。"
        ],
        "tips": "分裂阈值和合并阈值含义不同，别混在一起。",
    },
    {
        "title": "题19 彩色图像处理与颜色空间",
        "score": "10分",
        "test_points": "RGB、CMY、HSI、彩色增强、伪彩色",
        "question": [
            "归一化RGB颜色为R=0.2，G=0.7，B=0.4。",
            "1）求CMY。",
            "2）如果只想增强亮度而尽量保持色调，应在哪个颜色空间处理哪个分量？",
            "3）说明伪彩色增强与真彩色处理的区别。",
        ],
        "answer": [
            "C=1-R=0.8，M=1-G=0.3，Y=1-B=0.6。",
            "应转到HSI或类似亮度分离的颜色空间，处理I分量，再转换回RGB。这样能尽量保持H色调和S饱和度。",
            "伪彩色增强是把灰度值映射成颜色，目的是便于观察，不代表真实颜色；真彩色处理针对真实RGB彩色图像进行增强、滤波、分割或颜色空间变换。"
        ],
        "tips": "RGB面向显示，CMY/CMYK面向印刷，HSI更贴近人眼感知。",
    },
    {
        "title": "题20 图像去雾：大气散射模型",
        "score": "12分",
        "test_points": "大气散射模型、暗通道先验、A和t、恢复公式",
        "question": [
            "单通道大气散射模型I=Jt+A(1-t)。已知I=180，A=220，t=0.4。",
            "1）求清晰图像像素J。",
            "2）如果t很小，直接恢复有什么问题？",
            "3）简述暗通道先验去雾的基本流程。",
        ],
        "answer": [
            "由模型变形J=(I-A)/t+A。代入J=(180-220)/0.4+220=-100+220=120。",
            "如果t很小，公式中除以t会放大噪声和估计误差，所以实际常用max(t,t0)限制下限。",
            "暗通道先验流程：计算暗通道 -> 从暗通道最亮区域估计大气光A -> 根据暗通道估计透射率t -> 用恢复公式求J -> 必要时对t做导向滤波等细化。"
        ],
        "tips": "去雾不是简单锐化，而是估计A和t后按物理模型恢复。",
    },
]


def set_east_asian_font(run, font="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    title = styles.add_style("ReviewTitle", 1)
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles.add_style("SubtitleLine", 1)
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(10)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_header_footer(doc):
    header = doc.sections[0].header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("数字图像处理算法题专题")
    set_east_asian_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("综合题与算法应用复习")
    set_east_asian_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("777777")


def add_para(doc, text, label=None):
    p = doc.add_paragraph()
    if label:
        r = p.add_run(label)
        r.bold = True
        set_east_asian_font(r)
    r = p.add_run(text)
    set_east_asian_font(r)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_list(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_east_asian_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade_cell(hdr[i], "E8EEF5")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                set_east_asian_font(r)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    set_east_asian_font(r)
    return table


def add_title(doc):
    p = doc.add_paragraph(style="ReviewTitle")
    r = p.add_run("数字图像处理综合算法题专题复习")
    set_east_asian_font(r)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph(style="SubtitleLine")
    r = p.add_run("针对50分综合题：算法应用与实现。题目后直接附参考答案，覆盖PPT重点考点。")
    set_east_asian_font(r)


def add_question_block(doc, item, index):
    doc.add_heading(f"{index}. {item['title']}（{item['score']}）", level=2)
    add_para(doc, item["test_points"], "考点：")
    add_para(doc, "题目：", None)
    add_list(doc, item["question"], "List Number")
    add_para(doc, "参考答案：", None)
    add_list(doc, item["answer"], "List Number")
    add_para(doc, item["tips"], "拿分提醒：")


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_title(doc)

    doc.add_heading("一、先抓最高分：50分综合题怎么复习", level=1)
    add_para(
        doc,
        "从你发的考试题型看，综合题“算法应用与实现”占50分，是最值得优先复习的部分。"
        "它不是只考背诵，而是考你能不能把算法步骤、公式代入、结果分析和适用条件写完整。"
    )
    add_table(
        doc,
        ["复习层级", "你要练到什么程度", "对应考点"],
        [
            ("会算", "能手算小矩阵、小直方图、小阈值、小滤波模板。", "直方图均衡化、Otsu、Sobel、均值/中值、F(0,0)、数据量"),
            ("会写步骤", "能不看书写出算法流程，每步有输入、处理和输出。", "频域滤波、区域生长、分裂合并、复原、去雾"),
            ("会选方法", "看到图像现象能选算法，并说明理由和缺点。", "椒盐噪声选中值，周期噪声选陷波，光照不均选自适应阈值"),
            ("会分析", "能写优点、局限、易错点，拿简答和综合题尾分。", "逆滤波噪声放大，理想滤波振铃，Sobel抗噪优于Prewitt"),
        ],
        [1800, 3900, 3660],
    )
    add_para(
        doc,
        "复习顺序建议：先刷直方图、Otsu、Sobel、空间滤波这四类手算题；再背频域滤波、复原、区域分割三类流程题；最后用彩色处理、采样量化、去雾模型补齐选择题和简答题。"
    )

    doc.add_heading("二、算法题通用答题模板", level=1)
    for name, steps in ALGORITHM_TEMPLATES:
        doc.add_heading(name, level=2)
        add_list(doc, steps, "List Number")

    doc.add_heading("三、必背公式与判定表", level=1)
    add_table(doc, ["考点", "公式/判定"], FORMULA_ROWS, [2400, 6960])

    doc.add_heading("四、专题模拟试题与参考答案", level=1)
    add_para(
        doc,
        "下面的题目按综合题训练设计，每题后面直接附答案。你复习时先遮住答案做一遍，再按“公式、步骤、结论、分析”四项核对。"
    )
    for i, question in enumerate(QUESTIONS, 1):
        add_question_block(doc, question, i)

    doc.add_heading("五、最后一轮考前检查清单", level=1)
    add_list(
        doc,
        [
            "能否独立写出直方图均衡化的概率、CDF、映射步骤？",
            "能否独立算Otsu的P0、P1、m0、m1和类间方差？",
            "能否把3x3模板与邻域对应相乘求和？",
            "能否说明Prewitt、Sobel、Laplacian的区别？",
            "能否看频谱判断低通、高通、带阻、陷波？",
            "能否解释逆滤波为什么放大噪声，维纳滤波为什么更稳？",
            "能否根据噪声现象选择均值、中值、最大值、最小值、逆谐波或自适应滤波？",
            "能否写出区域生长和分裂合并的标准流程？",
            "能否解释RGB、CMY、HSI分别适合什么场景？",
            "能否写出暗通道去雾的大气散射模型和恢复公式？",
        ],
        "List Bullet",
    )
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
