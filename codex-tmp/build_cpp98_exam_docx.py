from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\moniC\project\learn\01-topics\algorithm-final-review\cpp98-exam-template-guide.docx"


def set_font(run, name="Calibri", size=11, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("C++98 算法考试板子速查")
    set_font(run, size=24, color="0B2545", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("STL / pair / DFS / BFS / 贪心 / DP / 背包 / Dijkstra / Prim / 二分")
    set_font(run, size=11, color="555555")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("使用方式：").bold = True
    p.add_run("先看“什么时候用”，再背代码，最后重点理解核心循环里的变量含义。所有代码按 C++98 风格整理。")


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.keep_with_next = True


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.keep_with_next = True


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.keep_with_next = True


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=10.5)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    run = p.add_run(code.strip("\n"))
    set_font(run, name="Consolas", size=8.5, color="1F2937")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")
    p_pr.append(ind)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "说明"
    shade_cell(table.rows[0].cells[0], "E8EEF5")
    shade_cell(table.rows[0].cells[1], "E8EEF5")
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, [1.45, 5.05])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=9.5)
    doc.add_paragraph()


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        shade_cell(table.rows[0].cells[idx], "E8EEF5")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=9.0)
    doc.add_paragraph()


def add_classic_table(doc, rows):
    add_h3(doc, "经典题型拓展")
    add_matrix_table(doc, ["题型", "板子作用", "边界处理"], rows, [1.55, 2.35, 2.60])


def add_explain_block(doc, variables, loop_text, mistakes):
    add_h3(doc, "变量说明")
    add_kv_table(doc, variables)
    add_h3(doc, "核心循环怎么跑")
    add_bullets(doc, loop_text)
    add_h3(doc, "易错点")
    add_bullets(doc, mistakes)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_reference_intro(doc):
    add_h1(doc, "一、考试统一写法")
    add_h2(doc, "1. 常用头文件")
    add_code(doc, """
#include <iostream>
#include <cstdio>
#include <cstring>
#include <cmath>
#include <algorithm>
#include <queue>
#include <stack>
#include <vector>
#include <string>
#include <utility>
#include <map>
#include <set>
#include <functional>
using namespace std;
""")
    add_bullets(doc, [
        "不使用 bits/stdc++.h，避免考试环境不支持。",
        "C++98 里嵌套模板要写空格，例如 vector< pair<int, int> >。",
        "常用 INF 写法：const int INF = 0x3f3f3f3f;",
        "如果边权或答案可能超过 int，距离数组和答案要改成 long long。",
    ])

    add_h2(doc, "2. 数组、排序和 algorithm")
    add_code(doc, """
sort(a, a + n);
reverse(a, a + n);
memset(a, 0, sizeof(a));
memset(a, -1, sizeof(a));
memset(a, 0x3f, sizeof(a));

int x = max(3, 5);
int y = min(3, 5);
swap(a[0], a[1]);

int p1 = lower_bound(a, a + n, x) - a;  // 第一个 >= x
int p2 = upper_bound(a, a + n, x) - a;  // 第一个 > x

sort(a, a + n);
int newn = unique(a, a + n) - a;        // 排序后去重

do {
    // 使用当前排列
} while (next_permutation(a, a + n));
""")
    add_kv_table(doc, [
        ("sort", "排序范围是左闭右开 [begin, end)，数组常写 a 到 a+n。"),
        ("lower_bound", "要求有序，返回第一个 >= x 的迭代器。"),
        ("upper_bound", "要求有序，返回第一个 > x 的迭代器。"),
        ("unique", "只删除相邻重复，真正去重前一般要先 sort。返回新结尾。"),
        ("next_permutation", "从当前排列找下一个字典序排列；如果要枚举全部，先 sort。"),
    ])

    add_h2(doc, "3. vector 和 string")
    add_code(doc, """
vector<int> v;
vector<int> a(n, 0);

v.push_back(10);
int len = (int)v.size();
int first = v[0];
int last = v.back();
v.pop_back();
v.clear();

string s = "abcde";
int slen = (int)s.size();
char c = s[0];
string t = s.substr(1, 3);        // 从下标 1 开始取 3 个字符
int pos = (int)s.find("cd");      // 找不到时返回 string::npos
sort(s.begin(), s.end());
""")
    add_kv_table(doc, [
        ("vector 下标", "v[i] 不检查越界，考试里先保证 0 <= i < size。"),
        ("back/pop_back", "空 vector 不能调用，先判断 !v.empty()。"),
        ("size()", "返回无符号类型，和 int 比较时建议强转成 int。"),
        ("substr(pos,len)", "pos 不能超过字符串长度；len 超出时会自动取到结尾。"),
        ("find", "找不到是 string::npos，不要直接当合法下标用。"),
    ])

    add_h2(doc, "4. pair、queue、stack、priority_queue")
    add_code(doc, """

pair<int, int> p = make_pair(3, 5);
cout << p.first << " " << p.second << endl;

queue< pair<int, int> > q;
q.push(make_pair(1, 2));
pair<int, int> cur = q.front();
q.pop();

stack<int> st;
st.push(1);
st.top();
st.pop();

priority_queue<int> big;  // 大根堆
priority_queue<int, vector<int>, greater<int> > small;  // 小根堆

priority_queue< pair<int, int>,
                vector< pair<int, int> >,
                greater< pair<int, int> > > pq;
""")
    add_kv_table(doc, [
        ("pair.first", "通常放第一个维度：点编号、x 坐标、区间左端点、排序主关键字。"),
        ("pair.second", "通常放第二个维度：距离、y 坐标、区间右端点、排序次关键字。"),
        ("make_pair(a,b)", "快速构造 pair，常和 queue、priority_queue、vector 配合。"),
        ("sort(pair数组)", "默认先按 first 升序，再按 second 升序。"),
        ("queue/stack", "front/top 之前必须确认非空。"),
        ("priority_queue", "默认大根堆；小根堆要写 greater，并包含 functional。"),
    ])

    add_h2(doc, "5. map 和 set")
    add_code(doc, """
map<string, int> mp;
mp["apple"]++;

if (mp.find("apple") != mp.end()) {
    cout << mp["apple"] << endl;
}

set<int> st;
st.insert(3);
st.insert(5);
st.count(3);

set<int>::iterator it = st.lower_bound(4);
if (it != st.end()) {
    cout << *it << endl;
}
""")
    add_kv_table(doc, [
        ("map", "键到值的映射，适合计数、编号、查表。"),
        ("mp[key]", "如果 key 不存在，会自动创建默认值；只判断存在时优先用 find。"),
        ("set", "有序且自动去重，适合维护不重复集合。"),
        ("count", "set 中只会返回 0 或 1。"),
        ("迭代器", "lower_bound 可能返回 end，解引用前必须判断。"),
    ])

    add_h2(doc, "6. STL 边界总提醒")
    add_kv_table(doc, [
        ("左闭右开", "STL 区间大多是 [l, r)，数组写 a 到 a+n，vector 写 begin 到 end。"),
        ("空容器", "front、back、top、pop 前先判断 empty。"),
        ("有序前提", "lower_bound、upper_bound、binary_search 必须在有序序列上用。"),
        ("比较函数", "sort 的 cmp 不能写 <=，相等时必须返回 false。"),
        ("迭代器失效", "vector 插入删除后，旧迭代器可能失效；考试中尽量少边遍历边删。"),
        ("C++98 空格", "vector< pair<int, int> > 中两个 > 之间要留空格。"),
    ])


def add_search(doc):
    add_h1(doc, "二、DFS 与 BFS")
    add_h2(doc, "1. 网格 DFS：连通块 / 能否到达")
    add_code(doc, """
int dx[4] = {-1, 1, 0, 0};
int dy[4] = {0, 0, -1, 1};

void dfs(int x, int y) {
    int k;
    vis[x][y] = 1;

    for (k = 0; k < 4; ++k) {
        int nx = x + dx[k];
        int ny = y + dy[k];
        if (nx < 0 || nx >= n || ny < 0 || ny >= m) continue;
        if (vis[nx][ny] || g[nx][ny] == '#') continue;
        dfs(nx, ny);
    }
}
""")
    add_explain_block(
        doc,
        [
            ("g[x][y]", "地图，通常 '#' 表示墙或不可走位置。"),
            ("vis[x][y]", "是否已经访问，防止递归来回走导致死循环。"),
            ("dx/dy", "四方向移动数组，把“上下左右”写成循环。"),
            ("nx/ny", "下一步尝试访问的位置。"),
        ],
        [
            "进入 dfs 后先把当前点标记为访问，表示这一格已经属于当前搜索区域。",
            "for 循环枚举四个方向，每次计算下一格 nx、ny。",
            "先判断越界，再判断是否访问过或是否为障碍。",
            "只有合法且没访问过的位置才递归进入。",
        ],
        [
            "忘记 vis[x][y] = 1 会无限递归。",
            "要求最短步数时不要用 DFS 当最短路板子。",
            "下标从 0 还是 1 要和输入保持一致。",
        ],
    )

    add_h2(doc, "2. 网格 BFS：无权最短路 / 最少步数")
    add_code(doc, """
int bfs(int sx, int sy, int tx, int ty) {
    queue< pair<int, int> > q;
    memset(dista, -1, sizeof(dista));

    dista[sx][sy] = 0;
    q.push(make_pair(sx, sy));

    while (!q.empty()) {
        pair<int, int> cur = q.front();
        q.pop();

        int x = cur.first;
        int y = cur.second;
        if (x == tx && y == ty) return dista[x][y];

        for (int k = 0; k < 4; ++k) {
            int nx = x + dx[k];
            int ny = y + dy[k];
            if (nx < 0 || nx >= n || ny < 0 || ny >= m) continue;
            if (g[nx][ny] == '#' || dista[nx][ny] != -1) continue;
            dista[nx][ny] = dista[x][y] + 1;
            q.push(make_pair(nx, ny));
        }
    }
    return -1;
}
""")
    add_explain_block(
        doc,
        [
            ("queue", "BFS 的队列，保证先到的点先扩展，形成一层一层搜索。"),
            ("dista[x][y]", "起点到当前格子的最短步数，-1 表示没到过。"),
            ("sx/sy", "起点坐标。"),
            ("tx/ty", "终点坐标。"),
        ],
        [
            "先把起点距离设为 0，并放入队列。",
            "while 循环每次取出队头，表示扩展当前最早到达的格子。",
            "四方向尝试新格子，合法且没到过才更新距离。",
            "新格子的距离一定是 dista[x][y] + 1，因为每一步代价相同。",
            "第一次取到终点时，距离就是最短距离。",
        ],
        [
            "访问标记要在入队时完成，避免同一个点重复入队。",
            "BFS 只适合无权图或每步代价相同的题。",
            "如果终点不可达，记得返回 -1 或按题目要求输出。",
        ],
    )

    add_h2(doc, "3. 链式前向星：不用 vector 的图存储")
    add_code(doc, """
int head[MAXN], to[MAXM], nxt[MAXM], idx;

void init_graph(int n) {
    for (int i = 1; i <= n; ++i) head[i] = -1;
    idx = 0;
}

void add_edge(int u, int v) {
    to[idx] = v;
    nxt[idx] = head[u];
    head[u] = idx;
    ++idx;
}

for (int i = head[u]; i != -1; i = nxt[i]) {
    int v = to[i];
}
""")
    add_kv_table(doc, [
        ("head[u]", "点 u 的第一条边编号。没有边时是 -1。"),
        ("to[i]", "第 i 条边指向的点。"),
        ("nxt[i]", "与第 i 条边同起点的下一条边。"),
        ("idx", "当前已经存了多少条边，也是下一条边的编号。"),
        ("遍历循环", "从 head[u] 开始，不断跳到 nxt[i]，直到 -1。"),
    ])
    add_classic_table(doc, [
        ("连通块 / 岛屿数量", "DFS 把一整片相连区域搜完，外层循环遇到未访问点就答案加 1。", "地图边界、障碍字符、八方向还是四方向要看题目。"),
        ("迷宫最短路", "BFS 按层扩展，第一次到达终点就是最短步数。", "起点等于终点答案为 0；不可达要输出 -1 或题目指定文字。"),
        ("多源 BFS", "把多个起点一起入队，适合最近出口、火焰扩散、多个源点最短距离。", "所有源点初始距离都是 0；入队时就标记。"),
        ("普通图遍历", "链式前向星保存邻接点，DFS/BFS 判断连通性或层数。", "无向图加两条边；MAXM 要按双向边开两倍。"),
    ])


def add_greedy_dp(doc):
    add_h1(doc, "三、贪心与 DP")
    add_h2(doc, "1. 贪心策略速记")
    add_kv_table(doc, [
        ("活动选择", "按结束时间从小到大排序，能选就选。"),
        ("区间覆盖", "按左端点排序，每次在可选区间里扩到最远右端点。"),
        ("排队接水", "耗时短的人先做，总等待时间更小。"),
        ("分数背包", "按单位价值从高到低拿。"),
        ("哈夫曼合并", "每次合并最小的两个。"),
    ])
    add_body(doc, "贪心题先问：排序依据是什么？当前选了以后会不会影响后面？能不能说明局部最优不会吃亏？")
    add_classic_table(doc, [
        ("活动安排", "按结束时间排序，尽量早结束，给后面留空间。", "区间是否允许端点相接：看题目是 a.r <= b.l 还是 a.r < b.l。"),
        ("区间覆盖", "当前能接上的区间里选右端点最远的，最少段覆盖目标。", "如果这一轮最远右端点没有推进，说明无法覆盖。"),
        ("排队接水 / 作业排序", "短任务优先通常能最小化总等待时间。", "输出平均值时注意保留小数；相同时间的顺序通常不影响最优值。"),
        ("哈夫曼合并", "每次取最小两个合并，总代价最小。", "使用小根堆；n=1 时通常代价为 0。"),
    ])

    add_h2(doc, "2. 线性 DP：最大子段和")
    add_code(doc, """
dp[1] = a[1];
ans = dp[1];
for (i = 2; i <= n; ++i) {
    dp[i] = max(a[i], dp[i - 1] + a[i]);
    ans = max(ans, dp[i]);
}
""")
    add_explain_block(
        doc,
        [
            ("dp[i]", "以第 i 个数结尾的最大连续子段和。必须包含 a[i]。"),
            ("ans", "所有 dp[i] 中的最大值，是最终答案。"),
            ("a[i]", "当前数字，可以单独开一段，也可以接在前一段后面。"),
        ],
        [
            "每次只考虑“以 i 结尾”的答案。",
            "如果 dp[i-1] 是负贡献，就从 a[i] 重新开始。",
            "如果 dp[i-1] 是正贡献，就接上 a[i]。",
        ],
        [
            "最大子段和是连续子段，不是可以跳着选的子序列。",
            "全是负数时，答案应该是最大的那个负数，不能默认 0。",
        ],
    )

    add_h2(doc, "3. LIS / LNDS")
    add_code(doc, """
for (i = 1; i <= n; ++i) dp[i] = 1;

ans = 1;
for (i = 1; i <= n; ++i) {
    for (j = 1; j < i; ++j) {
        if (a[j] <= a[i]) {
            dp[i] = max(dp[i], dp[j] + 1);
        }
    }
    ans = max(ans, dp[i]);
}
""")
    add_explain_block(
        doc,
        [
            ("dp[i]", "以 a[i] 结尾的最长不下降子序列长度。"),
            ("j < i", "只从 i 前面的元素转移，保证子序列顺序不乱。"),
            ("a[j] <= a[i]", "不下降条件。严格上升时改成 a[j] < a[i]。"),
        ],
        [
            "外层 i 枚举结尾位置。",
            "内层 j 枚举 i 前面所有可能接到 a[i] 前面的元素。",
            "只要 a[j] 能接到 a[i] 前面，就用 dp[j] + 1 更新 dp[i]。",
        ],
        [
            "严格上升和不下降的符号不同：< 与 <= 要看题目。",
            "答案通常不是 dp[n]，而是所有 dp[i] 的最大值。",
        ],
    )

    add_h2(doc, "4. LCS：最长公共子序列")
    add_code(doc, """
for (i = 1; i <= n; ++i) {
    for (j = 1; j <= m; ++j) {
        if (a[i] == b[j]) {
            dp[i][j] = dp[i - 1][j - 1] + 1;
        } else {
            dp[i][j] = max(dp[i - 1][j], dp[i][j - 1]);
        }
    }
}
""")
    add_explain_block(
        doc,
        [
            ("dp[i][j]", "a 的前 i 个字符和 b 的前 j 个字符的最长公共子序列长度。"),
            ("a[i] == b[j]", "两个末尾字符相同，可以一起选。"),
            ("dp[i-1][j]", "不选 a[i] 的情况。"),
            ("dp[i][j-1]", "不选 b[j] 的情况。"),
        ],
        [
            "二维循环按 i、j 从小到大填表。",
            "末尾相等时，答案来自左上角加 1。",
            "末尾不等时，只能丢掉其中一个末尾，取上方或左方最大值。",
        ],
        [
            "子序列不要求连续，子串才要求连续。",
            "字符串从 1 开始读入时，可以用 cin >> (a + 1)。",
        ],
    )

    add_h2(doc, "5. 背包 DP：0/1 与完全")
    add_code(doc, """
// 0/1 背包：每件物品最多一次，容量倒序
for (i = 1; i <= n; ++i) {
    for (j = m; j >= w[i]; --j) {
        dp[j] = max(dp[j], dp[j - w[i]] + v[i]);
    }
}

// 完全背包：每件物品无限次，容量正序
for (i = 1; i <= n; ++i) {
    for (j = w[i]; j <= m; ++j) {
        dp[j] = max(dp[j], dp[j - w[i]] + v[i]);
    }
}
""")
    add_explain_block(
        doc,
        [
            ("dp[j]", "容量不超过 j 时能取得的最大价值。"),
            ("w[i]", "第 i 件物品的重量或花费。"),
            ("v[i]", "第 i 件物品的价值。"),
            ("m", "背包最大容量。"),
        ],
        [
            "外层循环枚举物品，表示逐个考虑是否使用当前物品。",
            "0/1 背包容量倒序，是为了让 dp[j-w[i]] 仍然来自上一轮物品，当前物品不会被重复用。",
            "完全背包容量正序，是为了允许 dp[j-w[i]] 已经使用过当前物品，从而实现无限次选择。",
        ],
        [
            "0/1 背包写成正序会把一件物品用多次。",
            "完全背包写成倒序会变成每件最多一次。",
            "恰好装满时，除 dp[0] 外其他状态要初始化为负无穷。",
        ],
    )
    add_classic_table(doc, [
        ("最大子段和", "线性 DP 处理连续区间最大和。", "全负数不能把答案初始化为 0。"),
        ("LIS / LNDS", "dp[i] 表示以 i 结尾的最优长度。", "严格上升用 <；不下降用 <=；答案是所有 dp[i] 最大值。"),
        ("LCS / 最长公共子串", "LCS 不要求连续；公共子串不等时 dp 清 0。", "字符串下标从 0 还是 1 要统一。"),
        ("编辑距离", "二维 DP 处理插入、删除、替换。", "dp[i][0]=i，dp[0][j]=j 是关键初始化。"),
        ("0/1 背包", "每件物品最多用一次，容量倒序。", "恰好装满要把不可达状态设为负无穷。"),
        ("完全背包", "每件物品可无限次使用，容量正序。", "方案数题和最大价值题的转移不同，不要混用 max 和 +=。"),
        ("区间 DP", "先枚举长度，再枚举左端点和断点。", "初始化求最小值时用 INF；注意 cost(l,r) 是否可预处理。"),
    ])


def add_graph_binary(doc):
    add_h1(doc, "四、图论与二分")
    add_h2(doc, "1. Dijkstra：单源最短路")
    add_code(doc, """
void dijkstra(int s) {
    for (i = 1; i <= n; ++i) {
        dista[i] = INF;
        vis[i] = 0;
    }
    dista[s] = 0;

    for (i = 1; i <= n; ++i) {
        int u = 0;
        for (j = 1; j <= n; ++j) {
            if (!vis[j] && (u == 0 || dista[j] < dista[u])) {
                u = j;
            }
        }
        if (u == 0 || dista[u] == INF) break;
        vis[u] = 1;

        for (j = 1; j <= n; ++j) {
            if (!vis[j] && g[u][j] < INF) {
                dista[j] = min(dista[j], dista[u] + g[u][j]);
            }
        }
    }
}
""")
    add_explain_block(
        doc,
        [
            ("dista[v]", "源点 s 到 v 的当前最短距离。"),
            ("vis[v]", "v 的最短路是否已经确定。"),
            ("u", "当前未确定点中 dista 最小的点。"),
            ("g[u][j]", "u 到 j 的边权；INF 表示没有边。"),
        ],
        [
            "先把所有距离设成 INF，只把源点 s 设为 0。",
            "每轮在未确定的点里选 dista 最小的 u。",
            "如果 u 不存在或 dista[u] 还是 INF，说明剩下的点都不可达，直接结束。",
            "确定 u 后，用 u 去松弛其他点：经过 u 更短，就更新 dista[j]。",
        ],
        [
            "Dijkstra 不能处理负权边。",
            "邻接矩阵版适合点数不大；点数很大时要考虑邻接表和堆优化。",
            "初始化 g[i][i] = 0，其他为 INF。",
            "如果 dista[u] + g[u][j] 可能超过 int，要把 dista 和边权改成 long long。",
        ],
    )

    add_h2(doc, "2. Prim：最小生成树")
    add_code(doc, """
int prim() {
    int ans = 0, cnt = 1;
    for (i = 1; i <= n; ++i) {
        low[i] = g[1][i];
        vis[i] = 0;
    }
    vis[1] = 1;

    for (i = 1; i <= n - 1; ++i) {
        int u = -1, best = INF;
        for (j = 1; j <= n; ++j) {
            if (!vis[j] && low[j] < best) {
                best = low[j];
                u = j;
            }
        }
        if (u == -1) break;
        vis[u] = 1;
        ans += best;
        ++cnt;

        for (j = 1; j <= n; ++j) {
            if (!vis[j] && g[u][j] < low[j]) {
                low[j] = g[u][j];
            }
        }
    }
    if (cnt != n) return -1;
    return ans;
}
""")
    add_explain_block(
        doc,
        [
            ("low[v]", "当前生成树连接到 v 的最小边权。"),
            ("vis[v]", "v 是否已经加入生成树。"),
            ("ans", "最小生成树边权和。"),
            ("cnt", "已经加入生成树的点数。"),
        ],
        [
            "每轮找一个 low 最小的未加入点 u，把它接进生成树。",
            "best 是接入 u 的最小边权，所以加入 ans。",
            "u 加入后，用 u 的边更新其他未加入点的 low。",
            "循环结束后 cnt 不等于 n，说明图不连通。",
        ],
        [
            "Prim 求的是最小生成树，不是最短路。",
            "无向图建边必须双向赋值。",
            "Dijkstra 的 dista 是从源点到各点；Prim 的 low 是树到各点的一条边。",
        ],
    )
    add_classic_table(doc, [
        ("单源最短路", "Dijkstra 求一个起点到所有点的最短距离。", "边权必须非负；不可达点保持 INF。"),
        ("多次最短路询问", "点数很小时可考虑 Floyd；只问一个源点就用 Dijkstra。", "Floyd 是 O(n^3)，n 大时不能硬套。"),
        ("最小生成树", "Prim 或 Kruskal 连接所有点且总边权最小。", "图不连通时没有生成树，要输出 impossible 或 -1。"),
        ("稠密图", "邻接矩阵 Dijkstra/Prim 可读性高，适合点数较小。", "矩阵初始化必须处理重边：取更小边权。"),
        ("稀疏图", "边少点多时可换链式前向星 + 堆优化。", "堆优化 Dijkstra 仍然不能处理负权边。"),
    ])

    add_h2(doc, "3. lower_bound / upper_bound")
    add_code(doc, """
// 头文件：#include <algorithm>
// 前提：区间必须已经按升序排好

lower_bound(first, last, x);  // 第一个 >= x 的位置
upper_bound(first, last, x);  // 第一个 > x 的位置
""")
    add_kv_table(doc, [
        ("返回值", "返回迭代器，不是直接返回下标，也不是直接返回元素值。"),
        ("数组转下标", "lower_bound(a, a+n, x) - a。"),
        ("vector 转下标", "lower_bound(v.begin(), v.end(), x) - v.begin()。"),
        ("找不到", "可能返回 last；数组中下标表现为 n，vector 中表现为 v.end()。"),
        ("核心记法", "lower 是 >=，upper 是 >。等于 x 的区间是 [lower, upper)。"),
    ])

    add_h3(doc, "数组与 vector 基本用法")
    add_code(doc, """
int a[8] = {1, 2, 2, 2, 4, 7, 9, 9};
int n = 8;
int x = 2;

int l = lower_bound(a, a + n, x) - a;  // 1
int r = upper_bound(a, a + n, x) - a;  // 4
int cnt = r - l;                       // 3

if (l < n && a[l] == x) {
    cout << "found" << endl;
}

vector<int> v;
int pos = lower_bound(v.begin(), v.end(), x) - v.begin();
""")

    add_h3(doc, "前驱、后继和插入位置")
    add_code(doc, """
// 第一个 >= x 的数，也就是 x 的后继候选
int p = lower_bound(a, a + n, x) - a;
if (p < n) cout << a[p] << endl;

// 最后一个 < x 的数
p = lower_bound(a, a + n, x) - a;
if (p > 0) cout << a[p - 1] << endl;

// 最后一个 <= x 的数
p = upper_bound(a, a + n, x) - a;
if (p > 0) cout << a[p - 1] << endl;

// 插入 x 后仍保持有序的位置，一般用 lower_bound
p = lower_bound(a, a + n, x) - a;
""")

    add_h3(doc, "pair 的 lower_bound")
    add_code(doc, """
pair<int, int> a[1005];
sort(a, a + n);  // 先按 first，再按 second

// 找第一个 first >= x 的 pair
pair<int, int> key = make_pair(x, -INF);
int pos = lower_bound(a, a + n, key) - a;

// 找第一个 first > x 的 pair
key = make_pair(x, INF);
pos = upper_bound(a, a + n, key) - a;
""")

    add_h3(doc, "set 与降序比较器")
    add_code(doc, """
set<int> st;
set<int>::iterator it;

it = st.lower_bound(x);  // set 中第一个 >= x
it = st.upper_bound(x);  // set 中第一个 > x

// 如果数组按降序排序，比较器也必须一致
sort(a, a + n, greater<int>());
int pos = lower_bound(a, a + n, x, greater<int>()) - a;
""")

    add_h3(doc, "手写版")
    add_code(doc, """
int my_lower_bound(int a[], int n, int x) {
    int l = 0, r = n;
    while (l < r) {
        int mid = l + (r - l) / 2;
        if (a[mid] >= x) r = mid;
        else l = mid + 1;
    }
    return l;
}

int my_upper_bound(int a[], int n, int x) {
    int l = 0, r = n;
    while (l < r) {
        int mid = l + (r - l) / 2;
        if (a[mid] > x) r = mid;
        else l = mid + 1;
    }
    return l;
}
""")
    add_kv_table(doc, [
        ("统计 x 出现次数", "`upper_bound - lower_bound`，如果 x 不存在结果自然是 0。"),
        ("判断 x 是否存在", "`l < n && a[l] == x`。只判断 `l < n` 不够。"),
        ("最后一个 < x", "`lower_bound` 的位置减 1，必须先判断 `pos > 0`。"),
        ("最后一个 <= x", "`upper_bound` 的位置减 1，必须先判断 `pos > 0`。"),
        ("第一个 >= x", "`lower_bound` 本身，必须判断 `pos < n`。"),
        ("第一个 > x", "`upper_bound` 本身，必须判断 `pos < n`。"),
        ("易错边界", "`a+n` 是右开结尾，不要写成 `a+n-1`。"),
        ("set 用法", "`set.lower_bound(x)` 和 `set.upper_bound(x)` 是成员函数，返回迭代器。"),
        ("降序序列", "语法可加 `greater<int>()`，但考试不熟时尽量转成升序。"),
    ])

    add_h2(doc, "4. 二分查找与二分答案")
    add_code(doc, """
// 找第一个满足条件的位置：false false true true
while (l < r) {
    int mid = l + (r - l) / 2;
    if (check(mid)) r = mid;
    else l = mid + 1;
}

// 找最后一个满足条件的位置：true true false false
while (l < r) {
    int mid = l + (r - l + 1) / 2;
    if (check(mid)) l = mid;
    else r = mid - 1;
}
""")
    add_explain_block(
        doc,
        [
            ("l / r", "当前答案可能存在的左右边界。"),
            ("mid", "本轮尝试的中间答案。"),
            ("check(mid)", "判断 mid 是否满足题目要求。"),
            ("单调性", "二分答案成立的根本条件。"),
        ],
        [
            "先判断题目是找最小可行答案，还是最大可行答案。",
            "找最小可行时，check(mid) 成功就收缩右边界 r = mid。",
            "找最大可行时，check(mid) 成功就保留 mid，令 l = mid。",
            "找最后一个满足时 mid 要偏右，写成 l + (r-l+1)/2，防止死循环。",
        ],
        [
            "没有单调性不能二分答案。",
            "check 函数要只回答能不能，不要顺手改乱全局状态。",
            "边界 low/high 要覆盖所有可能答案。",
        ],
    )

    add_h3(doc, "二分答案 check 示例：最小化最大段和")
    add_code(doc, """
int check(int limit) {
    int cnt = 1;
    int sum = 0;
    for (i = 1; i <= n; ++i) {
        if (a[i] > limit) return 0;
        if (sum + a[i] <= limit) {
            sum += a[i];
        } else {
            ++cnt;
            sum = a[i];
        }
    }
    return cnt <= k;
}
""")
    add_classic_table(doc, [
        ("查某个数", "普通二分或 binary_search 判断 x 是否存在。", "数组必须有序；找不到返回 -1 或 false。"),
        ("统计出现次数", "upper_bound - lower_bound。", "结果区间是 [l, r)，x 不存在时次数为 0。"),
        ("第一个满足", "false false true true，用 l=mid+1 / r=mid。", "r 常设为 n，表示可能不存在时返回 n。"),
        ("最后一个满足", "true true false false，mid 偏右。", "写 mid=(l+r+1)/2，防止 l=mid 死循环。"),
        ("最小化最大值", "check(x) 判断能否让答案不超过 x。", "low 至少取最大单项，high 可取总和或题目上界。"),
        ("最大化最小值", "check(x) 判断能否让答案至少为 x。", "check 越大越难成功，成功时移动左边界。"),
    ])


def add_summary(doc):
    add_h1(doc, "五、考前最后一页")
    add_h2(doc, "循环方向速记")
    add_kv_table(doc, [
        ("DFS", "for 枚举方向或邻边，合法就递归。"),
        ("BFS", "while 队列不空，每次出队并扩展下一层。"),
        ("0/1 背包", "物品外层，容量倒序。"),
        ("完全背包", "物品外层，容量正序。"),
        ("Dijkstra", "找最小 dista 的未确定点，再松弛。"),
        ("Prim", "找最小 low 的未入树点，再更新 low。"),
        ("二分最小答案", "check 成功，r = mid。"),
        ("二分最大答案", "check 成功，l = mid，mid 偏右。"),
    ])
    add_h2(doc, "最容易丢分的地方")
    add_bullets(doc, [
        "BFS 入队时就标记，别等出队才标记。",
        "sort 比较函数不要写 <=。",
        "LIS 严格上升用 <，最长不下降用 <=。",
        "0/1 背包和完全背包唯一看起来差一行，但循环方向决定本质。",
        "Dijkstra 不处理负权边。",
        "Prim 要判断图是否连通。",
        "lower_bound / upper_bound 必须在有序序列上使用。",
        "二分前先画出 check 的真假序列。",
    ])


def build():
    doc = Document()
    setup_styles(doc)
    add_title(doc)
    add_reference_intro(doc)
    add_search(doc)
    add_greedy_dp(doc)
    add_graph_binary(doc)
    add_summary(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
