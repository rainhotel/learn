import os
import re
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\moniC\project\learn\03-outputs\algorithm-final-review\algorithm-final-review-cpp98-detailed.docx"

BLACK = RGBColor(0, 0, 0)
PAGE_WIDTH_DXA = 9360


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_first_row_header(table):
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    tbl.insert(0, grid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mark_first_row_header(table)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    return p


def add_para(doc, text="", size=11, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    before = {1: 18, 2: 12, 3: 8}
    after = {1: 8, 2: 5, 3: 4}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_code(doc, code):
    for raw in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(raw)
        set_run_font(run, name="Consolas", size=8.2)
    add_para(doc, "", after=5)


def add_kv(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(label + "：")
    set_run_font(r1, size=10.5, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        p = clear_cell(table.rows[0].cells[i])
        r = p.add_run(header)
        set_run_font(r, size=9.2, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = clear_cell(cells[i])
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=9.0)
    add_para(doc, "", after=5)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = BLACK

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("算法期末复习讲义 | C++98")
    set_run_font(r, size=9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("原理、板子、典型例题与动态规划精讲")
    set_run_font(r, size=9)
    return doc


def add_common_problem_types(doc, items):
    add_heading(doc, "常见题型", 3)
    for item in items:
        add_bullet(doc, item)


def add_example_block(doc, title, statement, sample_in, sample_out, idea_steps, mistakes):
    add_heading(doc, "典型例题精讲：" + title, 2)
    add_kv(doc, "题意", statement)
    add_kv(doc, "样例输入", sample_in)
    add_kv(doc, "样例输出", sample_out)
    add_heading(doc, "思路拆解", 3)
    for step in idea_steps:
        add_number(doc, step)
    add_heading(doc, "易错点", 3)
    for item in mistakes:
        add_bullet(doc, item)


def build():
    doc = setup_doc()

    add_para(doc, "算法课期末复习讲义", size=25, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "全黑打印版：原理讲解 + C++98 板子 + 典型例题精讲", size=13, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "覆盖 DFS、BFS、二分、贪心、动态规划、Floyd、最小生成树。重点加强动态规划。", size=10.5, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "使用建议", 1)
    add_para(doc, "这份讲义适合按板块复习。每一节先读原理，再手写板子，最后跟着例题把状态、边界、循环顺序说出来。")
    add_para(doc, "期末复习不要只看代码。真正得分的关键是：你能不能解释为什么这样转移、为什么这样排序、为什么 BFS 第一次到达就是最短路。")

    add_heading(doc, "一页考点总览", 1)
    add_simple_table(
        doc,
        ["考点", "判断信号", "核心结构", "典型题"],
        [
            ["DFS", "搜完整片区域", "递归 + visited", "连通块计数"],
            ["BFS", "最少步数", "队列 + dist", "迷宫最短路"],
            ["二分", "有序或答案单调", "lower_bound / upper_bound", "统计出现次数"],
            ["贪心", "局部选择不吃亏", "排序 + 扫描", "活动安排"],
            ["最大子段和", "连续子段最大和", "dp[i] 表示以 i 结尾", "连续区间求和"],
            ["最长不降子序列", "子序列可跳选且允许相等", "O(n^2) 枚举前驱", "最长非递减序列"],
            ["0/1 背包", "每件物品最多选一次", "容量倒序", "容量内最大价值"],
            ["LCS", "两个序列找公共部分", "二维前缀 DP", "最长公共子序列"],
            ["Floyd", "任意两点最短路", "k,i,j 三重循环", "多源最短路"],
            ["Kruskal", "连通所有点且总权最小", "边排序 + 并查集", "最小生成树"],
        ],
        [1650, 2600, 2600, 2510],
    )

    add_heading(doc, "0. C++98 统一写法", 1)
    add_para(doc, "虽然很多 OJ 支持 bits/stdc++.h，但按 C++98 准备时建议使用标准头文件。这样考试环境更稳，代码也更容易被老师看懂。")
    add_code(doc, r"""
#include <iostream>
#include <cstdio>
#include <cstring>
#include <algorithm>
#include <queue>
#include <vector>
#include <utility>
using namespace std;
""")

    add_heading(doc, "1. DFS：深度优先搜索", 1)
    add_para(doc, "DFS 的直观理解是“一条路走到底，走不动再退回来”。从一个点出发，先访问一个相邻点，再从这个相邻点继续深入，直到没有新的点可以走，再回溯到上一个点换方向。")
    add_para(doc, "DFS 最适合处理“把一整片能到达的区域全部找出来”的问题。它不保证最短路，但很适合连通块、岛屿数量、区域染色、可达性判断。")
    add_common_problem_types(doc, [
        "网格中有多少个连通块或岛屿。",
        "从某个点出发能不能到达目标点。",
        "给一张无向图，求连通块数量。",
        "把一个区域里所有相同字符染成另一个字符。",
    ])
    add_heading(doc, "C++98 板子：网格连通块 DFS", 2)
    add_code(doc, r"""
const int MAXN = 105;
int n, m;
char g[MAXN][MAXN];
int vis[MAXN][MAXN];
int dx[4] = {-1, 1, 0, 0};
int dy[4] = {0, 0, -1, 1};

void dfs(int x, int y) {
    vis[x][y] = 1;
    for (int k = 0; k < 4; ++k) {
        int nx = x + dx[k];
        int ny = y + dy[k];
        if (nx < 0 || nx >= n || ny < 0 || ny >= m) continue;
        if (vis[nx][ny] || g[nx][ny] == '#') continue;
        dfs(nx, ny);
    }
}

int main() {
    cin >> n >> m;
    for (int i = 0; i < n; ++i) cin >> g[i];
    int ans = 0;
    for (int i = 0; i < n; ++i) {
        for (int j = 0; j < m; ++j) {
            if (!vis[i][j] && g[i][j] == '.') {
                ++ans;
                dfs(i, j);
            }
        }
    }
    cout << ans << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "网格连通块计数",
        "给 n x m 网格，'.' 表示空地，'#' 表示障碍。上下左右相邻的空地属于同一块，求空地连通块数量。",
        "4 5 / ..#.. / .#.#. / ##..# / ...##",
        "3",
        [
            "双重循环扫描所有格子，遇到尚未访问过的 '.'，说明发现了一个新的连通块。",
            "答案加 1，然后从这个格子开始 DFS，把这一整块所有 '.' 都标记为已访问。",
            "后面扫描到同一块里的其他格子时，因为已经访问过，就不会重复计数。",
            "样例中左上角是一块，右上角是一块，下方连接起来的是一块，所以答案为 3。",
        ],
        [
            "答案加 1 的位置在 DFS 启动前，不是在 DFS 每访问一个格子时。",
            "越界判断必须在访问 g[nx][ny] 或 vis[nx][ny] 之前。",
            "如果题目说八方向连通，需要把方向数组改成 8 个方向。",
        ],
    )

    add_heading(doc, "2. BFS：广度优先搜索", 1)
    add_para(doc, "BFS 的直观理解是“一层一层向外扩散”。起点是第 0 层，离起点一步的点是第 1 层，离起点两步的点是第 2 层。队列保证了先到达的近距离点先被处理。")
    add_para(doc, "当每条边或每次移动的代价都一样时，BFS 第一次到达某个点时的距离就是最短距离。所以迷宫最少步数、无权图最短路、最少操作次数，优先考虑 BFS。")
    add_common_problem_types(doc, [
        "迷宫从 S 到 T 的最少步数。",
        "无权图从 1 号点到所有点的最短距离。",
        "最少几次操作能把一个状态变成目标状态。",
        "多起点同时扩散，例如多个火源或多个起点。",
    ])
    add_heading(doc, "C++98 板子：网格 BFS 最短路", 2)
    add_code(doc, r"""
const int MAXN = 105;
int n, m;
char g[MAXN][MAXN];
int dista[MAXN][MAXN];
int dx[4] = {-1, 1, 0, 0};
int dy[4] = {0, 0, -1, 1};

int bfs(int sx, int sy, int tx, int ty) {
    queue< pair<int, int> > q;
    memset(dista, -1, sizeof(dista));
    dista[sx][sy] = 0;
    q.push(make_pair(sx, sy));

    while (!q.empty()) {
        pair<int, int> cur = q.front();
        q.pop();
        int x = cur.first, y = cur.second;
        if (x == tx && y == ty) return dista[x][y];

        for (int k = 0; k < 4; ++k) {
            int nx = x + dx[k], ny = y + dy[k];
            if (nx < 0 || nx >= n || ny < 0 || ny >= m) continue;
            if (g[nx][ny] == '#' || dista[nx][ny] != -1) continue;
            dista[nx][ny] = dista[x][y] + 1;
            q.push(make_pair(nx, ny));
        }
    }
    return -1;
}
""")
    add_example_block(
        doc,
        "迷宫最短步数",
        "给 4 x 4 迷宫，S 是起点，T 是终点，# 是墙。每次上下左右走一步，求最少步数。",
        "4 4 / S... / ##.. / ...# / ..T.",
        "5",
        [
            "题目问最少步数，每一步代价都是 1，所以用 BFS。",
            "把 S 的距离设为 0 并入队。每弹出一个格子，就尝试扩展上下左右四个方向。",
            "第一次走到 T 时返回 dist[T]。样例中路径可为 S -> 右 -> 右 -> 下 -> 下 -> 下，到 T 共 5 步。",
            "如果队列空了还没到 T，说明不可达，返回 -1。",
        ],
        [
            "BFS 入队时就要设置 dist，避免同一个点重复入队。",
            "dist 数组既保存距离，也可以当 visited 使用。",
            "有边权时普通 BFS 不再适用。",
        ],
    )

    add_heading(doc, "3. 二分查找", 1)
    add_para(doc, "二分查找的本质是利用单调性，每次用中点排除一半范围。它不只可以在有序数组中找数，也可以在答案空间中找“最小可行值”或“最大可行值”。")
    add_para(doc, "期末最常见的是 lower_bound 和 upper_bound。lower_bound 找第一个 >= x 的位置，upper_bound 找第一个 > x 的位置，两者相减就是 x 的出现次数。")
    add_common_problem_types(doc, [
        "在有序数组中找某个数第一次出现的位置。",
        "统计有序数组中某个数出现了多少次。",
        "找第一个大于等于 x 或第一个大于 x 的位置。",
        "答案具有单调性时，二分答案。",
    ])
    add_heading(doc, "C++98 板子：STL lower_bound / upper_bound", 2)
    add_code(doc, r"""
int a[1005];

int main() {
    int n, x;
    cin >> n >> x;
    for (int i = 0; i < n; ++i) cin >> a[i];
    sort(a, a + n);

    int l = lower_bound(a, a + n, x) - a; // first >= x
    int r = upper_bound(a, a + n, x) - a; // first > x
    cout << r - l << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "统计数字出现次数",
        "给数组 1 2 2 2 4 7 9 和 x = 2，求 2 出现了多少次。",
        "7 2 / 1 2 2 2 4 7 9",
        "3",
        [
            "数组已经有序，可以直接二分。",
            "lower_bound 找第一个 >= 2 的位置，下标为 1。",
            "upper_bound 找第一个 > 2 的位置，下标为 4。",
            "区间 [1,4) 内一共有 4 - 1 = 3 个元素，所以答案是 3。",
        ],
        [
            "lower_bound 和 upper_bound 返回的是迭代器，数组里要减去 a 才是下标。",
            "如果数组无序，必须先排序，否则二分没有意义。",
            "如果 x 不存在，两个边界相等，出现次数为 0。",
        ],
    )

    add_heading(doc, "4. 贪心", 1)
    add_para(doc, "贪心不是一个固定公式，而是一种策略：每一步都做当前看来最好的选择，并且这个选择不会影响整体最优。考试里最常见的贪心题通常是“先排序，再扫描”。")
    add_para(doc, "活动安排是最典型的例子。为了选尽量多的活动，我们每次选结束最早的活动，因为它会给后续活动留下最多时间。")
    add_common_problem_types(doc, [
        "最多选择多少个互不重叠区间。",
        "按结束时间、截止时间、权值等排序后扫描。",
        "用一个变量维护当前最优边界。",
        "局部最优可以用交换思想解释。",
    ])
    add_heading(doc, "C++98 板子：活动选择", 2)
    add_code(doc, r"""
struct Node {
    int l, r;
} a[1005];

bool cmp(const Node &x, const Node &y) {
    if (x.r != y.r) return x.r < y.r;
    return x.l < y.l;
}

int main() {
    int n;
    cin >> n;
    for (int i = 0; i < n; ++i) cin >> a[i].l >> a[i].r;
    sort(a, a + n, cmp);

    int ans = 0;
    int last = -1000000000;
    for (int i = 0; i < n; ++i) {
        if (a[i].l >= last) {
            ++ans;
            last = a[i].r;
        }
    }
    cout << ans << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "最大不重叠区间数",
        "有区间 [1,3] [2,4] [3,5] [0,6] [5,7] [8,9]，求最多能选几个互不重叠区间。端点相接允许。",
        "6 / 1 3 / 2 4 / 3 5 / 0 6 / 5 7 / 8 9",
        "4",
        [
            "按右端点排序后依次考虑，优先选结束早的区间。",
            "先选 [1,3]，因为它最早结束；接着 [3,5] 可以接上；再选 [5,7]；最后选 [8,9]。",
            "共选 4 个区间。",
            "如果选 [0,6]，它占用太长，会挡住更多区间，所以不是好选择。",
        ],
        [
            "不要按左端点排序，开始早不代表结束早。",
            "a[i].l >= last 还是 a[i].l > last 要根据题目是否允许端点相接决定。",
            "贪心题要能讲清排序依据，不要只背代码。",
        ],
    )

    add_heading(doc, "5. 动态规划总方法", 1)
    add_para(doc, "动态规划最怕“只背公式”。真正稳定的写法是先说清 dp 的含义。dp 的含义一旦清楚，转移、初值、遍历顺序都会自然很多。")
    add_heading(doc, "DP 四问", 2)
    for item in [
        "状态是什么：dp[i] 或 dp[i][j] 具体表示什么。",
        "从哪里来：当前状态由哪些更小状态转移。",
        "初值是什么：最小规模问题的答案是什么。",
        "顺序是什么：怎样遍历才能保证依赖已经算好。",
    ]:
        add_number(doc, item)

    add_heading(doc, "5.1 最大子段和", 1)
    add_para(doc, "最大子段和要求连续，所以每个位置 i 只有两种选择：要么把 a[i] 接在前面的连续段后面，要么从 a[i] 重新开始。")
    add_kv(doc, "状态定义", "dp[i] 表示必须以 a[i] 结尾的最大连续子段和。")
    add_kv(doc, "转移方程", "dp[i] = max(a[i], dp[i - 1] + a[i])。")
    add_kv(doc, "答案", "max(dp[i])，因为最大子段不一定以最后一个元素结尾。")
    add_common_problem_types(doc, [
        "连续子段最大和。",
        "连续子段最大和并输出左右端点。",
        "环形数组最大子段和的基础版本。",
    ])
    add_heading(doc, "C++98 板子：最大子段和", 2)
    add_code(doc, r"""
int a[1005], dp[1005];

int main() {
    int n;
    cin >> n;
    for (int i = 1; i <= n; ++i) cin >> a[i];

    dp[1] = a[1];
    int ans = dp[1];
    for (int i = 2; i <= n; ++i) {
        dp[i] = max(a[i], dp[i - 1] + a[i]);
        ans = max(ans, dp[i]);
    }
    cout << ans << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "连续子段最大和",
        "给序列 -2 3 -1 5 -6 2，求最大连续子段和。",
        "6 / -2 3 -1 5 -6 2",
        "7",
        [
            "dp[1] = -2。",
            "到 3 时，接上前面会变成 1，不如从 3 重新开始，所以 dp[2] = 3。",
            "到 -1 时，接上 3 得到 2，比单独 -1 好，所以 dp[3] = 2。",
            "到 5 时，接上前面得到 7，所以 dp[4] = 7。",
            "最后最大值是 7，对应子段 3, -1, 5。",
        ],
        [
            "全是负数时答案不能写成 0，应是最大的负数。",
            "连续子段不能跳选。",
            "ans 要在每个 i 更新，不是只输出 dp[n]。",
        ],
    )
    add_simple_table(
        doc,
        ["i", "a[i]", "dp[i]", "解释"],
        [
            ["1", "-2", "-2", "只能选 -2"],
            ["2", "3", "3", "重新从 3 开始更好"],
            ["3", "-1", "2", "3 + (-1)"],
            ["4", "5", "7", "2 + 5"],
            ["5", "-6", "1", "7 + (-6)"],
            ["6", "2", "3", "1 + 2"],
        ],
        [900, 1200, 1200, 6060],
    )

    add_heading(doc, "5.2 最长不降子序列", 1)
    add_para(doc, "子序列可以跳着选，但不能改变原来的相对顺序。不降表示后一个数可以大于或等于前一个数，所以比较条件是 <=。")
    add_kv(doc, "状态定义", "dp[i] 表示以 a[i] 结尾的最长不降子序列长度。")
    add_kv(doc, "转移方程", "如果 j < i 且 a[j] <= a[i]，则 dp[i] = max(dp[i], dp[j] + 1)。")
    add_kv(doc, "初值", "每个 dp[i] 初始为 1，因为单个元素本身就是一个长度为 1 的子序列。")
    add_common_problem_types(doc, [
        "最长不降子序列长度。",
        "最长严格上升子序列，把 <= 改成 <。",
        "要求输出一种方案时，需要记录 pre[i] 前驱。",
    ])
    add_heading(doc, "C++98 板子：最长不降子序列 O(n^2)", 2)
    add_code(doc, r"""
int a[1005], dp[1005];

int main() {
    int n;
    cin >> n;
    for (int i = 1; i <= n; ++i) {
        cin >> a[i];
        dp[i] = 1;
    }

    int ans = 1;
    for (int i = 1; i <= n; ++i) {
        for (int j = 1; j < i; ++j) {
            if (a[j] <= a[i]) {
                dp[i] = max(dp[i], dp[j] + 1);
            }
        }
        ans = max(ans, dp[i]);
    }
    cout << ans << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "最长不降子序列",
        "给序列 3 1 2 2 4 3，求最长不降子序列长度。",
        "6 / 3 1 2 2 4 3",
        "4",
        [
            "以第 1 个数 3 结尾，只能选它自己，所以 dp[1] = 1。",
            "以 2 结尾时，可以接在 1 后面，所以长度变成 2。",
            "第二个 2 可以接在前一个 2 后面，因为是不降，允许相等，所以长度变成 3。",
            "4 可以接在 1,2,2 后面，长度为 4；最后的 3 也可以接成 1,2,2,3，长度为 4。",
            "答案是 4。",
        ],
        [
            "不降是 <=，严格上升才是 <。",
            "答案不是 dp[n]，最长序列不一定以最后一个元素结尾。",
            "不要把子序列看成子段，子序列可以跳过元素。",
        ],
    )
    add_simple_table(
        doc,
        ["i", "a[i]", "dp[i]", "可形成的代表序列"],
        [
            ["1", "3", "1", "3"],
            ["2", "1", "1", "1"],
            ["3", "2", "2", "1,2"],
            ["4", "2", "3", "1,2,2"],
            ["5", "4", "4", "1,2,2,4"],
            ["6", "3", "4", "1,2,2,3"],
        ],
        [900, 1200, 1200, 6060],
    )

    add_heading(doc, "5.3 0/1 背包", 1)
    add_para(doc, "0/1 背包的核心是每件物品只能选一次。对每个物品，你只有两个选择：选它或不选它。")
    add_kv(doc, "状态定义", "dp[j] 表示在当前已经考虑过的物品中，容量不超过 j 时能得到的最大价值。")
    add_kv(doc, "转移方程", "dp[j] = max(dp[j], dp[j - w[i]] + v[i])。")
    add_kv(doc, "循环方向", "容量 j 必须从大到小枚举，防止同一件物品在一轮中被重复使用。")
    add_common_problem_types(doc, [
        "每件物品最多选一次，容量内最大价值。",
        "恰好装满背包的最大价值。",
        "问是否能凑出某个容量或价值。",
        "二维写法 dp[i][j] 转一维优化。",
    ])
    add_heading(doc, "C++98 板子：一维 0/1 背包", 2)
    add_code(doc, r"""
int w[1005], v[1005], dp[10005];

int main() {
    int n, m;
    cin >> n >> m;
    for (int i = 1; i <= n; ++i) cin >> w[i] >> v[i];

    for (int i = 1; i <= n; ++i) {
        for (int j = m; j >= w[i]; --j) {
            dp[j] = max(dp[j], dp[j - w[i]] + v[i]);
        }
    }
    cout << dp[m] << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "容量内最大价值",
        "背包容量为 7，有 4 件物品：(重量,价值) 分别是 (2,3), (3,4), (4,5), (5,8)。每件最多选一次，求最大价值。",
        "4 7 / 2 3 / 3 4 / 4 5 / 5 8",
        "11",
        [
            "先考虑第 1 件，容量至少为 2 时都能得到价值 3。",
            "考虑第 2 件后，容量 5 可以选第 1 和第 2 件，价值 7。",
            "考虑第 3 件后，容量 7 可以选第 2 和第 3 件，价值 9。",
            "考虑第 4 件后，容量 7 可以选第 1 和第 4 件，价值 3 + 8 = 11。",
            "所以最大价值是 11。",
        ],
        [
            "0/1 背包必须倒序枚举容量。",
            "如果正序枚举，当前物品会被重复使用，变成完全背包。",
            "如果题目要求恰好装满，初始化要特别处理不可达状态。",
        ],
    )
    add_simple_table(
        doc,
        ["阶段", "dp[0]", "dp[1]", "dp[2]", "dp[3]", "dp[4]", "dp[5]", "dp[6]", "dp[7]"],
        [
            ["初始", "0", "0", "0", "0", "0", "0", "0", "0"],
            ["物品1", "0", "0", "3", "3", "3", "3", "3", "3"],
            ["物品2", "0", "0", "3", "4", "4", "7", "7", "7"],
            ["物品3", "0", "0", "3", "4", "5", "7", "8", "9"],
            ["物品4", "0", "0", "3", "4", "5", "8", "8", "11"],
        ],
        [1440, 990, 990, 990, 990, 990, 990, 990, 990],
    )

    add_heading(doc, "5.4 最长公共子序列 LCS", 1)
    add_para(doc, "最长公共子序列比较两个序列中共同出现且相对顺序一致的最长部分。它不要求连续，所以和最长公共子串不同。")
    add_kv(doc, "状态定义", "dp[i][j] 表示 a 的前 i 个字符和 b 的前 j 个字符的 LCS 长度。")
    add_kv(doc, "转移方程", "如果 a[i] == b[j]，dp[i][j] = dp[i - 1][j - 1] + 1；否则取 max(dp[i - 1][j], dp[i][j - 1])。")
    add_kv(doc, "初值", "dp[0][j] 和 dp[i][0] 都是 0，因为空串和任何串的公共子序列长度为 0。")
    add_common_problem_types(doc, [
        "两个字符串的最长公共子序列长度。",
        "两个序列的相似程度。",
        "最少删除多少字符能让两个串相同。",
        "如果要求连续，要改成最长公共子串做法。",
    ])
    add_heading(doc, "C++98 板子：LCS", 2)
    add_code(doc, r"""
char a[1005], b[1005];
int dp[1005][1005];

int main() {
    cin >> (a + 1) >> (b + 1);
    int n = strlen(a + 1);
    int m = strlen(b + 1);

    for (int i = 1; i <= n; ++i) {
        for (int j = 1; j <= m; ++j) {
            if (a[i] == b[j]) {
                dp[i][j] = dp[i - 1][j - 1] + 1;
            } else {
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1]);
            }
        }
    }
    cout << dp[n][m] << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "两个字符串的最长公共子序列",
        "给 a = ABCBDAB，b = BDCABA，求最长公共子序列长度。",
        "ABCBDAB / BDCABA",
        "4",
        [
            "dp[i][j] 只看两个前缀，所以问题规模会逐渐变小。",
            "如果当前字符相等，就让这两个字符配对，答案来自左上角加 1。",
            "如果当前字符不等，就尝试丢掉 a 的末尾或 b 的末尾，取更大的结果。",
            "最终 dp[7][6] = 4，代表可以找到长度为 4 的公共子序列，例如 BCBA 或 BDAB。",
        ],
        [
            "LCS 是子序列，不要求连续。",
            "最长公共子串要求连续，字符不等时通常要变成 0。",
            "从 1 开始存字符串，可以让第 0 行和第 0 列自然表示空串。",
        ],
    )
    add_simple_table(
        doc,
        ["a\\b", "B", "D", "C", "A", "B", "A"],
        [
            ["A", "0", "0", "0", "1", "1", "1"],
            ["B", "1", "1", "1", "1", "2", "2"],
            ["C", "1", "1", "2", "2", "2", "2"],
            ["B", "1", "1", "2", "2", "3", "3"],
            ["D", "1", "2", "2", "2", "3", "3"],
            ["A", "1", "2", "2", "3", "3", "4"],
            ["B", "1", "2", "2", "3", "4", "4"],
        ],
        [1200, 1300, 1300, 1300, 1300, 1300, 1300],
    )

    add_heading(doc, "6. Floyd：全源最短路", 1)
    add_para(doc, "Floyd 用来求任意两点之间的最短路。它的核心想法是：枚举一个中转点 k，看看 i 到 j 是否可以通过 i -> k -> j 变得更短。")
    add_para(doc, "循环顺序必须是 k 在最外层，因为第 k 轮的含义是：允许使用 1 到 k 这些点作为中转点。")
    add_common_problem_types(doc, [
        "任意两点最短路。",
        "点数不大，但询问很多。",
        "有重边时取最短边。",
        "判断两点是否可达。",
    ])
    add_heading(doc, "C++98 板子：Floyd", 2)
    add_code(doc, r"""
const int INF = 0x3f3f3f3f;
int d[105][105];

int main() {
    int n, m;
    cin >> n >> m;
    for (int i = 1; i <= n; ++i)
        for (int j = 1; j <= n; ++j)
            d[i][j] = (i == j ? 0 : INF);

    for (int i = 1; i <= m; ++i) {
        int u, v, w;
        cin >> u >> v >> w;
        d[u][v] = min(d[u][v], w);
        d[v][u] = min(d[v][u], w);
    }

    for (int k = 1; k <= n; ++k)
        for (int i = 1; i <= n; ++i)
            for (int j = 1; j <= n; ++j)
                if (d[i][k] < INF && d[k][j] < INF)
                    d[i][j] = min(d[i][j], d[i][k] + d[k][j]);

    cout << d[1][n] << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "多源最短路查询",
        "有 4 个点，边为 1-2(5), 2-3(2), 1-3(10), 3-4(1), 2-4(9)，求 1 到 4 的最短路。",
        "4 5 / 1 2 5 / 2 3 2 / 1 3 10 / 3 4 1 / 2 4 9",
        "8",
        [
            "直接边 1 到 3 是 10，但通过 1 -> 2 -> 3 可以变成 7。",
            "3 到 4 是 1，所以 1 -> 2 -> 3 -> 4 的总距离是 5 + 2 + 1 = 8。",
            "Floyd 会枚举所有中转点，自动把 d[1][3] 和 d[1][4] 更新成更短距离。",
        ],
        [
            "k 必须放在最外层。",
            "INF 要足够大，且更新前最好判断两段路都不是 INF。",
            "无向图建边时要同时更新 d[u][v] 和 d[v][u]。",
        ],
    )

    add_heading(doc, "7. Kruskal：最小生成树", 1)
    add_para(doc, "最小生成树是在无向带权图中选 n - 1 条边，把所有点连通，并让总边权最小。Kruskal 的做法是把所有边按权值从小到大排序，能加入就加入。")
    add_para(doc, "能加入的判断依赖并查集。如果一条边的两个端点已经在同一个集合里，加入它会形成环，不能选。")
    add_common_problem_types(doc, [
        "把 n 个点全部连起来，总代价最小。",
        "修路、铺网线、连接城市。",
        "判断图是否能形成生成树。",
        "输出最小生成树总权值。",
    ])
    add_heading(doc, "C++98 板子：Kruskal + 并查集", 2)
    add_code(doc, r"""
struct Edge {
    int u, v, w;
} e[5005];

int fa[1005];

int find_set(int x) {
    if (fa[x] == x) return x;
    fa[x] = find_set(fa[x]);
    return fa[x];
}

bool cmp(const Edge &a, const Edge &b) {
    return a.w < b.w;
}

int main() {
    int n, m;
    cin >> n >> m;
    for (int i = 1; i <= m; ++i) cin >> e[i].u >> e[i].v >> e[i].w;
    for (int i = 1; i <= n; ++i) fa[i] = i;
    sort(e + 1, e + m + 1, cmp);

    int ans = 0, cnt = 0;
    for (int i = 1; i <= m; ++i) {
        int fu = find_set(e[i].u);
        int fv = find_set(e[i].v);
        if (fu != fv) {
            fa[fu] = fv;
            ans += e[i].w;
            ++cnt;
        }
    }

    if (cnt == n - 1) cout << ans << endl;
    else cout << "impossible" << endl;
    return 0;
}
""")
    add_example_block(
        doc,
        "最小生成树总权值",
        "有 4 个点，边为 1-2(1), 2-3(2), 1-3(4), 3-4(3), 2-4(5)，求最小生成树总权值。",
        "4 5 / 1 2 1 / 2 3 2 / 1 3 4 / 3 4 3 / 2 4 5",
        "6",
        [
            "边排序后依次为 1-2(1), 2-3(2), 3-4(3), 1-3(4), 2-4(5)。",
            "先选 1-2，再选 2-3，此时 1,2,3 已经连通。",
            "再选 3-4，四个点全部连通，选了 n - 1 = 3 条边。",
            "总权值为 1 + 2 + 3 = 6。",
        ],
        [
            "必须先初始化 fa[i] = i。",
            "如果某条边两端已经在同一集合，加入会成环，不能选。",
            "最后 cnt 不等于 n - 1 时，说明图不连通。",
        ],
    )

    add_heading(doc, "考前最后检查清单", 1)
    for item in [
        "DFS：是否正确标记 visited，是否重复统计连通块。",
        "BFS：是否在入队时设置 dist，是否第一次到达就返回最短路。",
        "二分：数组是否有序，lower_bound 与 upper_bound 是否分清。",
        "贪心：排序依据是否正确，端点相接时比较符号是否正确。",
        "最大子段和：全负数时初值是否正确。",
        "最长不降子序列：不降用 <=，严格上升用 <。",
        "0/1 背包：容量倒序，正序是完全背包。",
        "LCS：子序列不要求连续，子串才要求连续。",
        "Floyd：k 在最外层，重边取最小，INF 不能太小。",
        "Kruskal：并查集初始化，最后检查是否选满 n - 1 条边。",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    force_docx_text_black(OUT)


def force_docx_text_black(path):
    tmp_dir = tempfile.mkdtemp(prefix="docx_black_")
    tmp_path = os.path.join(tmp_dir, "patched.docx")
    color_re = re.compile(rb'(w:color\b[^>]*\bw:val=")([^"]+)(")')
    fill_re = re.compile(rb'(w:highlight\b[^>]*\bw:val=")([^"]+)(")')

    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                data = color_re.sub(rb'\g<1>000000\3', data)
                # Remove text highlighting if a built-in style carried it.
                data = fill_re.sub(rb'\g<1>none\3', data)
            zout.writestr(item, data)

    shutil.copyfile(tmp_path, path)
    shutil.rmtree(tmp_dir)


if __name__ == "__main__":
    build()
