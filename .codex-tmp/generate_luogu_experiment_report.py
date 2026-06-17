from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\moniC\project\learn")
OUT_DIR = ROOT / "03-outputs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "洛谷算法实验报告.docx"


def set_run_font(run, font_name="Microsoft YaHei", size=11, bold=False, italic=False, color="000000"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, before=0, after=6, line=1.15, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="D9D9D9", size="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_layout(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    table.columns[0].width = Cm(widths_cm[0])
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            set_cell_margins(row.cells[idx])
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total * 567)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)

for style_name, size, bold in [("Title", 20, True), ("Heading 1", 15, True), ("Heading 2", 12.5, True), ("Heading 3", 11, True)]:
    style = styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(title, before=0, after=4, line=1.0)
run = title.add_run("洛谷算法实验报告")
set_run_font(run, size=20, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(sub, before=0, after=10, line=1.0)
run = sub.add_run("题目：P2137、P2105、P2039、P1913、P2141")
set_run_font(run, size=11, color="555555")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(meta, before=0, after=18, line=1.0)
run = meta.add_run("实验说明：本报告整理了 5 道典型算法题的核心思路、实现方法与复杂度分析。")
set_run_font(run, size=10.5, color="666666")

doc.add_heading("一、实验目的", level=1)
p = doc.add_paragraph()
set_paragraph_format(p)
run = p.add_run("通过对多道基础算法题的分析与实现，掌握贪心、双指针、动态规划/性质分析、图最短路等常见方法，"
                "并训练从题目条件中抽象出关键限制、选择合适算法并给出复杂度分析的能力。")
set_run_font(run)

doc.add_heading("二、题目总览", level=1)
table = doc.add_table(rows=1, cols=4)
set_table_layout(table, [2.4, 3.9, 2.05, 1.65])
set_table_borders(table)
hdr = table.rows[0].cells
headers = ["题目", "核心思路", "时间复杂度", "结论"]
for cell, text in zip(hdr, headers):
    cell.text = ""
    shade_cell(cell, "EAF0F6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=0, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True)
for row in [
    ("P2137 乘2与乘3", "判断 b/a 是否仅含 2 和 3 的质因子", "O(log b)", "Yes/No"),
    ("P2105 收集樱花", "正数数组上的滑动窗口", "O(n)", "最大不超 V 的和"),
    ("P2039 饭搭子", "排序 + 双指针贪心配对", "O(n log n)", "最少桌数"),
    ("P1913 最长相邻特别子序列", "按奇偶性统计连续段数", "O(n)", "最长交替子序列长度"),
    ("P2141 星际旅行家2", "Floyd 预处理最短路后累加", "O(n^3)", "最少传送次数"),
]:
    cells = table.add_row().cells
    for idx, text in enumerate(row):
        cells[idx].text = ""
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p, before=0, after=0, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=10.3)
for row in table.rows:
    for cell in row.cells:
        set_cell_margins(cell)

def add_problem(title_text, analysis, steps, code, answer, complexity):
    doc.add_heading(title_text, level=1)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=4)
    r = p.add_run("题目分析：")
    set_run_font(r, bold=True)
    r = p.add_run(analysis)
    set_run_font(r)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=4)
    r = p.add_run("解题步骤：")
    set_run_font(r, bold=True)

    for item in steps:
        bp = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(bp, after=2, line=1.12)
        rr = bp.add_run(item)
        set_run_font(rr)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=4)
    r = p.add_run("参考实现：")
    set_run_font(r, bold=True)

    for line in code.strip().splitlines():
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(0.6)
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.0
        rr = cp.add_run(line)
        set_run_font(rr, font_name="Consolas", size=9.5)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=4)
    r = p.add_run("结论：")
    set_run_font(r, bold=True)
    r = p.add_run(answer)
    set_run_font(r)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=8)
    r = p.add_run("复杂度：")
    set_run_font(r, bold=True)
    r = p.add_run(complexity)
    set_run_font(r)


add_problem(
    "三、P2137 乘2与乘3",
    "每个 case 只需判断能否从 a 经过若干次乘 2 和乘 3 得到 b。由于乘法只会改变因子 2 和 3，"
    "所以只要 b 能被 a 整除，并且商中不含除 2、3 以外的质因子，就可以到达。",
    [
        "先判断 b 是否能被 a 整除，若不能直接输出 No。",
        "计算 x = b / a，持续把 x 中的 2 和 3 除掉。",
        "若最后 x = 1，说明商只含 2 和 3，输出 Yes；否则输出 No。",
    ],
    """
int x = b / a;
while (x % 2 == 0) x /= 2;
while (x % 3 == 0) x /= 3;
cout << (x == 1 ? "Yes" : "No");
""",
    "该方法只需检查商的质因子构成，能准确判断是否可由乘 2 和乘 3 得到目标数字。",
    "每组数据 O(log b)，共 n 组时总复杂度 O(n log b)。",
)

add_problem(
    "四、P2105 收集樱花",
    "题目要求选取一段连续樱花树，使得樱花数量之和不超过 V 且尽可能大。由于 a_i 都是正整数，"
    "窗口右端扩展时和单调增加，因此可以使用双指针维护一个滑动窗口。",
    [
        "维护左指针 l、右指针 r 和当前窗口和 sum。",
        "不断向右扩展 r，把 a[r] 加入窗口。",
        "如果 sum > V，就向右移动 l，直到 sum <= V。",
        "每次更新答案 ans = max(ans, sum)。",
    ],
    """
int l = 0, sum = 0, ans = 0;
for (int r = 0; r < n; ++r) {
    sum += a[r];
    while (sum > V) sum -= a[l++];
    ans = max(ans, sum);
}
cout << ans;
""",
    "滑动窗口始终保持合法区间，并且只会向右移动，能够在线性时间内得到不超过 V 的最大连续和。",
    "每个元素最多进出窗口一次，时间复杂度 O(n)，空间复杂度 O(1)。",
)

add_problem(
    "五、P2039 饭搭子",
    "每张桌子最多坐两人，目标是最少使用桌子。典型贪心策略是把饭量最大的学生尽量和最小的学生配对；"
    "如果两者之和都超过 W，那么最大的学生只能单独坐一桌。",
    [
        "先把所有学生饭量从小到大排序。",
        "用双指针分别指向最小和最大饭量学生。",
        "若二者之和不超过 W，就让他们共用一桌，左右指针同时移动。",
        "否则最大的学生单独占一桌，只移动右指针。",
        "每处理掉一桌，桌子数加 1。",
    ],
    """
sort(f.begin(), f.end());
int l = 0, r = n - 1, ans = 0;
while (l <= r) {
    if (f[l] + f[r] <= W) ++l, --r;
    else --r;
    ++ans;
}
cout << ans;
""",
    "排序后使用双指针，始终优先让最重的人尽量与最轻的人配对，得到的桌子数最少。",
    "排序 O(n log n)，双指针 O(n)，总复杂度 O(n log n)。",
)

add_problem(
    "六、P1913 最长相邻特别子序列",
    "题目条件为 (b_i + b_{i+1})^2 是奇数。平方的奇偶性与原数相同，因此该条件等价于 b_i + b_{i+1} 为奇数，"
    "也就是相邻两个数一个奇、一个偶。于是问题转化为：在原序列中找一个最长的奇偶交替子序列。",
    [
        "把每个数只看成奇数或偶数两种状态。",
        "依次扫描序列，统计奇偶性发生变化的次数。",
        "最长子序列长度就是奇偶连续段的个数。",
    ],
    """
int ans = 1;
for (int i = 1; i < n; ++i) {
    if ((a[i] & 1) != (a[i - 1] & 1)) ++ans;
}
cout << ans;
""",
    "因为子序列中不能出现相邻同奇或同偶，最优做法就是从每个连续奇偶段中取 1 个元素，从而形成最长交替序列。",
    "只需一次线性扫描，时间复杂度 O(n)，空间复杂度 O(1)。",
)

add_problem(
    "七、P2141 星际旅行家2",
    "需要按固定顺序访问 1,2,3,...,n 并最终回到 1。每一段相邻目标星球之间，都应该走图上的最短路径，因此总传送次数就是这些最短路长度之和。"
    "由于 n <= 100，且输入为邻接矩阵，Floyd-Warshall 是最稳妥的做法。",
    [
        "把 9999 视为无穷大 INF，建立邻接矩阵 dist。",
        "用 Floyd-Warshall 预处理任意两点之间的最短路径。",
        "累加 dist[1][2]、dist[2][3]、...、dist[n][1]。",
    ],
    """
const int INF = 1e9;
for (int k = 1; k <= n; ++k)
  for (int i = 1; i <= n; ++i)
    for (int j = 1; j <= n; ++j)
      dist[i][j] = min(dist[i][j], dist[i][k] + dist[k][j]);

int ans = 0;
for (int i = 1; i < n; ++i) ans += dist[i][i + 1];
ans += dist[n][1];
cout << ans;
""",
    "最小传送次数等于固定访问顺序下各段最短路长度之和，Floyd 可以一次性求出所有点对最短路。",
    "Floyd 复杂度 O(n^3)，额外空间 O(n^2)。",
)

doc.add_heading("八、实验总结", level=1)
p = doc.add_paragraph()
set_paragraph_format(p, after=4)
run = p.add_run(
    "本次实验覆盖了性质判断、双指针、贪心、序列分析和最短路五类典型算法。"
    "其中，P2137 强调对因子结构的快速判断；P2105 展示了正整数数组上滑动窗口的适用条件；"
    "P2039 体现了“最重 + 最轻”的经典贪心配对思想；P1913 把题意转换为奇偶交替序列问题，"
    "说明抽象题面条件的重要性；P2141 则说明在 n 较小、图结构明确时，Floyd-Warshall 是高可靠的最短路预处理方法。"
)
set_run_font(run)

p = doc.add_paragraph()
set_paragraph_format(p, after=0)
run = p.add_run("总体而言，本组题目能够帮助梳理“从条件抽象模型 -> 选择算法 -> 证明正确性 -> 分析复杂度”的完整解题流程。")
set_run_font(run)

doc.save(OUT_PATH)
print(OUT_PATH)
