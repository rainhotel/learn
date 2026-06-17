from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path(r"D:\moniC\project\learn")
INPUT_PATH = ROOT / "实验报告.docx"
OUTPUT_PATH = ROOT / "03-outputs" / "luogu-experiment-report-filled.docx"


def set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)


def fill_paragraphs(paragraphs, lines):
    for i, line in enumerate(lines):
        if i >= len(paragraphs):
            break
        p = paragraphs[i]
        if p.runs:
            p.runs[0].text = line
            for r in p.runs[1:]:
                r.text = ""
            set_black(p.runs[0])
        else:
            run = p.add_run(line)
            set_black(run)


def all_runs(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        yield r


def fill_cell(cell, lines):
    fill_paragraphs(cell.paragraphs, lines)


doc = Document(str(INPUT_PATH))

for run in all_runs(doc):
    set_black(run)

table1, table2, table3, table4, table5 = doc.tables

fill_cell(table1.cell(1, 1), [
    "1. b 能否被 a 整除。若不能整除，直接输出 No。",
    "2. 令 x = b / a，反复删除 x 中的因子 2 和 3。",
    "3. 如果最后 x = 1，说明 b 只比 a 多出若干个 2 和 3 的乘积，输出 Yes；否则输出 No。",
    "4. 这个判断等价于检查商的质因子是否只含 2 和 3。",
])
fill_cell(table1.cell(2, 1), [
    "#include <iostream>",
    "using namespace std;",
    "int main(){",
    "ios::sync_with_stdio(false); cin.tie(0);",
    "int n; cin>>n;",
    "while(n--){",
    "long long a,b; cin>>a>>b;",
    "if(b%a){ cout<<\"No\\n\"; continue; }",
    "long long x=b/a; while(x%2==0) x/=2; while(x%3==0) x/=3;",
    "cout<<(x==1?\"Yes\":\"No\")<<\"\\n\";",
    "}",
    "return 0; }",
])
fill_cell(table1.cell(3, 1), [
    "每组数据只需进行若干次除法运算，循环次数与 x 中 2 和 3 的个数有关。",
    "因此单组复杂度为 O(log b)，n 组总复杂度为 O(n log b)，额外空间 O(1)。",
])
fill_cell(table1.cell(4, 1), [
    "也可以先分解质因数，再统计除去 2 和 3 后是否还有剩余质因子。",
    "本题的核心仍然是判断商是否只含 2 和 3 两类因子。",
])

fill_cell(table2.cell(1, 1), [
    "1. 因为 a_i 都是正整数，区间和会随着右端点移动而单调增加。",
    "2. 可以用双指针维护一个连续窗口，实时保证窗口和不超过 V。",
    "3. 每次右端点扩展后，如果和超了，就移动左端点收缩窗口。",
    "4. 在所有合法窗口中取最大和即可。",
])
fill_cell(table2.cell(2, 1), [
    "#include <iostream>",
    "#include <vector>",
    "using namespace std;",
    "int main(){",
    "ios::sync_with_stdio(false); cin.tie(0);",
    "int n; long long V; cin>>n>>V;",
    "vector<int> a(n); for(int i=0;i<n;i++) cin>>a[i];",
    "long long sum=0,ans=0; int l=0;",
    "for(int r=0;r<n;r++){ sum+=a[r]; while(sum>V) sum-=a[l++]; if(sum>ans) ans=sum; }",
    "cout<<ans<<\"\\n\"; return 0; }",
])
fill_cell(table2.cell(3, 1), [
    "每个元素最多进入窗口一次、离开窗口一次，因此双指针总共只移动 O(n) 次。",
    "排序并不需要，故时间复杂度为 O(n)，空间复杂度为 O(1)。",
])
fill_cell(table2.cell(4, 1), [
    "如果数组中存在非正数，这种单调滑动窗口就不成立，需要换成其他方法。",
])

fill_cell(table3.cell(1, 1), [
    "1. 两人同桌的条件是饭量和不超过 W。",
    "2. 把所有学生按饭量从小到大排序后，最重的人应尽量和最轻的人尝试配对。",
    "3. 如果两者仍超过 W，则最重的人只能单独一桌。",
    "4. 这是一个典型的排序 + 双指针贪心问题。",
])
fill_cell(table3.cell(2, 1), [
    "#include <iostream>",
    "#include <algorithm>",
    "#include <vector>",
    "using namespace std;",
    "int main(){",
    "ios::sync_with_stdio(false); cin.tie(0);",
    "int n,w; cin>>n>>w;",
    "vector<int> a(n); for(int i=0;i<n;i++) cin>>a[i];",
    "sort(a.begin(),a.end());",
    "int l=0,r=n-1,ans=0;",
    "while(l<=r){ if(a[l]+a[r]<=w){ l++; r--; } else r--; ans++; }",
    "cout<<ans<<\"\\n\"; return 0; }",
])
fill_cell(table3.cell(3, 1), [
    "排序需要 O(n log n)，双指针扫描只需要 O(n)。",
    "因此总复杂度为 O(n log n)，额外空间 O(1)（不计排序实现内部开销）。",
])
fill_cell(table3.cell(4, 1), [
    "如果允许一桌坐三人或更多人，贪心策略就要重新设计，不能直接沿用本题做法。",
    "本题只允许 1 人或 2 人，因此双指针能够稳定得到最优解。",
])

fill_cell(table4.cell(1, 1), [
    "1. (b_i + b_{i+1})^2 为奇数，等价于 b_i + b_{i+1} 为奇数。",
    "2. 两数和为奇数，说明它们一奇一偶。",
    "3. 问题就转化成：求一个最长的奇偶交替子序列。",
    "4. 只需统计原序列中奇偶性变化的次数即可。",
])
fill_cell(table4.cell(2, 1), [
    "#include <iostream>",
    "#include <vector>",
    "using namespace std;",
    "int main(){",
    "ios::sync_with_stdio(false); cin.tie(0);",
    "int n; cin>>n;",
    "vector<int> a(n); for(int i=0;i<n;i++) cin>>a[i];",
    "int ans=1;",
    "for(int i=1;i<n;i++) if((a[i]&1)!=(a[i-1]&1)) ans++;",
    "cout<<ans<<\"\\n\"; return 0; }",
])
fill_cell(table4.cell(3, 1), [
    "只需要一次线性扫描，因此时间复杂度是 O(n)，空间复杂度是 O(1)。",
    "这是一个纯性质题，关键在于把平方奇偶性转换为原数奇偶性。",
])
fill_cell(table4.cell(4, 1), [
    "也可以理解成把序列按奇偶分段，每一段取一个元素，得到的长度就是答案。",
])

fill_cell(table5.cell(1, 1), [
    "1. 需要按 1,2,3,...,n 的顺序访问星球，最后再回到 1。",
    "2. 每一段相邻目标星球之间都应使用最短路径。",
    "3. 因为 n <= 100，直接使用 Floyd-Warshall 预处理所有点对最短路最稳妥。",
    "4. 读入时把 9999 当作无穷大即可。",
])
fill_cell(table5.cell(2, 1), [
    "#include <iostream>",
    "using namespace std;",
    "const long long INF=1000000000LL;",
    "long long d[105][105];",
    "int main(){",
    "ios::sync_with_stdio(false); cin.tie(0);",
    "int n; cin>>n;",
    "for(int i=1;i<=n;i++)for(int j=1;j<=n;j++){ long long x; cin>>x; d[i][j]=(i==j?0:(x==9999?INF:x)); }",
    "for(int k=1;k<=n;k++)for(int i=1;i<=n;i++)for(int j=1;j<=n;j++) if(d[i][k]+d[k][j]<d[i][j]) d[i][j]=d[i][k]+d[k][j];",
    "long long ans=0; for(int i=1;i<n;i++) ans+=d[i][i+1]; ans+=d[n][1];",
    "cout<<ans<<\"\\n\"; return 0; }",
])
fill_cell(table5.cell(3, 1), [
    "Floyd-Warshall 的时间复杂度是 O(n^3)，额外空间是 O(n^2)。",
    "在 n 只有 100 的前提下，这个复杂度完全可以接受。",
])
fill_cell(table5.cell(4, 1), [
    "也可以分段使用 Dijkstra 逐次求最短路，但实现上不如 Floyd 统一。",
])

for run in all_runs(doc):
    set_black(run)

doc.save(str(OUTPUT_PATH))
print(str(OUTPUT_PATH))
