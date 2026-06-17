from docx import Document

path = r"D:\moniC\project\learn\实验报告.docx"
doc = Document(path)

print("paragraphs", len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs, 1):
    txt = para.text.replace("\t", "\\t")
    print(i, repr(txt))

print("tables", len(doc.tables))
for ti, table in enumerate(doc.tables, 1):
    print("TABLE", ti, "rows", len(table.rows), "cols", len(table.columns))
    for r in range(len(table.rows)):
        counts = []
        for c in range(len(table.columns)):
            cell = table.cell(r, c)
            counts.append(len(cell.paragraphs))
        print(" row", r + 1, counts)
        for c in range(len(table.columns)):
            cell = table.cell(r, c)
            print("  cell", r + 1, c + 1, "text", repr(cell.text), "paragraphs", len(cell.paragraphs))
