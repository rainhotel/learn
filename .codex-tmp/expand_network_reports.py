from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\moniC\project\learn")
OUT_DIR = ROOT / "03-outputs"
OUT_DIR.mkdir(exist_ok=True)


def set_run_font(run, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.bold = bold


def format_para(paragraph, first_line=True, spacing_after=6):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(spacing_after)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if first_line:
        fmt.first_line_indent = Pt(24)
    for run in paragraph.runs:
        set_run_font(run)


def tidy_text(text):
    replacements = {
        " MAC 地址": " MAC地址",
        "MAC 地址": "MAC地址",
        " IP 地址": " IP地址",
        "IP 地址": "IP地址",
        " PC ": " PC",
        " VLAN": " VLAN",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def make_paragraph_like(anchor, text="", bold=False, first_line=True):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if text:
        run = paragraph.add_run(tidy_text(text))
        set_run_font(run, bold=bold)
    format_para(paragraph, first_line=first_line, spacing_after=6)
    return paragraph


def insert_after(anchor, entries):
    current = anchor
    for entry in entries:
        if isinstance(entry, tuple):
            text, bold, first_line = entry
        else:
            text, bold, first_line = entry, False, True
        current = make_paragraph_like(current, text, bold=bold, first_line=first_line)
    return current


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(tidy_text(text))
    set_run_font(run)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_para(paragraph, first_line=True, spacing_after=6)


def find_first(doc, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    raise ValueError("paragraph not found")


def find_after_heading(doc, heading):
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() == heading:
            return paragraphs[i + 1]
    raise ValueError(f"heading not found: {heading}")


def remove_trailing_empty_paragraphs(doc, keep=0):
    body = doc.element.body
    children = list(body)
    sect_pr = None
    if children and children[-1].tag == qn("w:sectPr"):
        sect_pr = children[-1]
        children = children[:-1]
    trailing = []
    for child in reversed(children):
        if child.tag != qn("w:p"):
            break
        paragraph = Paragraph(child, doc)
        has_text = bool(paragraph.text.strip())
        has_drawing = bool(child.xpath(".//w:drawing"))
        if has_text or has_drawing:
            break
        trailing.append(child)
    for child in trailing[keep:]:
        body.remove(child)


def normalize_added_headings(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in {"实验原理补充", "结果记录与分析"}:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, bold=True)
                run.font.size = Pt(12)
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)


def expand_exp1(src, dst):
    doc = Document(str(src))
    content_image = doc.paragraphs[11]
    result_anchor = doc.paragraphs[21]
    summary_para = find_after_heading(doc, "四、体会和总结")

    insert_after(
        content_image,
        [
            ("实验原理补充", True, False),
            "交换机工作在数据链路层，主要依据以太网帧中的目的 MAC 地址进行转发。交换机端口收到帧后，会先读取源 MAC 地址，并把该地址与入端口的对应关系写入 MAC 地址表；随后再根据目的 MAC 地址决定转发方式。",
            "当目的 MAC 地址已经存在于 MAC 地址表中时，交换机会把帧只转发到对应端口，这就是单播转发；当目的地址未知，或者收到广播帧时，交换机会向除入端口外的其他端口泛洪转发。通过观察单播包和广播包的传播范围，可以验证交换机“学习、查表、转发、泛洪”的基本工作过程。",
        ],
    )

    insert_after(
        result_anchor,
        [
            ("结果记录与分析", True, False),
            "实验环境中 4 台 PC 连接在同一台 2960 交换机上，IP 地址位于同一网段，因此主机之间可以直接通过二层交换完成通信。配置终端 IP 地址后，通过终端命令查看各主机的物理地址，为后续观察交换机学习 MAC 地址提供依据。",
            "查看交换机初始 MAC 地址表时，表项为空或只有系统默认信息，说明交换机还没有通过数据帧学习到主机地址。当 PC0 与 PC1 等主机发生通信后，交换机从收到的数据帧中提取源 MAC 地址，并把该 MAC 地址记录到对应的 FastEthernet 端口上，MAC 地址表中出现 dynamic 类型的动态表项。",
            "从截图中的 show mac address-table 结果可以看出，交换机已经学习到部分主机的 MAC 地址与端口的对应关系，例如某个 MAC 地址对应 Fa0/1，另一个 MAC 地址对应 Fa0/2。这说明交换机并不是预先知道网络中所有设备的位置，而是在通信过程中逐步建立转发表。",
            "执行 clear mac-address-table dynamic 后，再次查看 MAC 地址表，动态学习到的表项被清除。此时如果继续发送数据包，交换机会重新根据源 MAC 地址学习表项。该现象说明动态 MAC 表具有可清除、可老化、可重新学习的特点。",
            "在广播包观察中，广播帧没有唯一目的主机，交换机会把广播帧转发到除接收端口外的所有端口，因此同一交换机下的所有主机都能收到广播包。与单播转发相比，广播转发范围更大，也说明单交换机网络整体处于同一个广播域内。",
        ],
    )

    replace_paragraph(
        summary_para,
        "通过本次单交换机组网实验，我对交换机的二层转发机制有了更直观的理解。交换机通过学习源 MAC 地址建立 MAC 地址表，再根据目的 MAC 地址选择转发端口；初始表项为空、通信后出现 dynamic 动态表项、清除后又能重新学习，验证了 MAC 地址学习机制的动态性。单播通信只转发到目标端口，广播通信会扩散到同一广播域内的其他端口。由此可见，交换机能够减少不必要的单播流量，但不能隔离广播域，后续可结合 VLAN、路由器等技术进一步控制广播范围。"
    )

    normalize_added_headings(doc)
    remove_trailing_empty_paragraphs(doc)
    doc.save(str(dst))


def expand_exp2(src, dst):
    doc = Document(str(src))
    content_image = doc.paragraphs[9]
    result_anchor = doc.paragraphs[14]
    summary_para = find_after_heading(doc, "四、体会和总结")

    insert_after(
        content_image,
        [
            ("实验原理补充", True, False),
            "多交换机组网是在单交换机网络的基础上，通过交换机之间的链路把多个二层网络连接在一起。这样可以增加可接入主机数量，扩展以太网覆盖范围，使不同交换机下的主机仍然能够在同一网段内通信。",
            "需要注意的是，普通二层交换机连接后，默认并不会隔离广播域。广播帧从一台交换机进入后，会继续通过交换机间链路转发到另一台交换机，再由另一台交换机泛洪到其余端口。因此，多交换机可以扩大网络覆盖范围，但也会扩大广播传播范围。广播域过大时，可能导致广播流量增加，影响网络性能。",
        ],
    )

    insert_after(
        result_anchor,
        [
            ("结果记录与分析", True, False),
            "实验拓扑中使用两台 2960 交换机互联，PC0 至 PC3 连接在 Switch0 上，PC4 和 PC5 连接在 Switch1 上。所有终端配置在同一 IP 网段后，主机之间理论上可以通过交换机间链路完成二层转发。",
            "连通性测试时，如果目标主机位于另一台交换机下，数据帧需要先到达本端交换机，再通过交换机间链路转发到对端交换机。若目标 MAC 地址未知，交换机会先进行泛洪；当通信双方产生响应后，两台交换机都会逐步学习到相应 MAC 地址与端口的对应关系。",
            "在广播域观察中，可以看到广播帧不仅会到达本交换机下的其他主机，也会沿着交换机互联链路到达另一台交换机，并继续转发到该交换机连接的主机端口。这说明两台交换机互联后仍属于同一个广播域，广播范围从单台交换机扩展到了整个二层交换网络。",
            "截图中部分 ping 测试出现超时，可能与目标主机 IP 配置、链路状态、模拟器收敛过程或测试命令使用方式有关。若要进一步确认，应检查每台 PC 的 IP 地址与子网掩码是否一致，确认交换机间链路处于绿色连通状态，并在通信前后观察 MAC 地址表是否学习到跨交换机的表项。",
            "本实验说明，交换机级联能够扩展以太网规模，但并不会天然划分广播域。若希望在扩大网络规模的同时控制广播传播范围，需要进一步使用 VLAN 划分广播域，并通过三层设备实现不同 VLAN 之间的通信。",
        ],
    )

    replace_paragraph(
        summary_para,
        "通过本次多交换机组网实验，我进一步理解了交换机扩展以太网覆盖范围的方式。多台交换机互联后，不同交换机下的主机仍可通过二层链路通信，交换机会根据源 MAC 地址动态学习端口位置，并根据目的 MAC 地址转发数据帧。但交换机数量增加并不会自动缩小广播域，广播帧会通过交换机间链路继续传播到另一台交换机及其连接的主机。因此，多交换机网络虽然提高了接入范围和端口数量，也可能扩大广播影响范围。后续设计较大规模局域网时，应结合 VLAN、生成树协议和三层转发等技术，既保证连通性，又控制广播流量和故障影响范围。"
    )

    normalize_added_headings(doc)
    remove_trailing_empty_paragraphs(doc)
    doc.save(str(dst))


def main():
    exp1_src = Path(r"D:\moniC\实验1_202483290054_李林浩.docx")
    exp2_src = Path(r"D:\moniC\实验2_202483290054_李林浩.docx")
    expand_exp1(exp1_src, OUT_DIR / "实验1_202483290054_李林浩_扩写完善版.docx")
    expand_exp2(exp2_src, OUT_DIR / "实验2_202483290054_李林浩_扩写完善版.docx")


if __name__ == "__main__":
    main()
