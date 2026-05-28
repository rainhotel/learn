from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\moniC\project\learn\实验1-数据库表建立-实验报告-填图版.docx")

TEXTBOOK_IMAGES = [
    Path(r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-04\3409faaa25a4bde60010a315c0c6e6bc\a6942b2ee445487cc41d777f20da644c.jpg"),
    Path(r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-04\3409faaa25a4bde60010a315c0c6e6bc\924663b2d12c0fde565c40a1eaf26f88.jpg"),
    Path(r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-04\3409faaa25a4bde60010a315c0c6e6bc\b56d87935276ef1e159603d3c21f9635.jpg"),
    Path(r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-04\3409faaa25a4bde60010a315c0c6e6bc\eb84b7813c0e64230db0f9ecf0232fc6.jpg"),
]

SCREENSHOT_DIR = Path(r"D:\xwechat_files\wxid_0vnbq4ga2qso22_630f\temp\RWTemp\2026-04\9e20f478899dc29eb19741386f9343c8")
LOCAL_SCREENSHOTS = {
    "goodsorder_query": SCREENSHOT_DIR / "9605069c1916e1ef26d6ea8e6498807d.png",
    "orderlist_fk": SCREENSHOT_DIR / "f1246f32ba5e562ffeb440935bc5c0ed.png",
    "orderlist_structure": SCREENSHOT_DIR / "9cfd7ece2e84bbccdf8fa2ebf1ef8a5d.png",
    "goodsinfo_structure": SCREENSHOT_DIR / "55a5af6c7d8ffd16f984b8e7dab9cc10.png",
    "customerinfo_structure": SCREENSHOT_DIR / "9f0ba8fbc7e20e35a31da20dc0f29678.png",
    "course_pk": SCREENSHOT_DIR / "6692eebfe914cb6f573188f502e44b24.png",
    "stucourse_fk": SCREENSHOT_DIR / "322b6d761805ca3a54bf118487baada7.png",
    "student_pk": SCREENSHOT_DIR / "39801b7418ea4b7ec4c8f647a33803fb.png",
    "stucourse_structure": SCREENSHOT_DIR / "213332ef5c19fe7b8312106e35780fb5.png",
    "course_structure": SCREENSHOT_DIR / "9d46d8c8b4e83a010d36d2453703213e.png",
    "student_structure": SCREENSHOT_DIR / "a8f15f1af3abcd87449cd5016f47a353.png",
    "studentcourse_db": SCREENSHOT_DIR / "39cefec4c1d0a4b1ffda292a4e67aa21.png",
    "goodsorder_db": SCREENSHOT_DIR / "c0d13a07cee330f01c7c1fb2e53e8a2e.png",
    "studentcourse_tables": SCREENSHOT_DIR / "ad7e73a8a9305d7771ea12735fbf5cc4.png",
    "goodsorder_tables": SCREENSHOT_DIR / "3d0c0b640569970ee57608e979315214.png",
}


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="SimSun", size=10.5, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, header_fill="D9EAF7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
    for cell in table.rows[0].cells:
        shade_cell(cell, header_fill)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, name="SimHei", size=16 - level * 2, color=(0, 0, 0))
    run.bold = True
    return p


def add_normal(doc, text="", first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code_block(doc, text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9.5)


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[截图占位：{text}]")
    set_run_font(run, size=10, color=(0, 0, 0))


def add_image_or_placeholder(doc, image_path, placeholder_text, width=6.0):
    if image_path and Path(image_path).exists():
        doc.add_picture(str(image_path), width=Inches(width))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"图：{Path(image_path).name}")
        set_run_font(run, size=9.5, color=(0, 0, 0))
    else:
        add_placeholder(doc, placeholder_text)


def add_data_table(doc, headers, rows, col_widths=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(" | ".join(headers))
    set_run_font(run, name="Consolas", size=9.5)
    run.bold = True
    for row in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        line = " | ".join(str(v) for v in row)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
styles["Normal"].font.name = "SimSun"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
styles["Normal"].font.size = Pt(10.5)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    style = styles[style_name]
    style.font.name = "SimHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    style.font.color.rgb = RGBColor(0, 0, 0)


# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("南京信息工程大学 数据库系统实验（实习）报告")
set_run_font(run, name="SimHei", size=20)
run.bold = True

for line in [
    "实验（实习）名称：数据库、表的建立    实验（实习）日期：2026-04-26",
    "学院：__________    专业：__________    年级/班次：__________",
    "姓名：__________    学号：__________    指导教师：__________    得分：__________",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(line)
    set_run_font(r, size=11)

doc.add_paragraph()

add_heading(doc, "一、实验目的", level=1)
for text in [
    "了解 SQL Server 的界面结构和基本功能。",
    "掌握 SQL Server 的基本使用方法。",
    "掌握在查询分析器中执行 SQL 语句的方法。",
    "掌握 SQL Server 的基本数据类型。",
    "掌握创建数据库和表的 SQL 语句。",
]:
    add_bullet(doc, text, numbered=True)

add_heading(doc, "二、实验内容与步骤", level=1)
add_heading(doc, "1. 使用界面操作方式建立 GoodsOrder 数据库及其基本表", level=2)
add_normal(
    doc,
    "本部分采用 SQL Server Management Studio 的图形界面完成数据库和表的建立，并对主键、外键、非空约束和默认约束进行检查。实验完成后，已在对象资源管理器中确认 GoodsOrder 数据库、CustomerInfo 表、GoodsInfo 表和 OderList 表创建成功。",
)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_db"], "GoodsOrder 数据库在对象资源管理器中的创建成功截图", width=2.2)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_tables"], "GoodsOrder -> 表 节点下显示 dbo.CustomerInfo、dbo.GoodsInfo、dbo.OderList 的截图", width=2.3)

add_heading(doc, "（1）CustomerInfo 表结构与约束", level=3)
customer_schema = [
    ["客户编号", "char(6)", "主键，非空"],
    ["客户姓名", "char(20)", "非空"],
    ["性别", "char(2)", "允许空值"],
    ["出生日期", "datetime", "允许空值"],
    ["所在省市", "varchar(50)", "允许空值"],
    ["联系电话", "varchar(12)", "允许空值"],
    ["微信号", "varchar(30)", "允许空值"],
    ["VIP", "bit", "允许空值，默认值 0"],
    ["备注", "text", "允许空值"],
]
add_data_table(doc, ["字段名", "数据类型", "约束说明"], customer_schema, [1.3, 1.1, 3.5])
add_normal(
    doc,
    "实际操作中，客户编号被设置为主键，客户姓名设置为非空，VIP 字段设置了默认值 0。经设计视图检查，表结构与教材要求一致。",
)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_tables"], "CustomerInfo 表创建成功截图", width=2.3)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["customerinfo_structure"], "CustomerInfo 表结构及约束截图", width=6.2)
add_placeholder(doc, "CustomerInfo 表中 VIP 字段默认值截图")

add_heading(doc, "（2）GoodsInfo 表结构与约束", level=3)
goods_schema = [
    ["商品编号", "char(8)", "主键，非空"],
    ["商品类别", "char(20)", "非空"],
    ["商品名称", "varchar(50)", "非空"],
    ["品牌", "varchar(30)", "允许空值"],
    ["单价", "float", "允许空值"],
    ["生产商", "varchar(50)", "允许空值"],
    ["保质期", "datetime", "默认值 2000-1-1"],
    ["库存量", "int", "允许空值"],
    ["备注", "text", "允许空值"],
]
add_data_table(doc, ["字段名", "数据类型", "约束说明"], goods_schema, [1.3, 1.2, 3.4])
add_normal(
    doc,
    "实际操作中，商品编号被设置为主键，商品类别和商品名称设置为非空，保质期字段设置了默认值 2000-1-1。经设计视图检查，表结构与教材要求一致。",
)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_tables"], "GoodsInfo 表创建成功截图", width=2.3)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsinfo_structure"], "GoodsInfo 表结构及约束截图", width=6.2)
add_placeholder(doc, "GoodsInfo 表中保质期字段默认值截图")

add_heading(doc, "（3）OderList 表结构与约束", level=3)
order_schema = [
    ["客户编号", "char(6)", "联合主键组成字段，外键，非空"],
    ["商品编号", "char(8)", "联合主键组成字段，外键，非空"],
    ["订购时间", "datetime", "联合主键组成字段，非空"],
    ["数量", "int", "允许空值"],
    ["需要日期", "datetime", "允许空值"],
    ["付款方式", "varchar(40)", "允许空值"],
    ["送货方式", "varchar(50)", "允许空值"],
]
add_data_table(doc, ["字段名", "数据类型", "约束说明"], order_schema, [1.3, 1.2, 3.4])
add_normal(
    doc,
    "OderList 表将客户编号、商品编号和订购时间联合设置为主键，同时建立了客户编号到 CustomerInfo 表、商品编号到 GoodsInfo 表的外键约束。经关系窗口检查，外键关系设置正确。",
)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_tables"], "OderList 表创建成功截图", width=2.3)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["orderlist_structure"], "OderList 表结构及联合主键截图", width=6.2)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["orderlist_fk"], "OderList 表外键关系截图")

add_heading(doc, "（4）GoodsOrder 数据录入结果", level=3)
add_normal(
    doc,
    "根据教材第 22 页表 2.6、表 2.7 和表 2.8 中的数据，对三张表进行了录入。录入完成后，通过查询和结果视图检查，CustomerInfo 共录入 6 条记录，GoodsInfo 共录入 13 条记录，OderList 共录入 10 条记录，数据内容与教材一致。",
)

customer_rows = [
    ["100001", "张小林", "1982-02-01", "男", "江苏南京", "02581234678", "13980030075", "1", "银牌客户"],
    ["100002", "李红红", "1991-03-22", "女", "江苏苏州", "13908899120", "13908899120", "1", "金牌客户"],
    ["100003", "王晓美", "1986-08-20", "女", "上海市", "02166552101", "wxid_0021001", "0", "新客户"],
    ["100004", "赵明", "1992-03-28", "男", "河南郑州", "13809900118", "NULL", "0", "新客户"],
    ["100005", "张帆一", "1990-08-10", "男", "山东烟台", "13880933201", "NULL", "0", "NULL"],
    ["100006", "王芳芳", "1996-05-01", "女", "江苏南京", "13709092011", "wxid_7890921", "0", "NULL"],
]
add_normal(doc, "CustomerInfo 表录入结果如下：")
add_data_table(
    doc,
    ["客户编号", "客户姓名", "出生日期", "性别", "所在省市", "联系电话", "微信号", "VIP", "备注"],
    customer_rows,
    [0.9, 0.8, 1.0, 0.5, 1.0, 1.0, 1.1, 0.5, 0.8],
)

goods_rows = [
    ["10010001", "食品", "咖啡", "宇一", "50", "宇一饮料公司", "2021-08-31", "100", "NULL"],
    ["10010002", "食品", "苹果汁", "宇一", "5.2", "宇一饮料公司", "2020-12-31", "500", "NULL"],
    ["10020001", "食品", "大米", "健康", "35", "健康粮食生产基地", "2020-12-20", "100", "NULL"],
    ["10020002", "食品", "面粉", "健康", "18", "健康粮食生产基地", "2021-01-20", "20", "NULL"],
    ["20180001", "服装", "运动服", "天天", "200", "天天服饰公司", "2000-01-01", "5", "有断码"],
    ["20180002", "服装", "T恤", "天天", "120", "天天服饰公司", "2000-01-01", "10", "NULL"],
    ["30010001", "文具", "签字笔", "新新", "3.5", "新新文化用品制造厂", "2000-01-01", "100", "NULL"],
    ["30010002", "文具", "文件夹", "新新", "5.6", "新新文化用品制造厂", "2000-01-01", "50", "NULL"],
    ["40010001", "图书", "营养菜谱", "新华", "38", "食品出版公司", "2000-01-01", "12", "NULL"],
    ["40010002", "图书", "豆浆的做法", "新华", "20", "食品出版公司", "2000-01-01", "20", "NULL"],
    ["50020001", "体育用品", "羽毛球拍", "美好", "30", "美好体育用品公司", "2000-01-01", "30", "NULL"],
    ["50020002", "体育用品", "篮球", "美好", "80", "美好体育用品公司", "2000-01-01", "20", "NULL"],
    ["50020003", "体育用品", "足球", "美好", "65", "美好体育用品公司", "2000-01-01", "20", "NULL"],
]
add_normal(doc, "GoodsInfo 表录入结果如下：")
add_data_table(
    doc,
    ["商品编号", "商品类别", "商品名称", "品牌", "单价", "生产商", "保质期", "库存量", "备注"],
    goods_rows,
    [0.9, 0.8, 0.9, 0.6, 0.5, 1.4, 0.9, 0.6, 0.6],
)

order_rows = [
    ["100001", "10010001", "2020-02-18 12:20:00", "2", "2020-02-20", "支付宝", "客户自提"],
    ["100001", "30010001", "2020-02-10 12:30:00", "10", "2020-02-20", "网银转账", "送货上门"],
    ["100002", "10010001", "2020-02-18 13:00:00", "1", "2020-02-21", "微信支付", "客户自提"],
    ["100002", "50020001", "2020-02-18 13:20:00", "1", "2020-02-21", "微信支付", "客户自提"],
    ["100004", "20180002", "2020-02-19 10:00:00", "1", "2020-02-28", "信用卡", "送货上门"],
    ["100004", "50020002", "2020-02-19 10:40:00", "2", "2020-02-28", "信用卡", "送货上门"],
    ["100004", "30010002", "2020-02-19 11:00:00", "10", "2020-02-28", "信用卡", "送货上门"],
    ["100005", "40010001", "2020-02-20 08:00:00", "2", "2020-02-27", "支付宝", "送货上门"],
    ["100005", "40010002", "2020-02-20 08:20:00", "3", "2020-02-27", "支付宝", "送货上门"],
    ["100006", "10020001", "2020-02-23 09:00:00", "5", "2020-02-26", "信用卡", "送货上门"],
]
add_normal(doc, "OderList 表录入结果如下：")
add_data_table(
    doc,
    ["客户编号", "商品编号", "订购时间", "数量", "需要日期", "付款方式", "送货方式"],
    order_rows,
    [0.8, 0.9, 1.5, 0.5, 0.9, 0.8, 0.9],
)
add_placeholder(doc, "CustomerInfo 查询结果截图")
add_placeholder(doc, "GoodsInfo 查询结果截图")
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["goodsorder_query"], "GoodsOrder 查询结果截图")

add_heading(doc, "2. 使用 SQL 命令方式建立 StudentCourse 数据库及其基本表", level=2)
add_normal(
    doc,
    "本部分通过查询分析器输入 SQL 语句，建立 StudentCourse 数据库，并创建 Student、Course 和 StuCourse 三张基本表。经对象资源管理器和设计视图检查，三张表均创建成功，主键和外键约束设置正确。",
)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["studentcourse_db"], "StudentCourse 数据库在对象资源管理器中的创建成功截图", width=2.2)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["studentcourse_tables"], "StudentCourse -> 表 节点下显示 dbo.Student、dbo.Course、dbo.StuCourse 的截图", width=2.8)

sql_db = """
CREATE DATABASE StudentCourse;
GO
"""
sql_student = """
USE StudentCourse;
GO
CREATE TABLE Student (
    学号 char(6) NOT NULL,
    姓名 char(12) NOT NULL,
    性别 char(2) NOT NULL,
    专业名 varchar(20),
    出生日期 smalldatetime,
    总学分 int,
    备注 text,
    CONSTRAINT PK_Student PRIMARY KEY (学号)
);
GO
"""
sql_course = """
CREATE TABLE Course (
    课程号 char(4) NOT NULL,
    课程名 varchar(40) NOT NULL,
    开课学期 int,
    学时 int,
    学分 int,
    CONSTRAINT PK_Course PRIMARY KEY (课程号)
);
GO
"""
sql_stucourse = """
CREATE TABLE StuCourse (
    学号 char(6) NOT NULL,
    课程号 char(4) NOT NULL,
    成绩 real,
    CONSTRAINT PK_StuCourse PRIMARY KEY (学号, 课程号),
    CONSTRAINT FK_StuCourse_Student FOREIGN KEY (学号) REFERENCES Student(学号),
    CONSTRAINT FK_StuCourse_Course FOREIGN KEY (课程号) REFERENCES Course(课程号)
);
GO
"""

add_heading(doc, "（1）建立 StudentCourse 数据库的 SQL 语句", level=3)
add_code_block(doc, sql_db)
add_placeholder(doc, "StudentCourse 数据库创建成功截图")

add_heading(doc, "（2）建立 Student 表的 SQL 语句", level=3)
add_code_block(doc, sql_student)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["studentcourse_tables"], "Student 表创建成功截图", width=2.8)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["student_structure"], "Student 表结构截图", width=6.0)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["student_pk"], "Student 表主键约束截图", width=6.0)

add_heading(doc, "（3）建立 Course 表的 SQL 语句", level=3)
add_code_block(doc, sql_course)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["studentcourse_tables"], "Course 表创建成功截图", width=2.8)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["course_structure"], "Course 表结构截图", width=6.0)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["course_pk"], "Course 表主键约束截图", width=6.0)

add_heading(doc, "（4）建立 StuCourse 表的 SQL 语句", level=3)
add_code_block(doc, sql_stucourse)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["studentcourse_tables"], "StuCourse 表创建成功截图", width=2.8)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["stucourse_structure"], "StuCourse 表结构及联合主键截图", width=6.0)
add_image_or_placeholder(doc, LOCAL_SCREENSHOTS["stucourse_fk"], "StuCourse 表外键约束截图", width=6.0)

add_heading(doc, "（5）StudentCourse 数据录入情况说明", level=3)
add_normal(
    doc,
    "根据实验要求，需要将教材附录实验 3 中表 A3.1、表 A3.2 和表 A3.3 的数据分别录入 Student、Course 和 StuCourse 三张表，并给出录入后的结果截图。当前整理材料中未包含这三张数据表的原始数据及录入后的本地截图，因此本文档保留了结果位置。若后续补充对应数据，可直接在下列位置插入查询结果截图。",
)
add_placeholder(doc, "Student 表数据录入结果截图")
add_placeholder(doc, "Course 表数据录入结果截图")
add_placeholder(doc, "StuCourse 表数据录入结果截图")

add_heading(doc, "三、实验心得", level=1)
add_normal(
    doc,
    "通过本次实验，我进一步熟悉了 SQL Server Management Studio 的基本界面和常用操作流程，掌握了在对象资源管理器中建立数据库、创建数据表、设置主键和外键约束的方法。与单纯阅读教材相比，实际动手创建 CustomerInfo、GoodsInfo、OderList、Student、Course 和 StuCourse 等表，使我对关系数据库中的实体、联系以及完整性约束有了更直观的理解。",
)
add_normal(
    doc,
    "本次实验还让我体会到图形界面建表和 SQL 命令建表各有特点。图形界面操作直观，适合初学者理解表结构和约束的设置过程；而 SQL 语句方式更高效、可复用，尤其适合在需要重复建立数据库或批量操作时使用。通过对比两种方式，我认识到掌握标准 SQL 语句的重要性，也加深了对主键、联合主键、外键、默认值和非空约束等概念的理解。",
)
add_normal(
    doc,
    "在实验过程中，较容易出错的地方主要有三类：一是字段类型和长度不一致会导致外键无法建立；二是重复执行建表或插入语句会出现对象已存在或主键重复错误；三是截图材料若不及时保存，后续整理报告会比较被动。因此，在今后的实验中，我会养成边操作边验证、边截图边整理的习惯，提高实验效率和结果的规范性。",
)

add_heading(doc, "附录：教材参考图片", level=1)
add_normal(
    doc,
    "以下图片为本次实验整理时使用的教材参考资料，包括 GoodsInfo、CustomerInfo、OderList 表的教材样例数据及相关 SQL 建表示例，用于说明数据来源和结构依据。",
)
for img in TEXTBOOK_IMAGES:
    if img.exists():
        doc.add_picture(str(img), width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"图：{img.name}")
        set_run_font(run, size=9.5)

doc.save(OUT)
print(OUT)
