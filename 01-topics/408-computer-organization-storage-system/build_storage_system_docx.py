from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "408-computer-organization-storage-system-guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(89, 89, 89)
FILL_BLUE = "E8EEF5"
FILL_GRAY = "F2F4F7"
FILL_NOTE = "F7FBF2"
FILL_WARN = "FFF7E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_note(doc, title, body, fill=FILL_NOTE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    r2 = p.add_run(" " + body)
    set_run_font(r2)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths_cm, header_fill=FILL_BLUE, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=font_size)
            if len(val) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_width(table, widths_cm)
    doc.add_paragraph()
    return table


def configure_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("408 计算机组成原理：存储系统考点与题型手册")
    set_run_font(r, size=9, color=MUTED)
    return doc


def build_doc():
    doc = configure_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("408 计算机组成原理\n存储系统完整考点与题型手册")
    set_run_font(r, size=24, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("结合王道课程体系与历年真题高频命题方式整理")
    set_run_font(r, size=12, color=MUTED)
    doc.add_paragraph()

    add_note(
        doc,
        "使用说明：",
        "本文按 408 计组“存储系统”常见复习顺序组织：先建立知识框架，再给真题常见题型的写法和解法。真题部分采用高频题型归纳，不强行标注年份题号，重点服务做题。"
    )

    add_heading(doc, "0. 本章总览：真题到底考什么", 1)
    add_para(doc, "存储系统这一章的核心矛盾是：CPU 很快、主存较慢、辅存容量大但更慢。考试围绕“如何用层次结构弥补速度和容量矛盾”展开。")
    add_matrix(
        doc,
        ["模块", "王道复习重点", "408 高频考法", "做题关键词"],
        [
            ["存储器概述", "层次结构、局部性、性能指标", "选择题、概念判断、平均访问时间", "速度/容量/价格、时间局部性、空间局部性"],
            ["主存储器", "SRAM/DRAM、刷新、芯片扩展、CPU 连接", "芯片数量、地址线/数据线、片选、刷新方式", "位扩展、字扩展、译码、刷新周期"],
            ["多模块主存", "单体多字、交叉编址、低位/高位交叉", "连续访存时间、带宽、模块号判断", "低位交叉、流水启动、模块冲突"],
            ["Cache", "映射、地址划分、替换、写策略、命中率", "大题和选择题核心区", "Tag、行号、组号、块内地址、LRU、写回"],
            ["虚拟存储器", "页表、TLB、地址变换、缺页", "跨计组/OS 综合题", "页号、页内偏移、页框号、TLB miss、缺页"],
        ],
        [2.4, 4.0, 4.4, 5.4],
        font_size=9,
    )
    add_note(doc, "总口诀：", "Cache 解决“快不快”，虚拟存储器解决“大不大”；Cache 和虚拟存储器都依赖局部性原理，但交换单位、管理主体、命中/缺失含义完全不同。", FILL_WARN)

    add_heading(doc, "1. 存储层次与局部性原理", 1)
    add_heading(doc, "1.1 必背知识点", 2)
    add_bullets(doc, [
        "存储层次：寄存器 -> Cache -> 主存 -> 辅存。越靠近 CPU，速度越快、容量越小、价格越高。",
        "时间局部性：刚访问过的信息近期可能再次访问。Cache 替换中的 LRU、虚拟存储中的工作集都依赖它。",
        "空间局部性：访问某地址后，其附近地址也可能被访问。Cache 按块调入、分页按页调入都依赖它。",
        "存储系统评价指标：存取时间、存取周期、带宽、命中率、缺失率、平均访问时间。",
    ])
    add_heading(doc, "1.2 常考写法", 2)
    add_para(doc, "选择题遇到“为什么 Cache/虚存有效”，优先答“程序访问具有局部性”。遇到“Cache 与虚存区别”，先从目的区分：Cache 为速度，虚存为容量。")

    add_heading(doc, "2. 主存储器：器件、刷新、扩展与 CPU 连接", 1)
    add_heading(doc, "2.1 SRAM、DRAM、ROM", 2)
    add_matrix(
        doc,
        ["类型", "特点", "典型用途", "考点"],
        [
            ["SRAM", "触发器保存信息，速度快，集成度低，成本高，不需刷新", "Cache", "常与 DRAM 对比"],
            ["DRAM", "电容保存信息，集成度高，成本低，速度较慢，需刷新", "主存", "刷新方式、刷新周期、行列地址"],
            ["ROM/Flash", "掉电后信息仍保留，写入方式受限", "固件、启动程序", "非易失性"],
        ],
        [2.0, 6.0, 3.0, 5.2],
        font_size=9,
    )
    add_heading(doc, "2.2 DRAM 刷新", 2)
    add_bullets(doc, [
        "刷新对象通常按行刷新，不是按单元逐个刷新。",
        "集中刷新：在一段时间内集中完成刷新，期间可能出现访存死区。",
        "分散刷新：把刷新分散到每个存取周期中，系统速度受影响但无集中死区。",
        "异步刷新：在刷新周期内均匀安排各行刷新，是真题最常考的折中方案。",
    ])
    add_heading(doc, "2.3 主存扩展题解法", 2)
    add_para(doc, "芯片扩展题的题干通常给出“芯片规格”和“目标存储器规格”。固定按下面三步写。")
    add_numbers(doc, [
        "看字长：位扩展倍数 = 目标字长 / 单片字长。",
        "看字数：字扩展倍数 = 目标字数 / 单片字数。",
        "算总片数：总片数 = 位扩展倍数 × 字扩展倍数。再根据字扩展倍数判断需要多少片选信号。",
    ])
    add_note(doc, "例题模板：", "用 16K × 8 位芯片组成 64K × 32 位存储器。位扩展倍数 = 32/8 = 4；字扩展倍数 = 64K/16K = 4；总片数 = 16。每 4 片组成一组扩展字长，共 4 组扩展字数。", FILL_WARN)
    add_heading(doc, "2.4 CPU 与主存连接", 2)
    add_bullets(doc, [
        "地址线决定可寻址单元数。若按字节编址，n 根地址线可寻址 2^n 个字节。",
        "数据线宽度决定一次可并行传输的数据位数，常与存储字长、总线宽度结合考。",
        "片选信号由高位地址译码产生，片内地址由低位地址进入芯片。",
        "写题时先分清“片内地址线”和“片选译码地址线”。这是芯片扩展题最容易丢分的点。",
    ])

    add_heading(doc, "3. 多模块主存与交叉编址", 1)
    add_heading(doc, "3.1 知识点", 2)
    add_bullets(doc, [
        "多模块主存把主存分成多个可独立工作的模块，目的是提高带宽。",
        "高位交叉编址：连续地址大多在同一模块中，适合扩大容量，但连续访存并行性差。",
        "低位交叉编址：连续地址分布在不同模块中，适合流水方式连续访问，真题更爱考。",
        "低位交叉模块号常由地址低位决定：模块号 = 地址 mod 模块数。",
    ])
    add_heading(doc, "3.2 题型写法：连续读 n 个字需要多久", 2)
    add_numbers(doc, [
        "判断是高位交叉还是低位交叉。看到“连续地址轮流进入不同模块”，就是低位交叉。",
        "找两个时间：模块存取周期 T，模块启动间隔/总线传输时间 t。",
        "若模块数 m 足够满足 T <= m × t，则可流水访问，n 个字时间约为 T + (n - 1)t。",
        "若模块数不足，会产生模块冲突，需要按模块再次可用的时间逐个排。",
    ])
    add_note(doc, "易错点：", "低位交叉不是把一次访问变快，而是让连续多次访问能够重叠进行，提高吞吐率。单个字的访问延迟通常没有神奇缩短。", FILL_WARN)

    add_heading(doc, "4. Cache：本章最大分值区", 1)
    add_heading(doc, "4.1 Cache 基础", 2)
    add_bullets(doc, [
        "Cache 位于 CPU 和主存之间，通常由 SRAM 构成。",
        "主存与 Cache 之间按块交换。块越大，空间局部性利用越充分，但也可能增加替换代价和污染。",
        "命中率 h，缺失率 1-h。平均访问时间常写为：AMAT = 命中时间 + 缺失率 × 缺失代价。",
        "Cache 行通常包含：有效位、Tag、数据块、脏位、替换相关位。",
    ])

    add_heading(doc, "4.2 三种映射方式", 2)
    add_matrix(
        doc,
        ["映射方式", "放置规则", "地址结构", "查找方式", "优缺点"],
        [
            ["直接映射", "主存块只能放到固定 Cache 行；行号 = 主存块号 mod Cache 行数", "Tag + 行号 + 块内地址", "按行号定位 1 行，再比 Tag", "最快最简单；冲突多"],
            ["全相联映射", "主存块可放任意 Cache 行", "Tag + 块内地址", "所有行同时比较 Tag", "命中率高；硬件复杂"],
            ["组相联映射", "主存块固定到某组，组内任意行；组号 = 主存块号 mod 组数", "Tag + 组号 + 块内地址", "先找组，再组内比较 Tag", "折中，实际常用"],
        ],
        [2.4, 4.5, 3.0, 3.4, 2.9],
        font_size=8,
    )
    add_heading(doc, "4.3 Cache 地址划分题解法", 2)
    add_para(doc, "这是历年真题反复出现的计算题。固定写法如下。")
    add_numbers(doc, [
        "先求块内地址位数：块大小 = 2^b 字节，则块内地址 b 位。",
        "再求行号或组号位数。直接映射用 Cache 行数；组相联用组数 = Cache 行数 / 路数；全相联没有行号/组号。",
        "Tag 位数 = 主存地址总位数 - 块内地址位数 - 行号/组号位数。",
        "写命中判断时一定写：有效位为 1 且 Tag 匹配。",
    ])
    add_note(doc, "例题模板：", "32 位字节地址，Cache 容量 16KB，块大小 64B，4 路组相联。块内地址 = log2 64 = 6 位；Cache 行数 = 16KB/64B = 256 行；组数 = 256/4 = 64 组；组号 = 6 位；Tag = 32 - 6 - 6 = 20 位。", FILL_WARN)

    add_heading(doc, "4.4 Cache 命中过程题", 2)
    add_para(doc, "题目给一串地址，让你判断命中/未命中、替换过程或最终 Cache 状态。先把地址转成主存块号，再看映射。")
    add_numbers(doc, [
        "计算主存块号：主存块号 = 地址 / 块大小，块内偏移 = 地址 mod 块大小。",
        "直接映射：行号 = 主存块号 mod Cache 行数，只检查这一行。",
        "组相联：组号 = 主存块号 mod 组数，只在该组内检查和替换。",
        "全相联：在整个 Cache 中检查和替换。",
        "每次访问后更新替换信息。LRU 要更新最近访问顺序；FIFO 只看进入顺序。",
    ])
    add_note(doc, "答题书写建议：", "画表最稳。列出“访问块号、映射位置、命中/缺失、替换对象、访问后状态”。LRU/FIFO 题不要只在脑子里滚动，容易错。")

    add_heading(doc, "4.5 替换算法", 2)
    add_matrix(
        doc,
        ["算法", "淘汰依据", "适用范围", "真题提醒"],
        [
            ["直接覆盖", "直接映射下固定行被覆盖", "直接映射", "严格说不需要选择替换对象"],
            ["FIFO", "最早进入 Cache 的块", "全相联、组相联", "看进入顺序，不看最近访问"],
            ["LRU", "最近最久未访问的块", "全相联、组相联", "看最近访问顺序，依据时间局部性"],
            ["随机", "随机选一块", "了解", "选择题偶尔出现"],
        ],
        [2.0, 4.0, 3.0, 7.2],
        font_size=9,
    )

    add_heading(doc, "4.6 写策略与一致性", 2)
    add_matrix(
        doc,
        ["场景", "策略", "含义", "特点"],
        [
            ["写命中", "写直达", "同时写 Cache 和主存", "一致性好，写流量大"],
            ["写命中", "写回", "只写 Cache，替换时写回主存", "速度快，需要脏位"],
            ["写不命中", "写分配", "先调入 Cache 再写", "常与写回搭配"],
            ["写不命中", "非写分配", "直接写主存，不调入 Cache", "常与写直达搭配"],
        ],
        [2.0, 2.2, 6.2, 5.8],
        font_size=9,
    )
    add_bullets(doc, [
        "Cache 标记 Tag 解决“这行是不是我要的主存块”。",
        "有效位解决“这行数据是否有效”。命中条件是有效位为 1 且 Tag 匹配。",
        "脏位解决“写回法下 Cache 是否比主存新”。脏位为 1 的块被替换前必须写回主存。",
        "多处理器场景还会出现多个 Cache 副本之间的一致性问题，408 基础题通常只要求理解问题来源。",
    ])

    add_heading(doc, "4.7 Cache 平均访问时间题", 2)
    add_para(doc, "最常见公式：")
    add_para(doc, "AMAT = 命中时间 + 缺失率 × 缺失代价")
    add_numbers(doc, [
        "如果题目给命中率 h、Cache 时间 Tc、主存时间 Tm，要看“未命中访问时间”是否包含先访问 Cache 的时间。",
        "若缺失代价是额外代价：AMAT = Tc + (1-h) × 缺失代价。",
        "若未命中总时间单独给出：AMAT = h × 命中时间 + (1-h) × 未命中总时间。",
        "多级 Cache 按层展开：L1 未命中才访问 L2，L2 未命中才访问主存。",
    ])
    add_note(doc, "易错点：", "不要机械套公式。先判断题干中的“主存访问时间”是缺失后的额外时间，还是未命中全过程时间。", FILL_WARN)

    add_heading(doc, "5. 虚拟存储器：页表、地址转换、TLB 与缺页", 1)
    add_heading(doc, "5.1 核心概念", 2)
    add_bullets(doc, [
        "虚拟存储器把主存和辅存结合，为进程提供比实际主存更大的虚拟地址空间。",
        "分页系统中，虚拟空间划分为页，物理主存划分为页框，页和页框大小相同。",
        "页表记录虚页号到物理页框号的映射。每个进程通常有自己的页表。",
        "TLB 是页表项的高速缓存，缓存的是地址映射关系，不是数据块。",
        "缺页是页表表明目标页当前不在主存，需要缺页中断处理。",
    ])
    add_heading(doc, "5.2 页表项 PTE", 2)
    add_matrix(
        doc,
        ["字段", "作用", "常见考法"],
        [
            ["页框号", "指出该虚页在主存中的物理页框", "地址转换必用"],
            ["存在位/有效位", "表示该页是否在主存", "判断是否缺页"],
            ["访问位", "表示近期是否被访问", "CLOCK 等置换算法"],
            ["修改位/脏位", "表示调入后是否被写过", "换出时是否写回辅存"],
            ["保护位", "读/写/执行权限控制", "概念题"],
        ],
        [2.4, 7.0, 6.8],
        font_size=9,
    )

    add_heading(doc, "5.3 虚拟地址到物理地址转换", 2)
    add_para(doc, "分页地址结构：")
    add_para(doc, "虚拟地址 = 虚页号 + 页内偏移")
    add_para(doc, "物理地址 = 物理页框号 + 页内偏移")
    add_note(doc, "必背结论：", "分页系统中，地址转换前后页内偏移不变。真正被替换的是高位部分：虚页号通过页表变成页框号。", FILL_WARN)
    add_numbers(doc, [
        "页面大小为 2^b 字节，则页内偏移 b 位。",
        "虚拟地址总位数为 n，则虚页号位数 n-b，虚页数 2^(n-b)。",
        "物理地址总位数由主存容量决定，若主存容量为 2^m 字节，则物理地址 m 位。",
        "物理页框号位数 = 物理地址位数 - 页内偏移位数。",
    ])
    add_heading(doc, "5.4 TLB、页表、缺页的访问流程", 2)
    add_numbers(doc, [
        "CPU 产生虚拟地址，拆成页号和页内偏移。",
        "先查 TLB。若 TLB 命中，直接得到页框号，拼接物理地址并访问主存。",
        "若 TLB 未命中，去主存查页表。",
        "若页表存在位为 1，说明页在主存，取页框号，必要时把页表项装入 TLB。",
        "若页表存在位为 0，发生缺页中断，由操作系统从辅存调页。",
    ])
    add_note(doc, "最高频辨析：", "TLB miss 不等于缺页。TLB miss 只是快表里没有映射；缺页是页表说明该页不在主存。", FILL_WARN)

    add_heading(doc, "5.5 地址转换题型写法", 2)
    add_para(doc, "题干常给虚拟地址、页大小、页表或 TLB 内容，要求判断物理地址或是否缺页。")
    add_numbers(doc, [
        "先根据页面大小划分虚拟地址：高位页号，低位页内偏移。",
        "用页号查 TLB。命中则直接得到页框号；未命中再查页表。",
        "查页表时先看存在位。存在位为 0，不要继续拼物理地址，直接写“缺页”。",
        "存在位为 1 时，用页框号替换虚页号，页内偏移原样保留。",
        "若题目要求十六进制地址，最后把页框号和偏移按位拼接，不要用十进制随意相加导致位数错。",
    ])
    add_note(doc, "例题模板：", "页面大小 4KB，则偏移 12 位。虚拟地址 0x12345 的页号是 0x12，页内偏移是 0x345。若页表显示虚页 0x12 -> 页框 0x08，则物理地址为 0x08345。")

    add_heading(doc, "5.6 TLB 有效访问时间 EAT", 2)
    add_para(doc, "设 TLB 访问时间为 a，主存访问时间为 t，TLB 命中率为 h，且不考虑缺页：")
    add_para(doc, "EAT = h(a + t) + (1 - h)(a + 2t)")
    add_bullets(doc, [
        "TLB 命中：查 TLB + 访问目标主存。",
        "TLB 未命中但不缺页：查 TLB + 查主存中的页表 + 访问目标主存。",
        "如果题目说明 TLB 时间可忽略，则命中约为 t，未命中约为 2t。",
        "如果考虑缺页率 p，则总时间约为 (1-p) × 正常访问时间 + p × 缺页处理时间。",
    ])
    add_note(doc, "易错点：", "缺页处理时间通常远大于普通访存时间，缺页率即使很小，也会显著拉高平均访问时间。", FILL_WARN)

    add_heading(doc, "5.7 页面置换与缺页过程", 2)
    add_bullets(doc, [
        "缺页中断处理：保存现场 -> 找到外存页面 -> 找空闲页框或置换旧页 -> 必要时写回脏页 -> 调入新页 -> 修改页表/TLB -> 重新执行指令。",
        "OPT：淘汰未来最长时间不用的页，理论最优，用来比较。",
        "FIFO：淘汰最早进入主存的页，可能出现 Belady 异常。",
        "LRU：淘汰最近最久未访问的页，一般不出现 Belady 异常。",
        "CLOCK：利用访问位近似实现 LRU，常作为工程折中。",
    ])
    add_heading(doc, "5.8 页面置换题写法", 2)
    add_numbers(doc, [
        "画页框表，列按访问序列推进，行按页框数排列。",
        "每访问一个页号，先判断是否已在页框中。在则命中，不增加缺页次数。",
        "不在则缺页。若有空页框，直接放入；若满了，按算法选牺牲页。",
        "FIFO 维护进入队列；LRU 维护最近访问顺序；OPT 向后看未来访问序列。",
        "最后统计缺页次数、缺页率，并按题目要求判断是否有 Belady 异常。",
    ])

    add_heading(doc, "6. 真题题型总表：看到题目就知道怎么下手", 1)
    add_matrix(
        doc,
        ["题型", "识别信号", "解题步骤", "常见坑"],
        [
            ["芯片扩展", "给芯片规格和目标容量", "先位扩展，再字扩展，最后算片选", "把 bit 和 Byte 混用"],
            ["DRAM 刷新", "给行数、刷新周期、存取周期", "算每行刷新间隔或刷新开销", "把按行刷新误当按单元刷新"],
            ["交叉编址", "给模块数、存取周期、连续地址", "判断低位/高位，再排流水时间", "忽略模块冲突"],
            ["Cache 地址划分", "给地址位数、Cache 容量、块大小、相联度", "块内位 -> 组/行位 -> Tag 位", "组数算成行数"],
            ["Cache 命中模拟", "给地址序列和映射方式", "转块号，按映射查位置，记录状态", "LRU/FIFO 混淆"],
            ["Cache 写策略", "出现写命中/写不命中/脏位", "判断写直达/写回、写分配/非写分配", "忘记脏块替换前写回"],
            ["AMAT", "给命中率、访问时间", "先判定缺失代价含义，再套期望公式", "重复或漏算 Cache 时间"],
            ["页表转换", "给虚拟地址、页大小、页表", "拆页号偏移，查存在位，拼物理地址", "TLB miss 当缺页"],
            ["TLB 时间", "给 TLB 命中率和主存时间", "命中 a+t，未命中 a+2t", "漏掉查页表那次主存访问"],
            ["页面置换", "给页号访问串和页框数", "画表模拟 FIFO/LRU/OPT", "FIFO 看进入顺序，不看访问顺序"],
        ],
        [2.2, 3.2, 6.2, 4.6],
        font_size=8,
    )

    add_heading(doc, "7. 综合对比：Cache、TLB、页表、虚拟存储器", 1)
    add_matrix(
        doc,
        ["对象", "缓存/映射什么", "位置", "未命中含义", "典型处理"],
        [
            ["Cache", "主存数据块", "CPU 与主存之间", "Cache 中没有目标块", "去主存取块，可能替换"],
            ["TLB", "页表项，即页号到页框号映射", "地址变换部件中", "快表没有该映射", "查页表，可能更新 TLB"],
            ["页表", "完整虚页到页框映射及状态", "主存中，由 OS 维护", "不叫页表未命中；看存在位", "存在位 0 则缺页"],
            ["虚拟存储器", "主存与辅存之间的页/段调度", "主存与辅存之间", "目标页不在主存", "缺页中断，从辅存调页"],
        ],
        [2.2, 4.8, 3.0, 3.4, 3.0],
        font_size=8,
    )

    add_heading(doc, "8. 考前公式与结论清单", 1)
    add_bullets(doc, [
        "存储层次：寄存器 -> Cache -> 主存 -> 辅存。",
        "局部性：时间局部性支持 LRU 和近期复用；空间局部性支持按块/页调入。",
        "芯片总数 = 位扩展倍数 × 字扩展倍数。",
        "Cache 行数 = Cache 容量 / 块大小。",
        "组数 = Cache 行数 / 相联路数。",
        "直接映射地址：Tag + 行号 + 块内地址。",
        "组相联地址：Tag + 组号 + 块内地址。",
        "全相联地址：Tag + 块内地址。",
        "Cache 命中条件：有效位为 1 且 Tag 匹配。",
        "AMAT = 命中时间 + 缺失率 × 缺失代价。",
        "页内偏移位数 = log2 页面大小。",
        "虚拟地址 = 页号 + 页内偏移；物理地址 = 页框号 + 页内偏移。",
        "TLB miss 不等于缺页；缺页要看页表存在位。",
        "TLB EAT = h(a+t) + (1-h)(a+2t)。",
        "FIFO 可能出现 Belady 异常，LRU/OPT 通常不会。",
    ])

    add_heading(doc, "9. 易错点集中纠偏", 1)
    add_bullets(doc, [
        "Cache 解决速度问题，不是容量问题；虚拟存储器解决逻辑容量问题，不是让物理主存真的变大。",
        "直接映射不是不替换，而是没有选择余地，固定覆盖。",
        "组相联替换只发生在对应组内，不是在整个 Cache 里任选。",
        "LRU 看最近访问，FIFO 看进入时间；两者模拟表不能共用同一套顺序。",
        "Tag 匹配还不够，必须有效位为 1 才命中。",
        "写回法下主存可能暂时是旧值，脏位为 1 的块换出前要写回。",
        "页内偏移永远不参与页表映射，转换前后保持不变。",
        "TLB 是页表项缓存，不是数据 Cache。",
        "缺页中断后通常重新执行引起缺页的指令。",
        "做平均访问时间题，先判断题干给的是额外缺失代价还是未命中总时间。",
    ])

    add_heading(doc, "10. 建议复习顺序", 1)
    add_numbers(doc, [
        "第一轮：把存储层次、SRAM/DRAM、主存扩展、Cache 映射、虚拟地址转换全部过一遍。",
        "第二轮：集中刷 Cache 地址划分、命中模拟、替换算法、平均访问时间。",
        "第三轮：把 TLB、页表、缺页、页面置换与操作系统章节一起联动复习。",
        "考前：只看公式清单、易错点和题型总表，再做 2 到 3 套真题中的存储系统题。",
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()

