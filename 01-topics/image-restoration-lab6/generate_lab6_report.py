import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(r"D:\moniC\project\learn")
OUTPUT_DIR = ROOT / "03-outputs" / "image-restoration-lab6"
SUMMARY_PATH = OUTPUT_DIR / "summary.json"
REPORT_PATH = OUTPUT_DIR / "实验6_图像复原_自动生成报告.docx"
MATLAB_SCRIPT_PATH = ROOT / "01-topics" / "image-restoration-lab6" / "matlab" / "run_lab6_image_restoration.m"


def set_font(run, name_cn="宋体", name_en="Times New Roman", size=Pt(11), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = name_en
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name_cn)
    r_fonts.set(qn("w:ascii"), name_en)
    r_fonts.set(qn("w:hAnsi"), name_en)


def add_paragraph(doc, text, *, font_cn="宋体", font_en="Times New Roman", size=Pt(11),
                  bold=False, alignment=None, spacing_after=Pt(4), first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = spacing_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def add_title(doc, text):
    return add_paragraph(
        doc,
        text,
        font_cn="黑体",
        size=Pt(18),
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=Pt(10),
    )


def add_heading(doc, text):
    return add_paragraph(doc, text, font_cn="黑体", size=Pt(12), bold=True, spacing_after=Pt(3))


def add_subheading(doc, text):
    return add_paragraph(doc, text, font_cn="黑体", size=Pt(11), bold=True, spacing_after=Pt(2))


def add_body(doc, text):
    return add_paragraph(doc, text, size=Pt(10.5), spacing_after=Pt(3), first_line_indent=Cm(0.74))


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    for idx, line in enumerate(text.strip("\n").splitlines()):
        run = p.add_run(line)
        set_font(run, name_cn="Consolas", name_en="Consolas", size=Pt(8.5), bold=False)
        if idx < len(text.strip("\n").splitlines()) - 1:
            run.add_break()
    return p


def add_image(doc, path, width=Inches(5.6)):
    if not path.exists():
        add_body(doc, f"[缺少图片：{path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)


def fmt_number(value, digits=3):
    if value is None:
        return "-"
    return f"{value:.{digits}f}"


def fmt_sci(value):
    if value is None:
        return "-"
    return f"{value:.0e}"


def inverse_better(inv, wnr):
    return inv["psnr"] > wnr["psnr"]


def matlab_snippets():
    return [
        (
            "主程序入口",
            """root_dir = 'D:\\moniC\\project\\learn';
input_dir = fullfile(root_dir, '01-topics', 'image-restoration-lab6', 'input-images');
output_dir = fullfile(root_dir, '03-outputs', 'image-restoration-lab6');

motion_original = load_gray_image(fullfile(input_dir, 'Fig0526(a)(original_DIP).tif'));
turb_original = load_gray_image(fullfile(input_dir, 'Fig0525(a)(aerial_view_no_turb).tif'));
turb_degraded = load_gray_image(fullfile(input_dir, 'Fig0525(b)(aerial_view_turb_c_0pt0025).tif'));

summary.task1 = task1_motion_blur_demo(motion_original, output_dir);
summary.task2 = task2_turbulence_demo(turb_original, output_dir);
summary.task3 = task3_restore_turbulence(turb_original, turb_degraded, output_dir);
summary.task4 = task4_restore_motion_noise(motion_original, motion_noisy, output_dir);""",
        ),
        (
            "运动模糊与大气湍流模型",
            """function H = motion_psf_freq(M, N, a, b, T)
[U, V] = centered_frequency_grid(M, N);
uv = (U * a + V * b) * T;
H = exp(-1i * pi * uv) .* sin(pi * uv) ./ (pi * uv + eps);
H(abs(uv) < 1e-12) = 1;
end

function H = turbulence_psf_freq(M, N, k)
[U, V] = pixel_frequency_grid(M, N);
H = exp(-k .* ((U .^ 2 + V .^ 2) .^ (5 / 6)));
end""",
        ),
        (
            "逆滤波与维纳滤波",
            """function restored = inverse_restore(img, H, cutoff, threshold)
[U, V] = centered_frequency_grid(size(img,1), size(img,2));
D = sqrt(U .^ 2 + V .^ 2);
F = fftshift(fft2(img));
mask = (abs(H) >= threshold) & (D <= cutoff);
invH = zeros(size(H));
invH(mask) = 1 ./ H(mask);
restored = real(ifft2(ifftshift(F .* invH)));
restored = clamp01(restored);
end

function restored = wiener_restore(img, H, K, cutoff)
[U, V] = centered_frequency_grid(size(img,1), size(img,2));
D = sqrt(U .^ 2 + V .^ 2);
W = conj(H) ./ (abs(H) .^ 2 + K);
W(D > cutoff) = 0;
restored = real(ifft2(ifftshift(fftshift(fft2(img)) .* W)));
restored = clamp01(restored);
end""",
        ),
        (
            "任务4 参数搜索",
            """length_values = [10, 15, 20, 25, 30, 35, 40, 45];
theta_values = [30, 45, 60, 120, 135, 150];
cutoffs = [80, 120, 160, 220];
Ks = [1e-6, 1e-5, 1e-4, 1e-3];

for len = length_values
    for theta = theta_values
        psf = fspecial('motion', len, theta);
        H = fftshift(psf2otf(psf, size(original_img)));
        for cutoff = cutoffs
            for K = Ks
                restored = wiener_restore(least_noisy_img, H, K, cutoff);
            end
        end
    end
end""",
        ),
    ]


def make_doc():
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    if MATLAB_SCRIPT_PATH.exists():
        shutil.copy2(MATLAB_SCRIPT_PATH, OUTPUT_DIR / MATLAB_SCRIPT_PATH.name)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_title(doc, "南京信息工程大学  实验（实习）报告")
    add_paragraph(
        doc,
        "实验（实习）名称：  图像复原              实验（实习）日期：               得分：        指导教师：",
        size=Pt(10.5),
        spacing_after=Pt(2),
    )
    add_paragraph(
        doc,
        "学院：  计算机学院          专业：               年级：           班级：          姓名：              学号：",
        size=Pt(10.5),
        spacing_after=Pt(10),
    )

    add_heading(doc, "一、实验目的：")
    add_body(doc, "1. 理解图像复原的原理，掌握图像退化模型估计方法。")
    add_body(doc, "2. 掌握逆滤波和维纳滤波图像复原方法。")

    add_heading(doc, "二、实验内容：")
    add_body(doc, "1. 编程模拟图像运动模糊，对原始图像进行运动模糊处理，并比较不同参数 a、b、T 对模糊效果的影响。")
    add_body(doc, "2. 编程模拟大气湍流模糊，对原始图像进行大气湍流退化，并比较不同参数 k 对模糊效果的影响。")
    add_body(doc, "3. 分别使用逆滤波和维纳滤波复原大气湍流模糊图像，选择合适的参数获取较好效果。")
    add_body(doc, "4. 使用逆滤波和维纳滤波复原运动模糊带噪图像，比较不同噪声强度下两种方法的表现。")

    add_heading(doc, "三、实验过程与结果分析：")

    add_subheading(doc, "1. 运动模糊模拟")
    add_body(doc, "运动模糊的频域退化模型可表示为 H(u,v)。其中 a、b 共同控制运动方向，T 控制运动持续时间。实验中固定原图不变，分别调节 a、b、T，观察模糊方向与模糊长度的变化。")
    add_body(doc, "由结果可见，当 a 与 b 同时增大时，图像沿对应方向的拖尾更加明显；当 T 增大时，边缘细节进一步被平滑，整幅图像清晰度下降。该实验说明运动模糊与运动轨迹参数之间存在直接对应关系。")
    add_image(doc, OUTPUT_DIR / "task1_motion_blur_demo.png", width=Inches(5.8))

    add_subheading(doc, "2. 大气湍流模糊模拟")
    add_body(doc, "大气湍流退化常用频域模型为 H(u,v)=exp[-k(u^2+v^2)^(5/6)]。参数 k 反映湍流强度，k 越大，图像高频衰减越强。")
    add_body(doc, "实验结果表明，随着 k 从 0.0005 增大到 0.05，图像由轻微模糊逐步过渡到严重模糊，地物边界和纹理细节大量丢失。这与大气湍流优先破坏高频信息的理论一致。")
    add_image(doc, OUTPUT_DIR / "task2_turbulence_demo.png", width=Inches(5.8))

    inv3 = summary["task3"]["best_inverse"]
    wnr3 = summary["task3"]["best_wiener"]
    add_subheading(doc, "3. 大气湍流模糊图像复原")
    add_body(
        doc,
        "对给定的湍流模糊图像分别使用逆滤波和维纳滤波复原。自动搜索结果显示，"
        f"逆滤波最佳参数为 cutoff={inv3['cutoff']}、threshold={fmt_sci(inv3['threshold'])}，"
        f"对应 PSNR={fmt_number(inv3['psnr'], 2)} dB；"
        f"维纳滤波最佳参数为 cutoff={wnr3['cutoff']}、K={fmt_sci(wnr3['K'])}，"
        f"对应 PSNR={fmt_number(wnr3['psnr'], 2)} dB。"
    )
    add_body(doc, "从结果可以看出，逆滤波能够恢复部分细节，但对频域零点附近非常敏感，容易产生不稳定放大；维纳滤波通过引入平衡项 K，在复原清晰度和抑制噪声之间取得了更好的折中，因此最终 PSNR 略高，视觉效果也更自然。")
    add_image(doc, OUTPUT_DIR / "task3_turbulence_restoration.png", width=Inches(5.8))

    add_subheading(doc, "4. 运动模糊带噪图像复原")
    est = summary["task4"]["estimated_degradation"]
    add_body(
        doc,
        f"由于题目给出的三幅运动模糊图像同时叠加了不同强度噪声，先利用低噪声图像估计线性运动模糊核。自动搜索得到的较优 PSF 参数为 len={est['len']}、theta={est['theta']}，辅助复原参数为 cutoff={est['cutoff']}、K={fmt_sci(est['K'])}。"
    )
    add_image(doc, OUTPUT_DIR / "task4_motion_parameter_estimation.png", width=Inches(5.6))
    labels = {
        "noisiest": "高噪声图像",
        "medium": "中噪声图像",
        "least": "低噪声图像",
    }
    for key in ("noisiest", "medium", "least"):
        inv = summary["task4"][key]["best_inverse"]
        wnr = summary["task4"][key]["best_wiener"]
        better = "逆滤波略优" if inverse_better(inv, wnr) else "维纳滤波略优"
        add_body(
            doc,
            f"{labels[key]}的自动搜索结果为：逆滤波最佳 PSNR={fmt_number(inv['psnr'], 2)} dB，"
            f"参数 cutoff={inv['cutoff']}，threshold={fmt_sci(inv['threshold'])}；"
            f"维纳滤波最佳 PSNR={fmt_number(wnr['psnr'], 2)} dB，"
            f"参数 cutoff={wnr['cutoff']}，K={fmt_sci(wnr['K'])}。该组结果中 {better}。"
        )
    add_body(doc, "综合三组结果可见，这批运动模糊带噪图像的退化程度较重，尤其是高噪声图像中噪声纹理占主导地位，因此两种方法都难以恢复到理想清晰度。低噪声图像上两者表现接近，而在中高噪声情况下，复原结果更依赖退化模型的准确性。这个现象说明：当噪声较强且模糊核估计存在偏差时，传统频域复原方法的提升空间会明显受限。")
    add_image(doc, OUTPUT_DIR / "task4_motion_noise_restoration.png", width=Inches(5.9))

    add_heading(doc, "四、关键程序说明：")
    add_body(doc, "本实验使用 Matlab 实现全部功能，程序结构分为四个部分：图像读取、退化模型构造、复原算法实现和结果保存。")
    add_body(doc, "其中运动模糊与大气湍流模糊均在频域中构造退化函数；逆滤波通过 G(u,v)/H(u,v) 恢复原图，但加入阈值和截止频率以避免 H(u,v) 过小导致数值发散；维纳滤波使用 H*(u,v)/(|H(u,v)|^2+K) 形式，在抑制噪声和恢复细节之间进行折中。")
    add_body(doc, "程序还对多组参数进行自动搜索，并将最佳参数、PSNR、MSE 以及结果图统一输出，便于直接用于报告撰写。")
    add_body(doc, "报告中给出关键 Matlab 代码片段如下，完整代码文件可与实验结果一起单独提交。")
    for title, snippet in matlab_snippets():
        add_subheading(doc, title)
        add_code_block(doc, snippet)

    add_heading(doc, "五、实验总结与体会：")
    add_body(doc, "本实验围绕图像退化模型和频域复原方法展开，通过运动模糊和大气湍流模糊的模拟，我进一步理解了图像退化可以表示为原图像与退化函数的卷积，并可在频域中转化为乘法运算。")
    add_body(doc, "在复原阶段，逆滤波方法理论上最直接，但对噪声和退化模型误差非常敏感；维纳滤波则通过引入平衡项提高了稳定性，因此在大气湍流模糊图像复原中取得了更好的结果。")
    add_body(doc, "对于运动模糊带噪图像，本实验也观察到一个很重要的现实问题：如果噪声过强或者模糊核估计不够准确，传统逆滤波和维纳滤波都难以获得理想复原效果。这说明图像复原不仅依赖算法本身，也强烈依赖退化模型估计的准确程度。")
    add_body(doc, "通过本次实验，我掌握了运动模糊、大气湍流模糊、逆滤波和维纳滤波的基本实现流程，并进一步理解了参数选择、截止频率控制和噪声抑制在图像复原中的重要作用。")

    doc.save(REPORT_PATH)
    print(f"Report saved to: {REPORT_PATH}")


if __name__ == "__main__":
    make_doc()
