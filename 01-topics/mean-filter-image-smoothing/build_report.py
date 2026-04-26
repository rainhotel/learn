from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from PIL import Image, ImageOps


TOPIC_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = TOPIC_DIR / "output"
FIGURE_DIR = OUTPUT_DIR / "figures"
REPORT_DIR = OUTPUT_DIR / "report-qa"
METRICS_PATH = OUTPUT_DIR / "metrics.json"
REPORT_MD = TOPIC_DIR / "report-draft.md"
REPORT_DOCX = TOPIC_DIR / "mean-filter-report.docx"


def load_metrics() -> dict:
    with METRICS_PATH.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def stack_images_vertically(paths: list[Path], out_path: Path, background: str = "white") -> None:
    images = [Image.open(path).convert("RGB") for path in paths]
    max_width = max(image.width for image in images)
    padded_images = []
    for image in images:
        if image.width != max_width:
            image = ImageOps.pad(image, (max_width, image.height), color=background)
        padded_images.append(image)
    total_height = sum(image.height for image in padded_images) + 18 * (len(padded_images) - 1)
    canvas = Image.new("RGB", (max_width, total_height), background)
    y = 0
    for index, image in enumerate(padded_images):
        canvas.paste(image, (0, y))
        y += image.height
        if index < len(padded_images) - 1:
            y += 18
    canvas.save(out_path)


def apply_font(run, font_name: str = "SimSun", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(3)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        apply_font(run, font_name="Microsoft YaHei", size=12.5, bold=True)
    else:
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        apply_font(run, font_name="Microsoft YaHei", size=11.2, bold=True)


def add_paragraph(document: Document, text: str, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = Mm(7) if indent else Mm(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    apply_font(run)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    apply_font(run, size=9.5)


def add_figure(document: Document, image_path: Path, caption: str, width_in_inches: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_in_inches))
    add_caption(document, caption)


def add_cover_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("均值类滤波器图像平滑实验报告")
    apply_font(run, font_name="Microsoft YaHei", size=15, bold=True)

    lines = [
        "实验（实习）名称：均值类滤波器图像平滑    实验（实习）日期：2026.4.26    指导教师：范春年",
        "学院：计算机学院    专业：计算机科学与技术    年级：2024    班级：4",
        "姓名：李林浩    学号：202483290054",
    ]
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        apply_font(run, size=9.8)


def add_code_snippets(document: Document) -> None:
    snippets = [
        (
            "1. 高斯滤波器系数计算",
            "def gaussian_kernel(size, sigma):\n"
            "    radius = size // 2\n"
            "    axis = np.arange(-radius, radius + 1)\n"
            "    xx, yy = np.meshgrid(axis, axis)\n"
            "    kernel = np.exp(-(xx**2 + yy**2) / (2 * sigma**2))\n"
            "    return kernel / kernel.sum()",
        ),
        (
            "2. 双边滤波核心权重",
            "weights = spatial * np.exp(-((windows - center)**2) / (2 * sigma_range**2))\n"
            "result = np.sum(weights * windows) / np.sum(weights)",
        ),
        (
            "3. 逆谐波均值滤波",
            "numerator = np.sum((windows + eps) ** (Q + 1), axis=(2, 3))\n"
            "denominator = np.sum((windows + eps) ** Q, axis=(2, 3))\n"
            "result = numerator / (denominator + eps)",
        ),
    ]

    for title, code in snippets:
        add_paragraph(document, title, indent=False)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(6)
        paragraph.paragraph_format.space_after = Pt(3)
        for line in code.splitlines():
            run = paragraph.add_run(line)
            apply_font(run, font_name="Consolas", size=8.2)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.add_break()


def build_markdown(metrics: dict) -> None:
    task1 = metrics["task1"]
    task2 = metrics["task2"]
    task3 = metrics["task3"]
    task4 = metrics["task4"]
    task5 = metrics["task5"]

    text = f"""# 均值类滤波器图像平滑实验报告草稿

## 封面信息

- 实验（实习）名称：均值类滤波器图像平滑
- 实验（实习）日期：2026.4.26
- 指导教师：范春年
- 学院：计算机学院
- 专业：计算机科学与技术
- 年级：2024
- 班级：4
- 姓名：李林浩
- 学号：202483290054

## 一、实验目的

1. 理解模板操作的基本原理。
2. 了解高斯噪声、盐噪声和椒噪声的特点。
3. 掌握算术均值、加权均值、高斯均值、几何均值、谐波均值和逆谐波均值滤波器的实现方法。
4. 能结合实验结果分析不同滤波器在去噪与保边方面的差异。

## 二、实验原理

均值类滤波器的核心思想是：在图像上滑动一个局部模板，用模板内像素的统计平均值替换中心像素，从而削弱随机噪声。模板越大，参与平均的像素越多，去噪能力通常越强，但图像边缘和细节也更容易被模糊。不同滤波器之间的主要差异，在于邻域像素的权重分配方式不同。

高斯滤波器采用距离中心越远、权重越小的高斯分布；双边滤波除了考虑空间距离，还考虑灰度差异，因此能在平滑噪声的同时更好地保护边缘。几何均值、谐波均值和逆谐波均值属于不同形式的统计平均，其中逆谐波均值滤波器通过参数 Q 调节对不同类型脉冲噪声的抑制方向：一般来说，Q>0 更适合抑制黑色椒噪声，Q<0 更适合抑制白色盐噪声。

## 三、实验内容与结果分析

### 1. 高斯滤波器系数计算与参数影响

在同一 sigma=1.0 下比较 3x3、5x5 和 7x7 模板时，PSNR 分别为 {task1['size_3_sigma_1.0']['psnr_vs_orig']:.2f} dB、{task1['size_5_sigma_1.0']['psnr_vs_orig']:.2f} dB 和 {task1['size_7_sigma_1.0']['psnr_vs_orig']:.2f} dB，可见模板越大，平滑更强，但细节略有下降。固定模板为 5x5 时，sigma=0.6、1.0 和 1.6 的 PSNR 分别为 {task1['size_5_sigma_0.6']['psnr_vs_orig']:.2f} dB、{task1['size_5_sigma_1.0']['psnr_vs_orig']:.2f} dB 和 {task1['size_5_sigma_1.6']['psnr_vs_orig']:.2f} dB，说明 sigma 过大时也会明显加重模糊。

### 2. 双边滤波与高斯滤波比较

对 Fig0507(b) 进行处理时，高斯滤波的 PSNR 为 {task2['gaussian_psnr_vs_orig']:.2f} dB，双边滤波的 PSNR 为 {task2['bilateral_psnr_vs_orig']:.2f} dB。实验结果表明，双边滤波在滤除噪声的同时能更好地保持电路板导线和芯片边缘，综合效果优于普通高斯滤波。

### 3. 三种均值滤波器对高斯噪声的去噪效果

当噪声方差为 0.001 时，原始噪声图像的 PSNR 已达到 {task3['variance_0.001']['noisy_psnr']:.2f} dB，此时继续平滑反而会损失细节；当噪声方差升高到 0.005 和 0.010 时，高斯均值滤波器分别达到 {task3['variance_0.005']['gaussian_psnr']:.2f} dB 和 {task3['variance_0.010']['gaussian_psnr']:.2f} dB，优于加权均值和算术均值滤波器。对于本实验图像，在方差 0.005 条件下继续比较模板大小时，3x3 模板在三类滤波器上都取得了最高 PSNR，说明该图像中较细的线路纹理不适合使用过大的平滑窗口。

### 4. 算术均值与几何均值滤波比较

对 Fig0507(b) 使用 5x5 模板处理时，算术均值滤波 PSNR 为 {task4['arithmetic_psnr_vs_orig']:.2f} dB，几何均值滤波 PSNR 为 {task4['geometric_psnr_vs_orig']:.2f} dB。两者差距不大，但几何均值滤波在芯片边缘和导线纹理处保留了略多的细节，因此主观视觉效果更自然。

### 5. 逆谐波均值滤波参数 Q 的选择

对于 Fig0508(a) 的椒噪声图像，Q=0.5、1.5 和 3.0 的 PSNR 依次为 {task5['pepper']['Q_0.5']:.2f} dB、{task5['pepper']['Q_1.5']:.2f} dB 和 {task5['pepper']['Q_3.0']:.2f} dB，表明正 Q 的方向是正确的，但 Q 过大时会导致图像发白和细节损失。对于 Fig0508(b) 的盐噪声图像，Q=-0.5、-1.5 和 -3.0 的 PSNR 依次为 {task5['salt']['Q_-0.5']:.2f} dB、{task5['salt']['Q_-1.5']:.2f} dB 和 {task5['salt']['Q_-3.0']:.2f} dB，说明负 Q 更适合抑制白色盐噪声，且中等绝对值的参数更稳妥。

### 6. Fig0512(b) 上多种均值滤波器的综合比较

对含有混合噪声的 Fig0512(b) 进行实验后可以看到：算术均值滤波对整体噪声抑制最稳定，但仍有一定模糊；几何均值和谐波均值在混合盐椒噪声条件下容易产生明显的亮度失衡；逆谐波均值在 Q=1.5 时能够明显抑制黑色噪点，但对白色噪点无能为力。这说明没有一种均值类滤波器能在所有噪声类型上都占优，参数和噪声模型需要匹配。

## 四、实验总结与体会

通过本次实验，我系统理解了均值类滤波器的共同点和差异。算术均值滤波器结构最简单，但容易模糊；加权均值和高斯均值在保持自然过渡方面更有优势；双边滤波通过同时利用空间距离和灰度相似性，在保边去噪方面效果最好。逆谐波均值滤波则进一步说明，滤波器参数必须结合噪声类型进行选择，不能脱离图像特征机械套用。

从实验结果来看，去噪和保边之间始终存在权衡。当噪声较弱时，过强的平滑会破坏细节；当噪声较重时，合适的加权方式和模板大小可以明显提高恢复效果。本次实验不仅帮助我掌握了几类常见平滑算法的原理，也让我学会了通过结果图和定量指标共同分析滤波效果。完整实验程序已单独保存在 `run_lab.py` 中，可作为提交附件。
"""
    REPORT_MD.write_text(text, encoding="utf-8")


def build_docx(metrics: dict) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    task1_combo = REPORT_DIR / "task1_report_combo.png"
    task5_combo = REPORT_DIR / "task5_report_combo.png"
    stack_images_vertically(
        [
            FIGURE_DIR / "task1_kernel_heatmap.png",
            FIGURE_DIR / "task1_filter_size_compare.png",
            FIGURE_DIR / "task1_filter_sigma_compare.png",
        ],
        task1_combo,
    )
    stack_images_vertically(
        [
            FIGURE_DIR / "task5_pepper_q_compare.png",
            FIGURE_DIR / "task5_salt_q_compare.png",
        ],
        task5_combo,
    )

    document = Document()
    configure_document(document)
    add_cover_block(document)

    add_heading(document, "一、实验目的", level=1)
    add_paragraph(document, "理解模板操作的基本原理，掌握均值类图像平滑算法的实现过程，并能够结合结果图像分析模板大小、权重分布以及参数选择对去噪与保边效果的影响。")
    add_paragraph(document, "在完成高斯滤波、双边滤波、算术均值、加权均值、几何均值、谐波均值和逆谐波均值滤波实验的基础上，进一步比较它们对高斯噪声、盐噪声、椒噪声和混合噪声的适应性。")

    add_heading(document, "二、实验原理", level=1)
    add_paragraph(document, "均值类滤波器的基本思想是利用局部邻域像素的统计平均值替换中心像素，以削弱随机噪声。模板越大，参与平均的像素越多，噪声抑制能力通常越强，但图像边缘和细节也更容易被模糊。")
    add_paragraph(document, "高斯滤波器采用随空间距离衰减的高斯权重；双边滤波同时考虑空间距离和灰度差异，因此在平滑噪声时能较好保持边缘。几何均值、谐波均值和逆谐波均值属于不同形式的统计平均，其中逆谐波均值滤波器可通过参数 Q 对盐噪声和椒噪声进行定向抑制。")

    add_heading(document, "三、实验内容与结果分析", level=1)

    add_heading(document, "1. 高斯滤波器系数与参数影响", level=2)
    add_paragraph(document, f"实验首先计算不同模板尺寸下的高斯滤波器系数，并对其进行可视化。固定 sigma=1.0 时，3x3、5x5 和 7x7 模板对应的 PSNR 分别为 {metrics['task1']['size_3_sigma_1.0']['psnr_vs_orig']:.2f} dB、{metrics['task1']['size_5_sigma_1.0']['psnr_vs_orig']:.2f} dB 和 {metrics['task1']['size_7_sigma_1.0']['psnr_vs_orig']:.2f} dB，说明模板增大后平滑增强，但细节略有损失。")
    add_paragraph(document, f"固定模板为 5x5 时，sigma 从 0.6 增大到 1.6，PSNR 从 {metrics['task1']['size_5_sigma_0.6']['psnr_vs_orig']:.2f} dB 下降到 {metrics['task1']['size_5_sigma_1.6']['psnr_vs_orig']:.2f} dB，说明过大的 sigma 会导致权重分布过于平缓，边缘细节被进一步模糊。")
    add_figure(document, task1_combo, "图1 高斯滤波器系数可视化及其模板大小、sigma 参数对 Fig0507(b) 滤波结果的影响", 6.15)

    add_heading(document, "2. 双边滤波与高斯滤波比较", level=2)
    add_paragraph(document, f"对 Fig0507(b) 进行去噪时，高斯滤波 5x5、sigma=1.0 的 PSNR 为 {metrics['task2']['gaussian_psnr_vs_orig']:.2f} dB，双边滤波 5x5 的 PSNR 为 {metrics['task2']['bilateral_psnr_vs_orig']:.2f} dB。结果表明，双边滤波在降低随机噪声的同时更好地保留了导线边界和芯片轮廓。")
    add_figure(document, FIGURE_DIR / "task2_bilateral_vs_gaussian.png", "图2 双边滤波与高斯滤波对 Fig0507(b) 的处理结果对比", 6.15)

    add_heading(document, "3. 三种均值滤波器对高斯噪声的去噪效果", level=2)
    add_paragraph(document, "以 Fig0507(a) 为原图，分别叠加方差为 0.001、0.005 和 0.010 的高斯噪声，再使用算术均值、加权均值和高斯均值滤波器进行去噪。实验结果说明：当噪声很弱时，过强的平滑会带来过度模糊；当噪声增强时，高斯均值和加权均值的效果明显优于算术均值。")
    add_paragraph(document, f"定量指标方面，噪声方差为 0.001 时，原始噪声图像的 PSNR 已达到 {metrics['task3']['variance_0.001']['noisy_psnr']:.2f} dB，此时继续平滑反而降低了重建质量。方差提高到 0.005 后，算术均值、加权均值和高斯均值的 PSNR 分别为 {metrics['task3']['variance_0.005']['arithmetic_psnr']:.2f} dB、{metrics['task3']['variance_0.005']['weighted_psnr']:.2f} dB 和 {metrics['task3']['variance_0.005']['gaussian_psnr']:.2f} dB；方差为 0.010 时，对应的 PSNR 分别为 {metrics['task3']['variance_0.010']['arithmetic_psnr']:.2f} dB、{metrics['task3']['variance_0.010']['weighted_psnr']:.2f} dB 和 {metrics['task3']['variance_0.010']['gaussian_psnr']:.2f} dB。")
    add_paragraph(document, f"在方差为 0.005 的中等噪声下继续比较模板大小时，算术均值、加权均值和高斯均值在 3x3 模板上的 PSNR 分别达到 {metrics['task3']['arithmetic_size_compare']['size_3']:.2f} dB、{metrics['task3']['weighted_size_compare']['size_3']:.2f} dB 和 {metrics['task3']['gaussian_size_compare']['size_3']:.2f} dB，均高于更大的模板尺寸，说明该电路板图像不适合采用过大的平滑窗口。")
    add_figure(document, FIGURE_DIR / "task3_filter_compare_var_0_005.png", "图3 方差为 0.005 时三种均值类滤波器的去噪结果对比", 6.15)

    add_heading(document, "4. 算术均值与几何均值滤波比较", level=2)
    add_paragraph(document, f"对 Fig0507(b) 采用 5x5 模板处理时，算术均值滤波的 PSNR 为 {metrics['task4']['arithmetic_psnr_vs_orig']:.2f} dB，几何均值滤波的 PSNR 为 {metrics['task4']['geometric_psnr_vs_orig']:.2f} dB。两者差距不大，但几何均值滤波在芯片边缘和细线结构处的层次更自然。")
    add_figure(document, FIGURE_DIR / "task4_arithmetic_vs_geometric.png", "图4 算术均值滤波与几何均值滤波对 Fig0507(b) 的结果对比", 6.15)

    add_heading(document, "5. 逆谐波均值滤波器的参数 Q 对比", level=2)
    add_paragraph(document, "对含椒噪声和盐噪声的电路板图像分别选取不同 Q 值进行逆谐波均值滤波。实验结果显示：椒噪声图像在正 Q 下恢复较好，盐噪声图像在负 Q 下恢复较好，且中等绝对值的 Q 更稳定。")
    add_paragraph(document, f"具体来看，Fig0508(a) 在 Q=0.5、1.5 和 3.0 下的 PSNR 分别为 {metrics['task5']['pepper']['Q_0.5']:.2f} dB、{metrics['task5']['pepper']['Q_1.5']:.2f} dB 和 {metrics['task5']['pepper']['Q_3.0']:.2f} dB，其中 Q=0.5 效果最好；Fig0508(b) 在 Q=-0.5、-1.5 和 -3.0 下的 PSNR 分别为 {metrics['task5']['salt']['Q_-0.5']:.2f} dB、{metrics['task5']['salt']['Q_-1.5']:.2f} dB 和 {metrics['task5']['salt']['Q_-3.0']:.2f} dB，其中 Q=-0.5 最稳定。")
    add_figure(document, task5_combo, "图5 逆谐波均值滤波器在椒噪声和盐噪声图像上的参数 Q 对比结果", 6.15)

    add_heading(document, "6. Fig0512(b) 上多种均值滤波器的综合比较", level=2)
    add_paragraph(document, "Fig0512(b) 同时含有均匀噪声和盐椒噪声，因此更能体现不同滤波器的适用边界。实验结果表明，算术均值滤波虽然仍存在一定模糊，但对整体噪声的抑制最稳定；几何均值和谐波均值在该混合噪声条件下容易造成明显的亮度失衡；逆谐波均值在 Q=1.5 时能压制黑色噪点，但对白色盐噪点无能为力。")
    add_figure(document, FIGURE_DIR / "task6_mixed_noise_compare.png", "图6 Fig0512(b) 在多种均值类滤波器下的处理效果对比", 6.15)

    add_heading(document, "四、实验总结与体会", level=1)
    add_paragraph(document, "通过本次实验，我系统理解了均值类滤波器在图像平滑中的共同点和差异。算术均值滤波实现最简单，但容易产生明显模糊；加权均值和高斯均值在平衡去噪与保边方面更稳定；双边滤波由于同时利用空间距离和灰度相似性，在本次实验中表现出最好的保边去噪综合效果。")
    add_paragraph(document, "逆谐波均值滤波的实验进一步说明，滤波器参数必须和噪声类型匹配。正 Q 更适合椒噪声，负 Q 更适合盐噪声，而且参数绝对值过大时容易带来亮度失真。总体来看，本次实验不仅帮助我掌握了几种常见平滑算法的原理和实现，也让我学会了结合结果图与定量指标共同分析滤波效果。")
    add_paragraph(document, "完整实验程序已单独保存在本主题目录中的 run_lab.py，可与本实验报告一起作为程序附件提交。", indent=False)

    document.save(REPORT_DOCX)


def main() -> None:
    metrics = load_metrics()
    build_markdown(metrics)
    build_docx(metrics)
    print(f"Saved markdown draft to: {REPORT_MD}")
    print(f"Saved report to: {REPORT_DOCX}")


if __name__ == "__main__":
    main()
