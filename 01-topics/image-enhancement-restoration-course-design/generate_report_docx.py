from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TOPIC_DIR = Path(__file__).resolve().parent
ROOT_DIR = TOPIC_DIR.parents[1]
TEMPLATE_PATH = Path(r"D:\qq_setup_31980\报告模板-课设3图像增强与复原算法综合应用（2024版）.docx")
OUTPUT_DIR = ROOT_DIR / "03-outputs" / "image-enhancement-restoration-course-design"
DOCX_PATH = OUTPUT_DIR / "course-design3-report-final.docx"

STUDENT_NAME = "李林浩"
STUDENT_ID = "202483290054"
DEPARTMENT = "计算机学院"
MAJOR = "计算机科学与技术"
TEACHER = "范春年"


def set_run_font(run, source_run=None, color=None):
    if source_run is not None:
        if source_run.font.name:
            run.font.name = source_run.font.name
        if source_run.font.size:
            run.font.size = source_run.font.size
        run.bold = source_run.bold
        run.italic = source_run.italic
        run.underline = source_run.underline
        r_color = source_run.font.color.rgb if source_run.font.color else None
        if color is None and r_color:
            color = str(r_color)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if run.font.name:
        r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), run.font.name)
        r_fonts.set(qn("w:hAnsi"), run.font.name)
        r_fonts.set(qn("w:eastAsia"), run.font.name)


def replace_paragraph_text(paragraph, text, color="000000"):
    source_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, source_run=source_run, color=color)


def replace_paragraph_lines(
    paragraph,
    lines,
    color="000000",
    font_name=None,
    font_size_pt=None,
    align=None,
):
    source_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align

    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, source_run=source_run, color=color)
        if font_name:
            run.font.name = font_name
            r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:eastAsia"), font_name)
        if font_size_pt:
            run.font.size = Pt(font_size_pt)
        if idx < len(lines) - 1:
            run.add_break()


def clear_paragraph_and_add_picture(paragraph, image_path, width_inches):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def insert_table_after(paragraph, headers, rows, widths_inches):
    doc = paragraph.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        set_run_font(run, color="000000")
        cell.width = Inches(widths_inches[idx])

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, color="000000")
            cell.width = Inches(widths_inches[c_idx])

    paragraph._p.addnext(table._tbl)
    return table


def font_path():
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simfang.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


def get_font(size):
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size, index=0)
    return ImageFont.load_default()


def draw_flowchart(output_path, title, labels):
    image = Image.new("RGB", (1600, 360), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    box_font = get_font(20)

    draw.text((800, 28), title, fill="black", font=title_font, anchor="ma")

    box_w = 180
    box_h = 92
    gap = 24
    total = len(labels) * box_w + (len(labels) - 1) * gap
    x = (1600 - total) // 2
    y = 145

    for i, label in enumerate(labels):
        x1 = x + i * (box_w + gap)
        y1 = y
        x2 = x1 + box_w
        y2 = y1 + box_h
        draw.rounded_rectangle([x1, y1, x2, y2], radius=14, outline="black", width=3, fill="white")
        bbox = draw.multiline_textbbox((0, 0), label, font=box_font, align="center", spacing=4)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        tx = x1 + (box_w - tw) / 2
        ty = y1 + (box_h - th) / 2 - 2
        draw.multiline_text((tx, ty), label, fill="black", font=box_font, align="center", spacing=4)
        if i < len(labels) - 1:
            ax1 = x2 + 6
            ax2 = x1 + box_w + gap - 10
            ay = y1 + box_h / 2
            draw.line((ax1, ay, ax2, ay), fill="black", width=3)
            draw.polygon([(ax2, ay), (ax2 - 12, ay - 8), (ax2 - 12, ay + 8)], fill="black")

    image.save(output_path)


def build_supporting_images():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_flowchart(
        OUTPUT_DIR / "enhancement_flowchart.png",
        "图像增强算法流程",
        [
            "读入图像\n归一化",
            "FFT\n频谱分析",
            "高斯陷波\n抑制周期噪声",
            "5x5中值\n滤波",
            "3x3均值\n滤波",
            "高提升\n锐化",
            "百分位\n拉伸",
        ],
    )
    draw_flowchart(
        OUTPUT_DIR / "restoration_flowchart.png",
        "图像复原算法流程",
        [
            "读入降质\n图像",
            "PSF\n近似建模",
            "Wiener\n复原",
            "细节\n补偿",
            "百分位\n拉伸",
            "输出复原\n结果",
        ],
    )


def fill_template():
    build_supporting_images()
    doc = Document(str(TEMPLATE_PATH))

    # Title page
    replace_paragraph_text(doc.paragraphs[13], f"学生姓名                             {STUDENT_NAME}")
    replace_paragraph_text(doc.paragraphs[15], f"学    号                             {STUDENT_ID}")
    replace_paragraph_text(doc.paragraphs[17], f"院    系        {DEPARTMENT}          ")
    replace_paragraph_text(doc.paragraphs[19], f"专    业     {MAJOR}        ")
    replace_paragraph_text(doc.paragraphs[21], f"任课教师        {TEACHER}        ")

    # Task 1
    replace_paragraph_text(doc.paragraphs[31], "本节针对题目给出的第二组带噪图像开展增强实验，研究目标是在保留目标轮廓与纹理细节的前提下，尽可能抑制混合噪声对图像质量的影响。")
    replace_paragraph_text(doc.paragraphs[32], "已知待处理图像同时受到随机噪声与周期噪声的共同干扰。随机噪声主要表现为局部灰度抖动与颗粒感，周期噪声则表现为具有固定方向和周期的规则条纹。")
    replace_paragraph_text(doc.paragraphs[33], "由于两类噪声在形成机理和空间表现上存在明显差异，若仅采用单一空间域平滑方法，往往难以同时获得较好的去噪效果和细节保持能力。")
    replace_paragraph_text(doc.paragraphs[34], "因此，本文采用频率域与空间域相结合的技术路线，首先在频率域抑制周期噪声分量，再在空间域进一步削弱随机噪声，并辅以细节补偿与对比度增强。")
    replace_paragraph_text(doc.paragraphs[35], "题目提供了对应的清晰参考图像，因此增强结果可以采用均方误差、信噪比、峰值信噪比和结构相似度等有参考指标进行客观评价。")
    replace_paragraph_text(doc.paragraphs[36], "根据学号尾数为偶数的要求，本文选取第二组 dog 图像作为实验对象。实验过程中同时保存中间处理结果，以便对各阶段的作用进行分析。")
    replace_paragraph_text(doc.paragraphs[37], "原始参考图像与待增强图像如图1所示。")
    clear_paragraph_and_add_picture(doc.paragraphs[39], OUTPUT_DIR / "enhancement_comparison.png", 5.8)
    replace_paragraph_text(doc.paragraphs[40], "图1  第二组带噪图像增强结果对比")
    replace_paragraph_text(doc.paragraphs[42], "空间域观察表明，待处理图像在整体上叠加了较明显的规则条纹，同时伴随细小颗粒状干扰，说明该图像属于典型的混合噪声退化情形。")
    replace_paragraph_text(doc.paragraphs[43], "对待处理图像进行二维傅里叶变换后，可以在频谱中心附近观察到四个关于原点对称的亮点。根据频谱峰值位置估计，其主要偏移约为 (0, ±71) 和 (±74, 0)，说明周期噪声具有较稳定的频率成分。")
    replace_paragraph_text(doc.paragraphs[44], "对于随机噪声部分，其主要表现为局部异常像素和小尺度灰度扰动。中值滤波对脉冲性干扰更为敏感，均值滤波则有利于降低残余高频波动，但单独使用均值滤波会造成边缘模糊。")
    replace_paragraph_text(doc.paragraphs[45], "据此，本文确定“频率域抑制周期噪声、空间域削弱随机噪声、后续锐化补偿细节”的总体设计思想，以兼顾去噪性能与图像清晰度。")
    replace_paragraph_text(doc.paragraphs[47], "算法设计流程如下：首先对退化图像进行频谱分析，并依据离散峰值位置构造高斯陷波滤波器；随后对频域滤波结果进行 5×5 中值滤波和 3×3 均值滤波；最后采用高提升锐化和百分位灰度拉伸恢复边缘细节与整体对比度。")
    clear_paragraph_and_add_picture(doc.paragraphs[48], OUTPUT_DIR / "enhancement_flowchart.png", 5.8)
    replace_paragraph_text(doc.paragraphs[49], "图2  图像增强算法框图")
    replace_paragraph_text(doc.paragraphs[50], "在参数设置方面，陷波中心取 (0, ±71) 和 (±74, 0)，陷波半径 D0 取 10；中值滤波窗口为 5×5，均值滤波窗口为 3×3；高提升锐化系数取 0.15；百分位拉伸区间取 0.5% 至 99.5%。")
    replace_paragraph_text(doc.paragraphs[51], "上述参数是在多次试验基础上确定的折中结果，其目的在于既保证周期条纹得到有效抑制，又避免对主体轮廓和纹理细节造成过度损伤。")
    replace_paragraph_text(doc.paragraphs[53], "依据上述算法框图，增强任务的关键 MATLAB 代码如下所示。")
    replace_paragraph_lines(
        doc.paragraphs[54],
        [
            "[freqFiltered, notchFilter] = gaussian_notch_reject(noisy, notchOffsets, notchRadius);",
            "medianFiltered = median_filter2(freqFiltered, 5);",
            "meanFiltered = mean_filter2(medianFiltered, 3);",
            "sharpened = high_boost(meanFiltered, 0.15);",
            "enhanced = percentile_stretch(sharpened, 0.5, 99.5);",
        ],
        font_name="Courier New",
        font_size_pt=8.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    replace_paragraph_lines(
        doc.paragraphs[55],
        [
            "notchOffsets = [0,-71; 0,71; -74,0; 74,0];",
            "d2 = (y - (cy + dy)).^2 + (x - (cx + dx)).^2;",
            "H = H .* (1 - exp(-d2 / (2 * D0^2)));",
            "G = F .* H;",
            "out = real(ifft2(ifftshift(G)));",
        ],
        font_name="Courier New",
        font_size_pt=8.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    replace_paragraph_text(doc.paragraphs[56], "程序运行后，脚本将输出频域陷波结果、空间域滤波结果、最终增强结果及频谱分析图，并进一步依据参考图像计算客观评价指标。")
    replace_paragraph_text(doc.paragraphs[57], "表1  图像增强结果客观评价指标对比")
    insert_table_after(
        doc.paragraphs[57],
        ["阶段", "MSE", "RMSE", "SNR/dB", "PSNR/dB", "SSIM"],
        [
            ["处理前", "0.071191", "0.266816", "4.664", "11.476", "0.1039"],
            ["处理后", "0.006541", "0.080877", "15.032", "21.844", "0.6558"],
        ],
        [1.0, 0.9, 0.95, 0.95, 1.0, 0.8],
    )
    replace_paragraph_text(doc.paragraphs[58], "由表1可知，处理后 MSE 由 0.071191 降至 0.006541，RMSE 由 0.266816 降至 0.080877，表明增强结果与参考图像之间的像素误差显著减小。与此同时，SNR 由 4.664 dB 提升至 15.032 dB，PSNR 由 11.476 dB 提升至 21.844 dB，说明噪声能量得到有效抑制。")
    replace_paragraph_text(doc.paragraphs[59], "SSIM 由 0.1039 提升至 0.6558，说明所设计算法不仅改善了灰度误差，而且在较大程度上保持并恢复了图像结构信息。该方法的优点是针对混合噪声具有较好的适应性；不足之处在于陷波中心仍需依据频谱人工估计，自动化程度有待提高。")

    # Task 2
    replace_paragraph_text(doc.paragraphs[62], "本节针对降质图像复原任务开展研究。依据学号尾数为偶数的要求，本文选择图1模糊树林图像作为实验对象。与增强任务不同，本任务未提供清晰参考图像。")
    replace_paragraph_text(doc.paragraphs[63], "因此，算法设计需要首先根据图像的视觉退化特征对退化机理进行估计，再选择适当的复原模型进行处理。")
    replace_paragraph_text(doc.paragraphs[64], "在客观评价方面，本文采用信息熵、标准差、平均梯度和拉普拉斯方差等无参考指标，从灰度层次、对比度和清晰度三个层面分析复原效果。")
    replace_paragraph_text(doc.paragraphs[66], "图3  模糊树林图像复原结果对比")
    clear_paragraph_and_add_picture(doc.paragraphs[65], OUTPUT_DIR / "restoration_fig1_comparison.png", 5.8)
    replace_paragraph_text(doc.paragraphs[67], "复原结果如图3所示。")
    replace_paragraph_text(doc.paragraphs[70], "从空间域观察可知，该图像的主要问题表现为整体边缘模糊、树木轮廓发软以及纹理细节被削弱，说明退化形式以模糊为主，而不是明显的脉冲噪声或周期干扰。")
    replace_paragraph_text(doc.paragraphs[71], "若将该退化过程表示为 g(x, y) = h(x, y) * f(x, y) + n(x, y)，其中 h(x, y) 为未知退化函数，则当前任务的关键在于对 h(x, y) 作合理近似。")
    replace_paragraph_text(doc.paragraphs[72], "结合图像视觉表现，可以将退化函数近似为高斯型点扩散函数。该模型能够较好描述成像系统失焦或轻微运动引起的低通模糊效应。")
    replace_paragraph_text(doc.paragraphs[73], "由于直接逆滤波容易在退化函数幅值较小时放大噪声，本文选择稳定性更好的维纳滤波作为基础复原方法，并在后续增加轻量细节补偿。")
    replace_paragraph_text(doc.paragraphs[75], "复原算法首先提取彩色图像的亮度分量，并利用高斯点扩散函数构造模糊模型；随后在频率域对亮度分量实施维纳复原；最后对复原结果进行亮度融合、细节补偿和百分位拉伸，以改善视觉清晰度与层次感。")
    clear_paragraph_and_add_picture(doc.paragraphs[76], OUTPUT_DIR / "restoration_flowchart.png", 5.8)
    replace_paragraph_text(doc.paragraphs[77], "图4  图像复原算法框图")
    replace_paragraph_text(doc.paragraphs[78], "参数设置如下：高斯 PSF 尺寸取 13×13，标准差取 1.35；维纳滤波参数 K 取 0.004；亮度融合权重设置为 0.72 与 0.28；细节补偿项权重取 0.28；百分位拉伸区间取 1.0% 至 99.0%。")
    replace_paragraph_text(doc.paragraphs[79], "上述参数的设置依据是，在控制噪声放大和振铃现象的同时，尽可能恢复树林边缘与枝叶纹理，使复原结果保持自然。")
    replace_paragraph_text(doc.paragraphs[81], "依据算法设计，复原任务的关键 MATLAB 代码如下所示。")
    replace_paragraph_lines(
        doc.paragraphs[82],
        [
            "psf = gaussian_kernel2(13, 1.35);",
            "deconvY = wiener_deconvolution(y, psf, 0.004);",
            "deconvY = clamp01(0.72 * deconvY + 0.28 * y);",
            "detail = deconvY - gaussian_blur2(deconvY, 1.0);",
            "restoredY = percentile_stretch(clamp01(deconvY + 0.28 * detail), 1.0, 99.0);",
        ],
        font_name="Courier New",
        font_size_pt=8.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    replace_paragraph_lines(
        doc.paragraphs[83],
        [
            "H = psf2otf_custom(psf, size(img));",
            "G = fft2(img);",
            "W = conj(H) ./ (abs(H).^2 + K);",
            "out = real(ifft2(G .* W));",
            "out = clamp01(out);",
        ],
        font_name="Courier New",
        font_size_pt=8.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    replace_paragraph_text(doc.paragraphs[84], "程序运行后，脚本输出原始降质图像、复原图像及其无参考指标对比结果，并据此对算法性能进行分析。")
    replace_paragraph_text(doc.paragraphs[85], "表2  图像复原结果无参考指标对比")
    insert_table_after(
        doc.paragraphs[85],
        ["阶段", "信息熵", "标准差", "平均梯度", "拉普拉斯方差"],
        [
            ["复原前", "7.320", "0.176", "0.03489", "0.01926"],
            ["复原后", "7.607", "0.228", "0.11841", "0.12463"],
        ],
        [1.0, 0.95, 0.95, 1.1, 1.2],
    )
    replace_paragraph_text(doc.paragraphs[86], "由表2可知，复原后图像的信息熵由 7.320 提升至 7.607，标准差由 0.176 提升至 0.228，说明图像灰度层次与整体对比度均得到改善；平均梯度由 0.03489 提升至 0.11841，拉普拉斯方差由 0.01926 提升至 0.12463，表明边缘过渡更加明显，纹理细节恢复较为充分。")
    replace_paragraph_text(doc.paragraphs[87], "综合视觉效果与无参考指标可以认为，所设计算法能够较好地恢复模糊树林图像的主体结构与局部细节。但由于退化函数采用的是近似模型，局部区域仍可能存在轻微过锐化现象，后续可结合盲去卷积方法进一步提高复原精度。")

    try:
        doc.save(str(DOCX_PATH))
        return DOCX_PATH
    except PermissionError:
        fallback_path = OUTPUT_DIR / "course-design3-report-final-v2.docx"
        doc.save(str(fallback_path))
        return fallback_path


if __name__ == "__main__":
    print(fill_template())
